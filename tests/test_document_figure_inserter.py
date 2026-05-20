"""
tests/test_document_figure_inserter.py
Tests para DOC-03 — Insercion de figuras en DOCX.
"""
import base64
import json
import sys
import tempfile
import unittest
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent / "src"))
sys.path.insert(0, str(Path(__file__).parent.parent))

from eia_agent.core.document_figure_inserter import (
    FIGURE_OUTPUT_DOCX_FILENAME,
    FIGURE_RESULT_JSON_FILENAME,
    FIGURE_RESULT_MD_FILENAME,
    SUPPORTED_IMAGE_EXTENSIONS,
    DocumentFigure,
    FigureInsertionResult,
    _ADMIN_NOTE,
    add_figures_annex_to_docx,
    build_figure_caption,
    build_figure_result_markdown,
    build_figure_title,
    detect_figure_type,
    discover_document_figures,
    insert_figures_into_document,
    validate_image_file,
    write_figure_insertion_outputs,
)

# ---------------------------------------------------------------------------
# PNG minimo sintetico (10x10 azul)
# ---------------------------------------------------------------------------

_MINIMAL_PNG_B64 = (
    "iVBORw0KGgoAAAANSUhEUgAAAAoAAAAKCAYAAACNMs+9AAAAFUlEQVR42mNk"
    "+M9Qz0AEYBxVSF+FABJADveQjN7eAAAAAElFTkSuQmCC"
)

def _minimal_png_bytes() -> bytes:
    return base64.b64decode(_MINIMAL_PNG_B64)


def _make_png(path: Path) -> Path:
    path.write_bytes(_minimal_png_bytes())
    return path


def _make_docx(path: Path, content: str = "Texto de prueba.") -> Path:
    from docx import Document
    doc = Document()
    doc.add_paragraph(content)
    doc.save(str(path))
    return path


# ---------------------------------------------------------------------------
# 1. Tests de DocumentFigure
# ---------------------------------------------------------------------------

class TestDocumentFigure(unittest.TestCase):

    def _make_fig(self, **kwargs) -> DocumentFigure:
        defaults = dict(
            figure_id="FIG-001",
            figure_type="MAPA",
            title="Mapa de situacion",
            source_path="/tmp/mapa.png",
            relative_path="cartografia/mapa.png",
            caption="Figura FIG-001. Mapa.",
            section_hint="Bloque K",
            file_size_bytes=1024,
        )
        defaults.update(kwargs)
        return DocumentFigure(**defaults)

    def test_to_dict_keys(self):
        fig = self._make_fig()
        d = fig.to_dict()
        for k in ("figure_id", "figure_type", "title", "source_path",
                  "relative_path", "caption", "section_hint",
                  "file_size_bytes", "warnings", "notes"):
            self.assertIn(k, d)

    def test_to_dict_values(self):
        fig = self._make_fig(figure_type="CLIMOGRAMA")
        d = fig.to_dict()
        self.assertEqual(d["figure_type"], "CLIMOGRAMA")

    def test_summary_contains_id(self):
        fig = self._make_fig()
        s = fig.summary()
        self.assertIn("FIG-001", s)

    def test_summary_contains_type(self):
        fig = self._make_fig(figure_type="MAPA")
        s = fig.summary()
        self.assertIn("MAPA", s)

    def test_summary_contains_title(self):
        fig = self._make_fig(title="Mapa de situacion")
        s = fig.summary()
        self.assertIn("Mapa de situacion", s)


# ---------------------------------------------------------------------------
# 2. Tests de FigureInsertionResult
# ---------------------------------------------------------------------------

class TestFigureInsertionResult(unittest.TestCase):

    def _make_result(self, **kwargs) -> FigureInsertionResult:
        defaults = dict(
            expediente_id="TEST",
            input_docx_path="test.docx",
        )
        defaults.update(kwargs)
        return FigureInsertionResult(**defaults)

    def test_found_count(self):
        from eia_agent.core.document_figure_inserter import DocumentFigure
        fig = DocumentFigure("FIG-001","MAPA","T","/p","p","C","H",100)
        r = self._make_result(figures_found=[fig, fig])
        self.assertEqual(r.found_count(), 2)

    def test_inserted_count(self):
        r = self._make_result(figures_inserted=["FIG-001", "FIG-002"])
        self.assertEqual(r.inserted_count(), 2)

    def test_skipped_count(self):
        r = self._make_result(figures_skipped=["FIG-003"])
        self.assertEqual(r.skipped_count(), 1)

    def test_warning_count(self):
        r = self._make_result(warnings=["w1", "w2"])
        self.assertEqual(r.warning_count(), 2)

    def test_is_success_false_when_not_generated(self):
        r = self._make_result(generated=False)
        self.assertFalse(r.is_success())

    def test_is_success_false_without_output_path(self):
        r = self._make_result(generated=True, output_docx_path=None)
        self.assertFalse(r.is_success())

    def test_is_success_true_with_generated_and_path(self):
        r = self._make_result(generated=True, output_docx_path="out.docx")
        self.assertTrue(r.is_success())

    def test_to_dict_keys(self):
        r = self._make_result()
        d = r.to_dict()
        for k in ("expediente_id", "generated", "found_count",
                  "inserted_count", "skipped_count", "figures_found",
                  "figures_inserted", "figures_skipped", "warnings"):
            self.assertIn(k, d)

    def test_summary_contains_expediente(self):
        r = self._make_result(expediente_id="EXP-001")
        self.assertIn("EXP-001", r.summary())

    def test_summary_dry_run(self):
        r = self._make_result(generated=False)
        self.assertIn("DRY-RUN", r.summary())

    def test_summary_ok_when_success(self):
        r = self._make_result(generated=True, output_docx_path="out.docx")
        self.assertIn("OK", r.summary())


# ---------------------------------------------------------------------------
# 3. Tests de detect_figure_type
# ---------------------------------------------------------------------------

class TestDetectFigureType(unittest.TestCase):

    def test_mapa_situacion(self):
        self.assertEqual(detect_figure_type("mapa_situacion.png"), "MAPA")

    def test_mapa_cartografia(self):
        self.assertEqual(detect_figure_type("cartografia/mapa_base.png"), "MAPA")

    def test_climograma_aemet(self):
        self.assertEqual(detect_figure_type("climograma_aemet_lanzarote.png"), "CLIMOGRAMA")

    def test_clima_koppen(self):
        self.assertEqual(detect_figure_type("clasificacion_koppen.png"), "CLIMOGRAMA")

    def test_fotografia_nave(self):
        self.assertEqual(detect_figure_type("foto_nave.jpg"), "FOTOGRAFIA")

    def test_logo(self):
        self.assertEqual(detect_figure_type("logo_ecogestion.png"), "LOGO")

    def test_brand_logo(self):
        self.assertEqual(detect_figure_type("assets/brand/logo.png"), "LOGO")

    def test_grafico(self):
        self.assertEqual(detect_figure_type("grafico_resultado.png"), "GRAFICO")

    def test_otro(self):
        t = detect_figure_type("archivo.png")
        self.assertIn(t, ["OTRO", "MAPA", "CLIMOGRAMA", "FOTOGRAFIA", "GRAFICO", "LOGO"])

    def test_inundabilidad(self):
        self.assertEqual(detect_figure_type("MAP006_inundabilidad.png"), "MAPA")

    def test_red_natura(self):
        self.assertEqual(detect_figure_type("map004_red_natura.png"), "MAPA")


# ---------------------------------------------------------------------------
# 4. Tests de build_figure_title y build_figure_caption
# ---------------------------------------------------------------------------

class TestBuildFigureTitleCaption(unittest.TestCase):

    def test_title_from_underscore(self):
        title = build_figure_title("mapa_situacion.png", "MAPA")
        self.assertIn("mapa", title.lower())
        self.assertIn("situacion", title.lower())

    def test_title_not_empty(self):
        title = build_figure_title("x.png", "MAPA")
        self.assertIsInstance(title, str)
        self.assertGreater(len(title), 0)

    def test_title_climograma(self):
        title = build_figure_title("climograma_AEMET_Lanzarote.png", "CLIMOGRAMA")
        self.assertIn("aemet", title.lower())

    def test_caption_contains_fig_id(self):
        caption = build_figure_caption("FIG-001", "Titulo de prueba", "MAPA")
        self.assertIn("FIG-001", caption)

    def test_caption_contains_type(self):
        caption = build_figure_caption("FIG-002", "Titulo", "CLIMOGRAMA")
        self.assertIn("CLIMOGRAMA", caption)

    def test_caption_contains_title(self):
        caption = build_figure_caption("FIG-003", "Mi Titulo Especial", "OTRO")
        self.assertIn("Mi Titulo Especial", caption)

    def test_caption_contains_fuente(self):
        caption = build_figure_caption("FIG-001", "T", "MAPA")
        self.assertIn("expediente", caption.lower())


# ---------------------------------------------------------------------------
# 5. Tests de validate_image_file
# ---------------------------------------------------------------------------

class TestValidateImageFile(unittest.TestCase):

    def test_nonexistent_returns_false(self):
        self.assertFalse(validate_image_file("/no/existe/imagen.png"))

    def test_unsupported_extension_returns_false(self):
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / "archivo.txt"
            p.write_text("text")
            self.assertFalse(validate_image_file(p))

    def test_empty_file_returns_false(self):
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / "empty.png"
            p.write_bytes(b"")
            self.assertFalse(validate_image_file(p))

    def test_valid_png_returns_true(self):
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / "test.png"
            p.write_bytes(_minimal_png_bytes())
            self.assertTrue(validate_image_file(p))

    def test_valid_jpg_returns_true(self):
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / "test.jpg"
            p.write_bytes(_minimal_png_bytes())  # jpg extension, png content ok for validation
            self.assertTrue(validate_image_file(p))

    def test_docx_extension_returns_false(self):
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / "archivo.docx"
            p.write_bytes(b"content")
            self.assertFalse(validate_image_file(p))


# ---------------------------------------------------------------------------
# 6. Tests de discover_document_figures
# ---------------------------------------------------------------------------

class TestDiscoverDocumentFigures(unittest.TestCase):

    def setUp(self):
        self.tmp_obj = tempfile.TemporaryDirectory()
        self.tmp = Path(self.tmp_obj.name)
        self.exp = self.tmp / "expediente-test"
        self.exp.mkdir()

    def tearDown(self):
        self.tmp_obj.cleanup()

    def test_empty_expediente_returns_empty(self):
        result = discover_document_figures(self.exp)
        self.assertEqual(result, [])

    def test_finds_png_in_cartografia_mapas(self):
        mapas = self.exp / "cartografia" / "mapas"
        mapas.mkdir(parents=True)
        _make_png(mapas / "mapa_situacion.png")
        figs = discover_document_figures(self.exp)
        self.assertEqual(len(figs), 1)
        self.assertEqual(figs[0].figure_type, "MAPA")

    def test_finds_climograma_in_clima(self):
        clima = self.exp / "clima"
        clima.mkdir()
        _make_png(clima / "climograma_lanzarote.png")
        figs = discover_document_figures(self.exp)
        self.assertEqual(len(figs), 1)
        self.assertEqual(figs[0].figure_type, "CLIMOGRAMA")

    def test_finds_foto_in_inputs_fotos(self):
        fotos = self.exp / "inputs" / "fotos"
        fotos.mkdir(parents=True)
        _make_png(fotos / "foto_nave_exterior.png")
        figs = discover_document_figures(self.exp)
        self.assertEqual(len(figs), 1)
        self.assertEqual(figs[0].figure_type, "FOTOGRAFIA")

    def test_ignores_txt_files(self):
        cart = self.exp / "cartografia"
        cart.mkdir()
        (cart / "readme.txt").write_text("info")
        figs = discover_document_figures(self.exp)
        self.assertEqual(len(figs), 0)

    def test_ignores_empty_png(self):
        cart = self.exp / "cartografia"
        cart.mkdir()
        (cart / "empty.png").write_bytes(b"")
        figs = discover_document_figures(self.exp)
        self.assertEqual(len(figs), 0)

    def test_assigns_correlative_ids(self):
        cart = self.exp / "cartografia"
        cart.mkdir()
        for i in range(3):
            _make_png(cart / f"mapa_{i:02d}.png")
        figs = discover_document_figures(self.exp)
        ids = [f.figure_id for f in figs]
        self.assertIn("FIG-001", ids)
        self.assertIn("FIG-002", ids)
        self.assertIn("FIG-003", ids)

    def test_mapa_before_climograma(self):
        cart = self.exp / "cartografia"
        cart.mkdir()
        _make_png(cart / "mapa_base.png")
        clima = self.exp / "clima"
        clima.mkdir()
        _make_png(clima / "climograma.png")
        figs = discover_document_figures(self.exp)
        tipos = [f.figure_type for f in figs]
        self.assertEqual(tipos[0], "MAPA")
        self.assertEqual(tipos[1], "CLIMOGRAMA")

    def test_no_duplicates(self):
        cart = self.exp / "cartografia"
        cart.mkdir()
        _make_png(cart / "mapa.png")
        figs = discover_document_figures(self.exp)
        ids = [f.figure_id for f in figs]
        self.assertEqual(len(ids), len(set(ids)))

    def test_relative_path_set(self):
        cart = self.exp / "cartografia"
        cart.mkdir()
        _make_png(cart / "mapa_test.png")
        figs = discover_document_figures(self.exp)
        self.assertIn("cartografia", figs[0].relative_path)

    def test_file_size_set(self):
        cart = self.exp / "cartografia"
        cart.mkdir()
        _make_png(cart / "mapa_test.png")
        figs = discover_document_figures(self.exp)
        self.assertGreater(figs[0].file_size_bytes, 0)


# ---------------------------------------------------------------------------
# 7. Tests de add_figures_annex_to_docx
# ---------------------------------------------------------------------------

class TestAddFiguresAnnexToDocx(unittest.TestCase):

    def setUp(self):
        self.tmp_obj = tempfile.TemporaryDirectory()
        self.tmp = Path(self.tmp_obj.name)

    def tearDown(self):
        self.tmp_obj.cleanup()

    def _make_figure(self, path: Path) -> DocumentFigure:
        return DocumentFigure(
            figure_id="FIG-001",
            figure_type="MAPA",
            title="Mapa de prueba",
            source_path=str(path),
            relative_path="cartografia/mapa.png",
            caption="Figura FIG-001. Mapa de prueba. Tipo: MAPA. Fuente: expediente.",
            section_hint="Bloque K",
            file_size_bytes=path.stat().st_size if path.exists() else 0,
        )

    def test_generates_output_docx(self):
        in_docx = self.tmp / "input.docx"
        out_docx = self.tmp / "output.docx"
        png = _make_png(self.tmp / "mapa.png")
        _make_docx(in_docx)
        fig = self._make_figure(png)
        result = add_figures_annex_to_docx(in_docx, out_docx, [fig])
        self.assertTrue(out_docx.exists())
        self.assertGreater(out_docx.stat().st_size, 0)

    def test_output_opens_with_docx(self):
        from docx import Document
        in_docx = self.tmp / "input.docx"
        out_docx = self.tmp / "output.docx"
        png = _make_png(self.tmp / "mapa.png")
        _make_docx(in_docx)
        fig = self._make_figure(png)
        add_figures_annex_to_docx(in_docx, out_docx, [fig])
        doc = Document(str(out_docx))
        self.assertIsNotNone(doc)

    def test_contains_annex_heading(self):
        from docx import Document
        in_docx = self.tmp / "input.docx"
        out_docx = self.tmp / "output.docx"
        png = _make_png(self.tmp / "mapa.png")
        _make_docx(in_docx)
        fig = self._make_figure(png)
        add_figures_annex_to_docx(in_docx, out_docx, [fig])
        doc = Document(str(out_docx))
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
        self.assertTrue(
            any("nexo" in h for h in headings),
            f"No se encontro heading de anexo. Headings: {headings}"
        )

    def test_contains_caption(self):
        from docx import Document
        in_docx = self.tmp / "input.docx"
        out_docx = self.tmp / "output.docx"
        png = _make_png(self.tmp / "mapa.png")
        _make_docx(in_docx)
        fig = self._make_figure(png)
        add_figures_annex_to_docx(in_docx, out_docx, [fig])
        doc = Document(str(out_docx))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        self.assertIn("FIG-001", all_text)

    def test_does_not_modify_input_docx(self):
        from docx import Document
        in_docx = self.tmp / "input.docx"
        out_docx = self.tmp / "output.docx"
        png = _make_png(self.tmp / "mapa.png")
        _make_docx(in_docx, "Contenido original unico.")
        original_size = in_docx.stat().st_size
        fig = self._make_figure(png)
        add_figures_annex_to_docx(in_docx, out_docx, [fig])
        self.assertEqual(in_docx.stat().st_size, original_size)

    def test_no_figures_generates_docx_with_notice(self):
        from docx import Document
        in_docx = self.tmp / "input.docx"
        out_docx = self.tmp / "output.docx"
        _make_docx(in_docx)
        add_figures_annex_to_docx(in_docx, out_docx, [])
        doc = Document(str(out_docx))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        self.assertTrue(
            "no se localizaron" in all_text.lower() or "nexo" in all_text.lower(),
            f"No se encontro aviso de no figuras. Texto: {all_text[:200]}"
        )

    def test_invalid_image_skipped_no_crash(self):
        in_docx = self.tmp / "input.docx"
        out_docx = self.tmp / "output.docx"
        _make_docx(in_docx)
        bad_fig = DocumentFigure(
            figure_id="FIG-BAD",
            figure_type="MAPA",
            title="Imagen invalida",
            source_path="/ruta/no/existe/imagen.png",
            relative_path="cartografia/no_existe.png",
            caption="Caption BAD",
            section_hint="K",
            file_size_bytes=0,
        )
        result = add_figures_annex_to_docx(in_docx, out_docx, [bad_fig])
        self.assertTrue(out_docx.exists())
        self.assertIn("FIG-BAD", result.figures_skipped)

    def test_inserted_count_correct(self):
        in_docx = self.tmp / "input.docx"
        out_docx = self.tmp / "output.docx"
        png = _make_png(self.tmp / "mapa.png")
        _make_docx(in_docx)
        fig = self._make_figure(png)
        result = add_figures_annex_to_docx(in_docx, out_docx, [fig])
        self.assertEqual(result.inserted_count(), 1)


# ---------------------------------------------------------------------------
# 8. Tests de insert_figures_into_document
# ---------------------------------------------------------------------------

class TestInsertFiguresIntoDocument(unittest.TestCase):

    def setUp(self):
        self.tmp_obj = tempfile.TemporaryDirectory()
        self.tmp = Path(self.tmp_obj.name)

    def tearDown(self):
        self.tmp_obj.cleanup()

    def _make_expediente(self, with_docx: bool = True, with_png: bool = False) -> Path:
        exp = self.tmp / "expediente-qa"
        exp.mkdir()
        doc_dir = exp / "documento"
        doc_dir.mkdir()
        if with_docx:
            _make_docx(doc_dir / "documento_ambiental_borrador.docx", "Contenido base.")
        if with_png:
            cart = exp / "cartografia"
            cart.mkdir()
            _make_png(cart / "mapa_situacion.png")
        return exp

    def test_missing_docx_raises_file_not_found(self):
        exp = self._make_expediente(with_docx=False)
        with self.assertRaises(FileNotFoundError):
            insert_figures_into_document(exp, write_outputs=False)

    def test_without_write_generated_false(self):
        exp = self._make_expediente()
        result = insert_figures_into_document(exp, write_outputs=False)
        self.assertFalse(result.generated)

    def test_without_write_no_enriched_docx(self):
        exp = self._make_expediente()
        insert_figures_into_document(exp, write_outputs=False)
        out = exp / "documento" / FIGURE_OUTPUT_DOCX_FILENAME
        self.assertFalse(out.exists())

    def test_with_write_generates_enriched_docx(self):
        exp = self._make_expediente()
        result = insert_figures_into_document(exp, write_outputs=True)
        out = exp / "documento" / FIGURE_OUTPUT_DOCX_FILENAME
        self.assertTrue(out.exists())
        self.assertTrue(result.generated)

    def test_with_write_generates_json(self):
        exp = self._make_expediente()
        insert_figures_into_document(exp, write_outputs=True)
        json_path = exp / "documento" / FIGURE_RESULT_JSON_FILENAME
        self.assertTrue(json_path.exists())

    def test_with_write_generates_md(self):
        exp = self._make_expediente()
        insert_figures_into_document(exp, write_outputs=True)
        md_path = exp / "documento" / FIGURE_RESULT_MD_FILENAME
        self.assertTrue(md_path.exists())

    def test_json_loadable(self):
        exp = self._make_expediente()
        insert_figures_into_document(exp, write_outputs=True)
        json_path = exp / "documento" / FIGURE_RESULT_JSON_FILENAME
        data = json.loads(json_path.read_text(encoding="utf-8"))
        self.assertIn("expediente_id", data)
        self.assertIn("generated", data)

    def test_with_png_found_figures_populated(self):
        exp = self._make_expediente(with_png=True)
        result = insert_figures_into_document(exp, write_outputs=False)
        self.assertGreater(result.found_count(), 0)

    def test_with_write_and_png_is_success(self):
        exp = self._make_expediente(with_png=True)
        result = insert_figures_into_document(exp, write_outputs=True)
        self.assertTrue(result.is_success())

    def test_invalid_image_in_directory_skipped(self):
        exp = self._make_expediente(with_docx=True)
        cart = exp / "cartografia"
        cart.mkdir()
        (cart / "corrupta.png").write_bytes(b"not a real image but valid name")
        # write_bytes makes it non-empty so validate_image_file returns True (basic check)
        # but docx insertion might fail - test that result still generates
        result = insert_figures_into_document(exp, write_outputs=True)
        self.assertTrue(result.generated)  # should not crash

    def test_no_figures_still_generates_docx(self):
        exp = self._make_expediente(with_docx=True, with_png=False)
        result = insert_figures_into_document(exp, write_outputs=True)
        out = exp / "documento" / FIGURE_OUTPUT_DOCX_FILENAME
        self.assertTrue(out.exists())
        self.assertTrue(result.generated)


# ---------------------------------------------------------------------------
# 9. Tests de build_figure_result_markdown
# ---------------------------------------------------------------------------

class TestBuildFigureResultMarkdown(unittest.TestCase):

    def _make_result(self) -> FigureInsertionResult:
        fig = DocumentFigure(
            "FIG-001", "MAPA", "Mapa de test", "/p", "c/m.png",
            "Caption", "K", 512,
        )
        return FigureInsertionResult(
            expediente_id="TEST-EXP",
            input_docx_path="test.docx",
            output_docx_path="out.docx",
            figures_found=[fig],
            figures_inserted=["FIG-001"],
            figures_skipped=[],
            generated=True,
        )

    def test_contains_titulo(self):
        md = build_figure_result_markdown(self._make_result())
        self.assertIn("insercion de figuras", md.lower())

    def test_contains_resumen(self):
        md = build_figure_result_markdown(self._make_result())
        self.assertIn("Resumen", md)

    def test_contains_figuras_encontradas(self):
        md = build_figure_result_markdown(self._make_result())
        self.assertIn("FIG-001", md)

    def test_contains_admin_note(self):
        md = build_figure_result_markdown(self._make_result())
        self.assertIn("aptitud administrativa", md.lower())

    def test_contains_no_aptitud_declaration(self):
        md = build_figure_result_markdown(self._make_result())
        self.assertIn("no declara aptitud administrativa", md.lower())

    def test_no_figures_still_generates_md(self):
        r = FigureInsertionResult(
            expediente_id="EMPTY",
            input_docx_path="test.docx",
        )
        md = build_figure_result_markdown(r)
        self.assertIsInstance(md, str)
        self.assertGreater(len(md), 0)

    def test_contains_advertencias_section(self):
        md = build_figure_result_markdown(self._make_result())
        self.assertIn("Advertencias", md)


# ---------------------------------------------------------------------------
# 10. Tests CLI document-insert-figures
# ---------------------------------------------------------------------------

class TestCLIDocumentInsertFigures(unittest.TestCase):

    def setUp(self):
        self.tmp_obj = tempfile.TemporaryDirectory()
        self.tmp = Path(self.tmp_obj.name)

    def tearDown(self):
        self.tmp_obj.cleanup()

    def _make_expediente(self, with_png: bool = False) -> Path:
        exp = self.tmp / "expediente-cli"
        exp.mkdir()
        doc_dir = exp / "documento"
        doc_dir.mkdir()
        _make_docx(doc_dir / "documento_ambiental_borrador.docx")
        if with_png:
            cart = exp / "cartografia"
            cart.mkdir()
            _make_png(cart / "mapa_situacion.png")
        return exp

    def _run_cli(self, args: list[str]) -> int:
        from run_expediente import main
        return main(args)

    def test_without_write_no_enriched_docx(self):
        exp = self._make_expediente()
        self._run_cli([str(exp), "document-insert-figures"])
        out = exp / "documento" / FIGURE_OUTPUT_DOCX_FILENAME
        self.assertFalse(out.exists())

    def test_with_write_creates_docx(self):
        exp = self._make_expediente()
        self._run_cli([str(exp), "document-insert-figures", "--write"])
        out = exp / "documento" / FIGURE_OUTPUT_DOCX_FILENAME
        self.assertTrue(out.exists())

    def test_with_write_creates_json(self):
        exp = self._make_expediente()
        self._run_cli([str(exp), "document-insert-figures", "--write"])
        json_path = exp / "documento" / FIGURE_RESULT_JSON_FILENAME
        self.assertTrue(json_path.exists())

    def test_missing_docx_exit_1(self):
        exp = self.tmp / "exp-sin-docx"
        exp.mkdir()
        (exp / "documento").mkdir()
        rc = self._run_cli([str(exp), "document-insert-figures"])
        self.assertEqual(rc, 1)

    def test_valid_docx_without_write_exit_0(self):
        exp = self._make_expediente()
        rc = self._run_cli([str(exp), "document-insert-figures"])
        self.assertEqual(rc, 0)

    def test_valid_docx_with_png_and_write_exit_0(self):
        exp = self._make_expediente(with_png=True)
        rc = self._run_cli([str(exp), "document-insert-figures", "--write"])
        self.assertEqual(rc, 0)

    def test_invalid_expediente_exit_1(self):
        rc = self._run_cli(["/ruta/no/existe", "document-insert-figures"])
        self.assertEqual(rc, 1)


# ---------------------------------------------------------------------------
# 11. Tests de write_figure_insertion_outputs
# ---------------------------------------------------------------------------

class TestWriteFigureInsertionOutputs(unittest.TestCase):

    def test_writes_json_and_md(self):
        with tempfile.TemporaryDirectory() as tmp:
            r = FigureInsertionResult(
                expediente_id="TEST",
                input_docx_path="test.docx",
                generated=True,
            )
            json_p, md_p = write_figure_insertion_outputs(r, tmp)
            self.assertTrue(json_p.exists())
            self.assertTrue(md_p.exists())

    def test_json_valid(self):
        with tempfile.TemporaryDirectory() as tmp:
            r = FigureInsertionResult(
                expediente_id="TEST",
                input_docx_path="test.docx",
            )
            json_p, _ = write_figure_insertion_outputs(r, tmp)
            data = json.loads(json_p.read_text(encoding="utf-8"))
            self.assertIn("expediente_id", data)


# ---------------------------------------------------------------------------
# 12. Tests de constantes
# ---------------------------------------------------------------------------

class TestConstants(unittest.TestCase):

    def test_supported_extensions(self):
        self.assertIn(".png", SUPPORTED_IMAGE_EXTENSIONS)
        self.assertIn(".jpg", SUPPORTED_IMAGE_EXTENSIONS)
        self.assertIn(".jpeg", SUPPORTED_IMAGE_EXTENSIONS)

    def test_output_filenames(self):
        self.assertEqual(FIGURE_OUTPUT_DOCX_FILENAME, "documento_ambiental_borrador_con_figuras.docx")
        self.assertEqual(FIGURE_RESULT_JSON_FILENAME, "document_figures_result.json")
        self.assertEqual(FIGURE_RESULT_MD_FILENAME, "document_figures_result.md")

    def test_admin_note_no_aptitud(self):
        self.assertIn("aptitud administrativa", _ADMIN_NOTE.lower())


if __name__ == "__main__":
    unittest.main()
