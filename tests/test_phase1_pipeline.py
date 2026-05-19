"""Tests para phase1_pipeline -- IN-06."""
from __future__ import annotations

import json
import sys
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch, MagicMock

sys.path.insert(0, str(Path(__file__).parent.parent / "src"))

from eia_agent.core.phase1_pipeline import (
    Phase1Result,
    detect_phase1_basic_conflicts,
    merge_candidate_facts,
    run_phase1,
)

# ---------------------------------------------------------------------------
# Rutas de pilots
# ---------------------------------------------------------------------------

_ROOT = Path(__file__).parent.parent
_PARCELA = _ROOT / "expediente-EIA-2026-RECIMETAL-PARCELA"
_NAVE = _ROOT / "expediente-EIA-2026-RECIMETAL-NAVE-222"


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _make_fact(**kwargs) -> dict:
    """Devuelve un hecho candidato con defaults razonables."""
    defaults = {
        "id": "FACT-001",
        "categoria": "identificacion",
        "campo": "nombre_promotor",
        "valor": "EMPRESA TEST SL",
        "estado": "DECLARADO",
        "fuentes": ["doc1.docx"],
        "entity_type": "PROMOTOR",
        "confidence": 0.9,
        "context": "texto de contexto",
        "normalized_value": None,
        "notes": [],
    }
    defaults.update(kwargs)
    return defaults


def _make_classification_result(facts=None):
    """Crea un ClassificationResult mínimo con lista de facts."""
    cr = MagicMock()
    cr.facts = facts or []
    cr.warnings = []
    return cr


def _make_candidate_fact_obj(
    campo="nombre_promotor",
    valor="EMPRESA TEST SL",
    fuentes=None,
    entity_type="PROMOTOR",
    confidence=0.9,
):
    """Crea un objeto con atributos compatibles con _fact_to_dict."""
    f = MagicMock()
    f.id = "FACT-001"
    f.categoria = "identificacion"
    f.campo = campo
    f.valor = valor
    f.estado = "DECLARADO"
    f.fuentes = set(fuentes or ["doc1.docx"])
    f.entity_type = entity_type
    f.confidence = confidence
    f.context = "ctx"
    f.normalized_value = None
    f.notes = []
    return f


# ---------------------------------------------------------------------------
# Clase 1: Phase1Result — estructura y métodos
# ---------------------------------------------------------------------------

class TestPhase1ResultSummary(unittest.TestCase):
    def _make_result(self, **kwargs):
        defaults = dict(
            expediente_id="EXP-TEST",
            inputs_index={},
            candidate_facts=[],
            documents_processed=0,
            docx_processed=0,
            pdf_pending=0,
            warnings=[],
        )
        defaults.update(kwargs)
        return Phase1Result(**defaults)

    def test_summary_contiene_expediente_id(self):
        r = self._make_result(expediente_id="EXP-DEMO")
        self.assertIn("EXP-DEMO", r.summary())

    def test_summary_muestra_documentos(self):
        r = self._make_result(documents_processed=5, docx_processed=3, pdf_pending=2)
        s = r.summary()
        self.assertIn("5", s)
        self.assertIn("3", s)
        self.assertIn("2", s)

    def test_summary_muestra_hechos_candidatos(self):
        r = self._make_result(candidate_facts=[_make_fact(), _make_fact(id="F2")])
        self.assertIn("2", r.summary())

    def test_summary_sin_avisos_no_muestra_seccion_avisos(self):
        r = self._make_result(warnings=[])
        self.assertNotIn("Avisos", r.summary())

    def test_summary_muestra_avisos(self):
        r = self._make_result(warnings=["aviso uno", "aviso dos"])
        s = r.summary()
        self.assertIn("Avisos", s)
        self.assertIn("aviso uno", s)

    def test_summary_trunca_avisos_a_5(self):
        r = self._make_result(warnings=[f"aviso {i}" for i in range(10)])
        s = r.summary()
        self.assertIn("5 aviso(s) más", s)

    def test_to_dict_incluye_todos_los_campos(self):
        r = self._make_result(expediente_id="EXP-X", candidate_facts=[_make_fact()])
        d = r.to_dict()
        for key in ("expediente_id", "inputs_index", "candidate_facts",
                    "documents_processed", "docx_processed", "pdf_pending",
                    "warnings", "notes"):
            self.assertIn(key, d)

    def test_to_dict_es_json_serializable(self):
        r = self._make_result(candidate_facts=[_make_fact()])
        dumped = json.dumps(r.to_dict())
        self.assertIsInstance(dumped, str)


# ---------------------------------------------------------------------------
# Clase 2: merge_candidate_facts
# ---------------------------------------------------------------------------

class TestMergeCandidateFacts(unittest.TestCase):
    def test_lista_vacia_devuelve_lista_vacia(self):
        self.assertEqual(merge_candidate_facts([]), [])

    def test_un_resultado_sin_facts(self):
        cr = _make_classification_result(facts=[])
        self.assertEqual(merge_candidate_facts([cr]), [])

    def test_concatena_facts_de_un_resultado(self):
        facts = [_make_candidate_fact_obj(campo="nombre_promotor")]
        cr = _make_classification_result(facts=facts)
        merged = merge_candidate_facts([cr])
        self.assertEqual(len(merged), 1)
        self.assertEqual(merged[0]["campo"], "nombre_promotor")

    def test_concatena_facts_de_multiples_resultados(self):
        cr1 = _make_classification_result(facts=[_make_candidate_fact_obj(campo="titular")])
        cr2 = _make_classification_result(facts=[
            _make_candidate_fact_obj(campo="capacidad"),
            _make_candidate_fact_obj(campo="referencia_catastral"),
        ])
        merged = merge_candidate_facts([cr1, cr2])
        self.assertEqual(len(merged), 3)

    def test_no_deduplica_facts_iguales(self):
        f = _make_candidate_fact_obj(campo="titular", valor="EMPRESA A")
        cr = _make_classification_result(facts=[f, f])
        merged = merge_candidate_facts([cr])
        self.assertEqual(len(merged), 2)

    def test_valor_none_serializado_como_none(self):
        f = _make_candidate_fact_obj(campo="titular", valor=None)
        cr = _make_classification_result(facts=[f])
        merged = merge_candidate_facts([cr])
        self.assertIsNone(merged[0]["valor"])

    def test_valor_numerico_serializado_como_string(self):
        f = _make_candidate_fact_obj(campo="capacidad", valor=12345)
        cr = _make_classification_result(facts=[f])
        merged = merge_candidate_facts([cr])
        self.assertEqual(merged[0]["valor"], "12345")


# ---------------------------------------------------------------------------
# Clase 3: detect_phase1_basic_conflicts
# ---------------------------------------------------------------------------

class TestDetectPhase1BasicConflicts(unittest.TestCase):
    def test_lista_vacia_no_conflictos(self):
        self.assertEqual(detect_phase1_basic_conflicts([]), [])

    def test_un_hecho_por_campo_no_conflicto(self):
        facts = [_make_fact(campo="referencia_catastral", valor="ABC123")]
        self.assertEqual(detect_phase1_basic_conflicts(facts), [])

    def test_dos_hechos_mismo_valor_no_conflicto(self):
        facts = [
            _make_fact(id="F1", campo="nombre_promotor", valor="EMPRESA A SL"),
            _make_fact(id="F2", campo="nombre_promotor", valor="EMPRESA A SL"),
        ]
        self.assertEqual(detect_phase1_basic_conflicts(facts), [])

    def test_rc_distinta_genera_conflicto(self):
        facts = [
            _make_fact(id="F1", campo="referencia_catastral", valor="RC001", fuentes=["a.docx"]),
            _make_fact(id="F2", campo="referencia_catastral", valor="RC002", fuentes=["b.docx"]),
        ]
        conflicts = detect_phase1_basic_conflicts(facts)
        self.assertEqual(len(conflicts), 1)
        self.assertEqual(conflicts[0]["campo"], "referencia_catastral")
        self.assertEqual(conflicts[0]["tipo"], "valor_multiple")

    def test_promotores_distintos_genera_conflicto(self):
        facts = [
            _make_fact(id="F1", campo="nombre_promotor", valor="EMPRESA A", fuentes=["a.docx"]),
            _make_fact(id="F2", campo="nombre_promotor", valor="EMPRESA B", fuentes=["b.docx"]),
        ]
        conflicts = detect_phase1_basic_conflicts(facts)
        self.assertEqual(len(conflicts), 1)
        self.assertEqual(conflicts[0]["campo"], "nombre_promotor")

    def test_conflicto_incluye_todas_las_fuentes(self):
        facts = [
            _make_fact(id="F1", campo="titular", valor="A", fuentes=["doc1.docx"]),
            _make_fact(id="F2", campo="titular", valor="B", fuentes=["doc2.docx"]),
        ]
        conflicts = detect_phase1_basic_conflicts(facts)
        self.assertIn("doc1.docx", conflicts[0]["fuentes"])
        self.assertIn("doc2.docx", conflicts[0]["fuentes"])

    def test_conflicto_incluye_valores_distintos(self):
        facts = [
            _make_fact(id="F1", campo="capacidad", valor="100 t/año"),
            _make_fact(id="F2", campo="capacidad", valor="200 t/año"),
        ]
        conflicts = detect_phase1_basic_conflicts(facts)
        valores = conflicts[0]["valores"]
        self.assertIn("100 T/AÑO", valores)
        self.assertIn("200 T/AÑO", valores)

    def test_conflicto_n_hechos_correcto(self):
        facts = [
            _make_fact(id="F1", campo="superficie_parcela", valor="500"),
            _make_fact(id="F2", campo="superficie_parcela", valor="600"),
            _make_fact(id="F3", campo="superficie_parcela", valor="700"),
        ]
        conflicts = detect_phase1_basic_conflicts(facts)
        self.assertEqual(conflicts[0]["n_hechos"], 3)

    def test_campos_no_conflictivos_ignorados(self):
        facts = [
            _make_fact(id="F1", campo="municipio", valor="Las Palmas"),
            _make_fact(id="F2", campo="municipio", valor="Telde"),
        ]
        self.assertEqual(detect_phase1_basic_conflicts(facts), [])

    def test_todos_los_campos_de_superficie_monitoreados(self):
        campos = [
            "superficie_catastral", "superficie_construida",
            "superficie_util", "superficie_nave", "superficie_no_clasificada",
        ]
        for campo in campos:
            facts = [
                _make_fact(id="F1", campo=campo, valor="100"),
                _make_fact(id="F2", campo=campo, valor="200"),
            ]
            conflicts = detect_phase1_basic_conflicts(facts)
            self.assertEqual(len(conflicts), 1, f"Fallo en campo: {campo}")

    def test_hechos_sin_valor_ignorados_en_conflicto(self):
        facts = [
            _make_fact(id="F1", campo="capacidad", valor=None),
            _make_fact(id="F2", campo="capacidad", valor=None),
        ]
        self.assertEqual(detect_phase1_basic_conflicts(facts), [])

    def test_comparacion_case_insensitive(self):
        facts = [
            _make_fact(id="F1", campo="nombre_promotor", valor="empresa a sl"),
            _make_fact(id="F2", campo="nombre_promotor", valor="EMPRESA A SL"),
        ]
        self.assertEqual(detect_phase1_basic_conflicts(facts), [])

    def test_fuentes_ordenadas_en_conflicto(self):
        facts = [
            _make_fact(id="F1", campo="titular", valor="X", fuentes=["z.docx", "a.docx"]),
            _make_fact(id="F2", campo="titular", valor="Y", fuentes=["m.docx"]),
        ]
        conflicts = detect_phase1_basic_conflicts(facts)
        fuentes = conflicts[0]["fuentes"]
        self.assertEqual(fuentes, sorted(fuentes))


# ---------------------------------------------------------------------------
# Clase 4: run_phase1 — expediente vacío
# ---------------------------------------------------------------------------

class TestRunPhase1ExpedienteVacio(unittest.TestCase):
    def test_directorio_sin_inputs_devuelve_resultado_vacio(self):
        with tempfile.TemporaryDirectory() as tmp:
            result = run_phase1(tmp)
        self.assertIsInstance(result, Phase1Result)
        self.assertEqual(result.candidate_facts, [])
        self.assertEqual(result.documents_processed, 0)

    def test_directorio_sin_inputs_genera_warning(self):
        with tempfile.TemporaryDirectory() as tmp:
            result = run_phase1(tmp)
        self.assertTrue(any("document" in w.lower() or "entr" in w.lower()
                            for w in result.warnings))

    def test_expediente_id_es_nombre_directorio(self):
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            exp_dir = tmp_path / "expediente-EIA-TEST-001"
            exp_dir.mkdir()
            result = run_phase1(exp_dir)
        self.assertEqual(result.expediente_id, "expediente-EIA-TEST-001")

    def test_write_false_no_escribe_archivos(self):
        with tempfile.TemporaryDirectory() as tmp:
            result = run_phase1(tmp, write_outputs=False)
            ci = Path(tmp) / "control_interno"
            self.assertFalse((ci / "phase1_result.json").exists())
            self.assertFalse((ci / "phase1_result.md").exists())


# ---------------------------------------------------------------------------
# Clase 5: run_phase1 — write_outputs
# ---------------------------------------------------------------------------

class TestRunPhase1WriteOutputs(unittest.TestCase):
    def test_write_true_crea_json(self):
        with tempfile.TemporaryDirectory() as tmp:
            run_phase1(tmp, write_outputs=True)
            json_path = Path(tmp) / "control_interno" / "phase1_result.json"
            self.assertTrue(json_path.exists())

    def test_write_true_crea_md(self):
        with tempfile.TemporaryDirectory() as tmp:
            run_phase1(tmp, write_outputs=True)
            md_path = Path(tmp) / "control_interno" / "phase1_result.md"
            self.assertTrue(md_path.exists())

    def test_json_es_valido(self):
        with tempfile.TemporaryDirectory() as tmp:
            run_phase1(tmp, write_outputs=True)
            json_path = Path(tmp) / "control_interno" / "phase1_result.json"
            with open(json_path, encoding="utf-8") as f:
                data = json.load(f)
            self.assertIn("expediente_id", data)
            self.assertIn("candidate_facts", data)

    def test_md_contiene_fase1(self):
        with tempfile.TemporaryDirectory() as tmp:
            run_phase1(tmp, write_outputs=True)
            md_path = Path(tmp) / "control_interno" / "phase1_result.md"
            content = md_path.read_text(encoding="utf-8")
            self.assertIn("Fase 1", content)

    def test_output_dir_personalizado(self):
        with tempfile.TemporaryDirectory() as tmp:
            run_phase1(tmp, write_outputs=True, output_dir="mi_salida")
            json_path = Path(tmp) / "mi_salida" / "phase1_result.json"
            self.assertTrue(json_path.exists())

    def test_write_false_no_escribe_nada(self):
        with tempfile.TemporaryDirectory() as tmp:
            run_phase1(tmp, write_outputs=False)
            ci = Path(tmp) / "control_interno"
            self.assertFalse(ci.exists())


# ---------------------------------------------------------------------------
# Clase 6: run_phase1 — con DOCX sintético
# ---------------------------------------------------------------------------

class TestRunPhase1ConDocxSintetico(unittest.TestCase):
    def _make_docx_expediente(self, tmp_dir):
        """Crea estructura mínima de expediente con un DOCX real (vacío pero válido)."""
        from docx import Document
        exp_dir = Path(tmp_dir) / "expediente-TEST"
        exp_dir.mkdir()
        inputs_dir = exp_dir / "inputs"
        inputs_dir.mkdir()
        doc = Document()
        doc.add_paragraph("Promotor: EMPRESA SINTÉTICA SL")
        doc.add_paragraph("Referencia catastral: 1234567AB1234S0001XY")
        doc.save(str(inputs_dir / "memoria.docx"))
        return exp_dir

    def test_docx_genera_resultado_con_facts(self):
        try:
            with tempfile.TemporaryDirectory() as tmp:
                exp_dir = self._make_docx_expediente(tmp)
                result = run_phase1(exp_dir)
            self.assertIsInstance(result, Phase1Result)
            self.assertEqual(result.docx_processed, 1)
        except ImportError:
            self.skipTest("python-docx no disponible")

    def test_docx_sin_pdf_pending_cero(self):
        try:
            with tempfile.TemporaryDirectory() as tmp:
                exp_dir = self._make_docx_expediente(tmp)
                result = run_phase1(exp_dir)
            self.assertEqual(result.pdf_pending, 0)
        except ImportError:
            self.skipTest("python-docx no disponible")

    def test_documents_processed_cuenta_docx(self):
        try:
            with tempfile.TemporaryDirectory() as tmp:
                exp_dir = self._make_docx_expediente(tmp)
                result = run_phase1(exp_dir)
            self.assertGreaterEqual(result.documents_processed, 1)
        except ImportError:
            self.skipTest("python-docx no disponible")


# ---------------------------------------------------------------------------
# Clase 7: run_phase1 — con PDF
# ---------------------------------------------------------------------------

class TestRunPhase1ConPdf(unittest.TestCase):
    def _make_pdf_expediente(self, tmp_dir):
        exp_dir = Path(tmp_dir) / "expediente-TEST-PDF"
        exp_dir.mkdir()
        inputs_dir = exp_dir / "inputs"
        inputs_dir.mkdir()
        pdf_path = inputs_dir / "proyecto.pdf"
        pdf_path.write_bytes(b"%PDF-1.4 dummy content")
        return exp_dir

    def test_pdf_cuenta_como_pendiente(self):
        with tempfile.TemporaryDirectory() as tmp:
            exp_dir = self._make_pdf_expediente(tmp)
            result = run_phase1(exp_dir)
        self.assertEqual(result.pdf_pending, 1)

    def test_pdf_genera_warning(self):
        with tempfile.TemporaryDirectory() as tmp:
            exp_dir = self._make_pdf_expediente(tmp)
            result = run_phase1(exp_dir)
        self.assertTrue(any("PDF" in w or "pdf" in w for w in result.warnings))

    def test_pdf_no_genera_candidate_facts(self):
        with tempfile.TemporaryDirectory() as tmp:
            exp_dir = self._make_pdf_expediente(tmp)
            result = run_phase1(exp_dir)
        self.assertEqual(result.candidate_facts, [])

    def test_pdf_no_se_cuenta_como_docx_procesado(self):
        with tempfile.TemporaryDirectory() as tmp:
            exp_dir = self._make_pdf_expediente(tmp)
            result = run_phase1(exp_dir)
        self.assertEqual(result.docx_processed, 0)


# ---------------------------------------------------------------------------
# Clase 8: CLI — integración mínima
# ---------------------------------------------------------------------------

class TestCLIPhase1(unittest.TestCase):
    def _run_cli(self, *args):
        import subprocess
        cli = Path(__file__).parent.parent / "run_expediente.py"
        cmd = [sys.executable, str(cli)] + list(args)
        return subprocess.run(cmd, capture_output=True, text=True, timeout=30)

    def test_cli_phase1_sin_write_exit_0(self):
        with tempfile.TemporaryDirectory() as tmp:
            r = self._run_cli(tmp, "phase1")
        self.assertEqual(r.returncode, 0)

    def test_cli_phase1_sin_write_no_crea_archivos(self):
        with tempfile.TemporaryDirectory() as tmp:
            self._run_cli(tmp, "phase1")
            ci = Path(tmp) / "control_interno"
            self.assertFalse(ci.exists())

    def test_cli_phase1_con_write_crea_json(self):
        with tempfile.TemporaryDirectory() as tmp:
            self._run_cli(tmp, "phase1", "--write")
            json_path = Path(tmp) / "control_interno" / "phase1_result.json"
            self.assertTrue(json_path.exists())

    def test_cli_phase1_imprime_resumen(self):
        with tempfile.TemporaryDirectory() as tmp:
            r = self._run_cli(tmp, "phase1")
        self.assertIn("Fase 1", r.stdout)

    def test_cli_expediente_inexistente_exit_1(self):
        r = self._run_cli("/ruta/que/no/existe/jamas", "phase1")
        self.assertEqual(r.returncode, 1)


# ---------------------------------------------------------------------------
# Clase 9: Pilots — solo lectura (PARCELA)
# ---------------------------------------------------------------------------

@unittest.skipUnless(_PARCELA.exists(), "Piloto PARCELA no disponible")
class TestRunPhase1PilotoParcela(unittest.TestCase):
    def test_no_modifica_inputs(self):
        inputs_before = set((_PARCELA / "inputs").glob("**/*")) if (_PARCELA / "inputs").exists() else set()
        run_phase1(_PARCELA, write_outputs=False)
        inputs_after = set((_PARCELA / "inputs").glob("**/*")) if (_PARCELA / "inputs").exists() else set()
        self.assertEqual(inputs_before, inputs_after)

    def test_devuelve_phase1result(self):
        result = run_phase1(_PARCELA, write_outputs=False)
        self.assertIsInstance(result, Phase1Result)

    def test_no_escribe_en_control_interno(self):
        ci = _PARCELA / "control_interno"
        before = set(ci.glob("phase1_result*")) if ci.exists() else set()
        run_phase1(_PARCELA, write_outputs=False)
        after = set(ci.glob("phase1_result*")) if ci.exists() else set()
        self.assertEqual(before, after)

    def test_summary_es_string_no_vacio(self):
        result = run_phase1(_PARCELA, write_outputs=False)
        s = result.summary()
        self.assertIsInstance(s, str)
        self.assertGreater(len(s), 0)


# ---------------------------------------------------------------------------
# Clase 10: Pilots — solo lectura (NAVE-222)
# ---------------------------------------------------------------------------

@unittest.skipUnless(_NAVE.exists(), "Piloto NAVE-222 no disponible")
class TestRunPhase1PilotoNave222(unittest.TestCase):
    def test_no_modifica_inputs(self):
        inputs_dir = _NAVE / "inputs"
        inputs_before = set(inputs_dir.glob("**/*")) if inputs_dir.exists() else set()
        run_phase1(_NAVE, write_outputs=False)
        inputs_after = set(inputs_dir.glob("**/*")) if inputs_dir.exists() else set()
        self.assertEqual(inputs_before, inputs_after)

    def test_devuelve_phase1result(self):
        result = run_phase1(_NAVE, write_outputs=False)
        self.assertIsInstance(result, Phase1Result)

    def test_no_escribe_en_control_interno(self):
        ci = _NAVE / "control_interno"
        before = set(ci.glob("phase1_result*")) if ci.exists() else set()
        run_phase1(_NAVE, write_outputs=False)
        after = set(ci.glob("phase1_result*")) if ci.exists() else set()
        self.assertEqual(before, after)


if __name__ == "__main__":
    unittest.main()
