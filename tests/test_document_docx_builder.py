"""
tests/test_document_docx_builder.py
Tests para DOC-02 — Generador DOCX del Documento Ambiental.

Cubre:
  1. safe_read_markdown
  2. parse_markdown_blocks
  3. create_docx_document
  4. add_cover_page
  5. add_table_of_contents_placeholder
  6. add_markdown_block_to_docx
  7. build_docx_from_markdown
  8. build_docx_from_expediente
  9. validate_docx_basic
  10. CLI document-build-docx
  11. Seguridad metodologica
"""
import json
import sys
import tempfile
import unittest
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent / "src"))
sys.path.insert(0, str(Path(__file__).parent.parent))

from eia_agent.core.document_docx_builder import (
    ADMIN_STYLE_PROFILE,
    DOCX_BUILD_RESULT_FILENAME,
    DOCX_OUTPUT_FILENAME,
    DEFAULT_SUBTITLE,
    DEFAULT_TITLE,
    SUPPORTED_MARKDOWN_ELEMENTS,
    DocxBuildWarning,
    DocumentDocxBuildResult,
    _ADMIN_DISCLAIMER,
    _strip_markdown_inline,
    add_cover_page,
    add_markdown_block_to_docx,
    add_running_header_footer,
    add_table_of_contents_placeholder,
    build_docx_from_expediente,
    build_docx_from_markdown,
    create_docx_document,
    parse_markdown_blocks,
    safe_read_markdown,
    validate_docx_basic,
    write_docx_build_result,
)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

_MINIMAL_MARKDOWN = """\
# Documento Ambiental — Borrador tecnico

> Documento generado automaticamente. No declara aptitud administrativa.

---

## Indice

- Bloque A
- Bloque B

---

## Bloque A — Identificacion

**Expediente:** TEST-001

Parrafo de descripcion del proyecto.

### A.1 Acciones

| ID | Nombre | Tipo |
|----|--------|------|
| AC-001 | Recepcion | RECEPCION |

### A.2 Otra seccion

Otro parrafo.

---

## Bloque I — Conclusiones

> Estado auditoria: CONFORME_CON_OBSERVACIONES

Conclusion tecnica prudente.
"""

_SIMPLE_TABLE_MARKDOWN = """\
| Col1 | Col2 |
|------|------|
| a    | b    |
| c    | d    |
"""


def _make_expediente_with_markdown(tmp: Path, md_content: str = _MINIMAL_MARKDOWN) -> Path:
    """Crea expediente temporal con documento_ambiental_borrador.md."""
    exp = tmp / "expediente-test-docx"
    exp.mkdir(exist_ok=True)
    doc_dir = exp / "documento"
    doc_dir.mkdir(exist_ok=True)
    (doc_dir / "documento_ambiental_borrador.md").write_text(
        md_content, encoding="utf-8"
    )
    return exp


def _make_expediente_empty(tmp: Path) -> Path:
    """Crea expediente temporal sin markdown."""
    exp = tmp / "expediente-empty-docx"
    exp.mkdir(exist_ok=True)
    return exp


# ---------------------------------------------------------------------------
# 1. Tests de safe_read_markdown
# ---------------------------------------------------------------------------

class TestSafeReadMarkdown(unittest.TestCase):

    def test_existing_file(self):
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / "test.md"
            p.write_text("# Titulo\nContenido.", encoding="utf-8")
            result = safe_read_markdown(p)
            self.assertEqual(result, "# Titulo\nContenido.")

    def test_nonexistent_raises_file_not_found(self):
        with self.assertRaises(FileNotFoundError):
            safe_read_markdown("/ruta/que/no/existe/archivo.md")

    def test_empty_file_raises_value_error(self):
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / "empty.md"
            p.write_text("", encoding="utf-8")
            with self.assertRaises(ValueError):
                safe_read_markdown(p)

    def test_whitespace_only_raises_value_error(self):
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / "ws.md"
            p.write_text("   \n\n   ", encoding="utf-8")
            with self.assertRaises(ValueError):
                safe_read_markdown(p)

    def test_string_path_works(self):
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / "test.md"
            p.write_text("Contenido", encoding="utf-8")
            result = safe_read_markdown(str(p))
            self.assertEqual(result, "Contenido")


# ---------------------------------------------------------------------------
# 2. Tests de parse_markdown_blocks
# ---------------------------------------------------------------------------

class TestParseMarkdownBlocks(unittest.TestCase):

    def _types(self, blocks):
        return [b.get("type") for b in blocks if b.get("type") != "blank_line"]

    def test_detects_headings(self):
        md = "# H1\n## H2\n### H3\n#### H4"
        blocks = parse_markdown_blocks(md)
        headings = [b for b in blocks if b.get("type") == "heading"]
        self.assertEqual(len(headings), 4)
        self.assertEqual(headings[0]["level"], 1)
        self.assertEqual(headings[1]["level"], 2)
        self.assertEqual(headings[2]["level"], 3)
        self.assertEqual(headings[3]["level"], 4)

    def test_heading_text_extracted(self):
        blocks = parse_markdown_blocks("## Bloque A — Titulo")
        h = next(b for b in blocks if b.get("type") == "heading")
        self.assertIn("Bloque A", h["text"])

    def test_detects_paragraph(self):
        blocks = parse_markdown_blocks("Este es un parrafo.")
        paras = [b for b in blocks if b.get("type") == "paragraph"]
        self.assertGreater(len(paras), 0)
        self.assertIn("parrafo", paras[0]["text"])

    def test_detects_bullet_list(self):
        blocks = parse_markdown_blocks("- item1\n- item2\n- item3")
        bullets = [b for b in blocks if b.get("type") == "bullet_list"]
        self.assertEqual(len(bullets), 1)
        self.assertEqual(len(bullets[0]["items"]), 3)

    def test_detects_bullet_with_asterisk(self):
        blocks = parse_markdown_blocks("* item1\n* item2")
        bullets = [b for b in blocks if b.get("type") == "bullet_list"]
        self.assertGreater(len(bullets), 0)

    def test_detects_numbered_list(self):
        blocks = parse_markdown_blocks("1. primero\n2. segundo\n3. tercero")
        numbered = [b for b in blocks if b.get("type") == "numbered_list"]
        self.assertEqual(len(numbered), 1)
        self.assertEqual(len(numbered[0]["items"]), 3)

    def test_detects_table(self):
        blocks = parse_markdown_blocks(_SIMPLE_TABLE_MARKDOWN)
        tables = [b for b in blocks if b.get("type") == "table"]
        self.assertEqual(len(tables), 1)
        t = tables[0]
        self.assertEqual(len(t["headers"]), 2)
        self.assertEqual(len(t["rows"]), 2)

    def test_table_header_content(self):
        blocks = parse_markdown_blocks(_SIMPLE_TABLE_MARKDOWN)
        t = next(b for b in blocks if b.get("type") == "table")
        self.assertIn("Col1", t["headers"])
        self.assertIn("Col2", t["headers"])

    def test_detects_blockquote(self):
        blocks = parse_markdown_blocks("> Este es un aviso importante.")
        bqs = [b for b in blocks if b.get("type") == "blockquote"]
        self.assertGreater(len(bqs), 0)
        self.assertIn("aviso", bqs[0]["text"])

    def test_detects_horizontal_rule(self):
        blocks = parse_markdown_blocks("Texto\n\n---\n\nOtro texto")
        hrs = [b for b in blocks if b.get("type") == "horizontal_rule"]
        self.assertGreater(len(hrs), 0)

    def test_detects_code_block(self):
        blocks = parse_markdown_blocks("```\npython code here\n```")
        code = [b for b in blocks if b.get("type") == "code_block"]
        self.assertGreater(len(code), 0)
        self.assertIn("python code", code[0]["text"])

    def test_does_not_break_on_empty_input(self):
        blocks = parse_markdown_blocks("")
        self.assertIsInstance(blocks, list)

    def test_does_not_break_on_weird_lines(self):
        weird = "!!! raroCaract3r\n||\n##\n# \n> \n---\n-\n1."
        try:
            blocks = parse_markdown_blocks(weird)
            self.assertIsInstance(blocks, list)
        except Exception as e:
            self.fail(f"parse_markdown_blocks lanzo excepcion con lineas raras: {e}")

    def test_heading_level_capped_at_4(self):
        blocks = parse_markdown_blocks("##### H5 heading")
        headings = [b for b in blocks if b.get("type") == "heading"]
        for h in headings:
            self.assertLessEqual(h["level"], 4)

    def test_full_minimal_markdown(self):
        blocks = parse_markdown_blocks(_MINIMAL_MARKDOWN)
        types = [b.get("type") for b in blocks]
        self.assertIn("heading", types)
        self.assertIn("paragraph", types)
        self.assertIn("table", types)
        self.assertIn("blockquote", types)
        self.assertIn("bullet_list", types)

    def test_blank_lines_detected(self):
        blocks = parse_markdown_blocks("texto\n\n\nmas texto")
        blank = [b for b in blocks if b.get("type") == "blank_line"]
        self.assertGreater(len(blank), 0)


class TestStripMarkdownInline(unittest.TestCase):

    def test_strips_bold(self):
        self.assertEqual(_strip_markdown_inline("**texto**"), "texto")

    def test_strips_italic(self):
        self.assertEqual(_strip_markdown_inline("*texto*"), "texto")

    def test_strips_code(self):
        self.assertEqual(_strip_markdown_inline("`codigo`"), "codigo")

    def test_strips_link(self):
        self.assertEqual(_strip_markdown_inline("[texto](https://url.com)"), "texto")

    def test_preserves_plain(self):
        self.assertEqual(_strip_markdown_inline("texto plano"), "texto plano")


# ---------------------------------------------------------------------------
# 3. Tests de create_docx_document
# ---------------------------------------------------------------------------

class TestCreateDocxDocument(unittest.TestCase):

    def test_returns_document(self):
        doc = create_docx_document()
        # python-docx Document() devuelve docx.document.Document
        self.assertIsNotNone(doc)
        self.assertTrue(hasattr(doc, "paragraphs"))
        self.assertTrue(hasattr(doc, "add_heading"))
        self.assertTrue(hasattr(doc, "add_paragraph"))

    def test_heading1_style_accessible(self):
        doc = create_docx_document()
        style = doc.styles["Heading 1"]
        self.assertIsNotNone(style)

    def test_heading2_style_accessible(self):
        doc = create_docx_document()
        self.assertIsNotNone(doc.styles["Heading 2"])

    def test_normal_style_accessible(self):
        doc = create_docx_document()
        self.assertIsNotNone(doc.styles["Normal"])

    def test_normal_style_uses_admin_reference_font(self):
        doc = create_docx_document()
        normal = doc.styles["Normal"]
        self.assertEqual(normal.font.name, "Calibri Light")
        self.assertEqual(normal.font.size.pt, 12)

    def test_margins_follow_admin_reference_profile(self):
        doc = create_docx_document()
        section = doc.sections[0]
        self.assertAlmostEqual(section.left_margin.cm, 2.5, places=1)
        self.assertAlmostEqual(section.right_margin.cm, 2.3, places=1)
        self.assertAlmostEqual(section.top_margin.cm, 2.2, places=1)
        self.assertAlmostEqual(section.bottom_margin.cm, 2.0, places=1)

    def test_has_sections(self):
        doc = create_docx_document()
        self.assertGreater(len(doc.sections), 0)

    def test_custom_title_not_crashes(self):
        doc = create_docx_document(title="Mi Titulo", subtitle="Mi Subtitulo")
        self.assertIsNotNone(doc)


# ---------------------------------------------------------------------------
# 4. Tests de add_cover_page
# ---------------------------------------------------------------------------

class TestAddCoverPage(unittest.TestCase):

    def _get_full_text(self, doc) -> str:
        return "\n".join(p.text for p in doc.paragraphs)

    def test_cover_contains_title(self):
        doc = create_docx_document()
        add_cover_page(doc, "EXP-001", "Mi Titulo", "Mi Subtitulo", "Aviso")
        text = self._get_full_text(doc)
        self.assertIn("MI TITULO", text)

    def test_cover_contains_expediente_id(self):
        doc = create_docx_document()
        add_cover_page(doc, "EXP-TEST-123", DEFAULT_TITLE, DEFAULT_SUBTITLE, _ADMIN_DISCLAIMER)
        text = self._get_full_text(doc)
        self.assertIn("EXP-TEST-123", text)

    def test_cover_contains_disclaimer(self):
        doc = create_docx_document()
        add_cover_page(doc, "EXP-001", DEFAULT_TITLE, DEFAULT_SUBTITLE, _ADMIN_DISCLAIMER)
        text = self._get_full_text(doc)
        self.assertIn("aptitud administrativa", text.lower())

    def test_cover_no_fail_without_logo(self):
        doc = create_docx_document()
        try:
            add_cover_page(doc, "EXP-001", DEFAULT_TITLE, DEFAULT_SUBTITLE, _ADMIN_DISCLAIMER,
                          logo_path=None)
        except Exception as e:
            self.fail(f"add_cover_page fallo sin logo: {e}")

    def test_cover_no_fail_nonexistent_logo(self):
        doc = create_docx_document()
        try:
            add_cover_page(doc, "EXP-001", DEFAULT_TITLE, DEFAULT_SUBTITLE, _ADMIN_DISCLAIMER,
                          logo_path="/ruta/logo/no/existe.png")
        except Exception as e:
            self.fail(f"add_cover_page fallo con logo inexistente: {e}")

    def test_cover_contains_fecha(self):
        doc = create_docx_document()
        add_cover_page(doc, "EXP-001", DEFAULT_TITLE, DEFAULT_SUBTITLE, _ADMIN_DISCLAIMER)
        text = self._get_full_text(doc)
        self.assertIn("generacion", text.lower())


class TestRunningHeaderFooter(unittest.TestCase):

    def test_adds_expediente_header(self):
        doc = create_docx_document()
        add_running_header_footer(doc, "EXP-HEADER-001")
        text = "\n".join(p.text for p in doc.sections[0].header.paragraphs)
        self.assertIn("EXP-HEADER-001", text)

    def test_adds_page_footer_label(self):
        doc = create_docx_document()
        add_running_header_footer(doc, "EXP-FOOTER-001")
        text = "\n".join(p.text for p in doc.sections[0].footer.paragraphs)
        self.assertIn("Pagina", text)


# ---------------------------------------------------------------------------
# 5. Tests de add_table_of_contents_placeholder
# ---------------------------------------------------------------------------

class TestAddTableOfContentsPlaceholder(unittest.TestCase):

    def test_adds_content(self):
        doc = create_docx_document()
        initial_para_count = len(doc.paragraphs)
        add_table_of_contents_placeholder(doc)
        self.assertGreater(len(doc.paragraphs), initial_para_count)

    def test_contains_indice_heading(self):
        doc = create_docx_document()
        add_table_of_contents_placeholder(doc)
        headings_text = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
        self.assertTrue(
            any("ndice" in t or "indice" in t.lower() for t in headings_text),
            f"No se encontro encabezado de indice. Headings: {headings_text}"
        )

    def test_contains_update_note(self):
        doc = create_docx_document()
        add_table_of_contents_placeholder(doc)
        all_text = "\n".join(p.text for p in doc.paragraphs)
        # Debe contener nota sobre actualizacion o TOC
        self.assertTrue(
            "ndice" in all_text or "TOC" in all_text or "actualiz" in all_text.lower(),
            f"No se encontro nota de TOC. Texto: {all_text[:200]}"
        )


# ---------------------------------------------------------------------------
# 6. Tests de add_markdown_block_to_docx
# ---------------------------------------------------------------------------

class TestAddMarkdownBlockToDocx(unittest.TestCase):

    def setUp(self):
        self.doc = create_docx_document()
        self.initial_paras = len(self.doc.paragraphs)

    def _para_count(self):
        return len(self.doc.paragraphs) - self.initial_paras

    def test_heading_adds_heading(self):
        block = {"type": "heading", "level": 2, "text": "Seccion de prueba"}
        counts = add_markdown_block_to_docx(self.doc, block)
        self.assertEqual(counts["heading_added"], 1)
        self.assertEqual(counts["paragraph_added"], 0)

    def test_heading_text_in_document(self):
        block = {"type": "heading", "level": 1, "text": "Mi Encabezado"}
        add_markdown_block_to_docx(self.doc, block)
        texts = [p.text for p in self.doc.paragraphs]
        self.assertTrue(any("Mi Encabezado" in t for t in texts))

    def test_paragraph_adds_paragraph(self):
        block = {"type": "paragraph", "text": "Un parrafo de prueba."}
        counts = add_markdown_block_to_docx(self.doc, block)
        self.assertEqual(counts["paragraph_added"], 1)
        self.assertEqual(counts["heading_added"], 0)

    def test_bullet_list_adds_paragraphs(self):
        block = {"type": "bullet_list", "items": ["item1", "item2", "item3"]}
        counts = add_markdown_block_to_docx(self.doc, block)
        self.assertEqual(counts["paragraph_added"], 3)

    def test_numbered_list_adds_paragraphs(self):
        block = {"type": "numbered_list", "items": ["uno", "dos"]}
        counts = add_markdown_block_to_docx(self.doc, block)
        self.assertEqual(counts["paragraph_added"], 2)

    def test_table_adds_table(self):
        block = {
            "type": "table",
            "headers": ["Col A", "Col B"],
            "rows": [["val1", "val2"], ["val3", "val4"]],
        }
        counts = add_markdown_block_to_docx(self.doc, block)
        self.assertEqual(counts["table_added"], 1)
        self.assertGreater(len(self.doc.tables), 0)

    def test_table_header_content(self):
        block = {
            "type": "table",
            "headers": ["ID", "Nombre"],
            "rows": [["001", "Test"]],
        }
        add_markdown_block_to_docx(self.doc, block)
        table = self.doc.tables[-1]
        cell_texts = [c.text for r in table.rows for c in r.cells]
        self.assertIn("ID", cell_texts)
        self.assertIn("Nombre", cell_texts)

    def test_blockquote_adds_paragraph(self):
        block = {"type": "blockquote", "text": "Texto de aviso importante."}
        counts = add_markdown_block_to_docx(self.doc, block)
        self.assertEqual(counts["paragraph_added"], 1)

    def test_blockquote_text_in_document(self):
        block = {"type": "blockquote", "text": "Aviso clave de prueba."}
        add_markdown_block_to_docx(self.doc, block)
        texts = [p.text for p in self.doc.paragraphs]
        self.assertTrue(any("Aviso clave" in t for t in texts))

    def test_code_block_adds_paragraph(self):
        block = {"type": "code_block", "text": "print('hello')"}
        counts = add_markdown_block_to_docx(self.doc, block)
        self.assertEqual(counts["paragraph_added"], 1)

    def test_horizontal_rule_no_exception(self):
        block = {"type": "horizontal_rule"}
        try:
            counts = add_markdown_block_to_docx(self.doc, block)
            self.assertEqual(counts["paragraph_added"], 1)
        except Exception as e:
            self.fail(f"horizontal_rule lanzo excepcion: {e}")

    def test_blank_line_adds_nothing(self):
        block = {"type": "blank_line"}
        counts = add_markdown_block_to_docx(self.doc, block)
        self.assertEqual(counts["paragraph_added"], 0)
        self.assertEqual(counts["heading_added"], 0)

    def test_unknown_block_type_no_exception(self):
        block = {"type": "tipo_desconocido_xyz", "text": "algo"}
        try:
            counts = add_markdown_block_to_docx(self.doc, block)
            self.assertIsInstance(counts, dict)
        except Exception as e:
            self.fail(f"Tipo desconocido lanzo excepcion: {e}")

    def test_empty_table_no_exception(self):
        block = {"type": "table", "headers": [], "rows": []}
        try:
            counts = add_markdown_block_to_docx(self.doc, block)
            self.assertIsInstance(counts, dict)
        except Exception as e:
            self.fail(f"Tabla vacia lanzo excepcion: {e}")


# ---------------------------------------------------------------------------
# 7. Tests de build_docx_from_markdown
# ---------------------------------------------------------------------------

class TestBuildDocxFromMarkdown(unittest.TestCase):

    def setUp(self):
        self.tmp_obj = tempfile.TemporaryDirectory()
        self.tmp = Path(self.tmp_obj.name)

    def tearDown(self):
        self.tmp_obj.cleanup()

    def test_generates_valid_docx(self):
        md_path = self.tmp / "test.md"
        md_path.write_text(_MINIMAL_MARKDOWN, encoding="utf-8")
        out_path = self.tmp / "output.docx"
        result = build_docx_from_markdown(md_path, out_path, expediente_id="TEST")
        self.assertTrue(result.generated)
        self.assertTrue(out_path.exists())
        self.assertGreater(out_path.stat().st_size, 0)

    def test_result_generated_true(self):
        md_path = self.tmp / "test.md"
        md_path.write_text(_MINIMAL_MARKDOWN, encoding="utf-8")
        out_path = self.tmp / "output.docx"
        result = build_docx_from_markdown(md_path, out_path)
        self.assertTrue(result.generated)
        self.assertTrue(result.is_success())

    def test_heading_count_greater_than_zero(self):
        md_path = self.tmp / "test.md"
        md_path.write_text(_MINIMAL_MARKDOWN, encoding="utf-8")
        out_path = self.tmp / "output.docx"
        result = build_docx_from_markdown(md_path, out_path)
        self.assertGreater(result.heading_count, 0)

    def test_paragraph_count_greater_than_zero(self):
        md_path = self.tmp / "test.md"
        md_path.write_text(_MINIMAL_MARKDOWN, encoding="utf-8")
        out_path = self.tmp / "output.docx"
        result = build_docx_from_markdown(md_path, out_path)
        self.assertGreater(result.paragraph_count, 0)

    def test_table_count_positive_when_table_present(self):
        md_path = self.tmp / "test.md"
        md_path.write_text(_MINIMAL_MARKDOWN, encoding="utf-8")
        out_path = self.tmp / "output.docx"
        result = build_docx_from_markdown(md_path, out_path)
        self.assertGreater(result.table_count, 0)

    def test_docx_is_openable(self):
        md_path = self.tmp / "test.md"
        md_path.write_text(_MINIMAL_MARKDOWN, encoding="utf-8")
        out_path = self.tmp / "output.docx"
        build_docx_from_markdown(md_path, out_path)
        self.assertTrue(validate_docx_basic(out_path))

    def test_does_not_modify_markdown_source(self):
        md_path = self.tmp / "test.md"
        md_path.write_text(_MINIMAL_MARKDOWN, encoding="utf-8")
        out_path = self.tmp / "output.docx"
        build_docx_from_markdown(md_path, out_path)
        content_after = md_path.read_text(encoding="utf-8")
        self.assertEqual(content_after, _MINIMAL_MARKDOWN)

    def test_expediente_id_in_result(self):
        md_path = self.tmp / "test.md"
        md_path.write_text(_MINIMAL_MARKDOWN, encoding="utf-8")
        out_path = self.tmp / "output.docx"
        result = build_docx_from_markdown(md_path, out_path, expediente_id="EXP-999")
        self.assertEqual(result.expediente_id, "EXP-999")

    def test_no_logo_does_not_fail(self):
        md_path = self.tmp / "test.md"
        md_path.write_text(_MINIMAL_MARKDOWN, encoding="utf-8")
        out_path = self.tmp / "output.docx"
        result = build_docx_from_markdown(md_path, out_path, logo_path=None)
        self.assertTrue(result.generated)

    def test_output_path_in_result(self):
        md_path = self.tmp / "test.md"
        md_path.write_text(_MINIMAL_MARKDOWN, encoding="utf-8")
        out_path = self.tmp / "output.docx"
        result = build_docx_from_markdown(md_path, out_path)
        self.assertEqual(Path(result.output_docx_path), out_path)


# ---------------------------------------------------------------------------
# 8. Tests de build_docx_from_expediente
# ---------------------------------------------------------------------------

class TestBuildDocxFromExpediente(unittest.TestCase):

    def setUp(self):
        self.tmp_obj = tempfile.TemporaryDirectory()
        self.tmp = Path(self.tmp_obj.name)

    def tearDown(self):
        self.tmp_obj.cleanup()

    def test_missing_markdown_raises_file_not_found(self):
        exp = _make_expediente_empty(self.tmp)
        with self.assertRaises(FileNotFoundError):
            build_docx_from_expediente(exp, write_outputs=False)

    def test_without_write_generated_false(self):
        exp = _make_expediente_with_markdown(self.tmp)
        result = build_docx_from_expediente(exp, write_outputs=False)
        self.assertFalse(result.generated)
        self.assertIsNone(result.output_docx_path)

    def test_without_write_no_docx_created(self):
        exp = _make_expediente_with_markdown(self.tmp)
        build_docx_from_expediente(exp, write_outputs=False)
        docx_path = exp / "documento" / DOCX_OUTPUT_FILENAME
        self.assertFalse(docx_path.exists())

    def test_without_write_returns_counts(self):
        exp = _make_expediente_with_markdown(self.tmp)
        result = build_docx_from_expediente(exp, write_outputs=False)
        self.assertGreater(result.heading_count, 0)

    def test_with_write_generates_docx(self):
        exp = _make_expediente_with_markdown(self.tmp)
        result = build_docx_from_expediente(exp, write_outputs=True)
        self.assertTrue(result.generated)
        docx_path = exp / "documento" / DOCX_OUTPUT_FILENAME
        self.assertTrue(docx_path.exists())

    def test_with_write_generates_json(self):
        exp = _make_expediente_with_markdown(self.tmp)
        build_docx_from_expediente(exp, write_outputs=True)
        json_path = exp / "documento" / DOCX_BUILD_RESULT_FILENAME
        self.assertTrue(json_path.exists())

    def test_json_is_loadable(self):
        exp = _make_expediente_with_markdown(self.tmp)
        build_docx_from_expediente(exp, write_outputs=True)
        json_path = exp / "documento" / DOCX_BUILD_RESULT_FILENAME
        data = json.loads(json_path.read_text(encoding="utf-8"))
        self.assertIn("expediente_id", data)
        self.assertIn("generated", data)
        self.assertTrue(data["generated"])

    def test_docx_openable_after_write(self):
        exp = _make_expediente_with_markdown(self.tmp)
        build_docx_from_expediente(exp, write_outputs=True)
        docx_path = exp / "documento" / DOCX_OUTPUT_FILENAME
        self.assertTrue(validate_docx_basic(docx_path))

    def test_result_expediente_id_correct(self):
        exp = _make_expediente_with_markdown(self.tmp)
        result = build_docx_from_expediente(exp, write_outputs=False)
        self.assertEqual(result.expediente_id, exp.name)


# ---------------------------------------------------------------------------
# 9. Tests de validate_docx_basic
# ---------------------------------------------------------------------------

class TestValidateDocxBasic(unittest.TestCase):

    def setUp(self):
        self.tmp_obj = tempfile.TemporaryDirectory()
        self.tmp = Path(self.tmp_obj.name)

    def tearDown(self):
        self.tmp_obj.cleanup()

    def _make_docx(self) -> Path:
        from docx import Document
        p = self.tmp / "valid.docx"
        doc = Document()
        doc.add_paragraph("Contenido de prueba.")
        doc.save(str(p))
        return p

    def test_valid_docx_returns_true(self):
        p = self._make_docx()
        self.assertTrue(validate_docx_basic(p))

    def test_nonexistent_returns_false(self):
        self.assertFalse(validate_docx_basic("/no/existe/archivo.docx"))

    def test_non_docx_extension_returns_false(self):
        p = self.tmp / "archivo.pdf"
        p.write_bytes(b"PDF content")
        self.assertFalse(validate_docx_basic(p))

    def test_empty_file_returns_false(self):
        p = self.tmp / "empty.docx"
        p.write_bytes(b"")
        self.assertFalse(validate_docx_basic(p))

    def test_invalid_content_returns_false(self):
        p = self.tmp / "invalid.docx"
        p.write_bytes(b"este no es un docx valido")
        self.assertFalse(validate_docx_basic(p))


# ---------------------------------------------------------------------------
# 10. Tests de CLI document-build-docx
# ---------------------------------------------------------------------------

class TestCLIDocumentBuildDocx(unittest.TestCase):

    def setUp(self):
        self.tmp_obj = tempfile.TemporaryDirectory()
        self.tmp = Path(self.tmp_obj.name)

    def tearDown(self):
        self.tmp_obj.cleanup()

    def _run_cli(self, args: list[str]) -> int:
        from run_expediente import main
        return main(args)

    def test_without_write_no_docx_created(self):
        exp = _make_expediente_with_markdown(self.tmp)
        self._run_cli([str(exp), "document-build-docx"])
        docx_path = exp / "documento" / DOCX_OUTPUT_FILENAME
        self.assertFalse(docx_path.exists())

    def test_with_write_creates_docx(self):
        exp = _make_expediente_with_markdown(self.tmp)
        self._run_cli([str(exp), "document-build-docx", "--write"])
        docx_path = exp / "documento" / DOCX_OUTPUT_FILENAME
        self.assertTrue(docx_path.exists())

    def test_with_write_creates_json(self):
        exp = _make_expediente_with_markdown(self.tmp)
        self._run_cli([str(exp), "document-build-docx", "--write"])
        json_path = exp / "documento" / DOCX_BUILD_RESULT_FILENAME
        self.assertTrue(json_path.exists())

    def test_missing_markdown_exit_1(self):
        exp = _make_expediente_empty(self.tmp)
        rc = self._run_cli([str(exp), "document-build-docx"])
        self.assertEqual(rc, 1)

    def test_valid_markdown_exit_0(self):
        exp = _make_expediente_with_markdown(self.tmp)
        rc = self._run_cli([str(exp), "document-build-docx"])
        self.assertEqual(rc, 0)

    def test_invalid_expediente_exit_1(self):
        rc = self._run_cli(["/ruta/no/existe", "document-build-docx"])
        self.assertEqual(rc, 1)


# ---------------------------------------------------------------------------
# 11. Tests de seguridad metodologica
# ---------------------------------------------------------------------------

class TestSeguridadMetodologica(unittest.TestCase):

    def setUp(self):
        self.tmp_obj = tempfile.TemporaryDirectory()
        self.tmp = Path(self.tmp_obj.name)

    def tearDown(self):
        self.tmp_obj.cleanup()

    def _get_docx_full_text(self, path: Path) -> str:
        from docx import Document
        doc = Document(str(path))
        return "\n".join(p.text for p in doc.paragraphs)

    def test_docx_contains_no_aptitud_administrativa(self):
        md_path = self.tmp / "test.md"
        md_path.write_text(_MINIMAL_MARKDOWN, encoding="utf-8")
        out_path = self.tmp / "output.docx"
        build_docx_from_markdown(md_path, out_path)
        text = self._get_docx_full_text(out_path).lower()
        self.assertIn("aptitud administrativa", text)

    def test_docx_contains_admin_disclaimer(self):
        md_path = self.tmp / "test.md"
        md_path.write_text(_MINIMAL_MARKDOWN, encoding="utf-8")
        out_path = self.tmp / "output.docx"
        build_docx_from_markdown(md_path, out_path)
        text = self._get_docx_full_text(out_path)
        # El disclaimer completo o fragmento debe aparecer
        self.assertIn("No declara aptitud administrativa", text)

    def test_does_not_remove_blockquote_warnings(self):
        md = "# Titulo\n\n> AVISO: Este bloque no puede completarse.\n\nParrafo final."
        md_path = self.tmp / "test.md"
        md_path.write_text(md, encoding="utf-8")
        out_path = self.tmp / "output.docx"
        build_docx_from_markdown(md_path, out_path)
        text = self._get_docx_full_text(out_path)
        self.assertIn("AVISO", text)

    def test_markdown_source_not_modified(self):
        md_path = self.tmp / "test.md"
        md_path.write_text(_MINIMAL_MARKDOWN, encoding="utf-8")
        out_path = self.tmp / "output.docx"
        build_docx_from_markdown(md_path, out_path)
        after = md_path.read_text(encoding="utf-8")
        self.assertEqual(after, _MINIMAL_MARKDOWN)

    def test_constants_defined(self):
        self.assertEqual(DOCX_OUTPUT_FILENAME, "documento_ambiental_borrador.docx")
        self.assertEqual(DOCX_BUILD_RESULT_FILENAME, "docx_build_result.json")
        self.assertEqual(ADMIN_STYLE_PROFILE, "recimetal_admin_reference_v1")
        self.assertIn("heading", SUPPORTED_MARKDOWN_ELEMENTS)
        self.assertIn("table", SUPPORTED_MARKDOWN_ELEMENTS)
        self.assertIn("blockquote", SUPPORTED_MARKDOWN_ELEMENTS)


# ---------------------------------------------------------------------------
# 12. Tests de DocxBuildWarning y DocumentDocxBuildResult
# ---------------------------------------------------------------------------

class TestDataclasses(unittest.TestCase):

    def test_warning_to_dict(self):
        w = DocxBuildWarning(code="TEST", message="msg", source_line=5, recommendation="rec")
        d = w.to_dict()
        self.assertEqual(d["code"], "TEST")
        self.assertEqual(d["message"], "msg")
        self.assertEqual(d["source_line"], 5)

    def test_warning_summary(self):
        w = DocxBuildWarning(code="ERR", message="error message")
        s = w.summary()
        self.assertIn("ERR", s)
        self.assertIn("error message", s)

    def test_result_warning_count(self):
        w = DocxBuildWarning(code="W1", message="aviso")
        r = DocumentDocxBuildResult(
            expediente_id="TEST",
            input_markdown_path="test.md",
            warnings=[w, w],
        )
        self.assertEqual(r.warning_count(), 2)

    def test_result_is_success_false_when_not_generated(self):
        r = DocumentDocxBuildResult(
            expediente_id="TEST",
            input_markdown_path="test.md",
            generated=False,
        )
        self.assertFalse(r.is_success())

    def test_result_is_success_true_when_generated(self):
        r = DocumentDocxBuildResult(
            expediente_id="TEST",
            input_markdown_path="test.md",
            generated=True,
        )
        self.assertTrue(r.is_success())

    def test_result_to_dict_keys(self):
        r = DocumentDocxBuildResult(
            expediente_id="TEST",
            input_markdown_path="test.md",
        )
        d = r.to_dict()
        for key in ("expediente_id", "generated", "paragraph_count", "heading_count",
                    "table_count", "warning_count", "warnings", "notes"):
            self.assertIn(key, d)

    def test_result_summary(self):
        r = DocumentDocxBuildResult(
            expediente_id="EXP-001",
            input_markdown_path="test.md",
            generated=True,
            heading_count=5,
            paragraph_count=20,
        )
        s = r.summary()
        self.assertIn("EXP-001", s)
        self.assertIn("5", s)


# ---------------------------------------------------------------------------
# 13. Tests de write_docx_build_result
# ---------------------------------------------------------------------------

class TestWriteDocxBuildResult(unittest.TestCase):

    def test_writes_json_file(self):
        with tempfile.TemporaryDirectory() as tmp:
            r = DocumentDocxBuildResult(
                expediente_id="TEST",
                input_markdown_path="test.md",
                generated=True,
            )
            out = Path(tmp) / "result.json"
            write_docx_build_result(r, out)
            self.assertTrue(out.exists())

    def test_json_is_valid(self):
        with tempfile.TemporaryDirectory() as tmp:
            r = DocumentDocxBuildResult(
                expediente_id="TEST",
                input_markdown_path="test.md",
                generated=True,
            )
            out = Path(tmp) / "result.json"
            write_docx_build_result(r, out)
            data = json.loads(out.read_text(encoding="utf-8"))
            self.assertTrue(data["generated"])


if __name__ == "__main__":
    unittest.main()
