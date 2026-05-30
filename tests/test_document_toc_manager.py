"""
tests/test_document_toc_manager.py — EN-05
Tests para document_toc_manager.py.

127 tests en 15 clases. 100% offline.
Usa DOCXs sinteticos generados con python-docx en memoria dentro de
tempfile.TemporaryDirectory.
"""
import json
import os
import sys
import tempfile
import zipfile
from pathlib import Path
from unittest import TestCase, main as unittest_main

# Asegurar que src/ esta en el path
sys.path.insert(0, str(Path(__file__).parent.parent / "src"))

from eia_agent.core.document_toc_manager import (
    DEFAULT_TOC_INSTRUCTION,
    TOC_OUTPUT_DOCX,
    TOC_RESULT_JSON,
    TOC_RESULT_MD,
    TOC_SEVERITY,
    TOC_STATUS,
    TocDetectionResult,
    TocIssue,
    TocProcessResult,
    _TOC_PLACEHOLDER_KEYWORDS,
    _compute_toc_status,
    _count_toc_fields,
    _has_update_fields_true,
    _inject_update_fields,
    analyze_toc,
    build_toc_field_paragraph,
    build_toc_report_markdown,
    detect_toc_in_docx,
    enable_update_fields_on_open,
    find_toc_placeholder_paragraphs,
    insert_or_replace_toc,
    process_document_toc,
    validate_docx_file,
    write_toc_outputs,
)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _make_docx(path: Path, paragraphs=None) -> Path:
    """Crea un DOCX sintetico con python-docx."""
    from docx import Document
    doc = Document()
    if paragraphs:
        for text in paragraphs:
            doc.add_paragraph(text)
    doc.save(str(path))
    return path


def _make_docx_with_toc(path: Path) -> Path:
    """Crea un DOCX sintetico con un campo TOC insertado."""
    from docx import Document
    doc = Document()
    doc.add_paragraph("Primer parrafo de prueba")
    build_toc_field_paragraph(doc)
    doc.add_paragraph("Ultimo parrafo de prueba")
    doc.save(str(path))
    return path


def _make_bad_zip(path: Path) -> Path:
    """Crea un archivo con extension .docx que no es ZIP valido."""
    path.write_bytes(b"NOT A ZIP FILE AT ALL")
    return path


def _make_expediente(tmp_dir: Path, docx_name: str = "documento_ambiental_borrador.docx") -> Path:
    """Crea estructura minima de expediente con un DOCX de prueba."""
    exp = tmp_dir / "expediente-TEST"
    doc_dir = exp / "documento"
    doc_dir.mkdir(parents=True)
    _make_docx(doc_dir / docx_name, ["Primer parrafo", "Segundo parrafo"])
    return exp


# ---------------------------------------------------------------------------
# 1. TestTocIssue
# ---------------------------------------------------------------------------


class TestTocIssue(TestCase):
    def test_basic_fields(self):
        issue = TocIssue(severity="ERROR", code="EN05-E001", message="msg")
        self.assertEqual(issue.severity, "ERROR")
        self.assertEqual(issue.code, "EN05-E001")
        self.assertEqual(issue.message, "msg")
        self.assertEqual(issue.recommendation, "")
        self.assertEqual(issue.evidence, [])

    def test_to_dict_keys(self):
        issue = TocIssue(severity="WARNING", code="EN05-W001", message="aviso",
                         recommendation="corregir", evidence=["ev1"])
        d = issue.to_dict()
        self.assertIn("severity", d)
        self.assertIn("code", d)
        self.assertIn("message", d)
        self.assertIn("recommendation", d)
        self.assertIn("evidence", d)

    def test_to_dict_values(self):
        issue = TocIssue(severity="ERROR", code="X", message="Y",
                         recommendation="R", evidence=["e"])
        d = issue.to_dict()
        self.assertEqual(d["severity"], "ERROR")
        self.assertEqual(d["code"], "X")
        self.assertEqual(d["message"], "Y")
        self.assertEqual(d["recommendation"], "R")
        self.assertEqual(d["evidence"], ["e"])

    def test_summary_format(self):
        issue = TocIssue(severity="ERROR", code="EN05-E001", message="texto")
        s = issue.summary()
        self.assertIn("ERROR", s)
        self.assertIn("EN05-E001", s)
        self.assertIn("texto", s)

    def test_warning_severity(self):
        issue = TocIssue(severity="WARNING", code="EN05-W001", message="w")
        self.assertEqual(issue.severity, "WARNING")
        self.assertIn("WARNING", issue.summary())

    def test_info_severity(self):
        issue = TocIssue(severity="INFO", code="EN05-I001", message="i")
        self.assertIn("INFO", issue.summary())

    def test_evidence_list(self):
        issue = TocIssue(severity="ERROR", code="X", message="Y",
                         evidence=["e1", "e2", "e3"])
        self.assertEqual(len(issue.evidence), 3)
        self.assertEqual(issue.to_dict()["evidence"], ["e1", "e2", "e3"])

    def test_default_evidence_is_empty_list(self):
        issue = TocIssue(severity="ERROR", code="X", message="Y")
        self.assertIsInstance(issue.evidence, list)
        self.assertEqual(issue.evidence, [])


# ---------------------------------------------------------------------------
# 2. TestTocDetectionResult
# ---------------------------------------------------------------------------


class TestTocDetectionResult(TestCase):
    def test_basic_fields(self):
        r = TocDetectionResult(
            docx_path="/a.docx",
            has_toc=True,
            update_fields_enabled=False,
            toc_paragraph_count=1,
            has_settings_xml=True,
        )
        self.assertTrue(r.has_toc)
        self.assertFalse(r.update_fields_enabled)
        self.assertEqual(r.toc_paragraph_count, 1)
        self.assertTrue(r.has_settings_xml)

    def test_to_dict_keys(self):
        r = TocDetectionResult(
            docx_path="x", has_toc=False,
            update_fields_enabled=False, toc_paragraph_count=0,
            has_settings_xml=False,
        )
        d = r.to_dict()
        for key in ["docx_path", "has_toc", "update_fields_enabled",
                    "toc_paragraph_count", "has_settings_xml", "issues", "notes"]:
            self.assertIn(key, d)

    def test_to_dict_values(self):
        r = TocDetectionResult(
            docx_path="x.docx", has_toc=True,
            update_fields_enabled=True, toc_paragraph_count=2,
            has_settings_xml=True,
        )
        d = r.to_dict()
        self.assertTrue(d["has_toc"])
        self.assertTrue(d["update_fields_enabled"])
        self.assertEqual(d["toc_paragraph_count"], 2)

    def test_summary_contains_toc(self):
        r = TocDetectionResult(
            docx_path="x", has_toc=True,
            update_fields_enabled=True, toc_paragraph_count=3,
            has_settings_xml=True,
        )
        s = r.summary()
        self.assertIn("si", s)
        self.assertIn("3", s)

    def test_summary_no_toc(self):
        r = TocDetectionResult(
            docx_path="x", has_toc=False,
            update_fields_enabled=False, toc_paragraph_count=0,
            has_settings_xml=False,
        )
        self.assertIn("no", r.summary())

    def test_issues_serialized(self):
        r = TocDetectionResult(
            docx_path="x", has_toc=False,
            update_fields_enabled=False, toc_paragraph_count=0,
            has_settings_xml=False,
            issues=[TocIssue("ERROR", "X", "Y")],
        )
        d = r.to_dict()
        self.assertEqual(len(d["issues"]), 1)

    def test_notes_in_dict(self):
        r = TocDetectionResult(
            docx_path="x", has_toc=False,
            update_fields_enabled=False, toc_paragraph_count=0,
            has_settings_xml=False, notes=["nota1"],
        )
        self.assertEqual(r.to_dict()["notes"], ["nota1"])

    def test_default_issues_empty(self):
        r = TocDetectionResult(
            docx_path="x", has_toc=False,
            update_fields_enabled=False, toc_paragraph_count=0,
            has_settings_xml=False,
        )
        self.assertEqual(r.issues, [])


# ---------------------------------------------------------------------------
# 3. TestTocProcessResult
# ---------------------------------------------------------------------------


class TestTocProcessResult(TestCase):
    def _make_result(self, **kwargs):
        defaults = dict(
            input_docx="in.docx", output_docx=None, status="OK",
            toc_inserted=False, toc_replaced=False,
            update_fields_set=False, placeholder_paragraphs_found=0,
        )
        defaults.update(kwargs)
        return TocProcessResult(**defaults)

    def test_basic_fields(self):
        r = self._make_result(toc_inserted=True, status="OK")
        self.assertTrue(r.toc_inserted)
        self.assertEqual(r.status, "OK")

    def test_error_count_zero(self):
        r = self._make_result()
        self.assertEqual(r.error_count(), 0)

    def test_error_count_with_errors(self):
        r = self._make_result(issues=[
            TocIssue("ERROR", "X", "m"),
            TocIssue("WARNING", "Y", "w"),
        ])
        self.assertEqual(r.error_count(), 1)

    def test_warning_count_from_issues(self):
        r = self._make_result(issues=[TocIssue("WARNING", "X", "m")])
        self.assertEqual(r.warning_count(), 1)

    def test_warning_count_from_warnings_list(self):
        r = self._make_result(warnings=["aviso1", "aviso2"])
        self.assertEqual(r.warning_count(), 2)

    def test_warning_count_combined(self):
        r = self._make_result(
            issues=[TocIssue("WARNING", "X", "m")],
            warnings=["aviso1"],
        )
        self.assertEqual(r.warning_count(), 2)

    def test_is_valid_no_errors(self):
        r = self._make_result()
        self.assertTrue(r.is_valid())

    def test_is_valid_with_error(self):
        r = self._make_result(issues=[TocIssue("ERROR", "X", "m")])
        self.assertFalse(r.is_valid())

    def test_is_valid_warning_only(self):
        r = self._make_result(issues=[TocIssue("WARNING", "X", "m")])
        self.assertTrue(r.is_valid())

    def test_to_dict_keys(self):
        r = self._make_result()
        d = r.to_dict()
        for key in ["input_docx", "output_docx", "status", "toc_inserted",
                    "toc_replaced", "update_fields_set", "placeholder_paragraphs_found",
                    "issues", "warnings", "notes"]:
            self.assertIn(key, d)

    def test_summary_insertado(self):
        r = self._make_result(toc_inserted=True, status="OK")
        self.assertIn("insertado", r.summary())

    def test_summary_reemplazado(self):
        r = self._make_result(toc_replaced=True, status="OK")
        self.assertIn("reemplazado", r.summary())

    def test_summary_no_insertado(self):
        r = self._make_result(status="OK")
        self.assertIn("no insertado", r.summary())

    def test_summary_contains_status(self):
        r = self._make_result(status="NO_CONFORME")
        self.assertIn("NO_CONFORME", r.summary())

    def test_to_dict_issues_serialized(self):
        r = self._make_result(issues=[TocIssue("ERROR", "X", "Y")])
        d = r.to_dict()
        self.assertEqual(len(d["issues"]), 1)
        self.assertEqual(d["issues"][0]["code"], "X")


# ---------------------------------------------------------------------------
# 4. TestValidateDocxFile
# ---------------------------------------------------------------------------


class TestValidateDocxFile(TestCase):
    def test_missing_file(self):
        self.assertFalse(validate_docx_file("/nonexistent/path/file.docx"))

    def test_empty_file(self):
        with tempfile.TemporaryDirectory() as td:
            p = Path(td) / "empty.docx"
            p.write_bytes(b"")
            self.assertFalse(validate_docx_file(p))

    def test_valid_docx(self):
        with tempfile.TemporaryDirectory() as td:
            p = Path(td) / "valid.docx"
            _make_docx(p)
            self.assertTrue(validate_docx_file(p))

    def test_bad_zip_file(self):
        with tempfile.TemporaryDirectory() as td:
            p = Path(td) / "bad.docx"
            _make_bad_zip(p)
            self.assertFalse(validate_docx_file(p))

    def test_nonexistent_directory(self):
        self.assertFalse(validate_docx_file("/no/such/dir/file.docx"))


# ---------------------------------------------------------------------------
# 5. TestCountTocFields
# ---------------------------------------------------------------------------


class TestCountTocFields(TestCase):
    def test_no_toc(self):
        xml = "<w:body><w:p><w:r><w:t>texto</w:t></w:r></w:p></w:body>"
        self.assertEqual(_count_toc_fields(xml), 0)

    def test_one_instrtext_toc(self):
        xml = '<w:instrText xml:space="preserve"> TOC \\o "1-3" \\h \\z </w:instrText>'
        self.assertEqual(_count_toc_fields(xml), 1)

    def test_one_fldSimple_toc(self):
        xml = '<w:fldSimple w:instr=" TOC \\o &quot;1-3&quot; ">'
        self.assertEqual(_count_toc_fields(xml), 1)

    def test_multiple_instrtext_toc(self):
        xml = (
            '<w:instrText>TOC \\o "1-3"</w:instrText>'
            '<w:instrText>TOC \\h</w:instrText>'
        )
        self.assertEqual(_count_toc_fields(xml), 2)

    def test_non_toc_instrtext(self):
        xml = '<w:instrText>DATE</w:instrText>'
        self.assertEqual(_count_toc_fields(xml), 0)


# ---------------------------------------------------------------------------
# 6. TestHasUpdateFieldsTrue
# ---------------------------------------------------------------------------


class TestHasUpdateFieldsTrue(TestCase):
    def test_true_value(self):
        xml = '<w:updateFields w:val="true"/>'
        self.assertTrue(_has_update_fields_true(xml))

    def test_1_value(self):
        xml = '<w:updateFields w:val="1"/>'
        self.assertTrue(_has_update_fields_true(xml))

    def test_on_value(self):
        xml = '<w:updateFields w:val="on"/>'
        self.assertTrue(_has_update_fields_true(xml))

    def test_false_value(self):
        xml = '<w:updateFields w:val="false"/>'
        self.assertFalse(_has_update_fields_true(xml))

    def test_no_update_fields(self):
        xml = '<w:settings><w:zoom w:percent="100"/></w:settings>'
        self.assertFalse(_has_update_fields_true(xml))

    def test_case_insensitive(self):
        xml = '<w:updateFields w:val="TRUE"/>'
        self.assertTrue(_has_update_fields_true(xml))


# ---------------------------------------------------------------------------
# 7. TestInjectUpdateFields
# ---------------------------------------------------------------------------


class TestInjectUpdateFields(TestCase):
    def test_adds_update_fields(self):
        xml = b'<?xml version=\'1.0\'?><w:settings xmlns:w="x"><w:zoom/></w:settings>'
        result = _inject_update_fields(xml).decode("utf-8")
        self.assertIn("w:updateFields", result)
        self.assertIn('w:val="true"', result)

    def test_replaces_existing_false(self):
        xml = b'<w:settings xmlns:w="x"><w:updateFields w:val="false"/></w:settings>'
        result = _inject_update_fields(xml).decode("utf-8")
        self.assertIn('w:val="true"', result)
        self.assertNotIn('w:val="false"', result)

    def test_preserves_existing_true(self):
        xml = b'<w:settings xmlns:w="x"><w:updateFields w:val="true"/></w:settings>'
        result = _inject_update_fields(xml).decode("utf-8")
        self.assertIn('w:val="true"', result)

    def test_returns_bytes(self):
        xml = b'<w:settings xmlns:w="x"></w:settings>'
        self.assertIsInstance(_inject_update_fields(xml), bytes)


# ---------------------------------------------------------------------------
# 8. TestDetectTocInDocx
# ---------------------------------------------------------------------------


class TestDetectTocInDocx(TestCase):
    def test_missing_file(self):
        r = detect_toc_in_docx("/no/such/file.docx")
        self.assertFalse(r.has_toc)
        self.assertFalse(r.update_fields_enabled)
        self.assertTrue(any(i.code == "EN05-E001" for i in r.issues))

    def test_bad_zip(self):
        with tempfile.TemporaryDirectory() as td:
            p = Path(td) / "bad.docx"
            _make_bad_zip(p)
            r = detect_toc_in_docx(p)
            self.assertFalse(r.has_toc)
            self.assertTrue(any("EN05-E001" in i.code for i in r.issues))

    def test_valid_docx_no_toc(self):
        with tempfile.TemporaryDirectory() as td:
            p = Path(td) / "notoc.docx"
            _make_docx(p, ["Hola mundo", "Segundo parrafo"])
            r = detect_toc_in_docx(p)
            self.assertFalse(r.has_toc)
            self.assertEqual(r.toc_paragraph_count, 0)
            self.assertEqual(len(r.issues), 0)

    def test_valid_docx_with_toc(self):
        with tempfile.TemporaryDirectory() as td:
            p = Path(td) / "withtoc.docx"
            _make_docx_with_toc(p)
            r = detect_toc_in_docx(p)
            self.assertTrue(r.has_toc)
            self.assertGreater(r.toc_paragraph_count, 0)

    def test_has_settings_xml(self):
        with tempfile.TemporaryDirectory() as td:
            p = Path(td) / "doc.docx"
            _make_docx(p)
            r = detect_toc_in_docx(p)
            self.assertTrue(r.has_settings_xml)

    def test_no_update_fields_by_default(self):
        with tempfile.TemporaryDirectory() as td:
            p = Path(td) / "doc.docx"
            _make_docx(p)
            r = detect_toc_in_docx(p)
            self.assertFalse(r.update_fields_enabled)

    def test_update_fields_after_enable(self):
        with tempfile.TemporaryDirectory() as td:
            src = Path(td) / "src.docx"
            dst = Path(td) / "dst.docx"
            _make_docx(src)
            enable_update_fields_on_open(src, dst)
            r = detect_toc_in_docx(dst)
            self.assertTrue(r.update_fields_enabled)

    def test_to_dict_serializable(self):
        with tempfile.TemporaryDirectory() as td:
            p = Path(td) / "doc.docx"
            _make_docx(p)
            r = detect_toc_in_docx(p)
            d = r.to_dict()
            json.dumps(d)  # no debe lanzar excepcion

    def test_docx_path_in_result(self):
        with tempfile.TemporaryDirectory() as td:
            p = Path(td) / "doc.docx"
            _make_docx(p)
            r = detect_toc_in_docx(p)
            self.assertEqual(r.docx_path, str(p))

    def test_notes_populated_no_toc(self):
        with tempfile.TemporaryDirectory() as td:
            p = Path(td) / "doc.docx"
            _make_docx(p, ["texto"])
            r = detect_toc_in_docx(p)
            self.assertGreater(len(r.notes), 0)


# ---------------------------------------------------------------------------
# 9. TestEnableUpdateFieldsOnOpen
# ---------------------------------------------------------------------------


class TestEnableUpdateFieldsOnOpen(TestCase):
    def test_missing_src(self):
        result = enable_update_fields_on_open("/no/file.docx", "/out/file.docx")
        self.assertFalse(result)

    def test_creates_output(self):
        with tempfile.TemporaryDirectory() as td:
            src = Path(td) / "src.docx"
            dst = Path(td) / "dst.docx"
            _make_docx(src)
            ok = enable_update_fields_on_open(src, dst)
            self.assertTrue(ok)
            self.assertTrue(dst.exists())

    def test_output_is_valid_docx(self):
        with tempfile.TemporaryDirectory() as td:
            src = Path(td) / "src.docx"
            dst = Path(td) / "dst.docx"
            _make_docx(src)
            enable_update_fields_on_open(src, dst)
            self.assertTrue(validate_docx_file(dst))

    def test_update_fields_detectable_after(self):
        with tempfile.TemporaryDirectory() as td:
            src = Path(td) / "src.docx"
            dst = Path(td) / "dst.docx"
            _make_docx(src)
            enable_update_fields_on_open(src, dst)
            r = detect_toc_in_docx(dst)
            self.assertTrue(r.update_fields_enabled)

    def test_same_file_in_place(self):
        with tempfile.TemporaryDirectory() as td:
            p = Path(td) / "inplace.docx"
            _make_docx(p)
            ok = enable_update_fields_on_open(p, p)
            self.assertTrue(ok)
            self.assertTrue(validate_docx_file(p))

    def test_idempotent(self):
        with tempfile.TemporaryDirectory() as td:
            src = Path(td) / "src.docx"
            dst = Path(td) / "dst.docx"
            _make_docx(src)
            enable_update_fields_on_open(src, dst)
            ok = enable_update_fields_on_open(dst, dst)
            self.assertTrue(ok)
            r = detect_toc_in_docx(dst)
            self.assertTrue(r.update_fields_enabled)

    def test_src_not_modified(self):
        with tempfile.TemporaryDirectory() as td:
            src = Path(td) / "src.docx"
            dst = Path(td) / "dst.docx"
            _make_docx(src)
            mtime_before = src.stat().st_mtime
            enable_update_fields_on_open(src, dst)
            mtime_after = src.stat().st_mtime
            self.assertAlmostEqual(mtime_before, mtime_after, places=0)

    def test_returns_true_on_success(self):
        with tempfile.TemporaryDirectory() as td:
            src = Path(td) / "src.docx"
            dst = Path(td) / "dst.docx"
            _make_docx(src)
            result = enable_update_fields_on_open(src, dst)
            self.assertTrue(result)


# ---------------------------------------------------------------------------
# 10. TestBuildTocFieldParagraph
# ---------------------------------------------------------------------------


class TestBuildTocFieldParagraph(TestCase):
    def _get_doc_with_toc(self):
        from docx import Document
        doc = Document()
        build_toc_field_paragraph(doc)
        return doc

    def test_adds_paragraph(self):
        from docx import Document
        doc = Document()
        initial = len(doc.paragraphs)
        build_toc_field_paragraph(doc)
        self.assertEqual(len(doc.paragraphs), initial + 1)

    def test_returns_paragraph(self):
        from docx import Document
        from docx.text.paragraph import Paragraph
        doc = Document()
        para = build_toc_field_paragraph(doc)
        self.assertIsInstance(para, Paragraph)

    def test_paragraph_has_run(self):
        from docx import Document
        doc = Document()
        para = build_toc_field_paragraph(doc)
        self.assertGreater(len(para.runs), 0)

    def test_toc_instruction_in_xml(self):
        from docx import Document
        doc = Document()
        para = build_toc_field_paragraph(doc)
        xml_str = para._p.xml
        self.assertIn("TOC", xml_str)

    def test_fldchar_begin_in_xml(self):
        from docx import Document
        doc = Document()
        para = build_toc_field_paragraph(doc)
        xml_str = para._p.xml
        self.assertIn("begin", xml_str)

    def test_fldchar_end_in_xml(self):
        from docx import Document
        doc = Document()
        para = build_toc_field_paragraph(doc)
        xml_str = para._p.xml
        self.assertIn("end", xml_str)

    def test_instrtext_in_xml(self):
        from docx import Document
        doc = Document()
        para = build_toc_field_paragraph(doc)
        xml_str = para._p.xml
        self.assertIn("instrText", xml_str)

    def test_detectable_after_save(self):
        with tempfile.TemporaryDirectory() as td:
            p = Path(td) / "toc.docx"
            _make_docx_with_toc(p)
            r = detect_toc_in_docx(p)
            self.assertTrue(r.has_toc)


# ---------------------------------------------------------------------------
# 11. TestFindTocPlaceholderParagraphs
# ---------------------------------------------------------------------------


class TestFindTocPlaceholderParagraphs(TestCase):
    def _make_doc(self, paragraphs):
        from docx import Document
        doc = Document()
        for text in paragraphs:
            doc.add_paragraph(text)
        return doc

    def test_no_paragraphs(self):
        from docx import Document
        doc = Document()
        result = find_toc_placeholder_paragraphs(doc)
        self.assertEqual(result, [])

    def test_indice_keyword(self):
        doc = self._make_doc(["Índice"])
        result = find_toc_placeholder_paragraphs(doc)
        self.assertGreater(len(result), 0)

    def test_indice_lowercase(self):
        doc = self._make_doc(["indice"])
        result = find_toc_placeholder_paragraphs(doc)
        self.assertGreater(len(result), 0)

    def test_toc_keyword(self):
        doc = self._make_doc(["TOC"])
        result = find_toc_placeholder_paragraphs(doc)
        self.assertGreater(len(result), 0)

    def test_tabla_de_contenido(self):
        doc = self._make_doc(["Tabla de contenido"])
        result = find_toc_placeholder_paragraphs(doc)
        self.assertGreater(len(result), 0)

    def test_long_paragraph_excluded(self):
        # Parrafo muy largo con "indice" no debe ser candidato
        long_text = "El indice de biodiversidad refleja la riqueza " + "x" * 100
        doc = self._make_doc([long_text])
        result = find_toc_placeholder_paragraphs(doc)
        self.assertEqual(result, [])

    def test_empty_paragraph_excluded(self):
        doc = self._make_doc([""])
        result = find_toc_placeholder_paragraphs(doc)
        self.assertEqual(result, [])

    def test_returns_correct_index(self):
        doc = self._make_doc(["Primer parrafo", "Índice", "Tercer parrafo"])
        result = find_toc_placeholder_paragraphs(doc)
        self.assertIn(1, result)

    def test_multiple_placeholders(self):
        doc = self._make_doc(["Índice", "contenido normal", "TOC"])
        result = find_toc_placeholder_paragraphs(doc)
        self.assertGreaterEqual(len(result), 1)

    def test_case_insensitive_keyword(self):
        doc = self._make_doc(["INDICE"])
        result = find_toc_placeholder_paragraphs(doc)
        self.assertGreater(len(result), 0)


# ---------------------------------------------------------------------------
# 12. TestInsertOrReplaceToc
# ---------------------------------------------------------------------------


class TestInsertOrReplaceToc(TestCase):
    def test_missing_input(self):
        with tempfile.TemporaryDirectory() as td:
            out = Path(td) / "out.docx"
            r = insert_or_replace_toc("/no/such/file.docx", out)
            self.assertFalse(r.is_valid())
            self.assertIsNone(r.output_docx)

    def test_insert_no_placeholder(self):
        with tempfile.TemporaryDirectory() as td:
            src = Path(td) / "src.docx"
            dst = Path(td) / "dst.docx"
            _make_docx(src, ["Primer parrafo", "Segundo parrafo"])
            r = insert_or_replace_toc(src, dst)
            self.assertTrue(r.is_valid())
            self.assertTrue(r.toc_inserted)
            self.assertFalse(r.toc_replaced)
            self.assertTrue(dst.exists())

    def test_replace_placeholder(self):
        with tempfile.TemporaryDirectory() as td:
            src = Path(td) / "src.docx"
            dst = Path(td) / "dst.docx"
            _make_docx(src, ["Primer parrafo", "Índice", "Tercer parrafo"])
            r = insert_or_replace_toc(src, dst)
            self.assertTrue(r.is_valid())
            self.assertTrue(r.toc_replaced)
            self.assertFalse(r.toc_inserted)

    def test_no_replace_flag(self):
        with tempfile.TemporaryDirectory() as td:
            src = Path(td) / "src.docx"
            dst = Path(td) / "dst.docx"
            _make_docx(src, ["Primer parrafo", "Índice"])
            r = insert_or_replace_toc(src, dst, replace_placeholder=False)
            self.assertTrue(r.toc_inserted)
            self.assertFalse(r.toc_replaced)

    def test_output_docx_contains_toc(self):
        with tempfile.TemporaryDirectory() as td:
            src = Path(td) / "src.docx"
            dst = Path(td) / "dst.docx"
            _make_docx(src, ["Primer parrafo"])
            insert_or_replace_toc(src, dst)
            detection = detect_toc_in_docx(dst)
            self.assertTrue(detection.has_toc)

    def test_src_not_modified(self):
        with tempfile.TemporaryDirectory() as td:
            src = Path(td) / "src.docx"
            dst = Path(td) / "dst.docx"
            _make_docx(src, ["texto"])
            mtime_before = src.stat().st_mtime
            insert_or_replace_toc(src, dst)
            self.assertAlmostEqual(src.stat().st_mtime, mtime_before, places=0)

    def test_update_fields_set_in_output(self):
        with tempfile.TemporaryDirectory() as td:
            src = Path(td) / "src.docx"
            dst = Path(td) / "dst.docx"
            _make_docx(src, ["texto"])
            r = insert_or_replace_toc(src, dst)
            self.assertTrue(r.update_fields_set)

    def test_output_docx_path_in_result(self):
        with tempfile.TemporaryDirectory() as td:
            src = Path(td) / "src.docx"
            dst = Path(td) / "dst.docx"
            _make_docx(src, ["texto"])
            r = insert_or_replace_toc(src, dst)
            self.assertEqual(r.output_docx, str(dst))

    def test_placeholder_count_in_result(self):
        with tempfile.TemporaryDirectory() as td:
            src = Path(td) / "src.docx"
            dst = Path(td) / "dst.docx"
            _make_docx(src, ["Índice", "otro"])
            r = insert_or_replace_toc(src, dst)
            self.assertGreaterEqual(r.placeholder_paragraphs_found, 1)

    def test_status_ok_on_success(self):
        with tempfile.TemporaryDirectory() as td:
            src = Path(td) / "src.docx"
            dst = Path(td) / "dst.docx"
            _make_docx(src, ["texto"])
            r = insert_or_replace_toc(src, dst)
            self.assertEqual(r.status, TOC_STATUS["OK"])


# ---------------------------------------------------------------------------
# 13. TestAnalyzeToc
# ---------------------------------------------------------------------------


class TestAnalyzeToc(TestCase):
    def test_missing_file(self):
        r = analyze_toc("/no/such/file.docx")
        self.assertFalse(r.is_valid())
        self.assertIsNone(r.output_docx)

    def test_valid_docx_no_toc(self):
        with tempfile.TemporaryDirectory() as td:
            p = Path(td) / "doc.docx"
            _make_docx(p, ["texto"])
            r = analyze_toc(p)
            self.assertTrue(r.is_valid())
            self.assertIsNone(r.output_docx)
            self.assertFalse(r.toc_inserted)
            self.assertFalse(r.toc_replaced)

    def test_valid_docx_with_placeholder(self):
        with tempfile.TemporaryDirectory() as td:
            p = Path(td) / "doc.docx"
            _make_docx(p, ["Índice", "texto"])
            r = analyze_toc(p)
            self.assertGreaterEqual(r.placeholder_paragraphs_found, 1)

    def test_valid_docx_with_toc(self):
        with tempfile.TemporaryDirectory() as td:
            p = Path(td) / "doc.docx"
            _make_docx_with_toc(p)
            r = analyze_toc(p)
            self.assertTrue(r.is_valid())

    def test_update_fields_false_by_default(self):
        with tempfile.TemporaryDirectory() as td:
            p = Path(td) / "doc.docx"
            _make_docx(p)
            r = analyze_toc(p)
            self.assertFalse(r.update_fields_set)

    def test_to_dict_serializable(self):
        with tempfile.TemporaryDirectory() as td:
            p = Path(td) / "doc.docx"
            _make_docx(p, ["texto"])
            r = analyze_toc(p)
            json.dumps(r.to_dict())  # no debe lanzar excepcion


# ---------------------------------------------------------------------------
# 14. TestComputeTocStatus
# ---------------------------------------------------------------------------


class TestComputeTocStatus(TestCase):
    def test_ok_no_issues(self):
        self.assertEqual(_compute_toc_status([], []), TOC_STATUS["OK"])

    def test_no_conforme_with_error(self):
        issues = [TocIssue("ERROR", "X", "m")]
        self.assertEqual(_compute_toc_status(issues, []), TOC_STATUS["NO_CONFORME"])

    def test_con_observaciones_warning_issue(self):
        issues = [TocIssue("WARNING", "X", "m")]
        self.assertEqual(_compute_toc_status(issues, []), TOC_STATUS["CON_OBSERVACIONES"])

    def test_con_observaciones_warnings_list(self):
        self.assertEqual(_compute_toc_status([], ["aviso"]), TOC_STATUS["CON_OBSERVACIONES"])

    def test_error_takes_precedence(self):
        issues = [TocIssue("ERROR", "X", "m"), TocIssue("WARNING", "Y", "w")]
        self.assertEqual(_compute_toc_status(issues, ["aviso"]), TOC_STATUS["NO_CONFORME"])


# ---------------------------------------------------------------------------
# 15. TestBuildTocReportMarkdown
# ---------------------------------------------------------------------------


class TestBuildTocReportMarkdown(TestCase):
    def _make_result(self, **kwargs):
        defaults = dict(
            input_docx="in.docx", output_docx=None, status="OK",
            toc_inserted=False, toc_replaced=False,
            update_fields_set=False, placeholder_paragraphs_found=0,
        )
        defaults.update(kwargs)
        return TocProcessResult(**defaults)

    def test_contains_en05(self):
        r = self._make_result()
        md = build_toc_report_markdown(r)
        self.assertIn("EN-05", md)

    def test_contains_status(self):
        r = self._make_result(status="NO_CONFORME")
        md = build_toc_report_markdown(r)
        self.assertIn("NO_CONFORME", md)

    def test_contains_input_docx(self):
        r = self._make_result(input_docx="/ruta/al/doc.docx")
        md = build_toc_report_markdown(r)
        self.assertIn("/ruta/al/doc.docx", md)

    def test_contains_issues_section(self):
        r = self._make_result(issues=[TocIssue("ERROR", "EN05-E001", "mensaje")])
        md = build_toc_report_markdown(r)
        self.assertIn("EN05-E001", md)
        self.assertIn("mensaje", md)

    def test_contains_advertencia(self):
        r = self._make_result()
        md = build_toc_report_markdown(r)
        self.assertIn("Advertencia", md)

    def test_contains_warnings_section(self):
        r = self._make_result(warnings=["aviso especial"])
        md = build_toc_report_markdown(r)
        self.assertIn("aviso especial", md)

    def test_returns_string(self):
        r = self._make_result()
        md = build_toc_report_markdown(r)
        self.assertIsInstance(md, str)

    def test_contains_placeholder_count(self):
        r = self._make_result(placeholder_paragraphs_found=3)
        md = build_toc_report_markdown(r)
        self.assertIn("3", md)


# ---------------------------------------------------------------------------
# 16. TestWriteTocOutputs
# ---------------------------------------------------------------------------


class TestWriteTocOutputs(TestCase):
    def _make_result(self):
        return TocProcessResult(
            input_docx="in.docx", output_docx=None, status="OK",
            toc_inserted=False, toc_replaced=False,
            update_fields_set=False, placeholder_paragraphs_found=0,
        )

    def test_writes_json(self):
        with tempfile.TemporaryDirectory() as td:
            out_dir = Path(td) / "documento"
            r = self._make_result()
            json_path, _ = write_toc_outputs(r, out_dir)
            self.assertTrue(json_path.exists())

    def test_writes_md(self):
        with tempfile.TemporaryDirectory() as td:
            out_dir = Path(td) / "documento"
            r = self._make_result()
            _, md_path = write_toc_outputs(r, out_dir)
            self.assertTrue(md_path.exists())

    def test_json_valid(self):
        with tempfile.TemporaryDirectory() as td:
            out_dir = Path(td) / "documento"
            r = self._make_result()
            json_path, _ = write_toc_outputs(r, out_dir)
            with open(json_path, encoding="utf-8") as f:
                data = json.load(f)
            self.assertIn("status", data)

    def test_filenames(self):
        with tempfile.TemporaryDirectory() as td:
            out_dir = Path(td) / "documento"
            r = self._make_result()
            json_path, md_path = write_toc_outputs(r, out_dir)
            self.assertEqual(json_path.name, TOC_RESULT_JSON)
            self.assertEqual(md_path.name, TOC_RESULT_MD)

    def test_creates_dir(self):
        with tempfile.TemporaryDirectory() as td:
            out_dir = Path(td) / "nuevo" / "documento"
            r = self._make_result()
            write_toc_outputs(r, out_dir)
            self.assertTrue(out_dir.exists())


# ---------------------------------------------------------------------------
# 17. TestProcessDocumentToc
# ---------------------------------------------------------------------------


class TestProcessDocumentToc(TestCase):
    def test_no_docx_returns_sin_datos(self):
        with tempfile.TemporaryDirectory() as td:
            exp = Path(td) / "vacio"
            exp.mkdir()
            r = process_document_toc(exp)
            self.assertEqual(r.status, TOC_STATUS["SIN_DATOS"])

    def test_analyze_only(self):
        with tempfile.TemporaryDirectory() as td:
            exp = _make_expediente(Path(td))
            r = process_document_toc(exp)
            self.assertFalse(r.toc_inserted)
            self.assertFalse(r.toc_replaced)
            self.assertIsNone(r.output_docx)

    def test_apply_toc(self):
        with tempfile.TemporaryDirectory() as td:
            exp = _make_expediente(Path(td))
            r = process_document_toc(exp, apply_toc=True)
            self.assertTrue(r.toc_inserted or r.toc_replaced)
            self.assertIsNotNone(r.output_docx)

    def test_apply_creates_output_file(self):
        with tempfile.TemporaryDirectory() as td:
            exp = _make_expediente(Path(td))
            r = process_document_toc(exp, apply_toc=True)
            if r.output_docx:
                self.assertTrue(Path(r.output_docx).exists())

    def test_write_outputs(self):
        with tempfile.TemporaryDirectory() as td:
            exp = _make_expediente(Path(td))
            process_document_toc(exp, write_outputs=True)
            json_path = exp / "documento" / TOC_RESULT_JSON
            md_path = exp / "documento" / TOC_RESULT_MD
            self.assertTrue(json_path.exists())
            self.assertTrue(md_path.exists())

    def test_write_apply_combined(self):
        with tempfile.TemporaryDirectory() as td:
            exp = _make_expediente(Path(td))
            r = process_document_toc(exp, write_outputs=True, apply_toc=True)
            self.assertTrue(r.toc_inserted or r.toc_replaced)
            json_path = exp / "documento" / TOC_RESULT_JSON
            self.assertTrue(json_path.exists())

    def test_prefers_numerado_docx(self):
        with tempfile.TemporaryDirectory() as td:
            exp = Path(td) / "exp"
            doc_dir = exp / "documento"
            doc_dir.mkdir(parents=True)
            _make_docx(doc_dir / "documento_ambiental_borrador.docx", ["borrador"])
            _make_docx(doc_dir / "documento_ambiental_numerado.docx", ["numerado"])
            r = process_document_toc(exp)
            self.assertIn("numerado", r.input_docx)

    def test_replace_placeholder_false(self):
        with tempfile.TemporaryDirectory() as td:
            exp = _make_expediente(Path(td))
            # Agregar un placeholder al DOCX
            doc_dir = exp / "documento"
            _make_docx(doc_dir / "documento_ambiental_borrador.docx", ["Índice", "texto"])
            r = process_document_toc(exp, apply_toc=True, replace_placeholder=False)
            if r.toc_inserted or r.toc_replaced:
                self.assertTrue(r.toc_inserted)


# ---------------------------------------------------------------------------
# 18. TestConstantes
# ---------------------------------------------------------------------------


class TestConstantes(TestCase):
    def test_toc_status_keys(self):
        for k in ["OK", "CON_OBSERVACIONES", "NO_CONFORME", "SIN_DATOS"]:
            self.assertIn(k, TOC_STATUS)

    def test_toc_severity_keys(self):
        for k in ["ERROR", "WARNING", "INFO"]:
            self.assertIn(k, TOC_SEVERITY)

    def test_default_toc_instruction_has_toc(self):
        self.assertIn("TOC", DEFAULT_TOC_INSTRUCTION)

    def test_output_docx_name(self):
        self.assertTrue(TOC_OUTPUT_DOCX.endswith(".docx"))

    def test_result_json_name(self):
        self.assertTrue(TOC_RESULT_JSON.endswith(".json"))

    def test_result_md_name(self):
        self.assertTrue(TOC_RESULT_MD.endswith(".md"))

    def test_placeholder_keywords_not_empty(self):
        self.assertGreater(len(_TOC_PLACEHOLDER_KEYWORDS), 0)


# ---------------------------------------------------------------------------
# 19. TestCLIDocumentToc
# ---------------------------------------------------------------------------


class TestCLIDocumentToc(TestCase):
    """Prueba la integracion con run_expediente.py via main()."""

    def _run_cli(self, args):
        import run_expediente
        return run_expediente.main(args)

    def test_analyze_no_docx_returns_1(self):
        with tempfile.TemporaryDirectory() as td:
            exp = Path(td) / "vacio"
            exp.mkdir()
            code = self._run_cli([str(exp), "document-toc"])
            self.assertEqual(code, 1)

    def test_analyze_valid_docx_returns_0(self):
        with tempfile.TemporaryDirectory() as td:
            exp = _make_expediente(Path(td))
            code = self._run_cli([str(exp), "document-toc"])
            self.assertEqual(code, 0)

    def test_write_flag_creates_outputs(self):
        with tempfile.TemporaryDirectory() as td:
            exp = _make_expediente(Path(td))
            self._run_cli([str(exp), "document-toc", "--write"])
            self.assertTrue((exp / "documento" / TOC_RESULT_JSON).exists())

    def test_apply_flag_creates_toc_docx(self):
        with tempfile.TemporaryDirectory() as td:
            exp = _make_expediente(Path(td))
            self._run_cli([str(exp), "document-toc", "--apply"])
            self.assertTrue((exp / "documento" / TOC_OUTPUT_DOCX).exists())

    def test_apply_write_combined(self):
        with tempfile.TemporaryDirectory() as td:
            exp = _make_expediente(Path(td))
            self._run_cli([str(exp), "document-toc", "--apply", "--write"])
            self.assertTrue((exp / "documento" / TOC_OUTPUT_DOCX).exists())
            self.assertTrue((exp / "documento" / TOC_RESULT_JSON).exists())

    def test_no_replace_flag(self):
        with tempfile.TemporaryDirectory() as td:
            exp = _make_expediente(Path(td))
            code = self._run_cli([str(exp), "document-toc", "--apply", "--no-replace"])
            self.assertEqual(code, 0)


if __name__ == "__main__":
    unittest_main()
