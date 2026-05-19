"""Tests para CL-05 — climogram_docx_inserter.py

Sin AEMET real. Sin red. DOCX y PNG sintéticos en directorios temporales.
"""
import struct
import sys
import tempfile
import unittest
import zlib
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent / "src"))

try:
    import matplotlib  # noqa: F401
    _MATPLOTLIB_OK = True
except ImportError:
    _MATPLOTLIB_OK = False

from docx import Document

from eia_agent.core.climogram_docx_inserter import (
    ClimogramDocxInsertConfig,
    ClimogramDocxInsertResult,
    count_docx_images,
    default_climogram_caption,
    insert_climogram_in_docx,
    validate_docx_contains_image,
)

# ---------------------------------------------------------------------------
# Helpers de fixtures
# ---------------------------------------------------------------------------

def _make_minimal_png(path: Path) -> Path:
    """PNG 1×1 blanco válido sin dependencias externas."""

    def _chunk(chunk_type: bytes, data: bytes) -> bytes:
        c = chunk_type + data
        return struct.pack(">I", len(data)) + c + struct.pack(">I", zlib.crc32(c) & 0xFFFFFFFF)

    sig = b"\x89PNG\r\n\x1a\n"
    ihdr = _chunk(b"IHDR", struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0))
    idat = _chunk(b"IDAT", zlib.compress(b"\x00\xff\xff\xff"))
    iend = _chunk(b"IEND", b"")
    path.write_bytes(sig + ihdr + idat + iend)
    return path


def _make_docx(path: Path, text: str = "Contenido de prueba.") -> Path:
    doc = Document()
    doc.add_paragraph(text)
    doc.save(str(path))
    return path


# ===========================================================================
# A. ClimogramDocxInsertConfig
# ===========================================================================

class TestClimogramDocxInsertConfig(unittest.TestCase):

    def test_default_heading(self):
        self.assertEqual(ClimogramDocxInsertConfig().heading, "Climograma")

    def test_default_caption_is_none(self):
        self.assertIsNone(ClimogramDocxInsertConfig().caption)

    def test_default_image_width(self):
        self.assertAlmostEqual(ClimogramDocxInsertConfig().image_width_inches, 5.8)

    def test_default_page_breaks_false(self):
        c = ClimogramDocxInsertConfig()
        self.assertFalse(c.insert_page_break_before)
        self.assertFalse(c.insert_page_break_after)

    def test_default_center_image_true(self):
        self.assertTrue(ClimogramDocxInsertConfig().center_image)

    def test_default_caption_style_none(self):
        self.assertIsNone(ClimogramDocxInsertConfig().caption_style)

    def test_to_dict_keys(self):
        d = ClimogramDocxInsertConfig().to_dict()
        for k in ("heading", "caption", "image_width_inches", "insert_page_break_before",
                  "insert_page_break_after", "caption_style", "center_image"):
            self.assertIn(k, d)

    def test_from_dict_roundtrip(self):
        cfg = ClimogramDocxInsertConfig(
            heading="Mi cabecera",
            caption="Mi caption",
            image_width_inches=4.5,
            insert_page_break_before=True,
            center_image=False,
        )
        cfg2 = ClimogramDocxInsertConfig.from_dict(cfg.to_dict())
        self.assertEqual(cfg2.heading, "Mi cabecera")
        self.assertEqual(cfg2.caption, "Mi caption")
        self.assertAlmostEqual(cfg2.image_width_inches, 4.5)
        self.assertTrue(cfg2.insert_page_break_before)
        self.assertFalse(cfg2.center_image)

    def test_from_dict_defaults_on_empty(self):
        cfg = ClimogramDocxInsertConfig.from_dict({})
        self.assertEqual(cfg.heading, "Climograma")
        self.assertAlmostEqual(cfg.image_width_inches, 5.8)
        self.assertFalse(cfg.insert_page_break_before)
        self.assertTrue(cfg.center_image)

    def test_from_dict_heading_none(self):
        cfg = ClimogramDocxInsertConfig.from_dict({"heading": None})
        self.assertIsNone(cfg.heading)


# ===========================================================================
# B. default_climogram_caption
# ===========================================================================

class TestDefaultCaption(unittest.TestCase):

    def test_no_args(self):
        cap = default_climogram_caption()
        self.assertIn("Figura", cap)
        self.assertIn("climática de referencia", cap)
        self.assertTrue(cap.endswith("."))

    def test_with_station_only(self):
        cap = default_climogram_caption("Lanzarote Aeropuerto")
        self.assertIn("Lanzarote Aeropuerto", cap)
        self.assertNotIn("periodo", cap)
        self.assertTrue(cap.endswith("."))

    def test_with_station_and_period(self):
        cap = default_climogram_caption("C029O Lanzarote Aeropuerto", "1991-2020")
        self.assertIn("C029O Lanzarote Aeropuerto", cap)
        self.assertIn("1991-2020", cap)
        self.assertIn("periodo", cap)
        self.assertTrue(cap.endswith("."))

    def test_none_station_returns_reference(self):
        cap = default_climogram_caption(station_name=None, period="1991-2020")
        self.assertIn("referencia", cap)
        self.assertNotIn("1991-2020", cap)

    def test_starts_with_figura(self):
        self.assertTrue(default_climogram_caption().startswith("Figura"))


# ===========================================================================
# C. Inserción básica
# ===========================================================================

class TestBasicInsertion(unittest.TestCase):

    def setUp(self):
        self.tmp = tempfile.mkdtemp()

    def _path(self, name):
        return Path(self.tmp) / name

    def test_output_docx_created(self):
        inp = _make_docx(self._path("in.docx"))
        png = _make_minimal_png(self._path("c.png"))
        out = self._path("out.docx")
        insert_climogram_in_docx(inp, png, out)
        self.assertTrue(out.exists())

    def test_validate_docx_contains_image_true(self):
        inp = _make_docx(self._path("in.docx"))
        png = _make_minimal_png(self._path("c.png"))
        out = self._path("out.docx")
        insert_climogram_in_docx(inp, png, out)
        self.assertTrue(validate_docx_contains_image(out))

    def test_count_images_is_one(self):
        inp = _make_docx(self._path("in.docx"))
        png = _make_minimal_png(self._path("c.png"))
        out = self._path("out.docx")
        insert_climogram_in_docx(inp, png, out)
        self.assertEqual(count_docx_images(out), 1)

    def test_result_is_ClimogramDocxInsertResult(self):
        inp = _make_docx(self._path("in.docx"))
        png = _make_minimal_png(self._path("c.png"))
        out = self._path("out.docx")
        result = insert_climogram_in_docx(inp, png, out)
        self.assertIsInstance(result, ClimogramDocxInsertResult)

    def test_result_inserted_true(self):
        inp = _make_docx(self._path("in.docx"))
        png = _make_minimal_png(self._path("c.png"))
        out = self._path("out.docx")
        result = insert_climogram_in_docx(inp, png, out)
        self.assertTrue(result.inserted)

    def test_result_output_path(self):
        inp = _make_docx(self._path("in.docx"))
        png = _make_minimal_png(self._path("c.png"))
        out = self._path("out.docx")
        result = insert_climogram_in_docx(inp, png, out)
        self.assertEqual(result.output_docx, str(out))

    def test_result_to_dict_keys(self):
        inp = _make_docx(self._path("in.docx"))
        png = _make_minimal_png(self._path("c.png"))
        out = self._path("out.docx")
        result = insert_climogram_in_docx(inp, png, out)
        d = result.to_dict()
        for k in ("input_docx", "output_docx", "png_path", "inserted",
                  "caption", "warnings", "notes"):
            self.assertIn(k, d)

    def test_result_summary_not_empty(self):
        inp = _make_docx(self._path("in.docx"))
        png = _make_minimal_png(self._path("c.png"))
        out = self._path("out.docx")
        result = insert_climogram_in_docx(inp, png, out)
        s = result.summary()
        self.assertGreater(len(s), 10)

    def test_validate_docx_false_for_missing(self):
        self.assertFalse(validate_docx_contains_image(self._path("missing.docx")))

    def test_count_images_zero_for_missing(self):
        self.assertEqual(count_docx_images(self._path("missing.docx")), 0)

    def test_count_images_zero_fresh_docx(self):
        inp = _make_docx(self._path("fresh.docx"))
        self.assertEqual(count_docx_images(inp), 0)

    def test_creates_nested_dirs(self):
        inp = _make_docx(self._path("in.docx"))
        png = _make_minimal_png(self._path("c.png"))
        nested = self._path("a") / "b" / "out.docx"
        insert_climogram_in_docx(inp, png, nested)
        self.assertTrue(nested.exists())


# ===========================================================================
# D. Heading y caption en el DOCX resultante
# ===========================================================================

class TestHeadingAndCaption(unittest.TestCase):

    def setUp(self):
        self.tmp = tempfile.mkdtemp()

    def _path(self, name):
        return Path(self.tmp) / name

    def _paragraph_texts(self, docx_path):
        return [p.text for p in Document(str(docx_path)).paragraphs]

    def test_heading_in_output(self):
        cfg = ClimogramDocxInsertConfig(heading="Clima de referencia")
        inp = _make_docx(self._path("in.docx"))
        png = _make_minimal_png(self._path("c.png"))
        out = self._path("out.docx")
        insert_climogram_in_docx(inp, png, out, config=cfg)
        texts = self._paragraph_texts(out)
        self.assertIn("Clima de referencia", texts)

    def test_caption_in_output(self):
        cap = "Figura. Climograma de prueba."
        cfg = ClimogramDocxInsertConfig(caption=cap)
        inp = _make_docx(self._path("in.docx"))
        png = _make_minimal_png(self._path("c.png"))
        out = self._path("out.docx")
        insert_climogram_in_docx(inp, png, out, config=cfg)
        texts = self._paragraph_texts(out)
        self.assertIn(cap, texts)

    def test_heading_none_not_in_output(self):
        cfg = ClimogramDocxInsertConfig(heading=None, caption=None)
        inp = _make_docx(self._path("in.docx"))
        png = _make_minimal_png(self._path("c.png"))
        out = self._path("out.docx")
        insert_climogram_in_docx(inp, png, out, config=cfg)
        texts = self._paragraph_texts(out)
        self.assertNotIn("Climograma", texts)

    def test_default_heading_present(self):
        inp = _make_docx(self._path("in.docx"))
        png = _make_minimal_png(self._path("c.png"))
        out = self._path("out.docx")
        insert_climogram_in_docx(inp, png, out)
        texts = self._paragraph_texts(out)
        self.assertIn("Climograma", texts)

    def test_result_caption_matches(self):
        cap = "Figura. Climograma de Telde."
        cfg = ClimogramDocxInsertConfig(caption=cap)
        inp = _make_docx(self._path("in.docx"))
        png = _make_minimal_png(self._path("c.png"))
        out = self._path("out.docx")
        result = insert_climogram_in_docx(inp, png, out, config=cfg)
        self.assertEqual(result.caption, cap)

    def test_result_caption_none_when_not_configured(self):
        cfg = ClimogramDocxInsertConfig(caption=None)
        inp = _make_docx(self._path("in.docx"))
        png = _make_minimal_png(self._path("c.png"))
        out = self._path("out.docx")
        result = insert_climogram_in_docx(inp, png, out, config=cfg)
        self.assertIsNone(result.caption)


# ===========================================================================
# E. Errores y seguridad
# ===========================================================================

class TestErrors(unittest.TestCase):

    def setUp(self):
        self.tmp = tempfile.mkdtemp()

    def _path(self, name):
        return Path(self.tmp) / name

    def test_missing_input_docx_raises(self):
        png = _make_minimal_png(self._path("c.png"))
        with self.assertRaises(FileNotFoundError):
            insert_climogram_in_docx(self._path("nope.docx"), png, self._path("out.docx"))

    def test_missing_png_raises(self):
        inp = _make_docx(self._path("in.docx"))
        with self.assertRaises(FileNotFoundError):
            insert_climogram_in_docx(inp, self._path("nope.png"), self._path("out.docx"))

    def test_invalid_png_raises_value_error(self):
        inp = _make_docx(self._path("in.docx"))
        bad_png = self._path("bad.png")
        bad_png.write_bytes(b"notapng1234567890")
        with self.assertRaises(ValueError):
            insert_climogram_in_docx(inp, bad_png, self._path("out.docx"))

    def test_empty_file_as_png_raises_value_error(self):
        inp = _make_docx(self._path("in.docx"))
        empty = self._path("empty.png")
        empty.write_bytes(b"")
        with self.assertRaises(ValueError):
            insert_climogram_in_docx(inp, empty, self._path("out.docx"))

    def test_non_docx_output_raises_value_error(self):
        inp = _make_docx(self._path("in.docx"))
        png = _make_minimal_png(self._path("c.png"))
        with self.assertRaises(ValueError):
            insert_climogram_in_docx(inp, png, self._path("out.pdf"))

    def test_no_extension_output_raises_value_error(self):
        inp = _make_docx(self._path("in.docx"))
        png = _make_minimal_png(self._path("c.png"))
        with self.assertRaises(ValueError):
            insert_climogram_in_docx(inp, png, self._path("out"))

    def test_original_docx_unchanged_when_output_differs(self):
        inp = _make_docx(self._path("in.docx"))
        png = _make_minimal_png(self._path("c.png"))
        out = self._path("out.docx")
        original_bytes = inp.read_bytes()
        insert_climogram_in_docx(inp, png, out)
        self.assertEqual(inp.read_bytes(), original_bytes)

    def test_non_png_extension_jpg_raises(self):
        inp = _make_docx(self._path("in.docx"))
        # File has PNG signature but .jpg extension — not the issue here
        # Test: file has no PNG signature
        bad = self._path("photo.jpg")
        bad.write_bytes(b"\xff\xd8\xff\xe0" + b"\x00" * 20)  # JPEG signature
        with self.assertRaises(ValueError):
            insert_climogram_in_docx(inp, bad, self._path("out.docx"))

    def test_validate_docx_bad_zip_returns_false(self):
        bad = self._path("bad.docx")
        bad.write_bytes(b"not a zip file at all")
        self.assertFalse(validate_docx_contains_image(bad))

    def test_count_images_bad_zip_returns_zero(self):
        bad = self._path("bad.docx")
        bad.write_bytes(b"not a zip")
        self.assertEqual(count_docx_images(bad), 0)


# ===========================================================================
# F. Configuración
# ===========================================================================

class TestConfiguration(unittest.TestCase):

    def setUp(self):
        self.tmp = tempfile.mkdtemp()

    def _path(self, name):
        return Path(self.tmp) / name

    def _insert(self, cfg=None, name="out.docx"):
        inp = _make_docx(self._path("in.docx"))
        png = _make_minimal_png(self._path("c.png"))
        out = self._path(name)
        return insert_climogram_in_docx(inp, png, out, config=cfg)

    def test_center_image_true_inserts_image(self):
        cfg = ClimogramDocxInsertConfig(center_image=True)
        result = self._insert(cfg, "ct.docx")
        self.assertTrue(validate_docx_contains_image(result.output_docx))

    def test_center_image_false_inserts_image(self):
        cfg = ClimogramDocxInsertConfig(center_image=False)
        result = self._insert(cfg, "cf.docx")
        self.assertTrue(validate_docx_contains_image(result.output_docx))

    def test_heading_none_still_inserts_image(self):
        cfg = ClimogramDocxInsertConfig(heading=None)
        result = self._insert(cfg, "hn.docx")
        self.assertTrue(validate_docx_contains_image(result.output_docx))

    def test_caption_none_still_inserts_image(self):
        cfg = ClimogramDocxInsertConfig(caption=None)
        result = self._insert(cfg, "cn.docx")
        self.assertTrue(validate_docx_contains_image(result.output_docx))

    def test_page_break_before_inserts_image(self):
        cfg = ClimogramDocxInsertConfig(insert_page_break_before=True)
        result = self._insert(cfg, "pbb.docx")
        self.assertTrue(validate_docx_contains_image(result.output_docx))

    def test_page_break_after_inserts_image(self):
        cfg = ClimogramDocxInsertConfig(insert_page_break_after=True)
        result = self._insert(cfg, "pba.docx")
        self.assertTrue(validate_docx_contains_image(result.output_docx))

    def test_none_config_uses_defaults(self):
        result = self._insert(None, "def.docx")
        self.assertEqual(result.inserted, True)
        self.assertIsNone(result.caption)

    def test_custom_width_still_inserts(self):
        cfg = ClimogramDocxInsertConfig(image_width_inches=3.0)
        result = self._insert(cfg, "w3.docx")
        self.assertTrue(validate_docx_contains_image(result.output_docx))

    def test_caption_style_invalid_adds_warning(self):
        cfg = ClimogramDocxInsertConfig(
            caption="Pie de figura.",
            caption_style="EstiloInexistente123",
        )
        result = self._insert(cfg, "ws.docx")
        self.assertTrue(len(result.warnings) > 0)

    def test_output_same_as_input_overwrites(self):
        inp = _make_docx(self._path("same.docx"))
        png = _make_minimal_png(self._path("c.png"))
        result = insert_climogram_in_docx(inp, png, inp)
        self.assertTrue(validate_docx_contains_image(result.output_docx))


# ===========================================================================
# G. Integración con CL-04
# ===========================================================================

@unittest.skipUnless(_MATPLOTLIB_OK, "matplotlib no disponible")
class TestCL04Integration(unittest.TestCase):

    def setUp(self):
        self.tmp = tempfile.mkdtemp()

    def _path(self, name):
        return Path(self.tmp) / name

    def test_real_climogram_inserted_in_docx(self):
        from eia_agent.core.climate_indices import MonthlyClimateData
        from eia_agent.core.climogram_generator import generate_climogram

        data = MonthlyClimateData(
            temperatures_c=[17.8, 18.1, 18.8, 19.4, 20.7, 22.7, 24.9, 25.7, 25.1, 23.5, 21.0, 18.6],
            precipitations_mm=[22.0, 19.0, 15.0, 7.0, 2.0, 1.0, 0.0, 1.0, 5.0, 14.0, 21.0, 24.0],
            station_id="C029O",
            station_name="Lanzarote Aeropuerto",
            period="1991-2020",
        )
        png_path = str(self._path("climo.png"))
        generate_climogram(data, png_path)

        inp = _make_docx(self._path("in.docx"), "Sección climática.")
        out = self._path("out_with_climo.docx")
        result = insert_climogram_in_docx(inp, png_path, out)

        self.assertTrue(result.inserted)
        self.assertTrue(validate_docx_contains_image(out))
        self.assertEqual(count_docx_images(out), 1)

    def test_real_climogram_with_auto_caption(self):
        from eia_agent.core.climate_indices import MonthlyClimateData
        from eia_agent.core.climogram_generator import generate_climogram

        data = MonthlyClimateData(
            temperatures_c=[5, 6, 8, 10, 13, 16, 18, 18, 15, 12, 8, 6],
            precipitations_mm=[80, 70, 65, 55, 60, 60, 55, 60, 65, 80, 85, 90],
            station_id="HUM01",
            station_name="Estación Húmeda",
            period="1991-2020",
        )
        png_path = str(self._path("humid.png"))
        generate_climogram(data, png_path)

        caption = default_climogram_caption("Estación Húmeda", "1991-2020")
        cfg = ClimogramDocxInsertConfig(caption=caption)
        inp = _make_docx(self._path("in2.docx"))
        out = self._path("humid_out.docx")
        result = insert_climogram_in_docx(inp, png_path, out, config=cfg)

        self.assertTrue(result.inserted)
        self.assertEqual(result.caption, caption)
        texts = [p.text for p in Document(str(out)).paragraphs]
        self.assertIn(caption, texts)

    def test_two_climograms_in_same_docx(self):
        from eia_agent.core.climate_indices import MonthlyClimateData
        from eia_agent.core.climogram_generator import generate_climogram

        # Datos distintos → bytes distintos → python-docx no deduplica
        data1 = MonthlyClimateData(
            temperatures_c=[10.0] * 12,
            precipitations_mm=[50.0] * 12,
        )
        data2 = MonthlyClimateData(
            temperatures_c=[17.8, 18.1, 18.8, 19.4, 20.7, 22.7, 24.9, 25.7, 25.1, 23.5, 21.0, 18.6],
            precipitations_mm=[22.0, 19.0, 15.0, 7.0, 2.0, 1.0, 0.0, 1.0, 5.0, 14.0, 21.0, 24.0],
            station_name="Lanzarote",
        )
        png1 = str(self._path("c1.png"))
        png2 = str(self._path("c2.png"))
        generate_climogram(data1, png1)
        generate_climogram(data2, png2)

        # Insertar primero, output → intermedio
        mid = self._path("mid.docx")
        inp = _make_docx(self._path("base.docx"))
        insert_climogram_in_docx(inp, png1, mid,
                                  config=ClimogramDocxInsertConfig(heading="Climo 1"))

        # Insertar segundo, mid → final
        final = self._path("final.docx")
        insert_climogram_in_docx(mid, png2, final,
                                  config=ClimogramDocxInsertConfig(heading="Climo 2"))

        self.assertEqual(count_docx_images(final), 2)


if __name__ == "__main__":
    unittest.main()
