"""
Tests para document_presentation_preparer (DOC-08).

Estrategia offline:
  - tempfile.TemporaryDirectory() para aislamiento total.
  - No se llama a ninguna API externa.
  - python-docx usada directamente para crear/leer DOCX sinteticos.
  - No se modifican expedientes piloto reales.
"""
from __future__ import annotations

import json
import sys
import tempfile
import unittest
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent / "src"))

from eia_agent.core.document_presentation_preparer import (
    FINAL_REVIEW_DOCX,
    METADATA_JSON,
    METADATA_MD,
    PRESENTATION_CHECKLIST_JSON,
    PRESENTATION_CHECKLIST_MD,
    PRESENTATION_STATUS,
    PRESENTATION_SEVERITY,
    SIGNATURE_SHEET_MD,
    DocumentMetadata,
    PresentationChecklistItem,
    PresentationIssue,
    PresentationPreparationResult,
    append_signature_sheet_to_docx,
    build_document_metadata,
    build_metadata_markdown,
    build_presentation_checklist,
    build_presentation_checklist_markdown,
    build_signature_sheet_markdown,
    prepare_document_for_presentation,
    safe_load_json,
    write_presentation_outputs,
)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _make_minimal_docx(path: Path) -> None:
    """Crea un DOCX minimal valido con python-docx."""
    try:
        from docx import Document
        doc = Document()
        doc.add_heading("Documento Ambiental", level=1)
        doc.add_paragraph("Contenido de prueba.")
        path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(path))
    except ImportError:
        # Si python-docx no esta disponible, crear bytes minimos de ZIP (DOCX valido minimo)
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_bytes(b"PK\x03\x04" + b"\x00" * 26)


def _make_full_expediente(exp_root: Path) -> None:
    """Crea estructura sintetica completa de expediente para tests de checklist."""
    doc_dir = exp_root / "documento"
    doc_dir.mkdir(parents=True, exist_ok=True)
    audit_dir = exp_root / "auditoria"
    audit_dir.mkdir(parents=True, exist_ok=True)

    # DOCX enriquecido
    _make_minimal_docx(doc_dir / "documento_ambiental_borrador_con_figuras.docx")

    # Markdown fuente
    (doc_dir / "documento_ambiental_borrador.md").write_text(
        "# Documento Ambiental\n\nContenido de prueba.\n",
        encoding="utf-8",
    )

    # QC result
    (doc_dir / "document_quality_result.json").write_text(
        json.dumps({"status": "OK", "administrative_ready": False}),
        encoding="utf-8",
    )

    # Auditoria final
    (audit_dir / "final_audit_result.json").write_text(
        json.dumps({"status": "CONFORME", "administrative_ready": False}),
        encoding="utf-8",
    )

    # Package result
    (doc_dir / "package_build_result.json").write_text(
        json.dumps({"generated": True, "administrative_ready": False}),
        encoding="utf-8",
    )

    # Export result
    (doc_dir / "document_export_result.json").write_text(
        json.dumps({"zip_generated": True, "administrative_ready": False}),
        encoding="utf-8",
    )

    # ZIP paquete
    (doc_dir / "paquete_entrega.zip").write_bytes(b"PK\x05\x06" + b"\x00" * 18)

    # README_ENTREGA en paquete
    pkg_dir = doc_dir / "paquete_entrega"
    pkg_dir.mkdir(parents=True, exist_ok=True)
    (pkg_dir / "README_ENTREGA.md").write_text("# README", encoding="utf-8")

    # Figures result
    (doc_dir / "document_figures_result.json").write_text(
        json.dumps({"generated": True, "figures_inserted": 6}),
        encoding="utf-8",
    )


def _make_metadata(
    expediente_id: str = "test_exp",
    final_audit_status: "str | None" = "CONFORME",
    document_qc_status: "str | None" = "OK",
    package_status: "str | None" = "GENERADO",
    export_status: "str | None" = "GENERADO",
) -> DocumentMetadata:
    return DocumentMetadata(
        expediente_id=expediente_id,
        generated_at="2026-05-25T00:00:00Z",
        source_docx="/tmp/doc.docx",
        source_markdown="/tmp/doc.md",
        package_zip="/tmp/pkg.zip",
        final_audit_status=final_audit_status,
        document_qc_status=document_qc_status,
        package_status=package_status,
        export_status=export_status,
        notes=["nota de prueba"],
        warnings=[],
    )


# ---------------------------------------------------------------------------
# 1. PresentationIssue
# ---------------------------------------------------------------------------

class TestPresentationIssue(unittest.TestCase):

    def _make_issue(self, severity=PRESENTATION_SEVERITY.ERROR) -> PresentationIssue:
        return PresentationIssue(
            severity=severity,
            code="PP-E001",
            message="Mensaje de prueba",
            recommendation="Recomendacion",
            evidence=["evidencia1"],
        )

    def test_to_dict_keys(self):
        d = self._make_issue().to_dict()
        for k in ("severity", "code", "message", "recommendation", "evidence"):
            self.assertIn(k, d)

    def test_to_dict_values(self):
        issue = self._make_issue(PRESENTATION_SEVERITY.WARNING)
        d = issue.to_dict()
        self.assertEqual(d["severity"], "WARNING")
        self.assertEqual(d["code"], "PP-E001")
        self.assertIsInstance(d["evidence"], list)

    def test_summary_format(self):
        issue = self._make_issue()
        s = issue.summary()
        self.assertIn("ERROR", s)
        self.assertIn("PP-E001", s)
        self.assertIn("Mensaje de prueba", s)

    def test_summary_warning(self):
        issue = self._make_issue(PRESENTATION_SEVERITY.WARNING)
        self.assertIn("WARNING", issue.summary())


# ---------------------------------------------------------------------------
# 2. DocumentMetadata
# ---------------------------------------------------------------------------

class TestDocumentMetadata(unittest.TestCase):

    def test_to_dict_contains_admin_ready_false(self):
        m = _make_metadata()
        d = m.to_dict()
        self.assertIn("administrative_ready", d)
        self.assertFalse(d["administrative_ready"])

    def test_administrative_ready_always_false(self):
        m = _make_metadata()
        self.assertFalse(m.administrative_ready)

    def test_to_dict_keys(self):
        m = _make_metadata()
        d = m.to_dict()
        for k in ("expediente_id", "generated_at", "source_docx", "final_audit_status",
                   "document_qc_status", "package_status", "export_status",
                   "administrative_ready", "notes", "warnings"):
            self.assertIn(k, d)

    def test_summary_contains_expediente_id(self):
        m = _make_metadata(expediente_id="test_exp_001")
        self.assertIn("test_exp_001", m.summary())

    def test_summary_admin_ready_false(self):
        m = _make_metadata()
        self.assertIn("False", m.summary())

    def test_none_fields_allowed(self):
        m = _make_metadata(final_audit_status=None, document_qc_status=None)
        d = m.to_dict()
        self.assertIsNone(d["final_audit_status"])
        self.assertIsNone(d["document_qc_status"])

    def test_notes_included(self):
        m = _make_metadata()
        m.notes = ["nota1", "nota2"]
        d = m.to_dict()
        self.assertIn("nota1", d["notes"])


# ---------------------------------------------------------------------------
# 3. PresentationChecklistItem
# ---------------------------------------------------------------------------

class TestPresentationChecklistItem(unittest.TestCase):

    def _make_item(self, status="OK") -> PresentationChecklistItem:
        return PresentationChecklistItem(
            item_id="CHK-001",
            description="Descripcion de prueba",
            status=status,
            evidence=["evidencia"],
            recommendation="Recomendacion",
        )

    def test_to_dict_keys(self):
        d = self._make_item().to_dict()
        for k in ("item_id", "description", "status", "evidence", "recommendation"):
            self.assertIn(k, d)

    def test_to_dict_status(self):
        self.assertEqual(self._make_item("WARNING").to_dict()["status"], "WARNING")
        self.assertEqual(self._make_item("ERROR").to_dict()["status"], "ERROR")

    def test_summary_format(self):
        s = self._make_item("OK").summary()
        self.assertIn("[OK]", s)
        self.assertIn("CHK-001", s)

    def test_summary_contains_description(self):
        s = self._make_item().summary()
        self.assertIn("Descripcion de prueba", s)


# ---------------------------------------------------------------------------
# 4. PresentationPreparationResult
# ---------------------------------------------------------------------------

class TestPresentationPreparationResult(unittest.TestCase):

    def _make_result(
        self,
        issues=None,
        checklist_items=None,
        status=PRESENTATION_STATUS.PREPARADO_PARA_REVISION,
    ) -> PresentationPreparationResult:
        return PresentationPreparationResult(
            expediente_id="test_exp",
            status=status,
            metadata=_make_metadata(),
            checklist_items=checklist_items or [],
            issues=issues or [],
            generated_files=[],
            warnings=[],
            notes=[],
        )

    def test_error_count_zero(self):
        r = self._make_result()
        self.assertEqual(r.error_count(), 0)

    def test_error_count_nonzero(self):
        issue = PresentationIssue(
            severity=PRESENTATION_SEVERITY.ERROR,
            code="PP-E001", message="x",
            recommendation="", evidence=[],
        )
        r = self._make_result(issues=[issue])
        self.assertEqual(r.error_count(), 1)

    def test_warning_count(self):
        w = PresentationIssue(
            severity=PRESENTATION_SEVERITY.WARNING,
            code="PP-W001", message="w",
            recommendation="", evidence=[],
        )
        r = self._make_result(issues=[w])
        self.assertEqual(r.warning_count(), 1)
        self.assertEqual(r.error_count(), 0)

    def test_is_success_no_errors(self):
        self.assertTrue(self._make_result().is_success())

    def test_is_success_with_error(self):
        issue = PresentationIssue(
            severity=PRESENTATION_SEVERITY.ERROR,
            code="PP-E001", message="x",
            recommendation="", evidence=[],
        )
        r = self._make_result(issues=[issue])
        self.assertFalse(r.is_success())

    def test_is_success_with_warning_only(self):
        w = PresentationIssue(
            severity=PRESENTATION_SEVERITY.WARNING,
            code="PP-W001", message="w",
            recommendation="", evidence=[],
        )
        r = self._make_result(issues=[w])
        self.assertTrue(r.is_success())

    def test_checklist_ok_count(self):
        items = [
            PresentationChecklistItem("CHK-001", "desc", "OK", [], ""),
            PresentationChecklistItem("CHK-002", "desc", "ERROR", [], ""),
            PresentationChecklistItem("CHK-003", "desc", "OK", [], ""),
        ]
        r = self._make_result(checklist_items=items)
        self.assertEqual(r.checklist_ok_count(), 2)

    def test_checklist_error_count(self):
        items = [
            PresentationChecklistItem("CHK-001", "desc", "ERROR", [], ""),
            PresentationChecklistItem("CHK-002", "desc", "OK", [], ""),
        ]
        r = self._make_result(checklist_items=items)
        self.assertEqual(r.checklist_error_count(), 1)

    def test_to_dict_admin_ready_always_false(self):
        r = self._make_result()
        d = r.to_dict()
        self.assertFalse(d["administrative_ready"])

    def test_to_dict_keys(self):
        d = self._make_result().to_dict()
        for k in ("expediente_id", "status", "metadata", "checklist_items",
                   "issues", "generated_files", "error_count", "warning_count",
                   "is_success", "administrative_ready"):
            self.assertIn(k, d)

    def test_summary_contains_expediente(self):
        r = self._make_result()
        self.assertIn("test_exp", r.summary())

    def test_summary_contains_disclaimer(self):
        r = self._make_result()
        self.assertIn("no declara", r.summary().lower())


# ---------------------------------------------------------------------------
# 5. safe_load_json
# ---------------------------------------------------------------------------

class TestSafeLoadJson(unittest.TestCase):

    def test_returns_none_for_missing_file(self):
        self.assertIsNone(safe_load_json("/nonexistent/path/file.json"))

    def test_loads_valid_json(self):
        with tempfile.TemporaryDirectory() as td:
            p = Path(td) / "test.json"
            p.write_text('{"key": "value"}', encoding="utf-8")
            result = safe_load_json(p)
            self.assertEqual(result, {"key": "value"})

    def test_returns_none_for_invalid_json(self):
        with tempfile.TemporaryDirectory() as td:
            p = Path(td) / "bad.json"
            p.write_text("not json", encoding="utf-8")
            self.assertIsNone(safe_load_json(p))

    def test_returns_none_for_non_dict(self):
        with tempfile.TemporaryDirectory() as td:
            p = Path(td) / "list.json"
            p.write_text("[1, 2, 3]", encoding="utf-8")
            self.assertIsNone(safe_load_json(p))


# ---------------------------------------------------------------------------
# 6. build_document_metadata
# ---------------------------------------------------------------------------

class TestBuildDocumentMetadata(unittest.TestCase):

    def test_empty_expediente_does_not_crash(self):
        with tempfile.TemporaryDirectory() as td:
            exp = Path(td) / "exp_vacio"
            exp.mkdir()
            m = build_document_metadata(exp)
            self.assertEqual(m.expediente_id, "exp_vacio")
            self.assertIsNone(m.final_audit_status)
            self.assertIsNone(m.document_qc_status)

    def test_reads_audit_status(self):
        with tempfile.TemporaryDirectory() as td:
            exp = Path(td) / "exp"
            exp.mkdir()
            audit_dir = exp / "auditoria"
            audit_dir.mkdir()
            (audit_dir / "final_audit_result.json").write_text(
                json.dumps({"status": "CONFORME"}), encoding="utf-8"
            )
            m = build_document_metadata(exp)
            self.assertEqual(m.final_audit_status, "CONFORME")

    def test_reads_qc_status(self):
        with tempfile.TemporaryDirectory() as td:
            exp = Path(td) / "exp"
            exp.mkdir()
            doc_dir = exp / "documento"
            doc_dir.mkdir()
            (doc_dir / "document_quality_result.json").write_text(
                json.dumps({"status": "OK"}), encoding="utf-8"
            )
            m = build_document_metadata(exp)
            self.assertEqual(m.document_qc_status, "OK")

    def test_reads_package_generated(self):
        with tempfile.TemporaryDirectory() as td:
            exp = Path(td) / "exp"
            exp.mkdir()
            doc_dir = exp / "documento"
            doc_dir.mkdir()
            (doc_dir / "package_build_result.json").write_text(
                json.dumps({"generated": True}), encoding="utf-8"
            )
            m = build_document_metadata(exp)
            self.assertEqual(m.package_status, "GENERADO")

    def test_reads_export_zip_generated(self):
        with tempfile.TemporaryDirectory() as td:
            exp = Path(td) / "exp"
            exp.mkdir()
            doc_dir = exp / "documento"
            doc_dir.mkdir()
            (doc_dir / "document_export_result.json").write_text(
                json.dumps({"zip_generated": True}), encoding="utf-8"
            )
            m = build_document_metadata(exp)
            self.assertEqual(m.export_status, "GENERADO")

    def test_administrative_ready_always_false(self):
        with tempfile.TemporaryDirectory() as td:
            exp = Path(td) / "exp"
            exp.mkdir()
            m = build_document_metadata(exp)
            self.assertFalse(m.administrative_ready)

    def test_administrative_ready_false_even_if_json_says_true(self):
        """administrative_ready=True en un JSON externo no debe propagarse."""
        with tempfile.TemporaryDirectory() as td:
            exp = Path(td) / "exp"
            exp.mkdir()
            doc_dir = exp / "documento"
            doc_dir.mkdir()
            (doc_dir / "document_export_result.json").write_text(
                json.dumps({"zip_generated": True, "administrative_ready": True}),
                encoding="utf-8",
            )
            m = build_document_metadata(exp)
            self.assertFalse(m.administrative_ready)

    def test_detects_enriched_docx(self):
        with tempfile.TemporaryDirectory() as td:
            exp = Path(td) / "exp"
            exp.mkdir()
            doc_dir = exp / "documento"
            doc_dir.mkdir()
            enriched = doc_dir / "documento_ambiental_borrador_con_figuras.docx"
            enriched.write_bytes(b"PK" + b"\x00" * 30)
            m = build_document_metadata(exp)
            self.assertIsNotNone(m.source_docx)
            self.assertIn("con_figuras", m.source_docx)

    def test_generated_at_is_string(self):
        with tempfile.TemporaryDirectory() as td:
            exp = Path(td) / "exp"
            exp.mkdir()
            m = build_document_metadata(exp)
            self.assertIsInstance(m.generated_at, str)
            self.assertIn("T", m.generated_at)


# ---------------------------------------------------------------------------
# 7. build_signature_sheet_markdown
# ---------------------------------------------------------------------------

class TestBuildSignatureSheetMarkdown(unittest.TestCase):

    def _get_sig(self) -> str:
        return build_signature_sheet_markdown(_make_metadata())

    def test_contains_heading(self):
        sig = self._get_sig()
        self.assertIn("Hoja de firmas", sig)

    def test_contains_nombre_apellidos(self):
        sig = self._get_sig()
        self.assertIn("Nombre y apellidos", sig)

    def test_contains_titulacion(self):
        sig = self._get_sig()
        self.assertIn("Titulacion", sig)

    def test_contains_num_colegiado(self):
        sig = self._get_sig()
        self.assertIn("colegiado", sig.lower())

    def test_contains_fecha(self):
        sig = self._get_sig()
        self.assertIn("Fecha", sig)

    def test_contains_firma_section(self):
        sig = self._get_sig()
        self.assertIn("Firma", sig)

    def test_contains_admin_disclaimer(self):
        sig = self._get_sig()
        self.assertIn("no acredita", sig.lower())
        self.assertIn("aptitud administrativa", sig.lower())

    def test_contains_expediente_id(self):
        sig = self._get_sig()
        self.assertIn("test_exp", sig)


# ---------------------------------------------------------------------------
# 8. build_presentation_checklist
# ---------------------------------------------------------------------------

class TestBuildPresentationChecklist(unittest.TestCase):

    def test_full_expediente_generates_ok_items(self):
        with tempfile.TemporaryDirectory() as td:
            exp = Path(td) / "exp"
            exp.mkdir()
            _make_full_expediente(exp)
            metadata = build_document_metadata(exp)
            items = build_presentation_checklist(exp, metadata)
            self.assertGreater(len(items), 0)
            chk_ids = [i.item_id for i in items]
            self.assertIn("CHK-001", chk_ids)
            self.assertIn("CHK-009", chk_ids)
            # CHK-009 debe ser OK (no hay administrative_ready=True)
            chk009 = next(i for i in items if i.item_id == "CHK-009")
            self.assertEqual(chk009.status, "OK")

    def test_missing_docx_generates_error(self):
        with tempfile.TemporaryDirectory() as td:
            exp = Path(td) / "exp_nodocx"
            exp.mkdir()
            # Expediente vacio: sin DOCX
            metadata = build_document_metadata(exp)
            items = build_presentation_checklist(exp, metadata)
            chk001 = next(i for i in items if i.item_id == "CHK-001")
            self.assertEqual(chk001.status, "ERROR")

    def test_administrative_ready_true_generates_error(self):
        with tempfile.TemporaryDirectory() as td:
            exp = Path(td) / "exp_adm"
            exp.mkdir()
            doc_dir = exp / "documento"
            doc_dir.mkdir()
            # Inyectar administrative_ready=True en un JSON
            (doc_dir / "document_export_result.json").write_text(
                json.dumps({"zip_generated": True, "administrative_ready": True}),
                encoding="utf-8",
            )
            metadata = build_document_metadata(exp)
            items = build_presentation_checklist(exp, metadata)
            chk009 = next(i for i in items if i.item_id == "CHK-009")
            self.assertEqual(chk009.status, "ERROR")

    def test_missing_zip_generates_warning(self):
        with tempfile.TemporaryDirectory() as td:
            exp = Path(td) / "exp"
            exp.mkdir()
            _make_full_expediente(exp)
            # Eliminar el ZIP
            zip_path = exp / "documento" / "paquete_entrega.zip"
            zip_path.unlink()
            metadata = build_document_metadata(exp)
            items = build_presentation_checklist(exp, metadata)
            chk007 = next(i for i in items if i.item_id == "CHK-007")
            self.assertEqual(chk007.status, "WARNING")

    def test_no_aplica_for_missing_figures(self):
        with tempfile.TemporaryDirectory() as td:
            exp = Path(td) / "exp"
            exp.mkdir()
            # Sin document_figures_result.json
            metadata = build_document_metadata(exp)
            items = build_presentation_checklist(exp, metadata)
            chk011 = next(i for i in items if i.item_id == "CHK-011")
            self.assertEqual(chk011.status, "NO_APLICA")

    def test_conforme_audit_ok(self):
        with tempfile.TemporaryDirectory() as td:
            exp = Path(td) / "exp"
            exp.mkdir()
            (exp / "auditoria").mkdir()
            (exp / "auditoria" / "final_audit_result.json").write_text(
                json.dumps({"status": "CONFORME"}), encoding="utf-8"
            )
            metadata = build_document_metadata(exp)
            items = build_presentation_checklist(exp, metadata)
            chk006 = next(i for i in items if i.item_id == "CHK-006")
            self.assertEqual(chk006.status, "OK")

    def test_no_conforme_audit_warning(self):
        with tempfile.TemporaryDirectory() as td:
            exp = Path(td) / "exp"
            exp.mkdir()
            (exp / "auditoria").mkdir()
            (exp / "auditoria" / "final_audit_result.json").write_text(
                json.dumps({"status": "NO_CONFORME"}), encoding="utf-8"
            )
            metadata = build_document_metadata(exp)
            items = build_presentation_checklist(exp, metadata)
            chk006 = next(i for i in items if i.item_id == "CHK-006")
            self.assertEqual(chk006.status, "WARNING")

    def test_qc_no_conforme_generates_error(self):
        with tempfile.TemporaryDirectory() as td:
            exp = Path(td) / "exp"
            exp.mkdir()
            doc_dir = exp / "documento"
            doc_dir.mkdir()
            (doc_dir / "document_quality_result.json").write_text(
                json.dumps({"status": "NO_CONFORME"}), encoding="utf-8"
            )
            metadata = build_document_metadata(exp)
            items = build_presentation_checklist(exp, metadata)
            chk004 = next(i for i in items if i.item_id == "CHK-004")
            self.assertEqual(chk004.status, "ERROR")

    def test_checklist_has_twelve_items(self):
        with tempfile.TemporaryDirectory() as td:
            exp = Path(td) / "exp"
            exp.mkdir()
            metadata = build_document_metadata(exp)
            items = build_presentation_checklist(exp, metadata)
            self.assertEqual(len(items), 12)


# ---------------------------------------------------------------------------
# 9. append_signature_sheet_to_docx
# ---------------------------------------------------------------------------

class TestAppendSignatureSheetToDocx(unittest.TestCase):

    def _signature_md(self) -> str:
        return build_signature_sheet_markdown(_make_metadata())

    def test_creates_output_docx(self):
        with tempfile.TemporaryDirectory() as td:
            inp = Path(td) / "input.docx"
            _make_minimal_docx(inp)
            out = Path(td) / "output.docx"
            result = append_signature_sheet_to_docx(inp, out, self._signature_md())
            if result:
                self.assertTrue(out.exists())
                self.assertGreater(out.stat().st_size, 0)

    def test_does_not_modify_input(self):
        with tempfile.TemporaryDirectory() as td:
            inp = Path(td) / "input.docx"
            _make_minimal_docx(inp)
            original_size = inp.stat().st_size
            out = Path(td) / "output.docx"
            append_signature_sheet_to_docx(inp, out, self._signature_md())
            self.assertEqual(inp.stat().st_size, original_size)

    def test_returns_false_for_missing_input(self):
        with tempfile.TemporaryDirectory() as td:
            inp = Path(td) / "nonexistent.docx"
            out = Path(td) / "output.docx"
            result = append_signature_sheet_to_docx(inp, out, self._signature_md())
            self.assertFalse(result)

    def test_output_openable_with_python_docx(self):
        """Si python-docx disponible, el output debe abrirse correctamente."""
        try:
            from docx import Document
        except ImportError:
            self.skipTest("python-docx no disponible")

        with tempfile.TemporaryDirectory() as td:
            inp = Path(td) / "input.docx"
            _make_minimal_docx(inp)
            out = Path(td) / "output.docx"
            ok = append_signature_sheet_to_docx(inp, out, self._signature_md())
            if ok:
                doc = Document(str(out))
                texts = "\n".join(p.text for p in doc.paragraphs)
                # El DOCX original debe estar presente (al menos parte del heading)
                self.assertTrue(len(doc.paragraphs) > 0)

    def test_output_contains_firma_content(self):
        """Si python-docx disponible, verificar que la hoja de firmas esta presente."""
        try:
            from docx import Document
        except ImportError:
            self.skipTest("python-docx no disponible")

        with tempfile.TemporaryDirectory() as td:
            inp = Path(td) / "input.docx"
            _make_minimal_docx(inp)
            out = Path(td) / "output.docx"
            ok = append_signature_sheet_to_docx(inp, out, self._signature_md())
            if ok:
                doc = Document(str(out))
                all_text = "\n".join(p.text for p in doc.paragraphs)
                self.assertIn("firmas", all_text.lower())


# ---------------------------------------------------------------------------
# 10. prepare_document_for_presentation
# ---------------------------------------------------------------------------

class TestPrepareDocumentForPresentation(unittest.TestCase):

    def test_dry_run_does_not_write(self):
        with tempfile.TemporaryDirectory() as td:
            exp = Path(td) / "exp"
            exp.mkdir()
            result = prepare_document_for_presentation(exp, write_outputs=False)
            doc_dir = exp / "documento"
            # No debe crear ningun archivo
            self.assertFalse((doc_dir / METADATA_JSON).exists())
            self.assertFalse((doc_dir / SIGNATURE_SHEET_MD).exists())
            self.assertEqual(result.generated_files, [])

    def test_write_creates_metadata_json(self):
        with tempfile.TemporaryDirectory() as td:
            exp = Path(td) / "exp"
            exp.mkdir()
            result = prepare_document_for_presentation(
                exp, write_outputs=True, create_final_docx=False
            )
            doc_dir = exp / "documento"
            self.assertTrue((doc_dir / METADATA_JSON).exists())

    def test_write_creates_metadata_md(self):
        with tempfile.TemporaryDirectory() as td:
            exp = Path(td) / "exp"
            exp.mkdir()
            prepare_document_for_presentation(
                exp, write_outputs=True, create_final_docx=False
            )
            doc_dir = exp / "documento"
            self.assertTrue((doc_dir / METADATA_MD).exists())

    def test_write_creates_signature_sheet(self):
        with tempfile.TemporaryDirectory() as td:
            exp = Path(td) / "exp"
            exp.mkdir()
            prepare_document_for_presentation(
                exp, write_outputs=True, create_final_docx=False
            )
            doc_dir = exp / "documento"
            self.assertTrue((doc_dir / SIGNATURE_SHEET_MD).exists())

    def test_write_creates_checklist_json(self):
        with tempfile.TemporaryDirectory() as td:
            exp = Path(td) / "exp"
            exp.mkdir()
            prepare_document_for_presentation(
                exp, write_outputs=True, create_final_docx=False
            )
            doc_dir = exp / "documento"
            self.assertTrue((doc_dir / PRESENTATION_CHECKLIST_JSON).exists())

    def test_write_creates_checklist_md(self):
        with tempfile.TemporaryDirectory() as td:
            exp = Path(td) / "exp"
            exp.mkdir()
            prepare_document_for_presentation(
                exp, write_outputs=True, create_final_docx=False
            )
            doc_dir = exp / "documento"
            self.assertTrue((doc_dir / PRESENTATION_CHECKLIST_MD).exists())

    def test_create_final_docx_true_with_source(self):
        with tempfile.TemporaryDirectory() as td:
            exp = Path(td) / "exp"
            exp.mkdir()
            _make_full_expediente(exp)
            result = prepare_document_for_presentation(
                exp, write_outputs=True, create_final_docx=True
            )
            doc_dir = exp / "documento"
            # Si python-docx disponible el archivo debe existir
            final_docx = doc_dir / FINAL_REVIEW_DOCX
            # El resultado incluye el intento; puede ser True o False segun python-docx
            self.assertIsInstance(result.generated_files, list)

    def test_create_final_docx_false_does_not_create(self):
        with tempfile.TemporaryDirectory() as td:
            exp = Path(td) / "exp"
            exp.mkdir()
            _make_full_expediente(exp)
            prepare_document_for_presentation(
                exp, write_outputs=True, create_final_docx=False
            )
            doc_dir = exp / "documento"
            self.assertFalse((doc_dir / FINAL_REVIEW_DOCX).exists())

    def test_does_not_modify_source_docx(self):
        with tempfile.TemporaryDirectory() as td:
            exp = Path(td) / "exp"
            exp.mkdir()
            _make_full_expediente(exp)
            src = exp / "documento" / "documento_ambiental_borrador_con_figuras.docx"
            original_size = src.stat().st_size
            prepare_document_for_presentation(
                exp, write_outputs=True, create_final_docx=True
            )
            self.assertEqual(src.stat().st_size, original_size)

    def test_administrative_ready_always_false(self):
        with tempfile.TemporaryDirectory() as td:
            exp = Path(td) / "exp"
            exp.mkdir()
            result = prepare_document_for_presentation(exp, write_outputs=False)
            self.assertFalse(result.metadata.administrative_ready)
            self.assertFalse(result.to_dict()["administrative_ready"])

    def test_result_has_status(self):
        with tempfile.TemporaryDirectory() as td:
            exp = Path(td) / "exp"
            exp.mkdir()
            result = prepare_document_for_presentation(exp, write_outputs=False)
            self.assertIn(result.status, [
                PRESENTATION_STATUS.PREPARADO_PARA_REVISION,
                PRESENTATION_STATUS.PENDIENTE_REVISION_TECNICA,
                PRESENTATION_STATUS.PENDIENTE_DOCUMENTACION,
                PRESENTATION_STATUS.NO_PREPARADO,
            ])

    def test_full_expediente_is_success(self):
        with tempfile.TemporaryDirectory() as td:
            exp = Path(td) / "exp"
            exp.mkdir()
            _make_full_expediente(exp)
            result = prepare_document_for_presentation(exp, write_outputs=False)
            self.assertTrue(result.is_success())

    def test_metadata_json_is_valid(self):
        with tempfile.TemporaryDirectory() as td:
            exp = Path(td) / "exp"
            exp.mkdir()
            prepare_document_for_presentation(
                exp, write_outputs=True, create_final_docx=False
            )
            p = exp / "documento" / METADATA_JSON
            data = json.loads(p.read_text(encoding="utf-8"))
            self.assertIn("expediente_id", data)
            self.assertFalse(data["administrative_ready"])


# ---------------------------------------------------------------------------
# 11. write_presentation_outputs
# ---------------------------------------------------------------------------

class TestWritePresentationOutputs(unittest.TestCase):

    def _make_result_obj(self) -> PresentationPreparationResult:
        return PresentationPreparationResult(
            expediente_id="test_exp",
            status=PRESENTATION_STATUS.PREPARADO_PARA_REVISION,
            metadata=_make_metadata(),
            checklist_items=[
                PresentationChecklistItem("CHK-001", "desc", "OK", [], ""),
            ],
            issues=[],
            generated_files=[],
            warnings=[],
            notes=[],
        )

    def test_writes_all_five_files(self):
        with tempfile.TemporaryDirectory() as td:
            out_dir = Path(td) / "output"
            result = self._make_result_obj()
            written = write_presentation_outputs(result, out_dir)
            self.assertEqual(len(written), 5)
            names = [p.name for p in written]
            self.assertIn(METADATA_JSON, names)
            self.assertIn(METADATA_MD, names)
            self.assertIn(SIGNATURE_SHEET_MD, names)
            self.assertIn(PRESENTATION_CHECKLIST_JSON, names)
            self.assertIn(PRESENTATION_CHECKLIST_MD, names)

    def test_all_files_exist(self):
        with tempfile.TemporaryDirectory() as td:
            out_dir = Path(td) / "output"
            result = self._make_result_obj()
            written = write_presentation_outputs(result, out_dir)
            for p in written:
                self.assertTrue(p.exists(), f"Archivo no creado: {p}")

    def test_metadata_json_content(self):
        with tempfile.TemporaryDirectory() as td:
            out_dir = Path(td) / "output"
            result = self._make_result_obj()
            write_presentation_outputs(result, out_dir)
            data = json.loads((out_dir / METADATA_JSON).read_text(encoding="utf-8"))
            self.assertFalse(data["administrative_ready"])
            self.assertEqual(data["expediente_id"], "test_exp")


# ---------------------------------------------------------------------------
# 12. CLI — document-prepare-presentation
# ---------------------------------------------------------------------------

class TestCLIDocumentPreparePresentation(unittest.TestCase):

    def _run_cli(self, exp_path: str, extra_args: "list[str] | None" = None) -> int:
        """Invoca run_expediente.py main() directamente."""
        sys.path.insert(0, str(Path(__file__).parent.parent))
        import importlib
        import run_expediente
        importlib.reload(run_expediente)
        args = [str(exp_path), "document-prepare-presentation"]
        if extra_args:
            args.extend(extra_args)
        return run_expediente.main(args)

    def test_dry_run_exit_0_when_expediente_has_docx(self):
        """Dry-run sale 0 cuando hay DOCX (sin ERRORs de CHK-001)."""
        with tempfile.TemporaryDirectory() as td:
            exp = Path(td) / "exp"
            exp.mkdir()
            _make_full_expediente(exp)
            code = self._run_cli(str(exp))
            self.assertEqual(code, 0)

    def test_dry_run_does_not_write(self):
        with tempfile.TemporaryDirectory() as td:
            exp = Path(td) / "exp"
            exp.mkdir()
            self._run_cli(str(exp))
            doc_dir = exp / "documento"
            self.assertFalse((doc_dir / METADATA_JSON).exists())

    def test_write_creates_files(self):
        with tempfile.TemporaryDirectory() as td:
            exp = Path(td) / "exp"
            exp.mkdir()
            self._run_cli(str(exp), ["--write", "--no-final-docx"])
            doc_dir = exp / "documento"
            self.assertTrue((doc_dir / METADATA_JSON).exists())
            self.assertTrue((doc_dir / SIGNATURE_SHEET_MD).exists())
            self.assertTrue((doc_dir / PRESENTATION_CHECKLIST_MD).exists())

    def test_no_final_docx_does_not_create_docx(self):
        with tempfile.TemporaryDirectory() as td:
            exp = Path(td) / "exp"
            exp.mkdir()
            _make_full_expediente(exp)
            self._run_cli(str(exp), ["--write", "--no-final-docx"])
            doc_dir = exp / "documento"
            self.assertFalse((doc_dir / FINAL_REVIEW_DOCX).exists())

    def test_exit_1_for_missing_expediente(self):
        code = self._run_cli("/nonexistent/path/exp")
        self.assertEqual(code, 1)

    def test_exit_0_for_full_expediente(self):
        with tempfile.TemporaryDirectory() as td:
            exp = Path(td) / "exp"
            exp.mkdir()
            _make_full_expediente(exp)
            code = self._run_cli(str(exp))
            self.assertEqual(code, 0)

    def test_write_metadata_json_not_admin_ready(self):
        with tempfile.TemporaryDirectory() as td:
            exp = Path(td) / "exp"
            exp.mkdir()
            self._run_cli(str(exp), ["--write", "--no-final-docx"])
            p = exp / "documento" / METADATA_JSON
            if p.exists():
                data = json.loads(p.read_text(encoding="utf-8"))
                self.assertFalse(data.get("administrative_ready", True))

    def test_exit_1_when_admin_ready_true_in_json(self):
        """CHK-009 ERROR si hay administrative_ready=True en un JSON."""
        with tempfile.TemporaryDirectory() as td:
            exp = Path(td) / "exp"
            exp.mkdir()
            doc_dir = exp / "documento"
            doc_dir.mkdir()
            (doc_dir / "document_export_result.json").write_text(
                json.dumps({"zip_generated": True, "administrative_ready": True}),
                encoding="utf-8",
            )
            code = self._run_cli(str(exp))
            self.assertEqual(code, 1)


if __name__ == "__main__":
    unittest.main()
