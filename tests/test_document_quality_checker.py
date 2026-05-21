"""Tests para document_quality_checker -- DOC-04."""
import json
import os
import tempfile
import unittest
from pathlib import Path

from eia_agent.core.document_quality_checker import (
    DOCUMENT_QC_SEVERITY,
    DOCUMENT_QC_STATUS,
    OPTIONAL_ENRICHED_FILES,
    REQUIRED_BLOCKS,
    REQUIRED_DOCUMENT_FILES,
    DocumentQualityIssue,
    DocumentQualityResult,
    build_document_quality_report_markdown,
    check_docx_structure,
    check_figures_and_captions,
    check_final_audit_visibility,
    check_no_administrative_ready_claim,
    check_required_document_files,
    detect_blocks_in_text,
    extract_docx_text,
    run_document_quality_check,
    safe_load_json,
    select_best_docx_for_qc,
    validate_docx_opens,
    write_document_quality_outputs,
)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _make_minimal_docx(path: Path, paragraphs: list = None, tables: list = None) -> None:
    """Crea un DOCX minimo con python-docx."""
    from docx import Document
    doc = Document()
    for text in (paragraphs or []):
        doc.add_paragraph(text)
    for rows in (tables or []):
        if rows:
            tbl = doc.add_table(rows=len(rows), cols=len(rows[0]))
            for r_idx, row_data in enumerate(rows):
                for c_idx, cell_text in enumerate(row_data):
                    tbl.rows[r_idx].cells[c_idx].text = str(cell_text)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


def _make_full_docx_with_blocks(path: Path) -> None:
    """DOCX con disclaimer, TOC y todos los bloques A-K."""
    paras = [
        "Documento Ambiental",
        "No declara aptitud administrativa",
        "Tabla de contenidos",
    ]
    for block in REQUIRED_BLOCKS:
        paras.append(f"{block} — Titulo del bloque {block}")
        paras.append(f"Contenido del bloque {block}.")
    _make_minimal_docx(path, paragraphs=paras)


def _make_exp(tmp: Path, sub: str = "expediente") -> Path:
    """Crea estructura minima de expediente temporal."""
    exp = tmp / sub
    exp.mkdir(parents=True, exist_ok=True)
    return exp


def _write_json(path: Path, data: dict) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with open(path, "w", encoding="utf-8") as fh:
        json.dump(data, fh, ensure_ascii=False)


# ---------------------------------------------------------------------------
# 1. DocumentQualityIssue
# ---------------------------------------------------------------------------

class TestDocumentQualityIssue(unittest.TestCase):

    def _issue(self, sev="ERROR", code="QC-E001"):
        return DocumentQualityIssue(
            severity=sev,
            code=code,
            file_path="documento/borrador.docx",
            message="Archivo no encontrado.",
            recommendation="Ejecutar build.",
            evidence=["Path: /tmp/x"],
        )

    def test_to_dict_keys(self):
        d = self._issue().to_dict()
        self.assertIn("severity", d)
        self.assertIn("code", d)
        self.assertIn("file_path", d)
        self.assertIn("message", d)
        self.assertIn("recommendation", d)
        self.assertIn("evidence", d)

    def test_to_dict_values(self):
        d = self._issue("WARNING", "QC-W001").to_dict()
        self.assertEqual(d["severity"], "WARNING")
        self.assertEqual(d["code"], "QC-W001")

    def test_summary_contains_severity_and_code(self):
        s = self._issue("ERROR", "QC-E003").summary()
        self.assertIn("ERROR", s)
        self.assertIn("QC-E003", s)

    def test_summary_contains_file_path(self):
        s = self._issue().summary()
        self.assertIn("documento/borrador.docx", s)

    def test_summary_no_file_path(self):
        issue = DocumentQualityIssue(
            severity="INFO", code="QC-I001", file_path=None,
            message="Test.", recommendation="", evidence=[],
        )
        s = issue.summary()
        self.assertIn("QC-I001", s)

    def test_evidence_empty_default(self):
        issue = DocumentQualityIssue(
            severity="INFO", code="QC-I001", file_path=None,
            message="Test.", recommendation="",
        )
        self.assertEqual(issue.evidence, [])


# ---------------------------------------------------------------------------
# 2. DocumentQualityResult
# ---------------------------------------------------------------------------

class TestDocumentQualityResult(unittest.TestCase):

    def _result(self, issues=None) -> DocumentQualityResult:
        return DocumentQualityResult(
            expediente_id="exp-test",
            status="OK",
            issues=issues or [],
        )

    def _issue(self, sev):
        return DocumentQualityIssue(
            severity=sev, code="QC-X001", file_path=None,
            message="msg", recommendation="",
        )

    def test_error_count(self):
        r = self._result([self._issue("ERROR"), self._issue("ERROR"), self._issue("WARNING")])
        self.assertEqual(r.error_count(), 2)

    def test_warning_count(self):
        r = self._result([self._issue("WARNING"), self._issue("INFO")])
        self.assertEqual(r.warning_count(), 1)

    def test_info_count(self):
        r = self._result([self._issue("INFO"), self._issue("INFO")])
        self.assertEqual(r.info_count(), 2)

    def test_is_valid_no_errors(self):
        r = self._result([self._issue("WARNING")])
        self.assertTrue(r.is_valid())

    def test_is_valid_with_errors(self):
        r = self._result([self._issue("ERROR")])
        self.assertFalse(r.is_valid())

    def test_is_valid_empty(self):
        r = self._result([])
        self.assertTrue(r.is_valid())

    def test_to_dict_keys(self):
        d = self._result().to_dict()
        for k in ("expediente_id", "status", "issues", "error_count", "warning_count"):
            self.assertIn(k, d)

    def test_to_dict_issues_serialized(self):
        r = self._result([self._issue("ERROR")])
        d = r.to_dict()
        self.assertIsInstance(d["issues"], list)
        self.assertEqual(d["issues"][0]["severity"], "ERROR")

    def test_summary_contains_status(self):
        r = self._result()
        r.status = "CON_OBSERVACIONES"
        s = r.summary()
        self.assertIn("CON_OBSERVACIONES", s)

    def test_summary_contains_no_admin_disclaimer(self):
        s = self._result().summary()
        self.assertIn("no declara", s.lower())

    def test_summary_valid_message(self):
        r = self._result([])
        self.assertIn("VALIDO", r.summary())

    def test_summary_invalid_message(self):
        r = self._result([self._issue("ERROR")])
        self.assertIn("NO VALIDO", r.summary())


# ---------------------------------------------------------------------------
# 3. check_required_document_files
# ---------------------------------------------------------------------------

class TestCheckRequiredDocumentFiles(unittest.TestCase):

    def setUp(self):
        self.tmp = Path(tempfile.mkdtemp())

    def tearDown(self):
        import shutil
        shutil.rmtree(self.tmp, ignore_errors=True)

    def _make_exp(self):
        return _make_exp(self.tmp)

    def test_empty_expediente_has_errors(self):
        exp = self._make_exp()
        issues = check_required_document_files(exp)
        errors = [i for i in issues if i.severity == "ERROR"]
        self.assertGreater(len(errors), 0)

    def test_missing_docx_base_is_error(self):
        exp = self._make_exp()
        issues = check_required_document_files(exp)
        codes = [i.code for i in issues if i.severity == "ERROR"]
        self.assertIn("QC-E001", codes)

    def test_missing_markdown_is_error(self):
        exp = self._make_exp()
        (exp / "documento").mkdir(exist_ok=True)
        issues = check_required_document_files(exp)
        file_paths = [i.file_path for i in issues if i.severity == "ERROR"]
        self.assertTrue(any("borrador.md" in (fp or "") for fp in file_paths))

    def test_with_minimum_required_files_no_critical_errors(self):
        exp = self._make_exp()
        doc_dir = exp / "documento"
        doc_dir.mkdir(exist_ok=True)
        _make_minimal_docx(doc_dir / "documento_ambiental_borrador.docx")
        (doc_dir / "documento_ambiental_borrador.md").write_text("# Doc", encoding="utf-8")
        issues = check_required_document_files(exp)
        errors = [i for i in issues if i.severity == "ERROR"]
        self.assertEqual(len(errors), 0)

    def test_figures_result_generated_true_without_enriched_docx_is_warning(self):
        exp = self._make_exp()
        doc_dir = exp / "documento"
        doc_dir.mkdir(exist_ok=True)
        _write_json(doc_dir / "document_figures_result.json", {"generated": True})
        issues = check_required_document_files(exp)
        warn_files = [i.file_path for i in issues if i.severity == "WARNING"]
        self.assertTrue(any("con_figuras" in (fp or "") for fp in warn_files))

    def test_no_figures_result_enriched_docx_is_info(self):
        exp = self._make_exp()
        (exp / "documento").mkdir(exist_ok=True)
        issues = check_required_document_files(exp)
        info = [i for i in issues if i.severity == "INFO" and "con_figuras" in (i.file_path or "")]
        self.assertTrue(len(info) >= 1)

    def test_enriched_docx_present_no_info_for_it(self):
        exp = self._make_exp()
        doc_dir = exp / "documento"
        doc_dir.mkdir(exist_ok=True)
        _make_minimal_docx(doc_dir / "documento_ambiental_borrador.docx")
        _make_minimal_docx(doc_dir / "documento_ambiental_borrador_con_figuras.docx")
        (doc_dir / "documento_ambiental_borrador.md").write_text("# Doc", encoding="utf-8")
        issues = check_required_document_files(exp)
        info_enriched = [
            i for i in issues
            if "con_figuras" in (i.file_path or "") and i.severity == "INFO"
        ]
        self.assertEqual(len(info_enriched), 0)


# ---------------------------------------------------------------------------
# 4. select_best_docx_for_qc
# ---------------------------------------------------------------------------

class TestSelectBestDocxForQc(unittest.TestCase):

    def setUp(self):
        self.tmp = Path(tempfile.mkdtemp())

    def tearDown(self):
        import shutil
        shutil.rmtree(self.tmp, ignore_errors=True)

    def _exp(self):
        return _make_exp(self.tmp)

    def test_prefers_enriched_docx(self):
        exp = self._exp()
        doc_dir = exp / "documento"
        doc_dir.mkdir(exist_ok=True)
        _make_minimal_docx(doc_dir / "documento_ambiental_borrador.docx")
        _make_minimal_docx(doc_dir / "documento_ambiental_borrador_con_figuras.docx")
        result = select_best_docx_for_qc(exp)
        self.assertIn("con_figuras", result.name)

    def test_fallback_to_base_docx(self):
        exp = self._exp()
        doc_dir = exp / "documento"
        doc_dir.mkdir(exist_ok=True)
        _make_minimal_docx(doc_dir / "documento_ambiental_borrador.docx")
        result = select_best_docx_for_qc(exp)
        self.assertEqual(result.name, "documento_ambiental_borrador.docx")

    def test_returns_none_if_no_docx(self):
        exp = self._exp()
        (exp / "documento").mkdir(exist_ok=True)
        result = select_best_docx_for_qc(exp)
        self.assertIsNone(result)

    def test_returns_none_if_no_documento_dir(self):
        exp = self._exp()
        result = select_best_docx_for_qc(exp)
        self.assertIsNone(result)


# ---------------------------------------------------------------------------
# 5. validate_docx_opens
# ---------------------------------------------------------------------------

class TestValidateDocxOpens(unittest.TestCase):

    def setUp(self):
        self.tmp = Path(tempfile.mkdtemp())

    def tearDown(self):
        import shutil
        shutil.rmtree(self.tmp, ignore_errors=True)

    def test_valid_docx_returns_true(self):
        p = self.tmp / "test.docx"
        _make_minimal_docx(p, paragraphs=["Hola"])
        self.assertTrue(validate_docx_opens(p))

    def test_nonexistent_file_returns_false(self):
        self.assertFalse(validate_docx_opens(self.tmp / "noexiste.docx"))

    def test_non_docx_file_returns_false(self):
        p = self.tmp / "fake.docx"
        p.write_bytes(b"esto no es un docx")
        self.assertFalse(validate_docx_opens(p))

    def test_empty_file_returns_false(self):
        p = self.tmp / "empty.docx"
        p.write_bytes(b"")
        self.assertFalse(validate_docx_opens(p))


# ---------------------------------------------------------------------------
# 6. extract_docx_text
# ---------------------------------------------------------------------------

class TestExtractDocxText(unittest.TestCase):

    def setUp(self):
        self.tmp = Path(tempfile.mkdtemp())

    def tearDown(self):
        import shutil
        shutil.rmtree(self.tmp, ignore_errors=True)

    def test_extracts_paragraphs(self):
        p = self.tmp / "doc.docx"
        _make_minimal_docx(p, paragraphs=["Parrafo uno", "Parrafo dos"])
        text = extract_docx_text(p)
        self.assertIn("Parrafo uno", text)
        self.assertIn("Parrafo dos", text)

    def test_extracts_table_cells(self):
        p = self.tmp / "doc.docx"
        _make_minimal_docx(p, tables=[[["Celda A", "Celda B"]]])
        text = extract_docx_text(p)
        self.assertIn("Celda A", text)
        self.assertIn("Celda B", text)

    def test_returns_empty_on_missing_file(self):
        text = extract_docx_text(self.tmp / "noexiste.docx")
        self.assertEqual(text, "")

    def test_returns_empty_on_non_docx(self):
        p = self.tmp / "fake.docx"
        p.write_bytes(b"not a docx")
        text = extract_docx_text(p)
        self.assertEqual(text, "")


# ---------------------------------------------------------------------------
# 7. detect_blocks_in_text
# ---------------------------------------------------------------------------

class TestDetectBlocksInText(unittest.TestCase):

    def test_detects_all_blocks_with_separator(self):
        text = "\n".join(f"{b} — Titulo" for b in REQUIRED_BLOCKS)
        found = detect_blocks_in_text(text)
        self.assertEqual(found, sorted(REQUIRED_BLOCKS))

    def test_detects_blocks_with_hash(self):
        text = "## A — Identificacion\n## B — Inventario"
        found = detect_blocks_in_text(text)
        self.assertIn("A", found)
        self.assertIn("B", found)

    def test_detects_bloque_pattern(self):
        text = "Bloque A de identificacion\nBloque K de anexos"
        found = detect_blocks_in_text(text)
        self.assertIn("A", found)
        self.assertIn("K", found)

    def test_detects_dot_separator(self):
        text = "A. Identificacion del proyecto"
        found = detect_blocks_in_text(text)
        self.assertIn("A", found)

    def test_no_duplicates(self):
        text = "A — Titulo\nBloque A\n## A"
        found = detect_blocks_in_text(text)
        self.assertEqual(found.count("A"), 1)

    def test_returns_sorted(self):
        text = "K — Anexos\nA — Ident\nF — Vulnerabilidad"
        found = detect_blocks_in_text(text)
        self.assertEqual(found, sorted(found))

    def test_empty_text(self):
        self.assertEqual(detect_blocks_in_text(""), [])

    def test_partial_blocks(self):
        text = "A — Identificacion\nC — Impactos\nK — Anexos"
        found = detect_blocks_in_text(text)
        self.assertIn("A", found)
        self.assertIn("C", found)
        self.assertIn("K", found)
        self.assertNotIn("B", found)

    def test_bloque_case_insensitive(self):
        text = "bloque a de identificacion"
        found = detect_blocks_in_text(text)
        self.assertIn("A", found)


# ---------------------------------------------------------------------------
# 8. check_docx_structure
# ---------------------------------------------------------------------------

class TestCheckDocxStructure(unittest.TestCase):

    def setUp(self):
        self.tmp = Path(tempfile.mkdtemp())

    def tearDown(self):
        import shutil
        shutil.rmtree(self.tmp, ignore_errors=True)

    def _exp(self):
        return _make_exp(self.tmp)

    def test_full_docx_no_errors(self):
        exp = self._exp()
        doc_dir = exp / "documento"
        doc_dir.mkdir(exist_ok=True)
        _make_full_docx_with_blocks(doc_dir / "documento_ambiental_borrador.docx")
        issues = check_docx_structure(exp)
        errors = [i for i in issues if i.severity == "ERROR"]
        self.assertEqual(len(errors), 0)

    def test_missing_disclaimer_is_error(self):
        exp = self._exp()
        doc_dir = exp / "documento"
        doc_dir.mkdir(exist_ok=True)
        paras = ["Tabla de contenidos"]
        for b in REQUIRED_BLOCKS:
            paras.append(f"{b} — Titulo")
        _make_minimal_docx(doc_dir / "documento_ambiental_borrador.docx", paragraphs=paras)
        issues = check_docx_structure(exp)
        errors = [i for i in issues if i.code == "QC-E004"]
        self.assertTrue(len(errors) >= 1)

    def test_missing_block_is_error(self):
        exp = self._exp()
        doc_dir = exp / "documento"
        doc_dir.mkdir(exist_ok=True)
        paras = ["No declara aptitud administrativa", "Tabla de contenidos"]
        for b in REQUIRED_BLOCKS:
            if b != "K":
                paras.append(f"{b} — Titulo")
        _make_minimal_docx(doc_dir / "documento_ambiental_borrador.docx", paragraphs=paras)
        issues = check_docx_structure(exp)
        block_errors = [i for i in issues if i.code == "QC-E003"]
        self.assertTrue(any("K" in i.message for i in block_errors))

    def test_missing_toc_is_warning(self):
        exp = self._exp()
        doc_dir = exp / "documento"
        doc_dir.mkdir(exist_ok=True)
        paras = ["No declara aptitud administrativa"]
        for b in REQUIRED_BLOCKS:
            paras.append(f"{b} — Titulo")
        _make_minimal_docx(doc_dir / "documento_ambiental_borrador.docx", paragraphs=paras)
        issues = check_docx_structure(exp)
        toc_warnings = [i for i in issues if i.code == "QC-W002"]
        self.assertTrue(len(toc_warnings) >= 1)

    def test_no_docx_returns_error(self):
        exp = self._exp()
        (exp / "documento").mkdir(exist_ok=True)
        issues = check_docx_structure(exp)
        self.assertTrue(any(i.severity == "ERROR" for i in issues))

    def test_unopenable_docx_returns_error(self):
        exp = self._exp()
        doc_dir = exp / "documento"
        doc_dir.mkdir(exist_ok=True)
        (doc_dir / "documento_ambiental_borrador.docx").write_bytes(b"not a docx")
        issues = check_docx_structure(exp)
        self.assertTrue(any(i.code == "QC-E002" for i in issues))

    def test_block_g_partial_is_warning(self):
        exp = self._exp()
        doc_dir = exp / "documento"
        doc_dir.mkdir(exist_ok=True)
        _make_full_docx_with_blocks(doc_dir / "documento_ambiental_borrador.docx")
        build_result = {
            "block_results": [
                {"block_id": "G", "status": "PARTIAL"},
            ]
        }
        _write_json(doc_dir / "document_build_result.json", build_result)
        issues = check_docx_structure(exp)
        partial_warns = [i for i in issues if i.code == "QC-W006"]
        self.assertTrue(len(partial_warns) >= 1)

    def test_prefers_enriched_docx(self):
        exp = self._exp()
        doc_dir = exp / "documento"
        doc_dir.mkdir(exist_ok=True)
        _make_full_docx_with_blocks(doc_dir / "documento_ambiental_borrador.docx")
        _make_full_docx_with_blocks(doc_dir / "documento_ambiental_borrador_con_figuras.docx")
        issues = check_docx_structure(exp)
        errors = [i for i in issues if i.severity == "ERROR"]
        self.assertEqual(len(errors), 0)


# ---------------------------------------------------------------------------
# 9. check_figures_and_captions
# ---------------------------------------------------------------------------

class TestCheckFiguresAndCaptions(unittest.TestCase):

    def setUp(self):
        self.tmp = Path(tempfile.mkdtemp())

    def tearDown(self):
        import shutil
        shutil.rmtree(self.tmp, ignore_errors=True)

    def _exp(self):
        return _make_exp(self.tmp)

    def test_no_figures_result_is_warning(self):
        exp = self._exp()
        (exp / "documento").mkdir(exist_ok=True)
        issues = check_figures_and_captions(exp)
        self.assertTrue(any(i.severity == "WARNING" for i in issues))

    def test_figures_result_generated_true_no_enriched_docx_is_error(self):
        exp = self._exp()
        doc_dir = exp / "documento"
        doc_dir.mkdir(exist_ok=True)
        _write_json(doc_dir / "document_figures_result.json", {
            "generated": True,
            "figures_inserted": ["FIG-001"],
            "figures_skipped": [],
        })
        issues = check_figures_and_captions(exp)
        errors = [i for i in issues if i.severity == "ERROR"]
        self.assertTrue(len(errors) >= 1)

    def test_figures_with_captions_in_docx_no_errors(self):
        exp = self._exp()
        doc_dir = exp / "documento"
        doc_dir.mkdir(exist_ok=True)
        _write_json(doc_dir / "document_figures_result.json", {
            "generated": True,
            "figures_inserted": ["FIG-001", "FIG-002"],
            "figures_skipped": [],
        })
        _make_minimal_docx(
            doc_dir / "documento_ambiental_borrador_con_figuras.docx",
            paragraphs=["Figura FIG-001. Mapa.", "Figura FIG-002. Mapa."],
        )
        issues = check_figures_and_captions(exp)
        errors = [i for i in issues if i.severity == "ERROR"]
        self.assertEqual(len(errors), 0)

    def test_missing_caption_in_docx_is_error(self):
        exp = self._exp()
        doc_dir = exp / "documento"
        doc_dir.mkdir(exist_ok=True)
        _write_json(doc_dir / "document_figures_result.json", {
            "generated": True,
            "figures_inserted": ["FIG-001"],
            "figures_skipped": [],
        })
        _make_minimal_docx(
            doc_dir / "documento_ambiental_borrador_con_figuras.docx",
            paragraphs=["Contenido sin caption"],
        )
        issues = check_figures_and_captions(exp)
        errors = [i for i in issues if i.code == "QC-E005"]
        self.assertTrue(len(errors) >= 1)

    def test_skipped_figures_is_warning(self):
        exp = self._exp()
        doc_dir = exp / "documento"
        doc_dir.mkdir(exist_ok=True)
        _write_json(doc_dir / "document_figures_result.json", {
            "generated": True,
            "figures_inserted": ["FIG-001"],
            "figures_skipped": ["FIG-002"],
        })
        _make_minimal_docx(
            doc_dir / "documento_ambiental_borrador_con_figuras.docx",
            paragraphs=["Figura FIG-001. Mapa."],
        )
        issues = check_figures_and_captions(exp)
        skip_warns = [i for i in issues if i.code == "QC-W004"]
        self.assertTrue(len(skip_warns) >= 1)

    def test_no_figures_inserted_is_warning(self):
        exp = self._exp()
        doc_dir = exp / "documento"
        doc_dir.mkdir(exist_ok=True)
        _write_json(doc_dir / "document_figures_result.json", {
            "generated": False,
            "figures_inserted": [],
            "figures_skipped": [],
        })
        issues = check_figures_and_captions(exp)
        self.assertTrue(any(i.severity == "WARNING" for i in issues))


# ---------------------------------------------------------------------------
# 10. check_final_audit_visibility
# ---------------------------------------------------------------------------

class TestCheckFinalAuditVisibility(unittest.TestCase):

    def setUp(self):
        self.tmp = Path(tempfile.mkdtemp())

    def tearDown(self):
        import shutil
        shutil.rmtree(self.tmp, ignore_errors=True)

    def _exp(self):
        return _make_exp(self.tmp)

    def test_no_audit_file_is_warning(self):
        exp = self._exp()
        issues = check_final_audit_visibility(exp)
        self.assertTrue(any(i.severity == "WARNING" for i in issues))

    def test_conforme_no_issues_about_visibility(self):
        exp = self._exp()
        (exp / "auditoria").mkdir(parents=True, exist_ok=True)
        _write_json(exp / "auditoria" / "final_audit_result.json", {
            "status": "CONFORME",
            "administrative_ready": False,
        })
        issues = check_final_audit_visibility(exp)
        errors = [i for i in issues if i.code in ("QC-E006", "QC-E007")]
        self.assertEqual(len(errors), 0)

    def test_no_conforme_not_visible_in_docx_is_error(self):
        exp = self._exp()
        (exp / "auditoria").mkdir(parents=True, exist_ok=True)
        doc_dir = exp / "documento"
        doc_dir.mkdir(exist_ok=True)
        _write_json(exp / "auditoria" / "final_audit_result.json", {
            "status": "NO_CONFORME",
            "administrative_ready": False,
        })
        _make_minimal_docx(
            doc_dir / "documento_ambiental_borrador.docx",
            paragraphs=["Texto sin mencion al estado."],
        )
        issues = check_final_audit_visibility(exp)
        errors = [i for i in issues if i.code == "QC-E006"]
        self.assertTrue(len(errors) >= 1)

    def test_no_conforme_visible_in_docx_no_error(self):
        exp = self._exp()
        (exp / "auditoria").mkdir(parents=True, exist_ok=True)
        doc_dir = exp / "documento"
        doc_dir.mkdir(exist_ok=True)
        _write_json(exp / "auditoria" / "final_audit_result.json", {
            "status": "NO_CONFORME",
            "administrative_ready": False,
        })
        _make_minimal_docx(
            doc_dir / "documento_ambiental_borrador.docx",
            paragraphs=["Estado: NO_CONFORME - con observaciones detectadas."],
        )
        issues = check_final_audit_visibility(exp)
        errors = [i for i in issues if i.code == "QC-E006"]
        self.assertEqual(len(errors), 0)

    def test_administrative_ready_true_is_error(self):
        exp = self._exp()
        (exp / "auditoria").mkdir(parents=True, exist_ok=True)
        _write_json(exp / "auditoria" / "final_audit_result.json", {
            "status": "CONFORME",
            "administrative_ready": True,
        })
        issues = check_final_audit_visibility(exp)
        errors = [i for i in issues if i.code == "QC-E007"]
        self.assertTrue(len(errors) >= 1)

    def test_administrative_ready_true_in_build_result_is_error(self):
        exp = self._exp()
        doc_dir = exp / "documento"
        doc_dir.mkdir(exist_ok=True)
        _write_json(doc_dir / "docx_build_result.json", {"administrative_ready": True})
        issues = check_final_audit_visibility(exp)
        errors = [i for i in issues if i.code == "QC-E007"]
        self.assertTrue(len(errors) >= 1)

    def test_no_conforme_visible_in_markdown_no_error(self):
        exp = self._exp()
        (exp / "auditoria").mkdir(parents=True, exist_ok=True)
        doc_dir = exp / "documento"
        doc_dir.mkdir(exist_ok=True)
        _write_json(exp / "auditoria" / "final_audit_result.json", {
            "status": "NO_CONFORME",
            "administrative_ready": False,
        })
        (doc_dir / "documento_ambiental_borrador.md").write_text(
            "Auditoria: con observaciones en modo test.", encoding="utf-8"
        )
        issues = check_final_audit_visibility(exp)
        errors = [i for i in issues if i.code == "QC-E006"]
        self.assertEqual(len(errors), 0)

    # --- DOC-05: tests de integracion con texto generado por DOC-01 ---

    def test_no_conforme_with_space_in_docx_no_qce006(self):
        """DOCX con 'NO CONFORME' (espacio) → QC lo detecta, sin QC-E006."""
        exp = self._exp()
        (exp / "auditoria").mkdir(parents=True, exist_ok=True)
        doc_dir = exp / "documento"
        doc_dir.mkdir(exist_ok=True)
        _write_json(exp / "auditoria" / "final_audit_result.json", {
            "status": "NO_CONFORME",
            "administrative_ready": False,
        })
        _make_minimal_docx(
            doc_dir / "documento_ambiental_borrador.docx",
            paragraphs=[
                "AVISO DE AUDITORIA FINAL: califica el expediente como NO CONFORME.",
                "Este borrador no debe considerarse apto para presentacion administrativa.",
            ],
        )
        issues = check_final_audit_visibility(exp)
        errors = [i for i in issues if i.code == "QC-E006"]
        self.assertEqual(len(errors), 0)

    def test_no_conforme_with_underscore_in_docx_no_qce006(self):
        """DOCX con 'NO_CONFORME' (guion bajo) tambien es detectado (DOC-04 resiliente)."""
        exp = self._exp()
        (exp / "auditoria").mkdir(parents=True, exist_ok=True)
        doc_dir = exp / "documento"
        doc_dir.mkdir(exist_ok=True)
        _write_json(exp / "auditoria" / "final_audit_result.json", {
            "status": "NO_CONFORME",
            "administrative_ready": False,
        })
        _make_minimal_docx(
            doc_dir / "documento_ambiental_borrador.docx",
            paragraphs=["Estado de auditoria (AU-04): NO_CONFORME"],
        )
        issues = check_final_audit_visibility(exp)
        errors = [i for i in issues if i.code == "QC-E006"]
        self.assertEqual(len(errors), 0)

    def test_no_conforme_with_observacion_in_docx_no_qce006(self):
        """'observacion' en DOCX tambien satisface la visibilidad."""
        exp = self._exp()
        (exp / "auditoria").mkdir(parents=True, exist_ok=True)
        doc_dir = exp / "documento"
        doc_dir.mkdir(exist_ok=True)
        _write_json(exp / "auditoria" / "final_audit_result.json", {
            "status": "NO_CONFORME",
            "administrative_ready": False,
        })
        _make_minimal_docx(
            doc_dir / "documento_ambiental_borrador.docx",
            paragraphs=["El expediente tiene observaciones pendientes de resolver."],
        )
        issues = check_final_audit_visibility(exp)
        errors = [i for i in issues if i.code == "QC-E006"]
        self.assertEqual(len(errors), 0)


# ---------------------------------------------------------------------------
# 11. check_no_administrative_ready_claim
# ---------------------------------------------------------------------------

class TestCheckNoAdministrativeReadyClaim(unittest.TestCase):

    def test_detects_apto_administrativamente(self):
        issues = check_no_administrative_ready_claim("El expediente es apto administrativamente.")
        self.assertTrue(any(i.code == "QC-E008" for i in issues))

    def test_detects_apto_para_presentacion_administrativa(self):
        issues = check_no_administrative_ready_claim(
            "Documento apto para presentacion administrativa."
        )
        self.assertTrue(any(i.code == "QC-E008" for i in issues))

    def test_allows_no_declara_aptitud(self):
        issues = check_no_administrative_ready_claim(
            "No declara aptitud administrativa."
        )
        admin_errors = [i for i in issues if i.code == "QC-E008"]
        self.assertEqual(len(admin_errors), 0)

    def test_allows_no_apto(self):
        issues = check_no_administrative_ready_claim(
            "Este documento no es apto administrativamente."
        )
        admin_errors = [i for i in issues if i.code == "QC-E008"]
        self.assertEqual(len(admin_errors), 0)

    def test_empty_text_no_issues(self):
        self.assertEqual(check_no_administrative_ready_claim(""), [])

    def test_normal_text_no_issues(self):
        issues = check_no_administrative_ready_claim(
            "El proyecto se ubica en zona industrial. Requiere evaluacion ambiental."
        )
        self.assertEqual(len(issues), 0)

    def test_detects_listo_para_presentar(self):
        issues = check_no_administrative_ready_claim("Paquete listo para presentar.")
        self.assertTrue(any(i.code == "QC-E008" for i in issues))

    def test_detects_validado_administrativamente(self):
        issues = check_no_administrative_ready_claim(
            "Expediente validado administrativamente segun normativa."
        )
        self.assertTrue(any(i.code == "QC-E008" for i in issues))

    def test_one_issue_per_phrase(self):
        text = "apto administrativamente " * 5
        issues = check_no_administrative_ready_claim(text)
        codes = [i.code for i in issues if i.code == "QC-E008"]
        self.assertEqual(len(codes), 1)

    def test_detects_conforme_para_presentar(self):
        issues = check_no_administrative_ready_claim("Conforme para presentar ante el organo.")
        self.assertTrue(any(i.code == "QC-E008" for i in issues))


# ---------------------------------------------------------------------------
# 12. run_document_quality_check
# ---------------------------------------------------------------------------

class TestRunDocumentQualityCheck(unittest.TestCase):

    def setUp(self):
        self.tmp = Path(tempfile.mkdtemp())

    def tearDown(self):
        import shutil
        shutil.rmtree(self.tmp, ignore_errors=True)

    def _exp(self):
        return _make_exp(self.tmp)

    def test_empty_expediente_is_sin_datos_or_no_conforme(self):
        exp = self._exp()
        result = run_document_quality_check(exp)
        self.assertIn(result.status, ("SIN_DATOS", "NO_CONFORME"))

    def test_empty_expediente_not_valid(self):
        exp = self._exp()
        result = run_document_quality_check(exp)
        self.assertFalse(result.is_valid())

    def test_complete_expediente_ok_or_con_observaciones(self):
        exp = self._exp()
        doc_dir = exp / "documento"
        doc_dir.mkdir(exist_ok=True)
        _make_full_docx_with_blocks(doc_dir / "documento_ambiental_borrador.docx")
        (doc_dir / "documento_ambiental_borrador.md").write_text(
            "# Doc\nNo declara aptitud administrativa\n"
            + "\n".join(f"{b} — Titulo" for b in REQUIRED_BLOCKS),
            encoding="utf-8",
        )
        _write_json(doc_dir / "document_manifest.json", {"status": "OK"})
        _write_json(doc_dir / "document_build_result.json", {"block_results": []})
        _write_json(doc_dir / "docx_build_result.json", {"generated": True})
        result = run_document_quality_check(exp)
        self.assertIn(result.status, ("OK", "CON_OBSERVACIONES"))

    def test_does_not_modify_files(self):
        exp = self._exp()
        doc_dir = exp / "documento"
        doc_dir.mkdir(exist_ok=True)
        docx_path = doc_dir / "documento_ambiental_borrador.docx"
        _make_full_docx_with_blocks(docx_path)
        mtime_before = docx_path.stat().st_mtime
        run_document_quality_check(exp)
        mtime_after = docx_path.stat().st_mtime
        self.assertEqual(mtime_before, mtime_after)

    def test_expediente_id_is_dirname(self):
        exp = self._exp()
        result = run_document_quality_check(exp)
        self.assertEqual(result.expediente_id, "expediente")

    def test_result_has_notes(self):
        exp = self._exp()
        result = run_document_quality_check(exp)
        self.assertTrue(len(result.notes) > 0)

    def test_notes_contain_no_admin_disclaimer(self):
        exp = self._exp()
        result = run_document_quality_check(exp)
        combined = " ".join(result.notes).lower()
        self.assertIn("no declara", combined)

    def test_figures_found_populated_when_present(self):
        exp = self._exp()
        doc_dir = exp / "documento"
        doc_dir.mkdir(exist_ok=True)
        _write_json(doc_dir / "document_figures_result.json", {
            "generated": True,
            "figures_inserted": ["FIG-001"],
            "figures_skipped": [],
        })
        _make_minimal_docx(
            doc_dir / "documento_ambiental_borrador_con_figuras.docx",
            paragraphs=["Figura FIG-001. Mapa de situacion."],
        )
        _make_full_docx_with_blocks(doc_dir / "documento_ambiental_borrador.docx")
        result = run_document_quality_check(exp)
        self.assertIn("FIG-001", result.figures_found)

    def test_captions_found_verified(self):
        exp = self._exp()
        doc_dir = exp / "documento"
        doc_dir.mkdir(exist_ok=True)
        _write_json(doc_dir / "document_figures_result.json", {
            "generated": True,
            "figures_inserted": ["FIG-001"],
            "figures_skipped": [],
        })
        _make_minimal_docx(
            doc_dir / "documento_ambiental_borrador_con_figuras.docx",
            paragraphs=["Figura FIG-001. Mapa."],
        )
        _make_full_docx_with_blocks(doc_dir / "documento_ambiental_borrador.docx")
        result = run_document_quality_check(exp)
        self.assertIn("FIG-001", result.captions_found)

    def test_status_no_conforme_when_errors(self):
        exp = self._exp()
        result = run_document_quality_check(exp)
        errors = result.error_count()
        if errors > 0 and len(result.missing_files) < 4:
            self.assertEqual(result.status, "NO_CONFORME")

    def test_status_sin_datos_when_almost_all_missing(self):
        exp = self._exp()
        result = run_document_quality_check(exp)
        missing_req = sum(
            1 for r in REQUIRED_DOCUMENT_FILES if not (exp / r).exists()
        )
        if missing_req >= 4:
            self.assertEqual(result.status, "SIN_DATOS")


# ---------------------------------------------------------------------------
# 13. Markdown report
# ---------------------------------------------------------------------------

class TestBuildDocumentQualityReportMarkdown(unittest.TestCase):

    def _result(self, status="OK", issues=None):
        return DocumentQualityResult(
            expediente_id="exp-test",
            status=status,
            issues=issues or [],
            blocks_found=list(REQUIRED_BLOCKS),
            blocks_missing=[],
        )

    def test_contains_title(self):
        md = build_document_quality_report_markdown(self._result())
        self.assertIn("Control de calidad", md)

    def test_contains_status(self):
        md = build_document_quality_report_markdown(self._result(status="CON_OBSERVACIONES"))
        self.assertIn("CON_OBSERVACIONES", md)

    def test_contains_blocks_section(self):
        md = build_document_quality_report_markdown(self._result())
        self.assertIn("Bloques A-K", md)

    def test_contains_figures_section(self):
        md = build_document_quality_report_markdown(self._result())
        self.assertIn("Figuras", md)

    def test_contains_auditoria_section(self):
        md = build_document_quality_report_markdown(self._result())
        self.assertIn("Auditoria", md)

    def test_contains_no_aptitud_administrativa(self):
        md = build_document_quality_report_markdown(self._result())
        self.assertIn("no declara", md.lower())

    def test_contains_disclaimer_section(self):
        md = build_document_quality_report_markdown(self._result())
        self.assertIn("Advertencia de alcance", md)

    def test_contains_recomendaciones_section(self):
        md = build_document_quality_report_markdown(self._result())
        self.assertIn("Recomendaciones", md)

    def test_errors_listed(self):
        issue = DocumentQualityIssue(
            severity="ERROR", code="QC-E003", file_path=None,
            message="Bloque K faltante.", recommendation="Regenerar.",
        )
        md = build_document_quality_report_markdown(self._result(issues=[issue]))
        self.assertIn("Bloque K faltante", md)

    def test_warnings_listed(self):
        issue = DocumentQualityIssue(
            severity="WARNING", code="QC-W002", file_path=None,
            message="Sin indice.", recommendation="Revisar.",
        )
        md = build_document_quality_report_markdown(self._result(issues=[issue]))
        self.assertIn("Sin indice", md)


# ---------------------------------------------------------------------------
# 14. write_document_quality_outputs
# ---------------------------------------------------------------------------

class TestWriteDocumentQualityOutputs(unittest.TestCase):

    def setUp(self):
        self.tmp = Path(tempfile.mkdtemp())

    def tearDown(self):
        import shutil
        shutil.rmtree(self.tmp, ignore_errors=True)

    def _result(self):
        return DocumentQualityResult(
            expediente_id="exp-write-test",
            status="OK",
        )

    def test_writes_json_and_md(self):
        json_path, md_path = write_document_quality_outputs(self._result(), self.tmp)
        self.assertTrue(json_path.exists())
        self.assertTrue(md_path.exists())

    def test_json_is_loadable(self):
        json_path, _ = write_document_quality_outputs(self._result(), self.tmp)
        with open(json_path, encoding="utf-8") as fh:
            data = json.load(fh)
        self.assertIn("expediente_id", data)

    def test_json_has_correct_expediente_id(self):
        json_path, _ = write_document_quality_outputs(self._result(), self.tmp)
        with open(json_path, encoding="utf-8") as fh:
            data = json.load(fh)
        self.assertEqual(data["expediente_id"], "exp-write-test")

    def test_md_contains_title(self):
        _, md_path = write_document_quality_outputs(self._result(), self.tmp)
        content = md_path.read_text(encoding="utf-8")
        self.assertIn("Control de calidad", content)

    def test_creates_output_dir_if_missing(self):
        out = self.tmp / "subdir" / "nested"
        write_document_quality_outputs(self._result(), out)
        self.assertTrue(out.exists())

    def test_returns_correct_paths(self):
        json_path, md_path = write_document_quality_outputs(self._result(), self.tmp)
        self.assertEqual(json_path.name, "document_quality_result.json")
        self.assertEqual(md_path.name, "document_quality_result.md")


# ---------------------------------------------------------------------------
# 15. CLI integration
# ---------------------------------------------------------------------------

class TestDocumentQcCLI(unittest.TestCase):

    def setUp(self):
        self.tmp = Path(tempfile.mkdtemp())

    def tearDown(self):
        import shutil
        shutil.rmtree(self.tmp, ignore_errors=True)

    def _run(self, exp_path: Path, extra_args: list = None) -> int:
        import sys
        sys.argv = ["run_expediente.py", str(exp_path), "document-qc"] + (extra_args or [])
        import importlib
        import run_expediente as re_mod
        importlib.reload(re_mod)
        return re_mod.main(sys.argv[1:])

    def test_dry_run_no_files_written(self):
        exp = _make_exp(self.tmp)
        doc_dir = exp / "documento"
        doc_dir.mkdir(exist_ok=True)
        _make_full_docx_with_blocks(doc_dir / "documento_ambiental_borrador.docx")
        self._run(exp)
        self.assertFalse((doc_dir / "document_quality_result.json").exists())

    def test_write_creates_files(self):
        exp = _make_exp(self.tmp)
        doc_dir = exp / "documento"
        doc_dir.mkdir(exist_ok=True)
        _make_full_docx_with_blocks(doc_dir / "documento_ambiental_borrador.docx")
        (doc_dir / "documento_ambiental_borrador.md").write_text("# Doc", encoding="utf-8")
        _write_json(doc_dir / "document_manifest.json", {})
        _write_json(doc_dir / "document_build_result.json", {"block_results": []})
        _write_json(doc_dir / "docx_build_result.json", {})
        self._run(exp, ["--write"])
        self.assertTrue((doc_dir / "document_quality_result.json").exists())
        self.assertTrue((doc_dir / "document_quality_result.md").exists())

    def test_exit_1_when_errors(self):
        exp = _make_exp(self.tmp)
        code = self._run(exp)
        self.assertEqual(code, 1)

    def test_exit_0_when_valid(self):
        exp = _make_exp(self.tmp)
        doc_dir = exp / "documento"
        doc_dir.mkdir(exist_ok=True)
        _make_full_docx_with_blocks(doc_dir / "documento_ambiental_borrador.docx")
        (doc_dir / "documento_ambiental_borrador.md").write_text("# Doc", encoding="utf-8")
        _write_json(doc_dir / "document_manifest.json", {})
        _write_json(doc_dir / "document_build_result.json", {"block_results": []})
        _write_json(doc_dir / "docx_build_result.json", {})
        code = self._run(exp)
        self.assertEqual(code, 0)


# ---------------------------------------------------------------------------
# Constantes
# ---------------------------------------------------------------------------

class TestConstants(unittest.TestCase):

    def test_required_document_files_count(self):
        self.assertEqual(len(REQUIRED_DOCUMENT_FILES), 6)

    def test_optional_enriched_files_count(self):
        self.assertEqual(len(OPTIONAL_ENRICHED_FILES), 3)

    def test_required_blocks_count(self):
        self.assertEqual(len(REQUIRED_BLOCKS), 11)

    def test_required_blocks_are_a_to_k(self):
        self.assertEqual(REQUIRED_BLOCKS, ["A", "B", "C", "D", "E", "F", "G", "H", "I", "J", "K"])

    def test_status_values(self):
        self.assertIn("OK", DOCUMENT_QC_STATUS)
        self.assertIn("NO_CONFORME", DOCUMENT_QC_STATUS)
        self.assertIn("SIN_DATOS", DOCUMENT_QC_STATUS)

    def test_severity_values(self):
        self.assertIn("ERROR", DOCUMENT_QC_SEVERITY)
        self.assertIn("WARNING", DOCUMENT_QC_SEVERITY)
        self.assertIn("INFO", DOCUMENT_QC_SEVERITY)


# ---------------------------------------------------------------------------
# safe_load_json
# ---------------------------------------------------------------------------

class TestSafeLoadJson(unittest.TestCase):

    def setUp(self):
        self.tmp = Path(tempfile.mkdtemp())

    def tearDown(self):
        import shutil
        shutil.rmtree(self.tmp, ignore_errors=True)

    def test_loads_valid_json(self):
        p = self.tmp / "data.json"
        _write_json(p, {"key": "value"})
        data = safe_load_json(p)
        self.assertEqual(data["key"], "value")

    def test_returns_none_if_missing(self):
        self.assertIsNone(safe_load_json(self.tmp / "noexiste.json"))

    def test_returns_none_if_corrupted(self):
        p = self.tmp / "bad.json"
        p.write_text("{ not valid json", encoding="utf-8")
        self.assertIsNone(safe_load_json(p))

    def test_loads_list_json(self):
        p = self.tmp / "list.json"
        p.write_text("[1, 2, 3]", encoding="utf-8")
        data = safe_load_json(p)
        self.assertEqual(data, [1, 2, 3])


if __name__ == "__main__":
    unittest.main()
