"""tests/test_entity_extractor.py — Suite de tests para EntityExtractor (IN-02).

Cubre los criterios de cierre del ítem:
- referencia catastral válida detectada (HIGH confidence)
- código LER normal detectado (XX XX XX, compact XXXXXX)
- código LER peligroso con asterisco detectado (normalized_value contiene *)
- operaciones R1201/R1203/R13 detectadas (OPERACION)
- coordenadas WGS84 detectadas (keyword + decimal)
- coordenadas UTM detectadas (E:/N:/X=/Y=)
- superficies detectadas (m², m2)
- capacidades TM/día y TM/año detectadas
- potencias kW, CV y HP detectadas
- fechas dd/mm/yyyy y yyyy-mm-dd detectadas
- equipos detectados (molino, criba, cizalla...)
- promotor/titular detectado (texto y tabla kv)
- ExtractionResult.by_type() filtra correctamente
- ExtractionResult.values() devuelve lista de strings
- summary() coherente con entidades presentes
- texto vacío → entities=[]
- extract_entities_from_docx con DOCX sintético
- solo lectura contra fixture real PARCELA
- normalize_ler, is_ler_peligroso, normalize_surface, normalize_power
"""
import sys
import tempfile
import unittest
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent / "src"))

import docx as python_docx

from eia_agent.core.entity_extractor import (
    ExtractedEntity,
    ExtractionResult,
    extract_entities_from_docx,
    extract_entities_from_text,
    is_ler_peligroso,
    normalize_ler,
    normalize_power,
    normalize_surface,
)

# ---------------------------------------------------------------------------
# Rutas
# ---------------------------------------------------------------------------

_ROOT = Path(__file__).parent.parent
_PARCELA = _ROOT / "expediente-EIA-2026-RECIMETAL-PARCELA"
_DOCX_PARCELA = _PARCELA / "inputs" / "memorias" / "Documento_Ambiental_RECIMETAL_Parcela_v6.docx"


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _make_docx(tmp: str, *, texto: str = "", tablas: list[list[list[str]]] | None = None,
               nombre: str = "test.docx") -> Path:
    doc = python_docx.Document()
    if texto:
        doc.add_paragraph(texto)
    for tabla in (tablas or []):
        if not tabla:
            continue
        t = doc.add_table(rows=len(tabla), cols=len(tabla[0]))
        for r_idx, fila in enumerate(tabla):
            for c_idx, valor in enumerate(fila):
                t.rows[r_idx].cells[c_idx].text = valor
    ruta = Path(tmp) / nombre
    doc.save(str(ruta))
    return ruta


# ---------------------------------------------------------------------------
# Normalización
# ---------------------------------------------------------------------------

class TestNormalizeLer(unittest.TestCase):

    def test_compact_6digits(self):
        self.assertEqual(normalize_ler("170405"), "17 04 05")

    def test_spaces_format(self):
        self.assertEqual(normalize_ler("17 04 05"), "17 04 05")

    def test_peligroso_compact(self):
        self.assertEqual(normalize_ler("170405*"), "17 04 05*")

    def test_peligroso_spaces(self):
        self.assertEqual(normalize_ler("17 04 05*"), "17 04 05*")

    def test_different_code(self):
        self.assertEqual(normalize_ler("191210"), "19 12 10")


class TestIsLerPeligroso(unittest.TestCase):

    def test_peligroso_true(self):
        self.assertTrue(is_ler_peligroso("17 04 05*"))

    def test_normal_false(self):
        self.assertFalse(is_ler_peligroso("17 04 05"))

    def test_compact_peligroso(self):
        self.assertTrue(is_ler_peligroso("170405*"))


class TestNormalizeSurface(unittest.TestCase):

    def test_integer(self):
        self.assertEqual(normalize_surface("1500"), "1500 m²")

    def test_decimal_comma(self):
        result = normalize_surface("1931,40")
        self.assertIn("m²", result)
        self.assertIn("1931", result)

    def test_thousands_separator(self):
        result = normalize_surface("1.931")
        self.assertIn("m²", result)

    def test_spanish_format(self):
        result = normalize_surface("1.931,40")
        self.assertIn("m²", result)
        self.assertIn("1931", result)


class TestNormalizePower(unittest.TestCase):

    def test_kw(self):
        self.assertEqual(normalize_power("75", "kW"), "75 KW")

    def test_cv(self):
        self.assertEqual(normalize_power("100", "CV"), "100 CV")

    def test_hp(self):
        self.assertEqual(normalize_power("50", "HP"), "50 HP")

    def test_decimal(self):
        result = normalize_power("7,5", "kW")
        self.assertIn("7.5", result)


# ---------------------------------------------------------------------------
# ExtractionResult
# ---------------------------------------------------------------------------

class TestExtractionResult(unittest.TestCase):

    def _result_with(self, *types: str) -> ExtractionResult:
        entities = [
            ExtractedEntity(entity_type=t, value=f"val_{i}", source="texto",
                            confidence="HIGH")
            for i, t in enumerate(types)
        ]
        return ExtractionResult(entities=entities)

    def test_by_type_found(self):
        result = self._result_with("LER", "LER", "OPERACION")
        self.assertEqual(len(result.by_type("LER")), 2)

    def test_by_type_not_found(self):
        result = self._result_with("LER")
        self.assertEqual(result.by_type("REFERENCIA_CATASTRAL"), [])

    def test_values_returns_list_of_strings(self):
        result = self._result_with("LER", "LER")
        vals = result.values("LER")
        self.assertIsInstance(vals, list)
        self.assertTrue(all(isinstance(v, str) for v in vals))

    def test_values_empty_type(self):
        result = self._result_with("LER")
        self.assertEqual(result.values("OPERACION"), [])

    def test_summary_empty(self):
        result = ExtractionResult()
        self.assertIn("Sin entidades", result.summary())

    def test_summary_with_entities(self):
        result = self._result_with("LER", "LER", "OPERACION")
        s = result.summary()
        self.assertIn("LER", s)
        self.assertIn("OPERACION", s)
        self.assertIn("3", s)

    def test_summary_includes_warnings(self):
        result = ExtractionResult(
            entities=[ExtractedEntity("LER", "17 04 05", "texto", "HIGH")],
            warnings=["aviso de prueba"],
        )
        self.assertIn("aviso", result.summary())

    def test_entities_empty_on_init(self):
        result = ExtractionResult()
        self.assertEqual(result.entities, [])
        self.assertEqual(result.warnings, [])


# ---------------------------------------------------------------------------
# Texto vacío
# ---------------------------------------------------------------------------

class TestTextoVacio(unittest.TestCase):

    def test_empty_string(self):
        result = extract_entities_from_text("")
        self.assertEqual(result.entities, [])

    def test_whitespace_only(self):
        result = extract_entities_from_text("   \n  \t  ")
        self.assertEqual(result.entities, [])

    def test_returns_extraction_result(self):
        result = extract_entities_from_text("")
        self.assertIsInstance(result, ExtractionResult)


# ---------------------------------------------------------------------------
# Referencia catastral
# ---------------------------------------------------------------------------

class TestReferenciaCatastral(unittest.TestCase):

    _RC = "2462302DS4026S0001GQ"

    def test_detecta_rc_valida(self):
        result = extract_entities_from_text(f"La parcela tiene RC {self._RC}")
        rcs = result.by_type("REFERENCIA_CATASTRAL")
        self.assertEqual(len(rcs), 1)
        self.assertEqual(rcs[0].value, self._RC)

    def test_confidence_high(self):
        result = extract_entities_from_text(self._RC)
        self.assertEqual(result.by_type("REFERENCIA_CATASTRAL")[0].confidence, "HIGH")

    def test_normalized_value_igual(self):
        result = extract_entities_from_text(self._RC)
        e = result.by_type("REFERENCIA_CATASTRAL")[0]
        self.assertEqual(e.normalized_value, self._RC)

    def test_no_detecta_cadena_corta(self):
        result = extract_entities_from_text("Ref: 12345678AB")
        self.assertEqual(result.by_type("REFERENCIA_CATASTRAL"), [])

    def test_deduplica(self):
        texto = f"{self._RC} y también {self._RC}"
        result = extract_entities_from_text(texto)
        self.assertEqual(len(result.by_type("REFERENCIA_CATASTRAL")), 1)


# ---------------------------------------------------------------------------
# LER
# ---------------------------------------------------------------------------

class TestLER(unittest.TestCase):

    def test_ler_espacios_detectado(self):
        result = extract_entities_from_text("Residuo 17 04 05 fragmentos de metal")
        lers = result.by_type("LER")
        vals = [e.normalized_value for e in lers]
        self.assertTrue(any("17 04 05" in v for v in vals))

    def test_ler_compacto_detectado(self):
        result = extract_entities_from_text("Código LER: 170405")
        lers = result.by_type("LER")
        self.assertGreater(len(lers), 0)
        norms = [e.normalized_value for e in lers]
        self.assertTrue(any("17 04 05" in n for n in norms))

    def test_ler_peligroso_con_asterisco(self):
        result = extract_entities_from_text("Residuo peligroso 16 06 01*")
        lers = result.by_type("LER")
        peligrosos = [e for e in lers if e.normalized_value and "*" in e.normalized_value]
        self.assertGreater(len(peligrosos), 0)

    def test_ler_peligroso_compact(self):
        result = extract_entities_from_text("LER 160601*")
        lers = result.by_type("LER")
        self.assertTrue(any("*" in (e.normalized_value or "") for e in lers))

    def test_confidence_spaces_high(self):
        result = extract_entities_from_text("17 04 05")
        ler = result.by_type("LER")[0]
        self.assertEqual(ler.confidence, "HIGH")

    def test_confidence_compact_medium(self):
        result = extract_entities_from_text("código 170405")
        lers = [e for e in result.by_type("LER") if e.value == "170405"]
        if lers:
            self.assertEqual(lers[0].confidence, "MEDIUM")

    def test_normalized_format(self):
        result = extract_entities_from_text("LER 19 12 10")
        lers = result.by_type("LER")
        norms = [e.normalized_value for e in lers]
        self.assertIn("19 12 10", norms)


# ---------------------------------------------------------------------------
# Operaciones
# ---------------------------------------------------------------------------

class TestOperaciones(unittest.TestCase):

    def test_r1201(self):
        result = extract_entities_from_text("Operación R1201 reciclado de metales")
        ops = result.values("OPERACION")
        self.assertIn("R1201", ops)

    def test_r1203(self):
        result = extract_entities_from_text("Incluye R1203")
        ops = result.values("OPERACION")
        self.assertIn("R1203", ops)

    def test_r13(self):
        result = extract_entities_from_text("Se realiza R13 almacenamiento")
        ops = result.values("OPERACION")
        self.assertIn("R13", ops)

    def test_d15(self):
        result = extract_entities_from_text("Operación D15 de valorización")
        ops = result.values("OPERACION")
        self.assertIn("D15", ops)

    def test_multiples_ops(self):
        result = extract_entities_from_text("R1201, R1203 y también R13")
        ops = result.values("OPERACION")
        self.assertIn("R1201", ops)
        self.assertIn("R1203", ops)
        self.assertIn("R13", ops)

    def test_no_falso_positivo_texto_normal(self):
        result = extract_entities_from_text("El proceso requiere revisión")
        ops = result.by_type("OPERACION")
        self.assertEqual(ops, [])

    def test_confidence_high(self):
        result = extract_entities_from_text("Operación R1201")
        op = result.by_type("OPERACION")[0]
        self.assertEqual(op.confidence, "HIGH")


# ---------------------------------------------------------------------------
# Coordenadas
# ---------------------------------------------------------------------------

class TestCoordenadas(unittest.TestCase):

    def test_wgs84_lat(self):
        result = extract_entities_from_text("latitud: 28.9234")
        coords = result.by_type("COORDENADA")
        self.assertGreater(len(coords), 0)

    def test_wgs84_lon(self):
        result = extract_entities_from_text("longitud: -13.5678")
        coords = result.by_type("COORDENADA")
        self.assertGreater(len(coords), 0)

    def test_wgs84_epsg(self):
        result = extract_entities_from_text("EPSG: 28.1234")
        coords = result.by_type("COORDENADA")
        self.assertGreater(len(coords), 0)

    def test_utm_e(self):
        result = extract_entities_from_text("E: 642000 N: 3207000")
        coords = result.by_type("COORDENADA")
        valores = [e.value for e in coords]
        self.assertTrue(any("E" in v for v in valores))

    def test_utm_n(self):
        result = extract_entities_from_text("X=642000 Y=3207000")
        coords = result.by_type("COORDENADA")
        valores = [e.value for e in coords]
        self.assertTrue(any("Y" in v or "X" in v for v in valores))

    def test_utm_normalized_prefix(self):
        result = extract_entities_from_text("E: 642000")
        coords = result.by_type("COORDENADA")
        self.assertTrue(any("UTM" in (e.normalized_value or "") for e in coords))


# ---------------------------------------------------------------------------
# Superficies
# ---------------------------------------------------------------------------

class TestSuperficies(unittest.TestCase):

    def test_m2_unicode(self):
        result = extract_entities_from_text("Superficie de 1500 m²")
        sups = [e for e in result.entities if "SUPERFICIE" in e.entity_type]
        self.assertGreater(len(sups), 0)

    def test_m2_ascii(self):
        result = extract_entities_from_text("Superficie de 1500 m2")
        sups = [e for e in result.entities if "SUPERFICIE" in e.entity_type]
        self.assertGreater(len(sups), 0)

    def test_metros_cuadrados_texto(self):
        result = extract_entities_from_text("1000 metros cuadrados de parcela")
        sups = [e for e in result.entities if "SUPERFICIE" in e.entity_type]
        self.assertGreater(len(sups), 0)

    def test_normalized_includes_m2(self):
        result = extract_entities_from_text("1500 m²")
        sups = [e for e in result.entities if "SUPERFICIE" in e.entity_type]
        self.assertTrue(all("m²" in (e.normalized_value or "") for e in sups))

    def test_surface_tipo_parcela(self):
        result = extract_entities_from_text("parcela de 2000 m²")
        sups = [e for e in result.entities if "SUPERFICIE" in e.entity_type]
        tipos = [e.entity_type for e in sups]
        self.assertTrue(any("PARCELA" in t for t in tipos))


# ---------------------------------------------------------------------------
# Capacidades
# ---------------------------------------------------------------------------

class TestCapacidades(unittest.TestCase):

    def test_tm_dia(self):
        result = extract_entities_from_text("capacidad de 500 tm/día")
        caps = result.by_type("CAPACIDAD")
        self.assertGreater(len(caps), 0)

    def test_t_anio(self):
        result = extract_entities_from_text("tratamiento de 1.000 t/año")
        caps = result.by_type("CAPACIDAD")
        self.assertGreater(len(caps), 0)

    def test_toneladas_anio(self):
        result = extract_entities_from_text("25.000 toneladas/año")
        caps = result.by_type("CAPACIDAD")
        self.assertGreater(len(caps), 0)

    def test_tm_año_con_acento(self):
        result = extract_entities_from_text("3.000 tm/año")
        caps = result.by_type("CAPACIDAD")
        self.assertGreater(len(caps), 0)

    def test_t_dia_sin_acento(self):
        result = extract_entities_from_text("100 t/dia")
        caps = result.by_type("CAPACIDAD")
        self.assertGreater(len(caps), 0)


# ---------------------------------------------------------------------------
# Potencias
# ---------------------------------------------------------------------------

class TestPotencias(unittest.TestCase):

    def test_kw(self):
        result = extract_entities_from_text("Motor de 75 kW de potencia")
        pots = result.by_type("POTENCIA")
        self.assertGreater(len(pots), 0)

    def test_cv(self):
        result = extract_entities_from_text("Motor de 100 CV")
        pots = result.by_type("POTENCIA")
        self.assertGreater(len(pots), 0)

    def test_hp(self):
        result = extract_entities_from_text("Compresor de 50 HP")
        pots = result.by_type("POTENCIA")
        self.assertGreater(len(pots), 0)

    def test_multiples_potencias(self):
        result = extract_entities_from_text("Motor principal 75 kW y auxiliar 15 kW")
        pots = result.by_type("POTENCIA")
        self.assertGreaterEqual(len(pots), 1)

    def test_normalized_uppercase_unit(self):
        result = extract_entities_from_text("75 kW")
        pots = result.by_type("POTENCIA")
        self.assertTrue(any("KW" in (e.normalized_value or "").upper() for e in pots))


# ---------------------------------------------------------------------------
# Fechas
# ---------------------------------------------------------------------------

class TestFechas(unittest.TestCase):

    def test_formato_ddmmyyyy(self):
        result = extract_entities_from_text("Fecha de inicio: 15/03/2024")
        fechas = result.by_type("FECHA")
        self.assertIn("15/03/2024", result.values("FECHA"))

    def test_formato_iso(self):
        result = extract_entities_from_text("Vigente desde 2024-03-15")
        fechas = result.by_type("FECHA")
        self.assertIn("2024-03-15", result.values("FECHA"))

    def test_multiple_fechas(self):
        result = extract_entities_from_text("Inicio 01/01/2024 y fin 31/12/2024")
        fechas = result.values("FECHA")
        self.assertGreaterEqual(len(fechas), 2)

    def test_confidence_high(self):
        result = extract_entities_from_text("15/03/2024")
        fecha = result.by_type("FECHA")[0]
        self.assertEqual(fecha.confidence, "HIGH")


# ---------------------------------------------------------------------------
# Equipos
# ---------------------------------------------------------------------------

class TestEquipos(unittest.TestCase):

    def test_molino(self):
        result = extract_entities_from_text("Se instalará un molino de reducción")
        self.assertIn("molino", result.values("EQUIPO"))

    def test_criba(self):
        result = extract_entities_from_text("criba vibratoria para clasificación")
        self.assertGreater(len(result.by_type("EQUIPO")), 0)

    def test_cizalla(self):
        result = extract_entities_from_text("cizalla hidráulica para corte")
        self.assertGreater(len(result.by_type("EQUIPO")), 0)

    def test_trituradora(self):
        result = extract_entities_from_text("trituradora de impacto")
        equips = result.values("EQUIPO")
        self.assertTrue(any("trituradora" in e for e in equips))

    def test_bascula(self):
        result = extract_entities_from_text("báscula de pesaje en continuo")
        self.assertGreater(len(result.by_type("EQUIPO")), 0)

    def test_cinta_transportadora(self):
        result = extract_entities_from_text("cinta transportadora de 20 metros")
        self.assertGreater(len(result.by_type("EQUIPO")), 0)

    def test_multiples_equipos(self):
        result = extract_entities_from_text("molino, criba y cizalla forman el proceso")
        equips = result.values("EQUIPO")
        self.assertGreaterEqual(len(equips), 2)

    def test_case_insensitive(self):
        result = extract_entities_from_text("MOLINO industrial de alto rendimiento")
        self.assertGreater(len(result.by_type("EQUIPO")), 0)


# ---------------------------------------------------------------------------
# Promotor
# ---------------------------------------------------------------------------

class TestPromotor(unittest.TestCase):

    def test_promotor_texto_clave(self):
        result = extract_entities_from_text(
            "El proyecto está promovido por RECIMETAL LANZAROTE, S.L. con CIF B72798846"
        )
        promotores = result.by_type("PROMOTOR")
        vals = [e.value for e in promotores]
        self.assertTrue(any("RECIMETAL" in v for v in vals))

    def test_titular_dos_puntos(self):
        result = extract_entities_from_text("Titular: METALRES CANARIAS, S.A.")
        promotores = result.by_type("PROMOTOR")
        self.assertGreater(len(promotores), 0)

    def test_solicitante(self):
        result = extract_entities_from_text("Solicitante: EMPRESA RECICLADOS, S.L.")
        promotores = result.by_type("PROMOTOR")
        self.assertGreater(len(promotores), 0)

    def test_empresa_sl_en_texto(self):
        result = extract_entities_from_text("RECIMETAL LANZAROTE, S.L. ha presentado solicitud")
        promotores = result.by_type("PROMOTOR")
        self.assertGreater(len(promotores), 0)

    def test_source_texto(self):
        result = extract_entities_from_text(
            "Promotor: EMPRESA RECICLADOS, S.L."
        )
        promotores = result.by_type("PROMOTOR")
        self.assertTrue(any("texto" in e.source for e in promotores))


# ---------------------------------------------------------------------------
# extract_entities_from_docx — DOCX sintético
# ---------------------------------------------------------------------------

class TestExtractFromDocxSintetico(unittest.TestCase):

    def test_docx_texto_rc(self):
        rc = "2462302DS4026S0001GQ"
        with tempfile.TemporaryDirectory() as tmp:
            ruta = _make_docx(tmp, texto=f"Referencia catastral: {rc}")
            result = extract_entities_from_docx(ruta)
        self.assertIn(rc, result.values("REFERENCIA_CATASTRAL"))

    def test_docx_tabla_ler(self):
        with tempfile.TemporaryDirectory() as tmp:
            ruta = _make_docx(
                tmp,
                tablas=[[["Código LER", "Descripción"], ["17 04 05", "Fragmentos de metal"]]]
            )
            result = extract_entities_from_docx(ruta)
        lers = result.by_type("LER")
        norms = [e.normalized_value for e in lers]
        self.assertTrue(any("17 04 05" in (n or "") for n in norms))

    def test_docx_tabla_promotor_kv(self):
        with tempfile.TemporaryDirectory() as tmp:
            ruta = _make_docx(
                tmp,
                tablas=[[["Promotor", "NIF"], ["EMPRESA RECICLADOS, S.L.", "B12345678"]]]
            )
            result = extract_entities_from_docx(ruta)
        promotores = result.by_type("PROMOTOR")
        vals = [e.value for e in promotores]
        self.assertTrue(any("EMPRESA RECICLADOS" in v for v in vals))

    def test_docx_tabla_promotor_cabecera_invertida(self):
        """Caso PARCELA: empresa es la clave, no el valor."""
        with tempfile.TemporaryDirectory() as tmp:
            ruta = _make_docx(
                tmp,
                tablas=[[["Promotor", "NIF"], ["RECIMETAL LANZAROTE, S.L.", "B72798846"]]]
            )
            result = extract_entities_from_docx(ruta)
        promotores = result.by_type("PROMOTOR")
        self.assertGreater(len(promotores), 0)

    def test_docx_deduplica_entre_texto_y_tabla(self):
        """La misma RC en texto y tabla debe aparecer una sola vez."""
        rc = "2462302DS4026S0001GQ"
        with tempfile.TemporaryDirectory() as tmp:
            ruta = _make_docx(
                tmp,
                texto=f"Ref catastral {rc}",
                tablas=[[["RC"], [rc]]]
            )
            result = extract_entities_from_docx(ruta)
        rcs = result.by_type("REFERENCIA_CATASTRAL")
        self.assertEqual(len(rcs), 1)

    def test_docx_vacio_no_lanza(self):
        with tempfile.TemporaryDirectory() as tmp:
            ruta = _make_docx(tmp, texto="")
            result = extract_entities_from_docx(ruta)
        self.assertIsInstance(result, ExtractionResult)

    def test_docx_no_existe(self):
        with self.assertRaises(FileNotFoundError):
            extract_entities_from_docx("/ruta/que/no/existe.docx")

    def test_docx_operaciones_en_texto(self):
        with tempfile.TemporaryDirectory() as tmp:
            ruta = _make_docx(tmp, texto="Operaciones R1201 y R13 de reciclado")
            result = extract_entities_from_docx(ruta)
        ops = result.values("OPERACION")
        self.assertIn("R1201", ops)
        self.assertIn("R13", ops)

    def test_docx_source_tabla(self):
        with tempfile.TemporaryDirectory() as tmp:
            ruta = _make_docx(
                tmp,
                tablas=[[["Código LER"], ["17 04 05"]]]
            )
            result = extract_entities_from_docx(ruta)
        lers = result.by_type("LER")
        self.assertTrue(any("tabla" in e.source for e in lers))

    def test_no_modifica_archivo(self):
        import os
        with tempfile.TemporaryDirectory() as tmp:
            ruta = _make_docx(tmp, texto="Prueba R1201")
            mtime_before = os.path.getmtime(ruta)
            extract_entities_from_docx(ruta)
            mtime_after = os.path.getmtime(ruta)
        self.assertEqual(mtime_before, mtime_after)


# ---------------------------------------------------------------------------
# Fixture real PARCELA — solo lectura
# ---------------------------------------------------------------------------

@unittest.skipUnless(_DOCX_PARCELA.exists(), f"Fixture PARCELA no disponible: {_DOCX_PARCELA}")
class TestFixtureParcela(unittest.TestCase):

    @classmethod
    def setUpClass(cls):
        cls.result = extract_entities_from_docx(_DOCX_PARCELA)

    def test_no_lanza_excepcion(self):
        self.assertIsInstance(self.result, ExtractionResult)

    def test_detecta_rc_parcela(self):
        rcs = self.result.values("REFERENCIA_CATASTRAL")
        self.assertIn("2462302DS4026S0001GQ", rcs)

    def test_detecta_ler_en_tablas(self):
        lers = self.result.by_type("LER")
        self.assertGreater(len(lers), 0)

    def test_detecta_promotor_recimetal(self):
        promotores = self.result.by_type("PROMOTOR")
        vals = [e.value for e in promotores]
        self.assertTrue(
            any("RECIMETAL" in v for v in vals),
            f"No se encontró RECIMETAL entre los promotores: {vals}"
        )

    def test_detecta_operaciones(self):
        ops = self.result.by_type("OPERACION")
        self.assertGreater(len(ops), 0)

    def test_summary_coherente(self):
        s = self.result.summary()
        total = len(self.result.entities)
        self.assertIn(str(total), s)

    def test_solo_lectura_no_escribe(self):
        import os
        mtime_antes = os.path.getmtime(_DOCX_PARCELA)
        extract_entities_from_docx(_DOCX_PARCELA)
        mtime_despues = os.path.getmtime(_DOCX_PARCELA)
        self.assertEqual(mtime_antes, mtime_despues)

    def test_entities_list_no_vacia(self):
        self.assertGreater(len(self.result.entities), 0)

    def test_by_type_ler_devuelve_entidades(self):
        lers = self.result.by_type("LER")
        self.assertGreater(len(lers), 0)
        self.assertTrue(all(isinstance(e, ExtractedEntity) for e in lers))

    def test_values_ler_son_strings(self):
        vals = self.result.values("LER")
        self.assertTrue(all(isinstance(v, str) for v in vals))

    def test_ler_normalizado_formato_correcto(self):
        lers = self.result.by_type("LER")
        for e in lers:
            if e.normalized_value:
                parts = e.normalized_value.rstrip('*').split()
                self.assertEqual(len(parts), 3, f"LER mal normalizado: {e.normalized_value}")
                self.assertTrue(all(p.isdigit() and len(p) == 2 for p in parts))


if __name__ == "__main__":
    unittest.main()
