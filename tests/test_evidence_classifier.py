"""tests/test_evidence_classifier.py — Suite de tests para EvidenceClassifier (IN-03).

Cubre los criterios de cierre del ítem:
- clasifica RC como emplazamiento/referencia_catastral
- clasifica LER como residuos/codigo_ler
- clasifica operacion R1203 como operaciones/operacion_residuos
- clasifica coordenadas WGS84 y UTM
- clasifica superficies según contexto (parcela, construida, etc.)
- clasifica capacidad
- clasifica potencia
- clasifica promotor/titular
- clasifica equipo
- entity_type desconocido → categoria otros
- default_state DECLARADO
- default_state ASUNCION_TEST con nota
- LOW confidence añade nota pero no cambia estado automáticamente
- no convierte nada en CONFIRMADO
- detecta conflicto simple de valores
- to_hecho_confirmado() genera dict compatible con schema
- to_hechos_confirmados() asigna IDs HC-001, HC-002...
- ClassificationResult.by_category() / by_field() / summary()
- classify_entities_from_docx con DOCX sintético
- solo lectura contra fixture real PARCELA
"""
import sys
import tempfile
import unittest
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent / "src"))

import docx as python_docx

from eia_agent.core.evidence_state import EvidenceState
from eia_agent.core.entity_extractor import ExtractedEntity, ExtractionResult
from eia_agent.core.evidence_classifier import (
    CandidateFact,
    ClassificationResult,
    classify_entities,
    classify_entities_from_docx,
    detect_simple_conflicts,
)

# ---------------------------------------------------------------------------
# Rutas
# ---------------------------------------------------------------------------

_ROOT = Path(__file__).parent.parent
_DOCX_PARCELA = (
    _ROOT
    / "expediente-EIA-2026-RECIMETAL-PARCELA"
    / "inputs"
    / "memorias"
    / "Documento_Ambiental_RECIMETAL_Parcela_v6.docx"
)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _make_entity(
    entity_type: str,
    value: str = "valor_test",
    confidence: str = "HIGH",
    context: str | None = None,
    normalized_value: str | None = None,
    source: str = "texto",
) -> ExtractedEntity:
    return ExtractedEntity(
        entity_type=entity_type,
        value=value,
        source=source,
        confidence=confidence,
        context=context,
        normalized_value=normalized_value,
    )


def _make_result(*entities: ExtractedEntity) -> ExtractionResult:
    return ExtractionResult(entities=list(entities))


def _classify(*entities: ExtractedEntity, state=EvidenceState.DECLARADO) -> ClassificationResult:
    return classify_entities(
        _make_result(*entities),
        source_doc_id="DOC-001",
        source_doc_name="test.docx",
        default_state=state,
    )


def _make_docx(tmp: str, *, texto: str = "", tablas=None, nombre="test.docx") -> Path:
    doc = python_docx.Document()
    if texto:
        doc.add_paragraph(texto)
    for tabla in (tablas or []):
        if not tabla:
            continue
        t = doc.add_table(rows=len(tabla), cols=len(tabla[0]))
        for r_idx, fila in enumerate(tabla):
            for c_idx, valor in enumerate(fila):
                t.rows[r_idx].cells[c_idx].text = valor
    ruta = Path(tmp) / nombre
    doc.save(str(ruta))
    return ruta


# ---------------------------------------------------------------------------
# Mapeo de tipos de entidad → categoria/campo
# ---------------------------------------------------------------------------

class TestMapeoReferenciaCatastral(unittest.TestCase):

    def test_categoria_emplazamiento(self):
        result = _classify(_make_entity("REFERENCIA_CATASTRAL", "2462302DS4026S0001GQ"))
        self.assertEqual(result.facts[0].categoria, "emplazamiento")

    def test_campo_referencia_catastral(self):
        result = _classify(_make_entity("REFERENCIA_CATASTRAL", "2462302DS4026S0001GQ"))
        self.assertEqual(result.facts[0].campo, "referencia_catastral")

    def test_valor_preservado(self):
        rc = "2462302DS4026S0001GQ"
        result = _classify(_make_entity("REFERENCIA_CATASTRAL", rc, normalized_value=rc))
        self.assertEqual(result.facts[0].valor, rc)


class TestMapeoLER(unittest.TestCase):

    def test_categoria_residuos(self):
        result = _classify(_make_entity("LER", "17 04 05", normalized_value="17 04 05"))
        self.assertEqual(result.facts[0].categoria, "residuos")

    def test_campo_codigo_ler(self):
        result = _classify(_make_entity("LER", "17 04 05", normalized_value="17 04 05"))
        self.assertEqual(result.facts[0].campo, "codigo_ler")

    def test_ler_peligroso_normalized(self):
        result = _classify(
            _make_entity("LER", "16 06 01*", normalized_value="16 06 01*")
        )
        self.assertIn("*", str(result.facts[0].valor))


class TestMapeoOperacion(unittest.TestCase):

    def test_categoria_operaciones(self):
        result = _classify(_make_entity("OPERACION", "R1203", normalized_value="R1203"))
        self.assertEqual(result.facts[0].categoria, "operaciones")

    def test_campo_operacion_residuos(self):
        result = _classify(_make_entity("OPERACION", "R1203", normalized_value="R1203"))
        self.assertEqual(result.facts[0].campo, "operacion_residuos")

    def test_r13(self):
        result = _classify(_make_entity("OPERACION", "R13", normalized_value="R13"))
        self.assertEqual(result.facts[0].campo, "operacion_residuos")

    def test_d15(self):
        result = _classify(_make_entity("OPERACION", "D15", normalized_value="D15"))
        self.assertEqual(result.facts[0].categoria, "operaciones")


class TestMapeoCoordenadasWGS84(unittest.TestCase):

    def test_campo_wgs84_dec(self):
        result = _classify(
            _make_entity("COORDENADA", "28.9234", normalized_value="DEC 28.9234")
        )
        f = result.facts[0]
        self.assertEqual(f.categoria, "emplazamiento")
        self.assertEqual(f.campo, "coordenadas_wgs84")

    def test_campo_wgs84_dec_lowercase(self):
        result = _classify(
            _make_entity("COORDENADA", "28.9234", normalized_value="dec 28.9234")
        )
        self.assertEqual(result.facts[0].campo, "coordenadas_wgs84")


class TestMapeoCoordenadasUTM(unittest.TestCase):

    def test_campo_utm(self):
        result = _classify(
            _make_entity("COORDENADA", "E: 642000", normalized_value="UTM E=642000")
        )
        f = result.facts[0]
        self.assertEqual(f.categoria, "emplazamiento")
        self.assertEqual(f.campo, "coordenadas_utm")

    def test_campo_utm_y(self):
        result = _classify(
            _make_entity("COORDENADA", "Y: 3207000", normalized_value="UTM Y=3207000")
        )
        self.assertEqual(result.facts[0].campo, "coordenadas_utm")

    def test_coordenada_sin_prefijo(self):
        result = _classify(
            _make_entity("COORDENADA", "12345", normalized_value=None)
        )
        # Sin prefijo DEC/UTM → coordenadas genéricas
        self.assertEqual(result.facts[0].campo, "coordenadas")


class TestMapeoSuperficies(unittest.TestCase):

    def test_superficie_parcela(self):
        result = _classify(
            _make_entity("SUPERFICIE_PARCELA", "2000 m²", normalized_value="2000 m²")
        )
        self.assertEqual(result.facts[0].campo, "superficie_parcela")
        self.assertEqual(result.facts[0].categoria, "superficies")

    def test_superficie_construida(self):
        result = _classify(
            _make_entity("SUPERFICIE_CONSTRUIDA", "500 m²", normalized_value="500 m²")
        )
        self.assertEqual(result.facts[0].campo, "superficie_construida")

    def test_superficie_util(self):
        result = _classify(
            _make_entity("SUPERFICIE_UTIL", "450 m²", normalized_value="450 m²")
        )
        self.assertEqual(result.facts[0].campo, "superficie_util")

    def test_superficie_catastral(self):
        result = _classify(
            _make_entity("SUPERFICIE_CATASTRAL", "2100 m²", normalized_value="2100 m²")
        )
        self.assertEqual(result.facts[0].campo, "superficie_catastral")

    def test_superficie_nave(self):
        result = _classify(
            _make_entity("SUPERFICIE_NAVE", "300 m²", normalized_value="300 m²")
        )
        self.assertEqual(result.facts[0].campo, "superficie_nave")

    def test_superficie_sin_subtipo(self):
        result = _classify(
            _make_entity("SUPERFICIE", "100 m²", normalized_value="100 m²")
        )
        self.assertEqual(result.facts[0].campo, "superficie_no_clasificada")


class TestMapeoCapacidad(unittest.TestCase):

    def test_categoria_capacidades(self):
        result = _classify(_make_entity("CAPACIDAD", "500 tm/día"))
        self.assertEqual(result.facts[0].categoria, "capacidades")

    def test_campo_capacidad(self):
        result = _classify(_make_entity("CAPACIDAD", "3.000 t/año"))
        self.assertEqual(result.facts[0].campo, "capacidad")


class TestMapeoPotencia(unittest.TestCase):

    def test_categoria_equipos(self):
        result = _classify(
            _make_entity("POTENCIA", "75 kW", normalized_value="75 KW")
        )
        self.assertEqual(result.facts[0].categoria, "equipos")

    def test_campo_potencia(self):
        result = _classify(
            _make_entity("POTENCIA", "100 CV", normalized_value="100 CV")
        )
        self.assertEqual(result.facts[0].campo, "potencia")


class TestMapeoFecha(unittest.TestCase):

    def test_categoria_fechas(self):
        result = _classify(_make_entity("FECHA", "15/03/2024"))
        self.assertEqual(result.facts[0].categoria, "fechas")

    def test_campo_fecha_documental(self):
        result = _classify(_make_entity("FECHA", "2024-03-15"))
        self.assertEqual(result.facts[0].campo, "fecha_documental")


class TestMapeoPromotor(unittest.TestCase):

    def test_categoria_promotor(self):
        result = _classify(
            _make_entity("PROMOTOR", "RECIMETAL LANZAROTE, S.L.")
        )
        self.assertEqual(result.facts[0].categoria, "promotor")

    def test_campo_nombre_promotor(self):
        result = _classify(
            _make_entity("PROMOTOR", "EMPRESA TEST, S.A.")
        )
        self.assertEqual(result.facts[0].campo, "nombre_promotor")

    def test_titular(self):
        result = _classify(_make_entity("TITULAR", "EMPRESA TEST, S.A."))
        self.assertEqual(result.facts[0].categoria, "titularidad")
        self.assertEqual(result.facts[0].campo, "titular")


class TestMapeoEquipo(unittest.TestCase):

    def test_categoria_equipos(self):
        result = _classify(_make_entity("EQUIPO", "molino"))
        self.assertEqual(result.facts[0].categoria, "equipos")

    def test_campo_equipo(self):
        result = _classify(_make_entity("EQUIPO", "criba"))
        self.assertEqual(result.facts[0].campo, "equipo")


class TestMapeoDesconocido(unittest.TestCase):

    def test_categoria_otros(self):
        result = _classify(_make_entity("TIPO_INVENTADO", "valor"))
        self.assertEqual(result.facts[0].categoria, "otros")

    def test_campo_entity_type_lower(self):
        result = _classify(_make_entity("TIPO_INVENTADO", "valor"))
        self.assertEqual(result.facts[0].campo, "tipo_inventado")


# ---------------------------------------------------------------------------
# Reglas de estado
# ---------------------------------------------------------------------------

class TestEstadoDeclarado(unittest.TestCase):

    def test_default_estado_declarado(self):
        result = _classify(_make_entity("LER", "17 04 05"))
        self.assertEqual(result.facts[0].estado, EvidenceState.DECLARADO.value)

    def test_estado_string_declarado(self):
        result = classify_entities(
            _make_result(_make_entity("LER", "17 04 05")),
            "DOC-001", "test.docx",
            default_state="DECLARADO",
        )
        self.assertEqual(result.facts[0].estado, EvidenceState.DECLARADO.value)

    def test_nunca_confirmado_desde_declarado(self):
        result = _classify(_make_entity("REFERENCIA_CATASTRAL", "2462302DS4026S0001GQ"))
        self.assertNotEqual(result.facts[0].estado, EvidenceState.CONFIRMADO.value)
        self.assertNotIn("CONFIRMADO", result.facts[0].estado)


class TestEstadoAsuncionTest(unittest.TestCase):

    def test_estado_asuncion_test(self):
        result = _classify(
            _make_entity("LER", "17 04 05"),
            state=EvidenceState.ASUNCION_TEST,
        )
        self.assertEqual(result.facts[0].estado, EvidenceState.ASUNCION_TEST.value)

    def test_nota_asuncion_test_añadida(self):
        result = _classify(
            _make_entity("LER", "17 04 05"),
            state=EvidenceState.ASUNCION_TEST,
        )
        notas = result.facts[0].notes
        self.assertTrue(any("asunción test" in n.lower() for n in notas))

    def test_nunca_confirmado_desde_asuncion(self):
        result = _classify(
            _make_entity("REFERENCIA_CATASTRAL", "2462302DS4026S0001GQ"),
            state=EvidenceState.ASUNCION_TEST,
        )
        self.assertNotIn("CONFIRMADO", result.facts[0].estado)

    def test_asuncion_por_string(self):
        result = classify_entities(
            _make_result(_make_entity("LER", "17 04 05")),
            "DOC-001", "test.docx",
            default_state="ASUNCION_TEST",
        )
        self.assertEqual(result.facts[0].estado, "ASUNCION_TEST")


class TestEstadoLowConfidence(unittest.TestCase):

    def test_low_confidence_no_cambia_estado(self):
        result = _classify(
            _make_entity("PROMOTOR", "EMPRESA, S.L.", confidence="LOW")
        )
        self.assertEqual(result.facts[0].estado, EvidenceState.DECLARADO.value)

    def test_low_confidence_añade_nota(self):
        result = _classify(
            _make_entity("PROMOTOR", "EMPRESA, S.L.", confidence="LOW")
        )
        notas = result.facts[0].notes
        self.assertTrue(any("confianza baja" in n.lower() for n in notas))

    def test_high_confidence_sin_nota_automatica(self):
        result = _classify(
            _make_entity("REFERENCIA_CATASTRAL", "2462302DS4026S0001GQ", confidence="HIGH")
        )
        # No debe haber nota de confianza baja
        notas = result.facts[0].notes
        self.assertFalse(any("confianza baja" in n.lower() for n in notas))


# ---------------------------------------------------------------------------
# Fuentes documentales
# ---------------------------------------------------------------------------

class TestFuentes(unittest.TestCase):

    def test_fuente_preservada(self):
        result = classify_entities(
            _make_result(_make_entity("LER", "17 04 05")),
            "DOC-042", "memoria.docx",
        )
        self.assertIn("DOC-042", result.facts[0].fuentes)

    def test_fuente_lista(self):
        result = classify_entities(
            _make_result(_make_entity("LER", "17 04 05")),
            "DOC-001", "test.docx",
        )
        self.assertIsInstance(result.facts[0].fuentes, list)
        self.assertGreater(len(result.facts[0].fuentes), 0)


# ---------------------------------------------------------------------------
# Conflictos
# ---------------------------------------------------------------------------

class TestConflictos(unittest.TestCase):

    def test_detecta_conflicto_mismo_campo(self):
        facts = [
            CandidateFact(
                id=None, categoria="emplazamiento", campo="referencia_catastral",
                valor="RC-001", estado="DECLARADO", fuentes=["DOC-001"],
                entity_type="REFERENCIA_CATASTRAL", confidence="HIGH"
            ),
            CandidateFact(
                id=None, categoria="emplazamiento", campo="referencia_catastral",
                valor="RC-002", estado="DECLARADO", fuentes=["DOC-001"],
                entity_type="REFERENCIA_CATASTRAL", confidence="HIGH"
            ),
        ]
        conflicts = detect_simple_conflicts(facts)
        self.assertEqual(len(conflicts), 1)
        self.assertEqual(conflicts[0]["campo"], "referencia_catastral")
        self.assertEqual(len(conflicts[0]["valores"]), 2)

    def test_no_conflicto_mismo_valor(self):
        facts = [
            CandidateFact(
                id=None, categoria="residuos", campo="codigo_ler",
                valor="17 04 05", estado="DECLARADO", fuentes=["DOC-001"],
                entity_type="LER", confidence="HIGH"
            ),
            CandidateFact(
                id=None, categoria="residuos", campo="codigo_ler",
                valor="17 04 05", estado="DECLARADO", fuentes=["DOC-001"],
                entity_type="LER", confidence="HIGH"
            ),
        ]
        conflicts = detect_simple_conflicts(facts)
        self.assertEqual(len(conflicts), 0)

    def test_conflicto_genera_warning(self):
        e1 = _make_entity("REFERENCIA_CATASTRAL", "2462302DS4026S0001GQ",
                           normalized_value="2462302DS4026S0001GQ")
        e2 = _make_entity("REFERENCIA_CATASTRAL", "9999999XX9999X9999XX",
                           normalized_value="9999999XX9999X9999XX")
        result = _classify(e1, e2)
        self.assertGreater(len(result.conflicts), 0)
        self.assertGreater(len(result.warnings), 0)

    def test_conflicto_no_resuelto(self):
        e1 = _make_entity("REFERENCIA_CATASTRAL", "2462302DS4026S0001GQ",
                           normalized_value="2462302DS4026S0001GQ")
        e2 = _make_entity("REFERENCIA_CATASTRAL", "9999999XX9999X9999XX",
                           normalized_value="9999999XX9999X9999XX")
        result = _classify(e1, e2)
        # Ambos hechos siguen DECLARADO; el conflicto no cambia su estado
        estados = [f.estado for f in result.by_field("referencia_catastral")]
        self.assertTrue(all(e == "DECLARADO" for e in estados))


# ---------------------------------------------------------------------------
# to_hecho_confirmado y to_hechos_confirmados
# ---------------------------------------------------------------------------

class TestToHechoConfirmado(unittest.TestCase):

    def _fact(self, **kwargs) -> CandidateFact:
        defaults = dict(
            id="HC-001",
            categoria="residuos",
            campo="codigo_ler",
            valor="17 04 05",
            estado="DECLARADO",
            fuentes=["DOC-001"],
            entity_type="LER",
            confidence="HIGH",
        )
        defaults.update(kwargs)
        return CandidateFact(**defaults)

    def test_campos_requeridos_presentes(self):
        d = self._fact().to_hecho_confirmado()
        for campo in ("id", "categoria", "campo", "valor", "estado", "fuentes"):
            self.assertIn(campo, d)

    def test_id_asignado(self):
        d = self._fact(id="HC-042").to_hecho_confirmado()
        self.assertEqual(d["id"], "HC-042")

    def test_nota_none_sin_notas(self):
        d = self._fact(notes=[]).to_hecho_confirmado()
        self.assertIsNone(d["nota"])

    def test_nota_unida_con_pipe(self):
        d = self._fact(notes=["nota A", "nota B"]).to_hecho_confirmado()
        self.assertIn("nota A", d["nota"])
        self.assertIn("nota B", d["nota"])
        self.assertIn("|", d["nota"])

    def test_fuentes_es_lista(self):
        d = self._fact().to_hecho_confirmado()
        self.assertIsInstance(d["fuentes"], list)

    def test_estado_es_string(self):
        d = self._fact().to_hecho_confirmado()
        self.assertIsInstance(d["estado"], str)


class TestToHechosConfirmados(unittest.TestCase):

    def test_asigna_ids_secuenciales(self):
        result = _classify(
            _make_entity("LER", "17 04 05", normalized_value="17 04 05"),
            _make_entity("OPERACION", "R1201", normalized_value="R1201"),
        )
        hechos = result.to_hechos_confirmados()
        ids = [h["id"] for h in hechos]
        self.assertEqual(ids, ["HC-001", "HC-002"])

    def test_start_index_configurable(self):
        result = _classify(_make_entity("LER", "17 04 05", normalized_value="17 04 05"))
        hechos = result.to_hechos_confirmados(start_index=5)
        self.assertEqual(hechos[0]["id"], "HC-005")

    def test_prefix_configurable(self):
        result = _classify(_make_entity("LER", "17 04 05", normalized_value="17 04 05"))
        hechos = result.to_hechos_confirmados(prefix="HC-T")
        self.assertTrue(hechos[0]["id"].startswith("HC-T"))

    def test_lista_de_dicts(self):
        result = _classify(_make_entity("LER", "17 04 05", normalized_value="17 04 05"))
        hechos = result.to_hechos_confirmados()
        self.assertIsInstance(hechos, list)
        self.assertIsInstance(hechos[0], dict)

    def test_sin_hechos_lista_vacia(self):
        result = ClassificationResult()
        hechos = result.to_hechos_confirmados()
        self.assertEqual(hechos, [])


# ---------------------------------------------------------------------------
# ClassificationResult API
# ---------------------------------------------------------------------------

class TestClassificationResultAPI(unittest.TestCase):

    def _build_result(self) -> ClassificationResult:
        return _classify(
            _make_entity("LER", "17 04 05", normalized_value="17 04 05"),
            _make_entity("LER", "19 12 10", normalized_value="19 12 10"),
            _make_entity("OPERACION", "R1201", normalized_value="R1201"),
            _make_entity("REFERENCIA_CATASTRAL", "2462302DS4026S0001GQ",
                         normalized_value="2462302DS4026S0001GQ"),
        )

    def test_by_category_residuos(self):
        result = self._build_result()
        residuos = result.by_category("residuos")
        self.assertEqual(len(residuos), 2)
        self.assertTrue(all(f.categoria == "residuos" for f in residuos))

    def test_by_category_vacio(self):
        result = self._build_result()
        self.assertEqual(result.by_category("climatic"), [])

    def test_by_field_codigo_ler(self):
        result = self._build_result()
        lers = result.by_field("codigo_ler")
        self.assertEqual(len(lers), 2)

    def test_by_field_vacio(self):
        result = self._build_result()
        self.assertEqual(result.by_field("campo_inexistente"), [])

    def test_values_retorna_lista(self):
        result = self._build_result()
        vals = result.values("codigo_ler")
        self.assertIsInstance(vals, list)
        self.assertEqual(len(vals), 2)

    def test_summary_contiene_total(self):
        result = self._build_result()
        s = result.summary()
        self.assertIn("4", s)

    def test_summary_contiene_categorias(self):
        result = self._build_result()
        s = result.summary()
        self.assertIn("residuos", s)
        self.assertIn("operaciones", s)

    def test_summary_vacio(self):
        result = ClassificationResult()
        self.assertIn("Sin hechos", result.summary())

    def test_summary_incluye_conflictos(self):
        e1 = _make_entity("REFERENCIA_CATASTRAL", "2462302DS4026S0001GQ",
                           normalized_value="2462302DS4026S0001GQ")
        e2 = _make_entity("REFERENCIA_CATASTRAL", "9999999XX9999X9999XX",
                           normalized_value="9999999XX9999X9999XX")
        result = _classify(e1, e2)
        s = result.summary()
        self.assertIn("conflicto", s.lower())


# ---------------------------------------------------------------------------
# classify_entities_from_docx — DOCX sintético
# ---------------------------------------------------------------------------

class TestClassifyFromDocxSintetico(unittest.TestCase):

    def test_docx_con_rc(self):
        rc = "2462302DS4026S0001GQ"
        with tempfile.TemporaryDirectory() as tmp:
            ruta = _make_docx(tmp, texto=f"Referencia catastral: {rc}")
            result = classify_entities_from_docx(ruta, "DOC-001")
        vals = result.values("referencia_catastral")
        self.assertTrue(any(rc in str(v) for v in vals))

    def test_docx_con_ler_en_tabla(self):
        with tempfile.TemporaryDirectory() as tmp:
            ruta = _make_docx(
                tmp,
                tablas=[[["Código LER", "Descripción"], ["17 04 05", "Metales"]]]
            )
            result = classify_entities_from_docx(ruta, "DOC-001")
        hechos_ler = result.by_field("codigo_ler")
        self.assertGreater(len(hechos_ler), 0)

    def test_docx_vacío_no_lanza(self):
        with tempfile.TemporaryDirectory() as tmp:
            ruta = _make_docx(tmp, texto="")
            result = classify_entities_from_docx(ruta, "DOC-001")
        self.assertIsInstance(result, ClassificationResult)

    def test_docx_no_existe(self):
        with self.assertRaises(FileNotFoundError):
            classify_entities_from_docx("/ruta/que/no/existe.docx")

    def test_source_doc_id_auto(self):
        with tempfile.TemporaryDirectory() as tmp:
            ruta = _make_docx(tmp, texto="Operación R1201")
            result = classify_entities_from_docx(ruta)  # sin source_doc_id
        for f in result.facts:
            self.assertIn("DOC-001", f.fuentes)

    def test_estado_por_defecto_declarado(self):
        with tempfile.TemporaryDirectory() as tmp:
            ruta = _make_docx(tmp, texto="Operación R1201")
            result = classify_entities_from_docx(ruta, "DOC-001")
        for f in result.facts:
            self.assertNotIn("CONFIRMADO", f.estado)

    def test_devuelve_classification_result(self):
        with tempfile.TemporaryDirectory() as tmp:
            ruta = _make_docx(tmp, texto="Test RC")
            result = classify_entities_from_docx(ruta)
        self.assertIsInstance(result, ClassificationResult)

    def test_no_modifica_docx(self):
        import os
        with tempfile.TemporaryDirectory() as tmp:
            ruta = _make_docx(tmp, texto="R1201 operación")
            mtime_antes = os.path.getmtime(ruta)
            classify_entities_from_docx(ruta, "DOC-001")
            mtime_despues = os.path.getmtime(ruta)
        self.assertEqual(mtime_antes, mtime_despues)

    def test_asuncion_test_via_classify(self):
        with tempfile.TemporaryDirectory() as tmp:
            ruta = _make_docx(tmp, texto="Operación R1201")
            result = classify_entities_from_docx(
                ruta, "DOC-001",
                default_state=EvidenceState.ASUNCION_TEST,
            )
        for f in result.facts:
            self.assertEqual(f.estado, "ASUNCION_TEST")


# ---------------------------------------------------------------------------
# Fixture real PARCELA — solo lectura
# ---------------------------------------------------------------------------

@unittest.skipUnless(_DOCX_PARCELA.exists(), f"Fixture PARCELA no disponible: {_DOCX_PARCELA}")
class TestFixtureParcela(unittest.TestCase):

    @classmethod
    def setUpClass(cls):
        cls.result = classify_entities_from_docx(_DOCX_PARCELA, "DOC-001")

    def test_no_lanza_excepcion(self):
        self.assertIsInstance(self.result, ClassificationResult)

    def test_genera_hechos_candidatos(self):
        self.assertGreater(len(self.result.facts), 0)

    def test_no_modifica_mtime(self):
        import os
        mtime_antes = os.path.getmtime(_DOCX_PARCELA)
        classify_entities_from_docx(_DOCX_PARCELA, "DOC-001")
        mtime_despues = os.path.getmtime(_DOCX_PARCELA)
        self.assertEqual(mtime_antes, mtime_despues)

    def test_rc_en_emplazamiento(self):
        emplazamiento = self.result.by_category("emplazamiento")
        self.assertTrue(len(emplazamiento) > 0)
        rc_vals = self.result.values("referencia_catastral")
        self.assertTrue(any("2462302DS4026S0001GQ" in str(v) for v in rc_vals))

    def test_ler_en_residuos(self):
        residuos = self.result.by_category("residuos")
        self.assertGreater(len(residuos), 0)

    def test_promotor_detectado(self):
        promotores = self.result.by_category("promotor")
        self.assertGreater(len(promotores), 0)
        vals = self.result.values("nombre_promotor")
        self.assertTrue(any("RECIMETAL" in str(v) for v in vals))

    def test_estado_no_confirmado(self):
        for f in self.result.facts:
            self.assertNotIn("CONFIRMADO", f.estado)

    def test_todos_tienen_fuente(self):
        for f in self.result.facts:
            self.assertIsInstance(f.fuentes, list)
            self.assertGreater(len(f.fuentes), 0)

    def test_to_hechos_confirmados_ids(self):
        hechos = self.result.to_hechos_confirmados()
        self.assertGreater(len(hechos), 0)
        for h in hechos:
            self.assertIsNotNone(h["id"])
            self.assertRegex(h["id"], r"^HC-\d{3,4}$")

    def test_summary_coherente(self):
        s = self.result.summary()
        total = len(self.result.facts)
        self.assertIn(str(total), s)


if __name__ == "__main__":
    unittest.main()
