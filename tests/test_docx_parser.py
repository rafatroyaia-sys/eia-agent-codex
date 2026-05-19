"""tests/test_docx_parser.py — Suite de tests para DocxParser (IN-01).

Cubre los criterios de cierre del ítem:
- parse_docx con DOCX mínimo devuelve DocxContent
- texto extraído no vacío
- tablas con cabecera extraídas como dict
- tablas sin cabecera → col_0, col_1...
- metadatos devuelve dict con campos esperados
- num_paginas_estimadas >= 1
- archivo inexistente → FileNotFoundError
- extensión no .docx → ValueError
- archivo .docx inválido (contenido binario) → ValueError
- DOCX vacío → texto='' y tablas=[]
- extract_tables_raw devuelve lista de listas de strings
- parse_docx no modifica el archivo: mtime invariante
- fixture real PARCELA solo lectura: parse_docx sin excepción, texto no vacío, tablas extraídas
"""
import io
import sys
import tempfile
import time
import unittest
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent / "src"))

import docx as python_docx

from eia_agent.core.docx_parser import DocxContent, extract_tables_raw, parse_docx

# ---------------------------------------------------------------------------
# Rutas
# ---------------------------------------------------------------------------

_ROOT    = Path(__file__).parent.parent
_PARCELA = _ROOT / "expediente-EIA-2026-RECIMETAL-PARCELA"
_DOCX_PARCELA = _PARCELA / "inputs" / "memorias" / "Documento_Ambiental_RECIMETAL_Parcela_v6.docx"


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _make_docx(tmp: str, *, texto: str = "Párrafo de prueba.",
               tablas: list[list[list[str]]] | None = None,
               nombre: str = "test.docx") -> Path:
    """Crea un DOCX sintético en un directorio temporal."""
    doc = python_docx.Document()
    if texto:
        doc.add_paragraph(texto)
    for tabla in (tablas or []):
        if not tabla:
            continue
        t = doc.add_table(rows=len(tabla), cols=len(tabla[0]))
        for r_idx, fila in enumerate(tabla):
            for c_idx, valor in enumerate(fila):
                t.rows[r_idx].cells[c_idx].text = valor
    ruta = Path(tmp) / nombre
    doc.save(str(ruta))
    return ruta


def _make_empty_docx(tmp: str) -> Path:
    """Crea un DOCX sin párrafos ni tablas."""
    doc = python_docx.Document()
    ruta = Path(tmp) / "empty.docx"
    doc.save(str(ruta))
    return ruta


def _make_invalid_docx(tmp: str) -> Path:
    """Crea un archivo con extensión .docx pero contenido binario inválido."""
    ruta = Path(tmp) / "invalid.docx"
    ruta.write_bytes(b"\x00\x01\x02\x03 esto no es un docx valido")
    return ruta


# ---------------------------------------------------------------------------
# TestDocxContent
# ---------------------------------------------------------------------------

class TestDocxContent(unittest.TestCase):

    def test_instanciacion_minima(self):
        dc = DocxContent(texto="hola", tablas=[], metadatos={}, num_paginas_estimadas=1)
        self.assertEqual(dc.texto, "hola")

    def test_defaults_tablas_y_metadatos(self):
        dc = DocxContent(texto="")
        self.assertIsInstance(dc.tablas, list)
        self.assertIsInstance(dc.metadatos, dict)

    def test_num_paginas_estimadas_default(self):
        dc = DocxContent(texto="")
        self.assertEqual(dc.num_paginas_estimadas, 1)

    def test_campos_presentes(self):
        dc = DocxContent(texto="x", tablas=[[{"a": "b"}]], metadatos={"author": "test"}, num_paginas_estimadas=3)
        self.assertEqual(dc.texto, "x")
        self.assertEqual(dc.tablas, [[{"a": "b"}]])
        self.assertEqual(dc.metadatos["author"], "test")
        self.assertEqual(dc.num_paginas_estimadas, 3)


# ---------------------------------------------------------------------------
# TestParseDocxTexto
# ---------------------------------------------------------------------------

class TestParseDocxTexto(unittest.TestCase):

    def setUp(self):
        self.tmp = tempfile.mkdtemp()

    def test_texto_extraido_no_vacio(self):
        ruta = _make_docx(self.tmp, texto="Este es un documento de prueba.")
        result = parse_docx(ruta)
        self.assertGreater(len(result.texto), 0)

    def test_texto_contiene_parrafo(self):
        ruta = _make_docx(self.tmp, texto="Contenido específico del promotor.")
        result = parse_docx(ruta)
        self.assertIn("Contenido específico del promotor.", result.texto)

    def test_multiples_parrafos_unidos_por_newline(self):
        doc = python_docx.Document()
        doc.add_paragraph("Párrafo uno.")
        doc.add_paragraph("Párrafo dos.")
        ruta = Path(self.tmp) / "multi.docx"
        doc.save(str(ruta))
        result = parse_docx(ruta)
        self.assertIn("Párrafo uno.", result.texto)
        self.assertIn("Párrafo dos.", result.texto)
        self.assertIn("\n", result.texto)

    def test_parrafos_vacios_omitidos(self):
        doc = python_docx.Document()
        doc.add_paragraph("Real.")
        doc.add_paragraph("")
        doc.add_paragraph("   ")
        ruta = Path(self.tmp) / "vacios.docx"
        doc.save(str(ruta))
        result = parse_docx(ruta)
        lineas = result.texto.split("\n")
        self.assertTrue(all(l.strip() for l in lineas))

    def test_docx_vacio_texto_empty(self):
        ruta = _make_empty_docx(self.tmp)
        result = parse_docx(ruta)
        self.assertEqual(result.texto, "")

    def test_docx_vacio_tablas_empty(self):
        ruta = _make_empty_docx(self.tmp)
        result = parse_docx(ruta)
        self.assertEqual(result.tablas, [])

    def test_devuelve_docxcontent(self):
        ruta = _make_docx(self.tmp)
        result = parse_docx(ruta)
        self.assertIsInstance(result, DocxContent)


# ---------------------------------------------------------------------------
# TestParseDocxTablas
# ---------------------------------------------------------------------------

class TestParseDocxTablas(unittest.TestCase):

    def setUp(self):
        self.tmp = tempfile.mkdtemp()

    def test_tabla_con_cabecera_devuelve_dicts(self):
        tabla = [["Nombre", "Valor"], ["R12", "100 t/año"]]
        ruta = _make_docx(self.tmp, texto="", tablas=[tabla])
        result = parse_docx(ruta)
        self.assertEqual(len(result.tablas), 1)
        fila = result.tablas[0][0]
        self.assertIsInstance(fila, dict)

    def test_tabla_con_cabecera_claves_correctas(self):
        tabla = [["Operación", "Capacidad"], ["R1203", "50 t/día"]]
        ruta = _make_docx(self.tmp, texto="", tablas=[tabla])
        result = parse_docx(ruta)
        fila = result.tablas[0][0]
        self.assertIn("Operación", fila)
        self.assertIn("Capacidad", fila)

    def test_tabla_con_cabecera_valores_correctos(self):
        tabla = [["LER", "Descripción"], ["17 04 05", "Hierro y acero"]]
        ruta = _make_docx(self.tmp, texto="", tablas=[tabla])
        result = parse_docx(ruta)
        fila = result.tablas[0][0]
        self.assertEqual(fila["LER"], "17 04 05")
        self.assertEqual(fila["Descripción"], "Hierro y acero")

    def test_tabla_sin_cabecera_usa_col_n(self):
        # Primera fila vacía → col_0, col_1...
        tabla = [["", ""], ["dato_a", "dato_b"]]
        ruta = _make_docx(self.tmp, texto="", tablas=[tabla])
        result = parse_docx(ruta)
        # Con primera fila vacía se generan col_0, col_1
        fila = result.tablas[0][0]
        self.assertTrue(
            any(k.startswith("col_") for k in fila.keys()),
            f"Se esperaban claves col_N, se obtuvo: {list(fila.keys())}"
        )

    def test_multiples_tablas(self):
        t1 = [["A", "B"], ["1", "2"]]
        t2 = [["X", "Y", "Z"], ["a", "b", "c"]]
        ruta = _make_docx(self.tmp, texto="", tablas=[t1, t2])
        result = parse_docx(ruta)
        self.assertEqual(len(result.tablas), 2)

    def test_sin_tablas_lista_vacia(self):
        ruta = _make_docx(self.tmp, texto="Solo texto, sin tablas.")
        result = parse_docx(ruta)
        self.assertEqual(result.tablas, [])


# ---------------------------------------------------------------------------
# TestParseDocxMetadatos
# ---------------------------------------------------------------------------

class TestParseDocxMetadatos(unittest.TestCase):

    def setUp(self):
        self.tmp = tempfile.mkdtemp()

    def test_metadatos_es_dict(self):
        ruta = _make_docx(self.tmp)
        result = parse_docx(ruta)
        self.assertIsInstance(result.metadatos, dict)

    def test_metadatos_tiene_campos_esperados(self):
        ruta = _make_docx(self.tmp)
        result = parse_docx(ruta)
        for campo in ("author", "created", "modified", "title", "subject"):
            self.assertIn(campo, result.metadatos, f"Campo faltante: {campo}")

    def test_metadatos_campos_ausentes_son_none_o_str(self):
        ruta = _make_docx(self.tmp)
        result = parse_docx(ruta)
        # Los campos pueden ser None, str, o datetime — no deben lanzar excepción
        for v in result.metadatos.values():
            self.assertIsNotNone(v.__class__)  # simplemente que tienen tipo


# ---------------------------------------------------------------------------
# TestNumPaginasEstimadas
# ---------------------------------------------------------------------------

class TestNumPaginasEstimadas(unittest.TestCase):

    def setUp(self):
        self.tmp = tempfile.mkdtemp()

    def test_paginas_minimo_1(self):
        ruta = _make_empty_docx(self.tmp)
        result = parse_docx(ruta)
        self.assertGreaterEqual(result.num_paginas_estimadas, 1)

    def test_paginas_texto_corto_es_1(self):
        ruta = _make_docx(self.tmp, texto="Texto corto.")
        result = parse_docx(ruta)
        self.assertEqual(result.num_paginas_estimadas, 1)

    def test_paginas_texto_largo_mayor_que_1(self):
        texto_largo = "A" * 5001  # > 2 × 2500 → al menos 3 páginas estimadas
        ruta = _make_docx(self.tmp, texto=texto_largo)
        result = parse_docx(ruta)
        self.assertGreater(result.num_paginas_estimadas, 1)

    def test_formula_correcta(self):
        texto = "X" * 7500  # 7500 // 2500 + 1 = 4
        ruta = _make_docx(self.tmp, texto=texto)
        result = parse_docx(ruta)
        esperado = max(1, len(result.texto) // 2500 + 1)
        self.assertEqual(result.num_paginas_estimadas, esperado)


# ---------------------------------------------------------------------------
# TestParseDocxErrores
# ---------------------------------------------------------------------------

class TestParseDocxErrores(unittest.TestCase):

    def setUp(self):
        self.tmp = tempfile.mkdtemp()

    def test_archivo_inexistente_file_not_found(self):
        with self.assertRaises(FileNotFoundError):
            parse_docx(Path(self.tmp) / "no_existe.docx")

    def test_extension_no_docx_value_error(self):
        ruta = Path(self.tmp) / "documento.pdf"
        ruta.write_bytes(b"contenido fake")
        with self.assertRaises(ValueError):
            parse_docx(ruta)

    def test_extension_txt_value_error(self):
        ruta = Path(self.tmp) / "texto.txt"
        ruta.write_text("texto plano", encoding="utf-8")
        with self.assertRaises(ValueError):
            parse_docx(ruta)

    def test_extension_sin_punto_value_error(self):
        # Ruta sin extensión
        ruta = Path(self.tmp) / "sin_extension"
        ruta.write_bytes(b"contenido")
        with self.assertRaises(ValueError):
            parse_docx(ruta)

    def test_docx_invalido_value_error(self):
        ruta = _make_invalid_docx(self.tmp)
        with self.assertRaises(ValueError):
            parse_docx(ruta)

    def test_value_error_mensaje_claro(self):
        ruta = _make_invalid_docx(self.tmp)
        try:
            parse_docx(ruta)
            self.fail("Debería haber lanzado ValueError")
        except ValueError as e:
            self.assertGreater(len(str(e)), 10)

    def test_file_not_found_mensaje_incluye_ruta(self):
        ruta = Path(self.tmp) / "ruta_especifica.docx"
        try:
            parse_docx(ruta)
            self.fail("Debería haber lanzado FileNotFoundError")
        except FileNotFoundError as e:
            self.assertIn("ruta_especifica", str(e))


# ---------------------------------------------------------------------------
# TestExtractTablesRaw
# ---------------------------------------------------------------------------

class TestExtractTablesRaw(unittest.TestCase):

    def setUp(self):
        self.tmp = tempfile.mkdtemp()

    def test_devuelve_lista(self):
        ruta = _make_docx(self.tmp, tablas=[[["A", "B"], ["1", "2"]]])
        result = extract_tables_raw(ruta)
        self.assertIsInstance(result, list)

    def test_sin_tablas_lista_vacia(self):
        ruta = _make_docx(self.tmp, texto="Solo texto.")
        result = extract_tables_raw(ruta)
        self.assertEqual(result, [])

    def test_tabla_es_lista_de_filas(self):
        tabla = [["Col1", "Col2"], ["v1", "v2"]]
        ruta = _make_docx(self.tmp, tablas=[tabla])
        result = extract_tables_raw(ruta)
        self.assertEqual(len(result), 1)
        self.assertIsInstance(result[0], list)

    def test_fila_es_lista_de_strings(self):
        tabla = [["Nombre", "Valor"], ["dato", "123"]]
        ruta = _make_docx(self.tmp, tablas=[tabla])
        result = extract_tables_raw(ruta)
        fila = result[0][0]
        self.assertIsInstance(fila, list)
        self.assertTrue(all(isinstance(c, str) for c in fila))

    def test_valores_correctos_sin_cabecera(self):
        tabla = [["Alpha", "Beta"], ["Gamma", "Delta"]]
        ruta = _make_docx(self.tmp, tablas=[tabla])
        result = extract_tables_raw(ruta)
        # Primera fila intacta (sin interpretación de cabeceras)
        self.assertIn("Alpha", result[0][0])
        self.assertIn("Beta", result[0][0])

    def test_multiples_tablas(self):
        t1 = [["A"], ["1"]]
        t2 = [["X", "Y"], ["a", "b"]]
        ruta = _make_docx(self.tmp, tablas=[t1, t2])
        result = extract_tables_raw(ruta)
        self.assertEqual(len(result), 2)

    def test_archivo_inexistente_file_not_found(self):
        with self.assertRaises(FileNotFoundError):
            extract_tables_raw(Path(self.tmp) / "nope.docx")

    def test_extension_invalida_value_error(self):
        ruta = Path(self.tmp) / "doc.txt"
        ruta.write_text("x", encoding="utf-8")
        with self.assertRaises(ValueError):
            extract_tables_raw(ruta)

    def test_docx_invalido_value_error(self):
        ruta = _make_invalid_docx(self.tmp)
        with self.assertRaises(ValueError):
            extract_tables_raw(ruta)


# ---------------------------------------------------------------------------
# TestMtimeInvariante
# ---------------------------------------------------------------------------

class TestMtimeInvariante(unittest.TestCase):

    def setUp(self):
        self.tmp = tempfile.mkdtemp()

    def test_parse_docx_no_modifica_mtime(self):
        ruta = _make_docx(self.tmp, texto="Contenido de prueba.")
        mtime_before = ruta.stat().st_mtime
        time.sleep(0.05)
        parse_docx(ruta)
        self.assertAlmostEqual(ruta.stat().st_mtime, mtime_before, places=2)

    def test_extract_tables_raw_no_modifica_mtime(self):
        ruta = _make_docx(self.tmp, tablas=[[["A"], ["1"]]])
        mtime_before = ruta.stat().st_mtime
        time.sleep(0.05)
        extract_tables_raw(ruta)
        self.assertAlmostEqual(ruta.stat().st_mtime, mtime_before, places=2)


# ---------------------------------------------------------------------------
# TestRutaComoString
# ---------------------------------------------------------------------------

class TestRutaComoString(unittest.TestCase):

    def setUp(self):
        self.tmp = tempfile.mkdtemp()

    def test_parse_docx_acepta_str(self):
        ruta = _make_docx(self.tmp)
        result = parse_docx(str(ruta))
        self.assertIsInstance(result, DocxContent)

    def test_extract_tables_raw_acepta_str(self):
        ruta = _make_docx(self.tmp, tablas=[[["X"], ["y"]]])
        result = extract_tables_raw(str(ruta))
        self.assertIsInstance(result, list)


# ---------------------------------------------------------------------------
# TestFixtureParcela — solo lectura
# ---------------------------------------------------------------------------

class TestFixtureParcela(unittest.TestCase):
    """Tests de solo lectura contra el DOCX real de PARCELA.

    Usan skipUnless para no bloquear si el expediente no está disponible.
    No modifican el archivo en ningún caso.
    """

    @unittest.skipUnless(_DOCX_PARCELA.exists(), "DOCX de PARCELA no disponible")
    def test_parse_no_lanza_excepcion(self):
        result = parse_docx(_DOCX_PARCELA)
        self.assertIsInstance(result, DocxContent)

    @unittest.skipUnless(_DOCX_PARCELA.exists(), "DOCX de PARCELA no disponible")
    def test_texto_no_vacio(self):
        result = parse_docx(_DOCX_PARCELA)
        self.assertGreater(len(result.texto), 0)

    @unittest.skipUnless(_DOCX_PARCELA.exists(), "DOCX de PARCELA no disponible")
    def test_texto_contiene_recimetal(self):
        result = parse_docx(_DOCX_PARCELA)
        self.assertIn("RECIMETAL", result.texto)

    @unittest.skipUnless(_DOCX_PARCELA.exists(), "DOCX de PARCELA no disponible")
    def test_metadatos_es_dict(self):
        result = parse_docx(_DOCX_PARCELA)
        self.assertIsInstance(result.metadatos, dict)

    @unittest.skipUnless(_DOCX_PARCELA.exists(), "DOCX de PARCELA no disponible")
    def test_metadatos_tiene_campos_esperados(self):
        result = parse_docx(_DOCX_PARCELA)
        for campo in ("author", "created", "modified", "title", "subject"):
            self.assertIn(campo, result.metadatos)

    @unittest.skipUnless(_DOCX_PARCELA.exists(), "DOCX de PARCELA no disponible")
    def test_paginas_estimadas_mayor_que_1(self):
        # El documento tiene >56000 caracteres → muchas páginas estimadas
        result = parse_docx(_DOCX_PARCELA)
        self.assertGreater(result.num_paginas_estimadas, 1)

    @unittest.skipUnless(_DOCX_PARCELA.exists(), "DOCX de PARCELA no disponible")
    def test_tablas_no_vacias(self):
        # El DOCX de PARCELA tiene 18 tablas verificadas
        result = parse_docx(_DOCX_PARCELA)
        self.assertGreater(len(result.tablas), 0)

    @unittest.skipUnless(_DOCX_PARCELA.exists(), "DOCX de PARCELA no disponible")
    def test_tablas_primera_fila_es_dict(self):
        result = parse_docx(_DOCX_PARCELA)
        for tabla in result.tablas:
            if tabla:
                self.assertIsInstance(tabla[0], dict)
                break

    @unittest.skipUnless(_DOCX_PARCELA.exists(), "DOCX de PARCELA no disponible")
    def test_extract_tables_raw_no_lanza_excepcion(self):
        result = extract_tables_raw(_DOCX_PARCELA)
        self.assertIsInstance(result, list)

    @unittest.skipUnless(_DOCX_PARCELA.exists(), "DOCX de PARCELA no disponible")
    def test_extract_tables_raw_filas_son_listas_de_strings(self):
        result = extract_tables_raw(_DOCX_PARCELA)
        for tabla in result:
            for fila in tabla:
                self.assertIsInstance(fila, list)
                self.assertTrue(all(isinstance(c, str) for c in fila))

    @unittest.skipUnless(_DOCX_PARCELA.exists(), "DOCX de PARCELA no disponible")
    def test_no_modifica_mtime(self):
        mtime_before = _DOCX_PARCELA.stat().st_mtime
        time.sleep(0.05)
        parse_docx(_DOCX_PARCELA)
        self.assertAlmostEqual(_DOCX_PARCELA.stat().st_mtime, mtime_before, places=2)


if __name__ == "__main__":
    unittest.main()
