"""
tests/test_document_numbering_manager.py — EN-04
Tests del modulo de numeracion DOCX.

100 % offline. Sin IA. Sin web. Sin APIs externas.
Usa DOCX sinteticos generados con python-docx en memoria.
"""
import sys
import json
import tempfile
import unittest
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent / "src"))
sys.path.insert(0, str(Path(__file__).parent.parent))

from eia_agent.core.document_numbering_manager import (
    DEFAULT_BULLET_STYLE_CANDIDATES,
    DEFAULT_NUMBERED_STYLE_CANDIDATES,
    NUMBERING_OUTPUT_DOCX,
    NUMBERING_RESULT_JSON,
    NUMBERING_RESULT_MD,
    NUMBERING_SEVERITY,
    NUMBERING_STATUS,
    NumberingIssue,
    NumberingResult,
    ParagraphNumberingStatus,
    analyze_docx_numbering,
    apply_list_styles_to_docx,
    build_numbering_report_markdown,
    docx_has_numbering_definitions,
    is_bullet_list_candidate,
    is_numbered_list_candidate,
    process_document_numbering,
    select_bullet_style,
    select_numbered_style,
    validate_docx_file,
    write_numbering_outputs,
)


# ---------------------------------------------------------------------------
# Helper: crear DOCX sinteticos
# ---------------------------------------------------------------------------

def _make_docx(path: Path, paragraphs: list = None) -> None:
    """
    Crea un DOCX minimo con los parrafos indicados.
    paragraphs: lista de (text, style_name) o strings.
    """
    import docx
    doc = docx.Document()
    if paragraphs:
        for item in paragraphs:
            if isinstance(item, str):
                doc.add_paragraph(item)
            else:
                text, style = item
                try:
                    doc.add_paragraph(text, style=style)
                except Exception:
                    doc.add_paragraph(text)
    else:
        doc.add_paragraph("Parrafo de prueba sin contenido especial.")
    doc.save(str(path))


def _make_minimal_docx(path: Path) -> None:
    """Crea un DOCX minimo con parrafos de lista para pruebas."""
    _make_docx(path, [
        "Texto normal introductorio.",
        "1. Primer elemento de lista numerada",
        "2. Segundo elemento de lista numerada",
        "- Viñeta uno",
        "- Viñeta dos",
        "a) Subelemento a",
        "Texto normal de cierre.",
    ])


# ---------------------------------------------------------------------------
# Clase 1: NumberingIssue
# ---------------------------------------------------------------------------

class TestNumberingIssue(unittest.TestCase):

    def test_creation_basic(self):
        issue = NumberingIssue(
            severity="ERROR",
            code="EN04-E001",
            paragraph_index=5,
            message="Error de prueba",
        )
        self.assertEqual(issue.severity, "ERROR")
        self.assertEqual(issue.code, "EN04-E001")
        self.assertEqual(issue.paragraph_index, 5)
        self.assertEqual(issue.message, "Error de prueba")

    def test_default_evidence_empty(self):
        issue = NumberingIssue(severity="WARNING", code="X", paragraph_index=None, message="m")
        self.assertEqual(issue.evidence, [])

    def test_default_recommendation_empty(self):
        issue = NumberingIssue(severity="INFO", code="X", paragraph_index=None, message="m")
        self.assertEqual(issue.recommendation, "")

    def test_to_dict_has_all_fields(self):
        issue = NumberingIssue(
            severity="WARNING",
            code="EN04-W001",
            paragraph_index=3,
            message="Estilo no encontrado",
            recommendation="Usar Word",
            evidence=["texto"],
        )
        d = issue.to_dict()
        self.assertIn("severity", d)
        self.assertIn("code", d)
        self.assertIn("paragraph_index", d)
        self.assertIn("message", d)
        self.assertIn("recommendation", d)
        self.assertIn("evidence", d)

    def test_to_dict_values_correct(self):
        issue = NumberingIssue(
            severity="ERROR",
            code="EN04-E002",
            paragraph_index=10,
            message="msg",
            recommendation="rec",
            evidence=["ev1", "ev2"],
        )
        d = issue.to_dict()
        self.assertEqual(d["severity"], "ERROR")
        self.assertEqual(d["paragraph_index"], 10)
        self.assertEqual(d["evidence"], ["ev1", "ev2"])

    def test_to_dict_none_paragraph_index(self):
        issue = NumberingIssue(severity="INFO", code="X", paragraph_index=None, message="m")
        d = issue.to_dict()
        self.assertIsNone(d["paragraph_index"])

    def test_summary_with_paragraph_index(self):
        issue = NumberingIssue(severity="ERROR", code="EN04-E001", paragraph_index=7, message="msg")
        s = issue.summary()
        self.assertIn("ERROR", s)
        self.assertIn("EN04-E001", s)
        self.assertIn("7", s)
        self.assertIn("msg", s)

    def test_summary_without_paragraph_index(self):
        issue = NumberingIssue(severity="WARNING", code="EN04-W001", paragraph_index=None, message="msg")
        s = issue.summary()
        self.assertIn("WARNING", s)
        self.assertNotIn("parrafo None", s)


# ---------------------------------------------------------------------------
# Clase 2: ParagraphNumberingStatus
# ---------------------------------------------------------------------------

class TestParagraphNumberingStatus(unittest.TestCase):

    def test_creation_basic(self):
        ps = ParagraphNumberingStatus(
            paragraph_index=0,
            text_preview="1. texto",
            style_name="Normal",
            detected_as_list_candidate=True,
            applied_style="List Number",
        )
        self.assertEqual(ps.paragraph_index, 0)
        self.assertTrue(ps.detected_as_list_candidate)
        self.assertEqual(ps.applied_style, "List Number")

    def test_default_warnings_empty(self):
        ps = ParagraphNumberingStatus(0, "t", None, False, None)
        self.assertEqual(ps.warnings, [])

    def test_default_notes_empty(self):
        ps = ParagraphNumberingStatus(0, "t", None, False, None)
        self.assertEqual(ps.notes, [])

    def test_to_dict_has_fields(self):
        ps = ParagraphNumberingStatus(2, "texto", "Normal", True, "List Number")
        d = ps.to_dict()
        self.assertIn("paragraph_index", d)
        self.assertIn("text_preview", d)
        self.assertIn("style_name", d)
        self.assertIn("detected_as_list_candidate", d)
        self.assertIn("applied_style", d)
        self.assertIn("warnings", d)
        self.assertIn("notes", d)

    def test_to_dict_values(self):
        ps = ParagraphNumberingStatus(5, "preview", "Heading 1", False, None)
        d = ps.to_dict()
        self.assertEqual(d["paragraph_index"], 5)
        self.assertFalse(d["detected_as_list_candidate"])
        self.assertIsNone(d["applied_style"])

    def test_summary_candidate(self):
        ps = ParagraphNumberingStatus(0, "1. texto", "Normal", True, "List Number")
        s = ps.summary()
        self.assertIn("CANDIDATO", s)
        self.assertIn("List Number", s)

    def test_summary_normal(self):
        ps = ParagraphNumberingStatus(3, "texto normal", "Normal", False, None)
        s = ps.summary()
        self.assertIn("normal", s)


# ---------------------------------------------------------------------------
# Clase 3: NumberingResult
# ---------------------------------------------------------------------------

class TestNumberingResult(unittest.TestCase):

    def _make_result(self, issues=None, warnings=None) -> NumberingResult:
        return NumberingResult(
            input_docx="doc.docx",
            output_docx=None,
            status=NUMBERING_STATUS["OK"],
            numbering_definitions_found=False,
            paragraphs_checked=10,
            list_candidates_found=3,
            styles_applied_count=0,
            issues=issues or [],
            warnings=warnings or [],
        )

    def test_error_count_zero(self):
        r = self._make_result()
        self.assertEqual(r.error_count(), 0)

    def test_error_count_with_errors(self):
        issues = [
            NumberingIssue(severity="ERROR", code="E", paragraph_index=None, message="m"),
            NumberingIssue(severity="WARNING", code="W", paragraph_index=None, message="m"),
        ]
        r = self._make_result(issues=issues)
        self.assertEqual(r.error_count(), 1)

    def test_warning_count_zero(self):
        r = self._make_result()
        self.assertEqual(r.warning_count(), 0)

    def test_warning_count_with_warnings(self):
        issues = [
            NumberingIssue(severity="WARNING", code="W", paragraph_index=None, message="m"),
            NumberingIssue(severity="WARNING", code="W2", paragraph_index=None, message="m"),
        ]
        r = self._make_result(issues=issues)
        self.assertEqual(r.warning_count(), 2)

    def test_is_valid_no_errors(self):
        r = self._make_result()
        self.assertTrue(r.is_valid())

    def test_is_valid_with_errors(self):
        issues = [NumberingIssue(severity="ERROR", code="E", paragraph_index=None, message="m")]
        r = self._make_result(issues=issues)
        self.assertFalse(r.is_valid())

    def test_is_valid_warnings_only(self):
        issues = [NumberingIssue(severity="WARNING", code="W", paragraph_index=None, message="m")]
        r = self._make_result(issues=issues)
        self.assertTrue(r.is_valid())

    def test_to_dict_has_all_fields(self):
        r = self._make_result()
        d = r.to_dict()
        expected = [
            "input_docx", "output_docx", "status", "numbering_definitions_found",
            "paragraphs_checked", "list_candidates_found", "styles_applied_count",
            "error_count", "warning_count", "is_valid",
            "paragraph_statuses", "issues", "warnings", "notes",
        ]
        for key in expected:
            self.assertIn(key, d, f"Falta campo '{key}' en to_dict()")

    def test_to_dict_values(self):
        r = self._make_result()
        d = r.to_dict()
        self.assertEqual(d["paragraphs_checked"], 10)
        self.assertEqual(d["list_candidates_found"], 3)
        self.assertEqual(d["error_count"], 0)
        self.assertTrue(d["is_valid"])

    def test_summary_contains_status(self):
        r = self._make_result()
        s = r.summary()
        self.assertIn("OK", s)

    def test_summary_contains_counts(self):
        r = self._make_result()
        s = r.summary()
        self.assertIn("10", s)
        self.assertIn("3", s)

    def test_status_sin_datos(self):
        r = NumberingResult(
            input_docx="x.docx", output_docx=None,
            status=NUMBERING_STATUS["SIN_DATOS"],
            numbering_definitions_found=False,
            paragraphs_checked=0, list_candidates_found=0, styles_applied_count=0,
        )
        self.assertEqual(r.status, "SIN_DATOS")

    def test_output_docx_in_to_dict(self):
        r = NumberingResult(
            input_docx="in.docx", output_docx="out.docx",
            status="OK", numbering_definitions_found=False,
            paragraphs_checked=0, list_candidates_found=0, styles_applied_count=0,
        )
        d = r.to_dict()
        self.assertEqual(d["output_docx"], "out.docx")


# ---------------------------------------------------------------------------
# Clase 4: validate_docx_file
# ---------------------------------------------------------------------------

class TestValidateDocxFile(unittest.TestCase):

    def test_valid_docx_returns_true(self):
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / "test.docx"
            _make_docx(p)
            self.assertTrue(validate_docx_file(p))

    def test_nonexistent_returns_false(self):
        self.assertFalse(validate_docx_file("/ruta/inexistente/doc.docx"))

    def test_empty_file_returns_false(self):
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / "empty.docx"
            p.write_bytes(b"")
            self.assertFalse(validate_docx_file(p))

    def test_non_docx_file_returns_false(self):
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / "archivo.txt"
            p.write_text("contenido de texto", encoding="utf-8")
            self.assertFalse(validate_docx_file(p))

    def test_path_as_string_works(self):
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / "test.docx"
            _make_docx(p)
            self.assertTrue(validate_docx_file(str(p)))

    def test_directory_returns_false(self):
        with tempfile.TemporaryDirectory() as tmp:
            self.assertFalse(validate_docx_file(tmp))


# ---------------------------------------------------------------------------
# Clase 5: docx_has_numbering_definitions
# ---------------------------------------------------------------------------

class TestDocxHasNumberingDefinitions(unittest.TestCase):

    def test_plain_docx_no_numbering(self):
        # python-docx crea DOCX sin numbering.xml por defecto
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / "plain.docx"
            _make_docx(p)
            result = docx_has_numbering_definitions(p)
            # Documenta el comportamiento: python-docx por defecto no crea numbering.xml
            self.assertIsInstance(result, bool)

    def test_nonexistent_returns_false(self):
        self.assertFalse(docx_has_numbering_definitions("/no/existe.docx"))

    def test_fake_docx_with_numbering_xml(self):
        import zipfile as _zf
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / "with_numbering.docx"
            # Crear un ZIP falso con word/numbering.xml
            with _zf.ZipFile(str(p), "w") as zf:
                zf.writestr("word/document.xml", "<root/>")
                zf.writestr("word/numbering.xml", "<numbering/>")
            self.assertTrue(docx_has_numbering_definitions(p))

    def test_fake_docx_without_numbering_xml(self):
        import zipfile as _zf
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / "without_numbering.docx"
            with _zf.ZipFile(str(p), "w") as zf:
                zf.writestr("word/document.xml", "<root/>")
            self.assertFalse(docx_has_numbering_definitions(p))

    def test_invalid_zip_returns_false(self):
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / "notzip.docx"
            p.write_bytes(b"no es un zip")
            self.assertFalse(docx_has_numbering_definitions(p))


# ---------------------------------------------------------------------------
# Clase 6: is_numbered_list_candidate
# ---------------------------------------------------------------------------

class TestIsNumberedListCandidate(unittest.TestCase):

    # --- Casos que DEBEN detectarse ---

    def test_numeric_dot(self):
        self.assertTrue(is_numbered_list_candidate("1. Primer elemento"))

    def test_numeric_paren(self):
        self.assertTrue(is_numbered_list_candidate("1) Primer elemento"))

    def test_numeric_dot_dash(self):
        self.assertTrue(is_numbered_list_candidate("1.- Primer elemento"))

    def test_lowercase_letter_paren(self):
        self.assertTrue(is_numbered_list_candidate("a) Subelemento"))

    def test_uppercase_letter_dot(self):
        self.assertTrue(is_numbered_list_candidate("A. Seccion A"))

    def test_uppercase_i_dot(self):
        self.assertTrue(is_numbered_list_candidate("I. Introduccion"))

    def test_lowercase_i_paren(self):
        self.assertTrue(is_numbered_list_candidate("i) subelemento"))

    def test_multidigit_numeric(self):
        self.assertTrue(is_numbered_list_candidate("10. Decimo elemento"))

    def test_uppercase_b_dot(self):
        self.assertTrue(is_numbered_list_candidate("B. Otro apartado"))

    def test_lowercase_b_paren(self):
        self.assertTrue(is_numbered_list_candidate("b) subelemento b"))

    def test_leading_spaces_stripped(self):
        self.assertTrue(is_numbered_list_candidate("   1. con espacios iniciales"))

    # --- Casos que NO deben detectarse ---

    def test_decimal_not_detected(self):
        self.assertFalse(is_numbered_list_candidate("1.5 kg de material"))

    def test_thousands_separator_not_detected(self):
        self.assertFalse(is_numbered_list_candidate("1.250 euros de presupuesto"))

    def test_slash_date_not_detected(self):
        self.assertFalse(is_numbered_list_candidate("01/01/2026"))

    def test_coordinate_not_detected(self):
        self.assertFalse(is_numbered_list_candidate("28.123456N"))

    def test_plain_text_not_detected(self):
        self.assertFalse(is_numbered_list_candidate("Texto normal sin patron de lista"))

    def test_empty_string_not_detected(self):
        self.assertFalse(is_numbered_list_candidate(""))

    def test_only_spaces_not_detected(self):
        self.assertFalse(is_numbered_list_candidate("   "))


# ---------------------------------------------------------------------------
# Clase 7: is_bullet_list_candidate
# ---------------------------------------------------------------------------

class TestIsBulletListCandidate(unittest.TestCase):

    # --- Casos que DEBEN detectarse ---

    def test_dash_bullet(self):
        self.assertTrue(is_bullet_list_candidate("- Elemento de lista"))

    def test_unicode_bullet(self):
        self.assertTrue(is_bullet_list_candidate("• Viñeta unicode"))

    def test_asterisk_bullet(self):
        self.assertTrue(is_bullet_list_candidate("* Elemento con asterisco"))

    def test_endash_bullet(self):
        self.assertTrue(is_bullet_list_candidate("– Elemento con guion largo"))

    def test_emdash_bullet(self):
        self.assertTrue(is_bullet_list_candidate("— Elemento con raya"))

    def test_leading_spaces_stripped(self):
        self.assertTrue(is_bullet_list_candidate("  - con espacios"))

    # --- Casos que NO deben detectarse ---

    def test_internal_hyphen_not_detected(self):
        self.assertFalse(is_bullet_list_candidate("texto-con-guion-interno"))

    def test_phrase_with_middle_dash_not_detected(self):
        self.assertFalse(is_bullet_list_candidate("El promotor — como señala la documentacion — propone"))

    def test_plain_text_not_detected(self):
        self.assertFalse(is_bullet_list_candidate("Texto normal introductorio"))

    def test_empty_not_detected(self):
        self.assertFalse(is_bullet_list_candidate(""))


# ---------------------------------------------------------------------------
# Clase 8: select_numbered_style y select_bullet_style
# ---------------------------------------------------------------------------

class TestSelectStyles(unittest.TestCase):

    def test_numbered_level1(self):
        self.assertEqual(select_numbered_style(1), "List Number")

    def test_numbered_level2(self):
        self.assertEqual(select_numbered_style(2), "List Number 2")

    def test_numbered_level3(self):
        self.assertEqual(select_numbered_style(3), "List Number 3")

    def test_numbered_level0_clamped_to_1(self):
        self.assertEqual(select_numbered_style(0), "List Number")

    def test_numbered_level_above_3_clamped_to_3(self):
        self.assertEqual(select_numbered_style(5), "List Number 3")
        self.assertEqual(select_numbered_style(10), "List Number 3")

    def test_bullet_level1(self):
        self.assertEqual(select_bullet_style(1), "List Bullet")

    def test_bullet_level2(self):
        self.assertEqual(select_bullet_style(2), "List Bullet 2")

    def test_bullet_level3(self):
        self.assertEqual(select_bullet_style(3), "List Bullet 3")

    def test_bullet_level0_clamped_to_1(self):
        self.assertEqual(select_bullet_style(0), "List Bullet")

    def test_bullet_level_above_3_clamped_to_3(self):
        self.assertEqual(select_bullet_style(7), "List Bullet 3")

    def test_numbered_default_level(self):
        self.assertEqual(select_numbered_style(), "List Number")

    def test_bullet_default_level(self):
        self.assertEqual(select_bullet_style(), "List Bullet")


# ---------------------------------------------------------------------------
# Clase 9: analyze_docx_numbering
# ---------------------------------------------------------------------------

class TestAnalyzeDocxNumbering(unittest.TestCase):

    def test_analyze_plain_docx(self):
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / "doc.docx"
            _make_docx(p)
            result = analyze_docx_numbering(p)
            self.assertIsInstance(result, NumberingResult)
            self.assertGreater(result.paragraphs_checked, 0)

    def test_analyze_detects_numbered_candidates(self):
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / "doc.docx"
            _make_minimal_docx(p)
            result = analyze_docx_numbering(p)
            self.assertGreater(result.list_candidates_found, 0)

    def test_analyze_no_output_docx(self):
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / "doc.docx"
            _make_docx(p)
            result = analyze_docx_numbering(p)
            self.assertIsNone(result.output_docx)

    def test_analyze_no_file_created(self):
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / "doc.docx"
            _make_minimal_docx(p)
            before = list(Path(tmp).iterdir())
            analyze_docx_numbering(p)
            after = list(Path(tmp).iterdir())
            self.assertEqual(len(before), len(after))

    def test_analyze_nonexistent_returns_sin_datos(self):
        result = analyze_docx_numbering("/no/existe/doc.docx")
        self.assertEqual(result.status, NUMBERING_STATUS["SIN_DATOS"])

    def test_analyze_valid_status(self):
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / "doc.docx"
            _make_docx(p)
            result = analyze_docx_numbering(p)
            self.assertIn(result.status, list(NUMBERING_STATUS.values()))

    def test_analyze_paragraph_statuses_count(self):
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / "doc.docx"
            _make_minimal_docx(p)
            result = analyze_docx_numbering(p)
            self.assertEqual(len(result.paragraph_statuses), result.paragraphs_checked)

    def test_analyze_path_as_string(self):
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / "doc.docx"
            _make_docx(p)
            result = analyze_docx_numbering(str(p))
            self.assertIsInstance(result, NumberingResult)


# ---------------------------------------------------------------------------
# Clase 10: apply_list_styles_to_docx
# ---------------------------------------------------------------------------

class TestApplyListStylesToDocx(unittest.TestCase):

    def test_creates_output_docx(self):
        with tempfile.TemporaryDirectory() as tmp:
            src = Path(tmp) / "in.docx"
            dst = Path(tmp) / "out.docx"
            _make_minimal_docx(src)
            apply_list_styles_to_docx(src, dst)
            self.assertTrue(dst.exists())

    def test_input_not_modified(self):
        with tempfile.TemporaryDirectory() as tmp:
            src = Path(tmp) / "in.docx"
            dst = Path(tmp) / "out.docx"
            _make_minimal_docx(src)
            size_before = src.stat().st_size
            apply_list_styles_to_docx(src, dst)
            size_after = src.stat().st_size
            self.assertEqual(size_before, size_after)

    def test_detects_candidates(self):
        with tempfile.TemporaryDirectory() as tmp:
            src = Path(tmp) / "in.docx"
            dst = Path(tmp) / "out.docx"
            _make_minimal_docx(src)
            result = apply_list_styles_to_docx(src, dst)
            self.assertGreater(result.list_candidates_found, 0)

    def test_output_opens_with_python_docx(self):
        with tempfile.TemporaryDirectory() as tmp:
            src = Path(tmp) / "in.docx"
            dst = Path(tmp) / "out.docx"
            _make_minimal_docx(src)
            apply_list_styles_to_docx(src, dst)
            import docx
            doc = docx.Document(str(dst))
            self.assertGreater(len(doc.paragraphs), 0)

    def test_missing_style_creates_warning_not_crash(self):
        with tempfile.TemporaryDirectory() as tmp:
            src = Path(tmp) / "in.docx"
            dst = Path(tmp) / "out.docx"
            _make_minimal_docx(src)
            # Siempre debe terminar sin lanzar excepcion
            result = apply_list_styles_to_docx(src, dst)
            self.assertIsInstance(result, NumberingResult)

    def test_nonexistent_input_returns_sin_datos(self):
        with tempfile.TemporaryDirectory() as tmp:
            dst = Path(tmp) / "out.docx"
            result = apply_list_styles_to_docx("/no/existe.docx", dst)
            self.assertEqual(result.status, NUMBERING_STATUS["SIN_DATOS"])

    def test_result_has_input_and_output_docx(self):
        with tempfile.TemporaryDirectory() as tmp:
            src = Path(tmp) / "in.docx"
            dst = Path(tmp) / "out.docx"
            _make_minimal_docx(src)
            result = apply_list_styles_to_docx(src, dst)
            self.assertIsNotNone(result.input_docx)
            self.assertIsNotNone(result.output_docx)

    def test_apply_numbered_false_skips_numbered(self):
        with tempfile.TemporaryDirectory() as tmp:
            src = Path(tmp) / "in.docx"
            dst = Path(tmp) / "out.docx"
            _make_minimal_docx(src)
            result_with = apply_list_styles_to_docx(src, dst, apply_numbered=True, apply_bullets=False)
            dst2 = Path(tmp) / "out2.docx"
            result_without = apply_list_styles_to_docx(src, dst2, apply_numbered=False, apply_bullets=False)
            # Sin numbered ni bullets, candidates = 0
            self.assertEqual(result_without.list_candidates_found, 0)

    def test_note_about_no_modification_present(self):
        with tempfile.TemporaryDirectory() as tmp:
            src = Path(tmp) / "in.docx"
            dst = Path(tmp) / "out.docx"
            _make_minimal_docx(src)
            result = apply_list_styles_to_docx(src, dst)
            # La nota de alcance debe estar presente en los notes
            notes_text = " ".join(result.notes)
            self.assertIn("no modifica el contenido", notes_text)


# ---------------------------------------------------------------------------
# Clase 11: build_numbering_report_markdown
# ---------------------------------------------------------------------------

class TestBuildNumberingReportMarkdown(unittest.TestCase):

    def _make_result(self) -> NumberingResult:
        return NumberingResult(
            input_docx="doc.docx",
            output_docx="doc_num.docx",
            status=NUMBERING_STATUS["OK"],
            numbering_definitions_found=True,
            paragraphs_checked=10,
            list_candidates_found=4,
            styles_applied_count=4,
            paragraph_statuses=[
                ParagraphNumberingStatus(0, "1. item", "Normal", True, "List Number"),
                ParagraphNumberingStatus(1, "texto", "Normal", False, None),
            ],
        )

    def test_returns_string(self):
        r = self._make_result()
        md = build_numbering_report_markdown(r)
        self.assertIsInstance(md, str)

    def test_has_title(self):
        r = self._make_result()
        md = build_numbering_report_markdown(r)
        self.assertIn("# Resultado de numeracion DOCX", md)

    def test_has_6_sections(self):
        r = self._make_result()
        md = build_numbering_report_markdown(r)
        for i in range(1, 7):
            self.assertIn(f"## {i}.", md)

    def test_has_scope_warning(self):
        r = self._make_result()
        md = build_numbering_report_markdown(r)
        self.assertIn("no modifica el contenido tecnico", md)

    def test_has_status(self):
        r = self._make_result()
        md = build_numbering_report_markdown(r)
        self.assertIn("OK", md)

    def test_sin_datos_result(self):
        r = NumberingResult(
            input_docx="x.docx", output_docx=None,
            status="SIN_DATOS",
            numbering_definitions_found=False,
            paragraphs_checked=0, list_candidates_found=0, styles_applied_count=0,
        )
        md = build_numbering_report_markdown(r)
        self.assertIn("SIN_DATOS", md)


# ---------------------------------------------------------------------------
# Clase 12: write_numbering_outputs
# ---------------------------------------------------------------------------

class TestWriteNumberingOutputs(unittest.TestCase):

    def _make_result(self) -> NumberingResult:
        return NumberingResult(
            input_docx="doc.docx",
            output_docx=None,
            status=NUMBERING_STATUS["OK"],
            numbering_definitions_found=False,
            paragraphs_checked=5,
            list_candidates_found=2,
            styles_applied_count=0,
        )

    def test_creates_json_and_md(self):
        with tempfile.TemporaryDirectory() as tmp:
            r = self._make_result()
            json_p, md_p = write_numbering_outputs(r, tmp)
            self.assertTrue(json_p.exists())
            self.assertTrue(md_p.exists())

    def test_json_is_valid(self):
        with tempfile.TemporaryDirectory() as tmp:
            r = self._make_result()
            json_p, _ = write_numbering_outputs(r, tmp)
            data = json.loads(json_p.read_text(encoding="utf-8"))
            self.assertIn("status", data)
            self.assertIn("paragraphs_checked", data)

    def test_md_has_content(self):
        with tempfile.TemporaryDirectory() as tmp:
            r = self._make_result()
            _, md_p = write_numbering_outputs(r, tmp)
            content = md_p.read_text(encoding="utf-8")
            self.assertIn("#", content)

    def test_returns_paths(self):
        with tempfile.TemporaryDirectory() as tmp:
            r = self._make_result()
            json_p, md_p = write_numbering_outputs(r, tmp)
            self.assertIsInstance(json_p, Path)
            self.assertIsInstance(md_p, Path)

    def test_filenames(self):
        with tempfile.TemporaryDirectory() as tmp:
            r = self._make_result()
            json_p, md_p = write_numbering_outputs(r, tmp)
            self.assertEqual(json_p.name, NUMBERING_RESULT_JSON)
            self.assertEqual(md_p.name, NUMBERING_RESULT_MD)


# ---------------------------------------------------------------------------
# Clase 13: process_document_numbering
# ---------------------------------------------------------------------------

class TestProcessDocumentNumbering(unittest.TestCase):

    def _make_expediente(self, tmp: str, with_docx: bool = True) -> Path:
        exp = Path(tmp) / "expediente"
        (exp / "documento").mkdir(parents=True)
        if with_docx:
            docx_path = exp / "documento" / "documento_ambiental_borrador.docx"
            _make_minimal_docx(docx_path)
        return exp

    def test_no_docx_returns_sin_datos(self):
        with tempfile.TemporaryDirectory() as tmp:
            exp = self._make_expediente(tmp, with_docx=False)
            result = process_document_numbering(exp)
            self.assertEqual(result.status, NUMBERING_STATUS["SIN_DATOS"])

    def test_apply_false_no_numerado_docx(self):
        with tempfile.TemporaryDirectory() as tmp:
            exp = self._make_expediente(tmp)
            process_document_numbering(exp, apply_styles=False)
            numerado = exp / "documento" / NUMBERING_OUTPUT_DOCX
            self.assertFalse(numerado.exists())

    def test_apply_true_creates_numerado_docx(self):
        with tempfile.TemporaryDirectory() as tmp:
            exp = self._make_expediente(tmp)
            process_document_numbering(exp, apply_styles=True)
            numerado = exp / "documento" / NUMBERING_OUTPUT_DOCX
            self.assertTrue(numerado.exists())

    def test_write_outputs_true_creates_json_md(self):
        with tempfile.TemporaryDirectory() as tmp:
            exp = self._make_expediente(tmp)
            process_document_numbering(exp, write_outputs=True)
            json_p = exp / "documento" / NUMBERING_RESULT_JSON
            md_p = exp / "documento" / NUMBERING_RESULT_MD
            self.assertTrue(json_p.exists())
            self.assertTrue(md_p.exists())

    def test_write_outputs_false_no_json_md(self):
        with tempfile.TemporaryDirectory() as tmp:
            exp = self._make_expediente(tmp)
            process_document_numbering(exp, write_outputs=False)
            json_p = exp / "documento" / NUMBERING_RESULT_JSON
            self.assertFalse(json_p.exists())

    def test_returns_numbering_result(self):
        with tempfile.TemporaryDirectory() as tmp:
            exp = self._make_expediente(tmp)
            result = process_document_numbering(exp)
            self.assertIsInstance(result, NumberingResult)

    def test_prioritizes_best_docx(self):
        with tempfile.TemporaryDirectory() as tmp:
            exp = self._make_expediente(tmp)
            # Crear un DOCX de mayor prioridad
            final_docx = exp / "documento" / "documento_ambiental_final_revisable.docx"
            _make_docx(final_docx, ["Contenido del documento final revisable."])
            result = process_document_numbering(exp)
            # Debe usar el de mayor prioridad
            self.assertIn("final_revisable", result.input_docx)

    def test_original_not_modified(self):
        with tempfile.TemporaryDirectory() as tmp:
            exp = self._make_expediente(tmp)
            borrador = exp / "documento" / "documento_ambiental_borrador.docx"
            size_before = borrador.stat().st_size
            process_document_numbering(exp, apply_styles=True)
            size_after = borrador.stat().st_size
            self.assertEqual(size_before, size_after)


# ---------------------------------------------------------------------------
# Clase 14: Constantes
# ---------------------------------------------------------------------------

class TestConstantes(unittest.TestCase):

    def test_numbering_status_has_4_keys(self):
        self.assertEqual(len(NUMBERING_STATUS), 4)

    def test_numbering_status_values(self):
        self.assertIn("OK", NUMBERING_STATUS)
        self.assertIn("CON_OBSERVACIONES", NUMBERING_STATUS)
        self.assertIn("NO_CONFORME", NUMBERING_STATUS)
        self.assertIn("SIN_DATOS", NUMBERING_STATUS)

    def test_numbering_severity_has_3_keys(self):
        self.assertEqual(len(NUMBERING_SEVERITY), 3)

    def test_numbered_style_candidates_count(self):
        self.assertEqual(len(DEFAULT_NUMBERED_STYLE_CANDIDATES), 3)
        self.assertIn("List Number", DEFAULT_NUMBERED_STYLE_CANDIDATES)

    def test_bullet_style_candidates_count(self):
        self.assertEqual(len(DEFAULT_BULLET_STYLE_CANDIDATES), 3)
        self.assertIn("List Bullet", DEFAULT_BULLET_STYLE_CANDIDATES)

    def test_output_filenames_defined(self):
        self.assertTrue(NUMBERING_OUTPUT_DOCX.endswith(".docx"))
        self.assertTrue(NUMBERING_RESULT_JSON.endswith(".json"))
        self.assertTrue(NUMBERING_RESULT_MD.endswith(".md"))


# ---------------------------------------------------------------------------
# Clase 15: CLI document-numbering
# ---------------------------------------------------------------------------

class TestCLIDocumentNumbering(unittest.TestCase):

    def _make_expediente(self, tmp: str) -> Path:
        exp = Path(tmp) / "expediente-test"
        exp.mkdir()
        (exp / "documento").mkdir()
        docx_path = exp / "documento" / "documento_ambiental_borrador.docx"
        _make_minimal_docx(docx_path)
        return exp

    def test_command_exists_and_runs(self):
        from run_expediente import main
        with tempfile.TemporaryDirectory() as tmp:
            exp = self._make_expediente(tmp)
            rc = main([str(exp), "document-numbering"])
            self.assertIn(rc, [0, 1])

    def test_no_flags_no_outputs(self):
        from run_expediente import main
        with tempfile.TemporaryDirectory() as tmp:
            exp = self._make_expediente(tmp)
            main([str(exp), "document-numbering"])
            json_p = exp / "documento" / NUMBERING_RESULT_JSON
            self.assertFalse(json_p.exists())

    def test_write_flag_creates_json_md(self):
        from run_expediente import main
        with tempfile.TemporaryDirectory() as tmp:
            exp = self._make_expediente(tmp)
            main([str(exp), "document-numbering", "--write"])
            json_p = exp / "documento" / NUMBERING_RESULT_JSON
            md_p = exp / "documento" / NUMBERING_RESULT_MD
            self.assertTrue(json_p.exists())
            self.assertTrue(md_p.exists())

    def test_apply_flag_creates_numerado_docx(self):
        from run_expediente import main
        with tempfile.TemporaryDirectory() as tmp:
            exp = self._make_expediente(tmp)
            main([str(exp), "document-numbering", "--apply"])
            numerado = exp / "documento" / NUMBERING_OUTPUT_DOCX
            self.assertTrue(numerado.exists())

    def test_write_and_apply_flags(self):
        from run_expediente import main
        with tempfile.TemporaryDirectory() as tmp:
            exp = self._make_expediente(tmp)
            main([str(exp), "document-numbering", "--write", "--apply"])
            numerado = exp / "documento" / NUMBERING_OUTPUT_DOCX
            json_p = exp / "documento" / NUMBERING_RESULT_JSON
            self.assertTrue(numerado.exists())
            self.assertTrue(json_p.exists())

    def test_exit_code_valid_docx(self):
        from run_expediente import main
        with tempfile.TemporaryDirectory() as tmp:
            exp = self._make_expediente(tmp)
            rc = main([str(exp), "document-numbering"])
            # DOCX valido → exit 0
            self.assertEqual(rc, 0)


if __name__ == "__main__":
    unittest.main()
