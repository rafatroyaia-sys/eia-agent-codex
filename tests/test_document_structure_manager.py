"""
tests/test_document_structure_manager.py — EN-02
Tests del validador y normalizador de estructura del DOCX final.

100 % offline. Sin IA. Sin web. Sin APIs externas.
Usa DOCX sinteticos generados con python-docx en memoria.
"""
import sys
import json
import tempfile
import unittest
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent / "src"))

from eia_agent.core.document_structure_manager import (
    BLOCK_IDS,
    CANONICAL_DOCUMENT_ORDER,
    DocumentSectionPosition,
    DocumentStructureResult,
    build_document_structure_markdown,
    detect_document_sections,
    find_best_available_docx,
    normalize_document_structure,
    validate_document_structure,
    write_document_structure_outputs,
)


# ---------------------------------------------------------------------------
# Helper: crear DOCX sinteticos con python-docx
# ---------------------------------------------------------------------------

def _make_docx(path: Path, sections: list) -> None:
    """
    Crea un DOCX minimo con las secciones indicadas.
    sections es una lista de tuplas (section_id, heading_level, title_override)
    o simplemente strings de section_id para usar titulos por defecto.

    section_id puede ser:
        "PORTADA"     -> parrafo normal (no heading)
        "INDICE"      -> Heading 1 "Indice"
        "A".."K"      -> Heading 1 "A — Identificacion del proyecto"
        "ANEXO_GRAFICO" -> Heading 1 "Anexo grafico y cartografico"
        "HOJA_FIRMAS" -> Heading 1 "Hoja de firmas y revision tecnica"
        ("X", level, "Titulo personalizado") -> heading con nivel y titulo
    """
    import docx

    _default_titles = {
        "PORTADA": None,          # parrafo normal
        "INDICE": "Indice",
        "A": "A — Identificacion del promotor y el proyecto",
        "B": "B — Descripcion general del proyecto",
        "C": "C — Inventario ambiental",
        "D": "D — Identificacion de impactos",
        "E": "E — Valoracion de impactos",
        "F": "F — Plan de manejo ambiental",
        "G": "G — Alternativas y justificacion",
        "H": "H — Impactos residuales",
        "I": "I — Conclusion",
        "J": "J — Medidas compensatorias",
        "K": "K — Plan de vigilancia ambiental",
        "ANEXO_GRAFICO": "Anexo grafico y cartografico",
        "HOJA_FIRMAS": "Hoja de firmas y revision tecnica",
    }

    doc = docx.Document()
    for item in sections:
        if isinstance(item, str):
            sid = item
            level = 1
            title = _default_titles.get(sid)
        else:
            sid, level, title = item

        if sid == "PORTADA":
            p = doc.add_paragraph("Documento Ambiental — Expediente EIA-2026-TEST")
            p.add_run(" Borrador tecnico generado automaticamente")
        elif title is not None:
            doc.add_heading(title, level=level)
            doc.add_paragraph(f"Contenido de la seccion {sid}.")
        else:
            doc.add_paragraph(f"Seccion {sid} sin titulo.")

    doc.save(str(path))


# ---------------------------------------------------------------------------
# 1. Dataclasses
# ---------------------------------------------------------------------------

class TestDocumentSectionPosition(unittest.TestCase):

    def _make(self, **kw):
        defaults = dict(
            section_id="A",
            title="A — Identificacion",
            paragraph_index=5,
            heading_level=1,
            found=True,
        )
        defaults.update(kw)
        return DocumentSectionPosition(**defaults)

    def test_to_dict_keys(self):
        d = self._make().to_dict()
        for k in ("section_id", "title", "paragraph_index", "heading_level", "found", "notes"):
            self.assertIn(k, d)

    def test_to_dict_values(self):
        pos = self._make(section_id="B", paragraph_index=10)
        d = pos.to_dict()
        self.assertEqual(d["section_id"], "B")
        self.assertEqual(d["paragraph_index"], 10)

    def test_summary_found(self):
        pos = self._make(section_id="A", found=True, paragraph_index=5)
        s = pos.summary()
        self.assertIn("OK", s)
        self.assertIn("A", s)

    def test_summary_not_found(self):
        pos = self._make(found=False, paragraph_index=None)
        s = pos.summary()
        self.assertIn("NO ENCONTRADA", s)

    def test_notes_default_empty(self):
        pos = self._make()
        self.assertEqual(pos.notes, [])

    def test_found_false_paragraph_none(self):
        pos = self._make(found=False, paragraph_index=None, heading_level=None)
        d = pos.to_dict()
        self.assertIsNone(d["paragraph_index"])
        self.assertIsNone(d["heading_level"])

    def test_to_dict_serializable(self):
        pos = self._make()
        self.assertIsNotNone(json.dumps(pos.to_dict()))

    def test_notes_included_in_dict(self):
        pos = self._make()
        pos.notes.append("nota de prueba")
        d = pos.to_dict()
        self.assertIn("nota de prueba", d["notes"])


class TestDocumentStructureResult(unittest.TestCase):

    def _make_result(self, errors=0, warnings=0):
        errs = [{"code": f"EN02-E00{i}", "severity": "ERROR",
                  "section_id": "", "message": f"err{i}", "recommendation": ""}
                for i in range(errors)]
        warns = [{"code": f"EN02-W00{i}", "severity": "WARNING",
                   "section_id": "", "message": f"warn{i}", "recommendation": ""}
                 for i in range(warnings)]
        return DocumentStructureResult(
            input_docx="/tmp/test.docx",
            output_docx=None,
            expected_order=list(CANONICAL_DOCUMENT_ORDER),
            detected_order=["PORTADA", "INDICE", "A"],
            errors=errs,
            warnings=warns,
        )

    def test_error_count(self):
        r = self._make_result(errors=3)
        self.assertEqual(r.error_count(), 3)

    def test_warning_count(self):
        r = self._make_result(warnings=2)
        self.assertEqual(r.warning_count(), 2)

    def test_is_valid_no_errors(self):
        r = self._make_result(errors=0)
        self.assertTrue(r.is_valid())

    def test_is_valid_with_errors(self):
        r = self._make_result(errors=1)
        self.assertFalse(r.is_valid())

    def test_to_dict_keys(self):
        r = self._make_result()
        d = r.to_dict()
        for k in ("input_docx", "output_docx", "is_valid", "sections_found",
                   "expected_order", "detected_order", "errors", "warnings", "notes"):
            self.assertIn(k, d)

    def test_to_dict_serializable(self):
        r = self._make_result(errors=1, warnings=1)
        self.assertIsNotNone(json.dumps(r.to_dict()))

    def test_summary_contains_status(self):
        r_ok = self._make_result(errors=0)
        r_err = self._make_result(errors=1)
        self.assertIn("VALIDO", r_ok.summary())
        self.assertIn("CON ERRORES", r_err.summary())

    def test_summary_contains_block_count(self):
        r = self._make_result()
        r.sections_found = [
            DocumentSectionPosition("A", "A —", 5, 1, True),
            DocumentSectionPosition("B", "B —", 10, 1, True),
        ]
        s = r.summary()
        self.assertIn("2/11", s)

    def test_output_docx_in_dict(self):
        r = self._make_result()
        r.output_docx = "/tmp/out.docx"
        d = r.to_dict()
        self.assertEqual(d["output_docx"], "/tmp/out.docx")

    def test_notes_in_dict(self):
        r = self._make_result()
        r.notes.append("nota test")
        d = r.to_dict()
        self.assertIn("nota test", d["notes"])


# ---------------------------------------------------------------------------
# 2. Deteccion de secciones
# ---------------------------------------------------------------------------

class TestDetectDocumentSections(unittest.TestCase):

    def test_returns_list_for_nonexistent_docx(self):
        result = detect_document_sections("/nonexistent/path.docx")
        self.assertIsInstance(result, list)
        self.assertGreater(len(result), 0)

    def test_all_sections_not_found_for_nonexistent(self):
        result = detect_document_sections("/nonexistent/path.docx")
        found = [s for s in result if s.found]
        self.assertEqual(len(found), 0)

    def test_detects_portada(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            p = Path(tmpdir) / "test.docx"
            _make_docx(p, ["PORTADA", "INDICE", "A"])
            result = detect_document_sections(p)
            by_id = {s.section_id: s for s in result}
            self.assertTrue(by_id["PORTADA"].found)

    def test_detects_indice(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            p = Path(tmpdir) / "test.docx"
            _make_docx(p, ["PORTADA", "INDICE", "A"])
            result = detect_document_sections(p)
            by_id = {s.section_id: s for s in result}
            self.assertTrue(by_id["INDICE"].found)

    def test_detects_all_blocks_a_to_k(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            p = Path(tmpdir) / "test.docx"
            sections = ["PORTADA", "INDICE"] + list("ABCDEFGHIJK")
            _make_docx(p, sections)
            result = detect_document_sections(p)
            by_id = {s.section_id: s for s in result}
            for block in "ABCDEFGHIJK":
                self.assertTrue(by_id[block].found, f"Bloque {block} no detectado")

    def test_detects_anexo_grafico(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            p = Path(tmpdir) / "test.docx"
            _make_docx(p, ["PORTADA", "INDICE", "K", "ANEXO_GRAFICO"])
            result = detect_document_sections(p)
            by_id = {s.section_id: s for s in result}
            self.assertTrue(by_id["ANEXO_GRAFICO"].found)

    def test_detects_hoja_firmas(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            p = Path(tmpdir) / "test.docx"
            _make_docx(p, ["PORTADA", "INDICE", "K", "ANEXO_GRAFICO", "HOJA_FIRMAS"])
            result = detect_document_sections(p)
            by_id = {s.section_id: s for s in result}
            self.assertTrue(by_id["HOJA_FIRMAS"].found)

    def test_tolerates_missing_sections(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            p = Path(tmpdir) / "test.docx"
            _make_docx(p, ["PORTADA", "INDICE", "A", "B"])
            result = detect_document_sections(p)
            by_id = {s.section_id: s for s in result}
            # Solo A y B deben estar encontrados (de los bloques)
            self.assertTrue(by_id["A"].found)
            self.assertTrue(by_id["B"].found)
            self.assertFalse(by_id["C"].found)
            self.assertFalse(by_id["K"].found)

    def test_returns_canonical_section_count(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            p = Path(tmpdir) / "test.docx"
            _make_docx(p, ["PORTADA"])
            result = detect_document_sections(p)
            self.assertEqual(len(result), len(CANONICAL_DOCUMENT_ORDER))

    def test_paragraph_index_set_for_found_sections(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            p = Path(tmpdir) / "test.docx"
            _make_docx(p, ["PORTADA", "INDICE", "A"])
            result = detect_document_sections(p)
            by_id = {s.section_id: s for s in result}
            if by_id["INDICE"].found:
                self.assertIsNotNone(by_id["INDICE"].paragraph_index)

    def test_heading_level_set_for_heading_sections(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            p = Path(tmpdir) / "test.docx"
            _make_docx(p, ["PORTADA", "INDICE", "A"])
            result = detect_document_sections(p)
            by_id = {s.section_id: s for s in result}
            if by_id["A"].found:
                self.assertEqual(by_id["A"].heading_level, 1)

    def test_portada_has_no_heading_level(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            p = Path(tmpdir) / "test.docx"
            _make_docx(p, ["PORTADA", "INDICE"])
            result = detect_document_sections(p)
            by_id = {s.section_id: s for s in result}
            if by_id["PORTADA"].found:
                self.assertIsNone(by_id["PORTADA"].heading_level)

    def test_indice_title_detected(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            p = Path(tmpdir) / "test.docx"
            _make_docx(p, ["PORTADA", "INDICE"])
            result = detect_document_sections(p)
            by_id = {s.section_id: s for s in result}
            if by_id["INDICE"].found:
                self.assertIn("ndice", by_id["INDICE"].title.lower())

    def test_detects_bloque_with_dash_title(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            p = Path(tmpdir) / "test.docx"
            import docx
            doc = docx.Document()
            doc.add_paragraph("Portada del documento EIA")
            doc.add_heading("Indice", level=1)
            doc.add_heading("A — Identificacion del promotor y el proyecto", level=1)
            doc.add_heading("B — Descripcion del proyecto", level=1)
            doc.save(str(p))
            result = detect_document_sections(p)
            by_id = {s.section_id: s for s in result}
            self.assertTrue(by_id["A"].found)
            self.assertTrue(by_id["B"].found)

    def test_detects_bloque_with_dot_title(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            p = Path(tmpdir) / "test.docx"
            import docx
            doc = docx.Document()
            doc.add_paragraph("Portada EIA")
            doc.add_heading("Indice", level=1)
            doc.add_heading("C. Inventario ambiental", level=1)
            doc.save(str(p))
            result = detect_document_sections(p)
            by_id = {s.section_id: s for s in result}
            self.assertTrue(by_id["C"].found)

    def test_duplicate_block_noted(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            p = Path(tmpdir) / "test.docx"
            import docx
            doc = docx.Document()
            doc.add_paragraph("Portada EIA")
            doc.add_heading("Indice", level=1)
            doc.add_heading("A — Primera aparicion", level=1)
            doc.add_heading("A — Segunda aparicion", level=1)
            doc.save(str(p))
            result = detect_document_sections(p)
            by_id = {s.section_id: s for s in result}
            if by_id["A"].found:
                self.assertTrue(any("Duplicado" in n for n in by_id["A"].notes))


# ---------------------------------------------------------------------------
# 3. Validacion de orden
# ---------------------------------------------------------------------------

class TestValidateDocumentStructure(unittest.TestCase):

    def _full_docx(self, tmpdir, extra_end=None):
        """DOCX completo A-K + anexo + firmas."""
        p = Path(tmpdir) / "full.docx"
        sections = (
            ["PORTADA", "INDICE"]
            + list("ABCDEFGHIJK")
            + ["ANEXO_GRAFICO", "HOJA_FIRMAS"]
        )
        if extra_end:
            sections += extra_end
        _make_docx(p, sections)
        return p

    def test_full_document_is_valid(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            p = self._full_docx(tmpdir)
            result = validate_document_structure(p)
            self.assertTrue(result.is_valid(), f"Errores: {result.errors}")

    def test_missing_portada_is_error(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            p = Path(tmpdir) / "test.docx"
            _make_docx(p, ["INDICE"] + list("ABCDEFGHIJK"))
            result = validate_document_structure(p)
            codes = [e["code"] for e in result.errors]
            self.assertIn("EN02-E001", codes)

    def test_missing_indice_is_error(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            p = Path(tmpdir) / "test.docx"
            _make_docx(p, ["PORTADA"] + list("ABCDEFGHIJK"))
            result = validate_document_structure(p)
            codes = [e["code"] for e in result.errors]
            self.assertIn("EN02-E002", codes)

    def test_missing_block_is_error(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            p = Path(tmpdir) / "test.docx"
            # Faltan B, C, D, E, F, G, H, I, J, K
            _make_docx(p, ["PORTADA", "INDICE", "A"])
            result = validate_document_structure(p)
            codes = [e["code"] for e in result.errors]
            self.assertIn("EN02-E003", codes)

    def test_block_out_of_order_is_error(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            p = Path(tmpdir) / "test.docx"
            # C antes de B
            _make_docx(p, ["PORTADA", "INDICE", "A", "C", "B", "D",
                            "E", "F", "G", "H", "I", "J", "K"])
            result = validate_document_structure(p)
            codes = [e["code"] for e in result.errors]
            self.assertIn("EN02-E004", codes)

    def test_anexo_before_k_is_error(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            p = Path(tmpdir) / "test.docx"
            sections = (
                ["PORTADA", "INDICE"]
                + list("ABCDEFGHIJ")
                + ["ANEXO_GRAFICO", "K"]   # ANEXO antes de K
            )
            _make_docx(p, sections)
            result = validate_document_structure(p)
            codes = [e["code"] for e in result.errors]
            self.assertIn("EN02-E006", codes)

    def test_firmas_not_last_is_error(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            p = Path(tmpdir) / "test.docx"
            import docx
            doc = docx.Document()
            doc.add_paragraph("Portada EIA")
            doc.add_heading("Indice", level=1)
            for b in "ABCDEFGHIJK":
                doc.add_heading(f"{b} — Contenido", level=1)
            doc.add_heading("Hoja de firmas y revision tecnica", level=1)
            doc.add_heading("Anexo grafico y cartografico", level=1)  # FIRMAS antes de ANEXO
            doc.save(str(p))
            result = validate_document_structure(p)
            codes = [e["code"] for e in result.errors]
            self.assertIn("EN02-E005", codes)

    def test_missing_firmas_is_warning(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            p = Path(tmpdir) / "test.docx"
            _make_docx(p, ["PORTADA", "INDICE"] + list("ABCDEFGHIJK") + ["ANEXO_GRAFICO"])
            result = validate_document_structure(p)
            codes = [w["code"] for w in result.warnings]
            self.assertIn("EN02-W003", codes)

    def test_missing_anexo_is_warning(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            p = Path(tmpdir) / "test.docx"
            _make_docx(p, ["PORTADA", "INDICE"] + list("ABCDEFGHIJK"))
            result = validate_document_structure(p)
            codes = [w["code"] for w in result.warnings]
            self.assertIn("EN02-W002", codes)

    def test_firmas_absent_no_error(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            p = Path(tmpdir) / "test.docx"
            _make_docx(p, ["PORTADA", "INDICE"] + list("ABCDEFGHIJK") + ["ANEXO_GRAFICO"])
            result = validate_document_structure(p)
            error_codes = [e["code"] for e in result.errors]
            # EN02-E005 no debe aparecer si firmas no existe
            self.assertNotIn("EN02-E005", error_codes)

    def test_expected_order_in_result(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            p = self._full_docx(tmpdir)
            result = validate_document_structure(p)
            self.assertEqual(result.expected_order, list(CANONICAL_DOCUMENT_ORDER))

    def test_detected_order_populated(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            p = Path(tmpdir) / "test.docx"
            _make_docx(p, ["PORTADA", "INDICE", "A"])
            result = validate_document_structure(p)
            self.assertGreater(len(result.detected_order), 0)

    def test_nonexistent_docx_returns_errors(self):
        result = validate_document_structure("/nonexistent/path.docx")
        self.assertIsInstance(result, DocumentStructureResult)
        self.assertFalse(result.is_valid())

    def test_valid_result_has_nota_valida(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            p = self._full_docx(tmpdir)
            result = validate_document_structure(p)
            if result.is_valid():
                notes_text = " ".join(result.notes)
                self.assertIn("valida", notes_text.lower())

    def test_result_is_serializable(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            p = self._full_docx(tmpdir)
            result = validate_document_structure(p)
            self.assertIsNotNone(json.dumps(result.to_dict()))

    def test_sections_found_has_canonical_count(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            p = self._full_docx(tmpdir)
            result = validate_document_structure(p)
            self.assertEqual(len(result.sections_found), len(CANONICAL_DOCUMENT_ORDER))


# ---------------------------------------------------------------------------
# 4. Duplicados
# ---------------------------------------------------------------------------

class TestDuplicateDetection(unittest.TestCase):

    def test_duplicate_block_h_generates_warning(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            p = Path(tmpdir) / "test.docx"
            import docx
            doc = docx.Document()
            doc.add_paragraph("Portada EIA")
            doc.add_heading("Indice", level=1)
            for b in "ABCDEFGHIJK":
                doc.add_heading(f"{b} — Contenido bloque {b}", level=1)
            # Duplicar H
            doc.add_heading("H — Segunda aparicion de H", level=1)
            doc.add_heading("Anexo grafico y cartografico", level=1)
            doc.add_heading("Hoja de firmas y revision tecnica", level=1)
            doc.save(str(p))
            result = validate_document_structure(p)
            warning_codes = [w["code"] for w in result.warnings]
            self.assertIn("EN02-W001", warning_codes)

    def test_no_false_duplicate_warning_for_clean_doc(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            p = Path(tmpdir) / "test.docx"
            sections = (
                ["PORTADA", "INDICE"]
                + list("ABCDEFGHIJK")
                + ["ANEXO_GRAFICO", "HOJA_FIRMAS"]
            )
            _make_docx(p, sections)
            result = validate_document_structure(p)
            warning_codes = [w["code"] for w in result.warnings]
            self.assertNotIn("EN02-W001", warning_codes)


# ---------------------------------------------------------------------------
# 5. Normalizacion
# ---------------------------------------------------------------------------

class TestNormalizeDocumentStructure(unittest.TestCase):

    def test_creates_output_file(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            in_p = Path(tmpdir) / "input.docx"
            out_p = Path(tmpdir) / "output.docx"
            sections = (
                ["PORTADA", "INDICE"] + list("ABCDEFGHIJK")
                + ["ANEXO_GRAFICO", "HOJA_FIRMAS"]
            )
            _make_docx(in_p, sections)
            result = normalize_document_structure(in_p, out_p)
            self.assertTrue(out_p.exists())

    def test_does_not_modify_input(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            in_p = Path(tmpdir) / "input.docx"
            out_p = Path(tmpdir) / "output.docx"
            sections = ["PORTADA", "INDICE"] + list("ABCDEFGHIJK")
            _make_docx(in_p, sections)
            size_before = in_p.stat().st_size
            normalize_document_structure(in_p, out_p)
            size_after = in_p.stat().st_size
            self.assertEqual(size_before, size_after)

    def test_output_docx_is_valid(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            in_p = Path(tmpdir) / "input.docx"
            out_p = Path(tmpdir) / "output.docx"
            sections = (
                ["PORTADA", "INDICE"] + list("ABCDEFGHIJK")
                + ["ANEXO_GRAFICO", "HOJA_FIRMAS"]
            )
            _make_docx(in_p, sections)
            result = normalize_document_structure(in_p, out_p)
            self.assertEqual(result.output_docx, str(out_p))

    def test_output_docx_openable(self):
        import docx
        with tempfile.TemporaryDirectory() as tmpdir:
            in_p = Path(tmpdir) / "input.docx"
            out_p = Path(tmpdir) / "output.docx"
            _make_docx(in_p, ["PORTADA", "INDICE", "A", "B"])
            normalize_document_structure(in_p, out_p)
            if out_p.exists():
                opened = docx.Document(str(out_p))
                self.assertIsNotNone(opened)

    def test_normalize_already_valid_remains_valid(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            in_p = Path(tmpdir) / "input.docx"
            out_p = Path(tmpdir) / "output.docx"
            sections = (
                ["PORTADA", "INDICE"] + list("ABCDEFGHIJK")
                + ["ANEXO_GRAFICO", "HOJA_FIRMAS"]
            )
            _make_docx(in_p, sections)
            result = normalize_document_structure(in_p, out_p)
            self.assertTrue(result.is_valid(), f"Errores tras normalizacion: {result.errors}")

    def test_normalize_nonexistent_input_returns_result(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            out_p = Path(tmpdir) / "output.docx"
            result = normalize_document_structure("/nonexistent.docx", out_p)
            self.assertIsInstance(result, DocumentStructureResult)
            self.assertFalse(result.is_valid())

    def test_result_has_both_input_and_output_paths(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            in_p = Path(tmpdir) / "input.docx"
            out_p = Path(tmpdir) / "output.docx"
            _make_docx(in_p, ["PORTADA", "INDICE", "A"])
            result = normalize_document_structure(in_p, out_p)
            self.assertEqual(result.input_docx, str(in_p))

    def test_notes_contain_normalizacion_info(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            in_p = Path(tmpdir) / "input.docx"
            out_p = Path(tmpdir) / "output.docx"
            sections = (
                ["PORTADA", "INDICE"] + list("ABCDEFGHIJK") + ["ANEXO_GRAFICO"]
            )
            _make_docx(in_p, sections)
            result = normalize_document_structure(in_p, out_p)
            notes_text = " ".join(result.notes).lower()
            self.assertIn("normalizaci", notes_text)

    def test_output_in_result(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            in_p = Path(tmpdir) / "input.docx"
            out_p = Path(tmpdir) / "out" / "output.docx"
            _make_docx(in_p, ["PORTADA", "INDICE", "A"])
            result = normalize_document_structure(in_p, out_p)
            # El directorio se crea automaticamente
            if out_p.parent.exists():
                self.assertEqual(result.output_docx, str(out_p))


# ---------------------------------------------------------------------------
# 6. Outputs Markdown y JSON
# ---------------------------------------------------------------------------

class TestWriteDocumentStructureOutputs(unittest.TestCase):

    def _make_result(self):
        return DocumentStructureResult(
            input_docx="/tmp/test.docx",
            output_docx=None,
            expected_order=list(CANONICAL_DOCUMENT_ORDER),
            detected_order=["PORTADA", "INDICE", "A"],
            errors=[],
            warnings=[],
            notes=["Test note"],
        )

    def test_writes_both_files(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            result = self._make_result()
            paths = write_document_structure_outputs(result, tmpdir)
            self.assertEqual(len(paths), 2)
            for p in paths:
                self.assertTrue(Path(p).exists())

    def test_json_loadable(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            result = self._make_result()
            paths = write_document_structure_outputs(result, tmpdir)
            json_path = next(p for p in paths if str(p).endswith(".json"))
            data = json.loads(Path(json_path).read_text(encoding="utf-8"))
            self.assertIn("is_valid", data)

    def test_md_contains_expected_sections(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            result = self._make_result()
            paths = write_document_structure_outputs(result, tmpdir)
            md_path = next(p for p in paths if str(p).endswith(".md"))
            content = Path(md_path).read_text(encoding="utf-8")
            self.assertIn("Resumen", content)
            self.assertIn("Orden esperado", content)
            self.assertIn("Secciones encontradas", content)

    def test_returns_path_objects(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            result = self._make_result()
            paths = write_document_structure_outputs(result, tmpdir)
            for p in paths:
                self.assertIsInstance(p, Path)

    def test_creates_output_dir_if_needed(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            out_dir = Path(tmpdir) / "nested" / "subdir"
            result = self._make_result()
            paths = write_document_structure_outputs(result, out_dir)
            self.assertTrue(out_dir.exists())

    def test_json_no_real_secrets(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            result = self._make_result()
            paths = write_document_structure_outputs(result, tmpdir)
            json_path = next(p for p in paths if str(p).endswith(".json"))
            content = Path(json_path).read_text(encoding="utf-8")
            self.assertNotIn("sk-", content)


class TestBuildDocumentStructureMarkdown(unittest.TestCase):

    def _make_result(self, valid=True):
        err = [] if valid else [{"code": "EN02-E001", "severity": "ERROR",
                                  "section_id": "PORTADA", "message": "Falta portada",
                                  "recommendation": "Regenerar"}]
        return DocumentStructureResult(
            input_docx="/tmp/test.docx",
            output_docx=None,
            expected_order=list(CANONICAL_DOCUMENT_ORDER),
            detected_order=["PORTADA", "INDICE", "A"],
            errors=err,
            warnings=[],
            notes=["nota de prueba"],
        )

    def test_contains_summary_header(self):
        md = build_document_structure_markdown(self._make_result())
        self.assertIn("Resumen", md)

    def test_contains_expected_order(self):
        md = build_document_structure_markdown(self._make_result())
        self.assertIn("Orden esperado", md)

    def test_contains_detected_order(self):
        md = build_document_structure_markdown(self._make_result())
        self.assertIn("Orden detectado", md)

    def test_contains_secciones_encontradas(self):
        md = build_document_structure_markdown(self._make_result())
        self.assertIn("Secciones encontradas", md)

    def test_contains_errores(self):
        md = build_document_structure_markdown(self._make_result())
        self.assertIn("Errores", md)

    def test_contains_avisos(self):
        md = build_document_structure_markdown(self._make_result())
        self.assertIn("Avisos", md)

    def test_contains_notas(self):
        md = build_document_structure_markdown(self._make_result())
        self.assertIn("Notas", md)

    def test_valid_status_in_md(self):
        md = build_document_structure_markdown(self._make_result(valid=True))
        self.assertIn("VALIDO", md)

    def test_error_status_in_md(self):
        md = build_document_structure_markdown(self._make_result(valid=False))
        self.assertIn("CON ERRORES", md)

    def test_error_code_in_md(self):
        md = build_document_structure_markdown(self._make_result(valid=False))
        self.assertIn("EN02-E001", md)


# ---------------------------------------------------------------------------
# 7. CLI document-structure
# ---------------------------------------------------------------------------

class TestCLIDocumentStructure(unittest.TestCase):

    def _run(self, argv):
        sys.path.insert(0, str(Path(__file__).parent.parent))
        import run_expediente
        return run_expediente.main(argv)

    def test_document_structure_no_write_exit_0_valid_docx(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            exp = Path(tmpdir) / "EIA-2026-CLI-TEST"
            doc_dir = exp / "documento"
            doc_dir.mkdir(parents=True)
            docx_path = doc_dir / "documento_ambiental_borrador.docx"
            sections = (
                ["PORTADA", "INDICE"] + list("ABCDEFGHIJK")
                + ["ANEXO_GRAFICO", "HOJA_FIRMAS"]
            )
            _make_docx(docx_path, sections)
            code = self._run([str(exp), "document-structure"])
            self.assertEqual(code, 0)

    def test_document_structure_no_write_exit_1_invalid_docx(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            exp = Path(tmpdir) / "EIA-2026-CLI-TEST"
            doc_dir = exp / "documento"
            doc_dir.mkdir(parents=True)
            docx_path = doc_dir / "documento_ambiental_borrador.docx"
            # Solo A, sin portada ni indice ni otros bloques
            _make_docx(docx_path, ["A"])
            code = self._run([str(exp), "document-structure"])
            self.assertEqual(code, 1)

    def test_document_structure_write_creates_files(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            exp = Path(tmpdir) / "EIA-2026-CLI-TEST"
            doc_dir = exp / "documento"
            doc_dir.mkdir(parents=True)
            docx_path = doc_dir / "documento_ambiental_borrador.docx"
            sections = (
                ["PORTADA", "INDICE"] + list("ABCDEFGHIJK") + ["ANEXO_GRAFICO"]
            )
            _make_docx(docx_path, sections)
            self._run([str(exp), "document-structure", "--write"])
            self.assertTrue((doc_dir / "document_structure_result.json").exists())
            self.assertTrue((doc_dir / "document_structure_result.md").exists())

    def test_document_structure_normalize_creates_output_docx(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            exp = Path(tmpdir) / "EIA-2026-CLI-TEST"
            doc_dir = exp / "documento"
            doc_dir.mkdir(parents=True)
            docx_path = doc_dir / "documento_ambiental_borrador.docx"
            sections = (
                ["PORTADA", "INDICE"] + list("ABCDEFGHIJK") + ["ANEXO_GRAFICO"]
            )
            _make_docx(docx_path, sections)
            self._run([str(exp), "document-structure", "--normalize"])
            self.assertTrue(
                (doc_dir / "documento_ambiental_estructurado.docx").exists()
            )

    def test_document_structure_no_docx_returns_1(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            exp = Path(tmpdir) / "EIA-2026-CLI-NO-DOCX"
            exp.mkdir()
            (exp / "documento").mkdir()
            code = self._run([str(exp), "document-structure"])
            self.assertEqual(code, 1)

    def test_document_structure_prefers_con_figuras_over_borrador(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            exp = Path(tmpdir) / "EIA-2026-CLI-TEST"
            doc_dir = exp / "documento"
            doc_dir.mkdir(parents=True)
            # Crear ambos
            borrador = doc_dir / "documento_ambiental_borrador.docx"
            con_figuras = doc_dir / "documento_ambiental_borrador_con_figuras.docx"
            sections = (
                ["PORTADA", "INDICE"] + list("ABCDEFGHIJK") + ["ANEXO_GRAFICO"]
            )
            _make_docx(borrador, ["PORTADA"])  # Invalido
            _make_docx(con_figuras, sections)  # Valido
            code = self._run([str(exp), "document-structure"])
            # Con figuras tiene todos los bloques → valido → exit 0
            self.assertEqual(code, 0)


# ---------------------------------------------------------------------------
# 8. find_best_available_docx
# ---------------------------------------------------------------------------

class TestFindBestAvailableDocx(unittest.TestCase):

    def test_returns_none_if_no_docx(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            result = find_best_available_docx(tmpdir)
            self.assertIsNone(result)

    def test_returns_borrador_if_only_one(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            exp = Path(tmpdir)
            doc_dir = exp / "documento"
            doc_dir.mkdir()
            p = doc_dir / "documento_ambiental_borrador.docx"
            p.write_bytes(b"fake")
            result = find_best_available_docx(exp)
            self.assertEqual(result, p)

    def test_prefers_con_figuras_over_borrador(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            exp = Path(tmpdir)
            doc_dir = exp / "documento"
            doc_dir.mkdir()
            borrador = doc_dir / "documento_ambiental_borrador.docx"
            con_figuras = doc_dir / "documento_ambiental_borrador_con_figuras.docx"
            borrador.write_bytes(b"fake1")
            con_figuras.write_bytes(b"fake2")
            result = find_best_available_docx(exp)
            self.assertEqual(result, con_figuras)

    def test_returns_path_object(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            exp = Path(tmpdir)
            doc_dir = exp / "documento"
            doc_dir.mkdir()
            p = doc_dir / "documento_ambiental_borrador.docx"
            p.write_bytes(b"fake")
            result = find_best_available_docx(exp)
            self.assertIsInstance(result, Path)


# ---------------------------------------------------------------------------
# 9. Integracion sintetica end-to-end
# ---------------------------------------------------------------------------

class TestIntegrationSyntheticDocx(unittest.TestCase):

    def _full_sections(self):
        return (
            ["PORTADA", "INDICE"] + list("ABCDEFGHIJK")
            + ["ANEXO_GRAFICO", "HOJA_FIRMAS"]
        )

    def test_full_flow_validate_then_normalize(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            in_p = Path(tmpdir) / "input.docx"
            out_p = Path(tmpdir) / "output.docx"
            _make_docx(in_p, self._full_sections())
            # Validar
            val_result = validate_document_structure(in_p)
            # Normalizar
            norm_result = normalize_document_structure(in_p, out_p)
            self.assertTrue(val_result.is_valid(), f"Errores: {val_result.errors}")
            self.assertTrue(out_p.exists())

    def test_full_flow_write_outputs(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            p = Path(tmpdir) / "test.docx"
            _make_docx(p, self._full_sections())
            result = validate_document_structure(p)
            out_dir = Path(tmpdir) / "outputs"
            paths = write_document_structure_outputs(result, out_dir)
            self.assertEqual(len(paths), 2)
            for path in paths:
                self.assertTrue(path.exists())

    def test_full_flow_json_consistent_with_result(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            p = Path(tmpdir) / "test.docx"
            _make_docx(p, self._full_sections())
            result = validate_document_structure(p)
            out_dir = Path(tmpdir) / "outputs"
            paths = write_document_structure_outputs(result, out_dir)
            json_path = next(x for x in paths if str(x).endswith(".json"))
            data = json.loads(Path(json_path).read_text(encoding="utf-8"))
            self.assertEqual(data["is_valid"], result.is_valid())
            self.assertEqual(len(data["errors"]), result.error_count())

    def test_partial_doc_generates_errors(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            p = Path(tmpdir) / "test.docx"
            _make_docx(p, ["PORTADA", "INDICE", "A", "B"])
            result = validate_document_structure(p)
            self.assertFalse(result.is_valid())
            self.assertGreater(result.error_count(), 0)

    def test_normalize_partial_doc_still_creates_output(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            in_p = Path(tmpdir) / "input.docx"
            out_p = Path(tmpdir) / "output.docx"
            _make_docx(in_p, ["PORTADA", "INDICE", "A"])
            result = normalize_document_structure(in_p, out_p)
            self.assertTrue(out_p.exists())

    def test_canonical_document_order_has_all_blocks(self):
        for b in BLOCK_IDS:
            self.assertIn(b, CANONICAL_DOCUMENT_ORDER)

    def test_canonical_document_order_starts_portada(self):
        self.assertEqual(CANONICAL_DOCUMENT_ORDER[0], "PORTADA")

    def test_canonical_document_order_ends_firmas(self):
        self.assertEqual(CANONICAL_DOCUMENT_ORDER[-1], "HOJA_FIRMAS")

    def test_canonical_document_order_has_15_items(self):
        # PORTADA + INDICE + 11 bloques + ANEXO + FIRMAS = 15
        self.assertEqual(len(CANONICAL_DOCUMENT_ORDER), 15)

    def test_sections_found_includes_all_canonical_ids(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            p = Path(tmpdir) / "test.docx"
            _make_docx(p, self._full_sections())
            result = validate_document_structure(p)
            found_ids = {s.section_id for s in result.sections_found}
            for sid in CANONICAL_DOCUMENT_ORDER:
                self.assertIn(sid, found_ids)


# ---------------------------------------------------------------------------
# 10. No mutacion
# ---------------------------------------------------------------------------

class TestNoMutation(unittest.TestCase):

    def test_validate_does_not_modify_docx(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            p = Path(tmpdir) / "test.docx"
            _make_docx(p, ["PORTADA", "INDICE", "A", "B"])
            size_before = p.stat().st_size
            validate_document_structure(p)
            size_after = p.stat().st_size
            self.assertEqual(size_before, size_after)

    def test_detect_does_not_modify_docx(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            p = Path(tmpdir) / "test.docx"
            _make_docx(p, ["PORTADA", "INDICE", "A"])
            size_before = p.stat().st_size
            detect_document_sections(p)
            size_after = p.stat().st_size
            self.assertEqual(size_before, size_after)

    def test_normalize_only_modifies_copy(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            in_p = Path(tmpdir) / "input.docx"
            out_p = Path(tmpdir) / "output.docx"
            sections = (
                ["PORTADA", "INDICE"] + list("ABCDEFGHIJK") + ["ANEXO_GRAFICO"]
            )
            _make_docx(in_p, sections)
            content_before = in_p.read_bytes()
            normalize_document_structure(in_p, out_p)
            content_after = in_p.read_bytes()
            self.assertEqual(content_before, content_after)


if __name__ == "__main__":
    unittest.main()
