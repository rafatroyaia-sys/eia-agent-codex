"""
document_docx_builder -- DOC-02
Generador determinista del Documento Ambiental en formato DOCX.

Convierte el borrador Markdown generado por DOC-01 en un documento DOCX
profesional con portada, indice, estilos y tablas.

Principios no negociables:
  - No usa IA.
  - No consulta fuentes externas.
  - No genera Markdown (eso es DOC-01).
  - No inventa datos.
  - No corrige outputs tecnicos del pipeline.
  - No modifica el Markdown fuente.
  - No declara aptitud administrativa.
  - No modifica impactos, medidas, PVA ni auditorias.
  - No genera PDF.
"""
from __future__ import annotations

import json
import re
from dataclasses import dataclass, field
from pathlib import Path

# ---------------------------------------------------------------------------
# Constantes publicas
# ---------------------------------------------------------------------------

DOCX_OUTPUT_FILENAME = "documento_ambiental_borrador.docx"
DOCX_BUILD_RESULT_FILENAME = "docx_build_result.json"

DEFAULT_TITLE = "Documento Ambiental"
DEFAULT_SUBTITLE = "Borrador tecnico generado automaticamente"
ADMIN_STYLE_PROFILE = "recimetal_admin_reference_v1"

SUPPORTED_MARKDOWN_ELEMENTS: list[str] = [
    "heading",
    "paragraph",
    "bullet_list",
    "numbered_list",
    "table",
    "blockquote",
    "horizontal_rule",
    "code_block",
    "blank_line",
]

_ADMIN_DISCLAIMER = (
    "Documento generado automaticamente a partir de outputs tecnicos. "
    "Requiere revision tecnica/juridica. "
    "No declara aptitud administrativa."
)

_ADMIN_STYLE_NOTE = (
    "Perfil formal basado en referencia administrativa RECIMETAL: portada sobria, "
    "Calibri Light 12, jerarquia compacta, indice y tablas con cuadricula."
)

# ---------------------------------------------------------------------------
# Dataclasses
# ---------------------------------------------------------------------------


@dataclass
class DocxBuildWarning:
    """Advertencia generada durante la construccion del DOCX."""

    code: str
    message: str
    source_line: "int | None" = None
    recommendation: str = ""

    def to_dict(self) -> dict:
        return {
            "code": self.code,
            "message": self.message,
            "source_line": self.source_line,
            "recommendation": self.recommendation,
        }

    def summary(self) -> str:
        loc = f" (linea {self.source_line})" if self.source_line is not None else ""
        return f"[{self.code}]{loc} {self.message}"


@dataclass
class DocumentDocxBuildResult:
    """Resultado de la generacion del DOCX."""

    expediente_id: str
    input_markdown_path: str
    output_docx_path: "str | None" = None
    generated: bool = False
    paragraph_count: int = 0
    heading_count: int = 0
    table_count: int = 0
    image_count: int = 0
    warnings: list[DocxBuildWarning] = field(default_factory=list)
    notes: list[str] = field(default_factory=list)

    def warning_count(self) -> int:
        return len(self.warnings)

    def is_success(self) -> bool:
        return self.generated

    def to_dict(self) -> dict:
        return {
            "expediente_id": self.expediente_id,
            "input_markdown_path": self.input_markdown_path,
            "output_docx_path": self.output_docx_path,
            "generated": self.generated,
            "paragraph_count": self.paragraph_count,
            "heading_count": self.heading_count,
            "table_count": self.table_count,
            "image_count": self.image_count,
            "warning_count": self.warning_count(),
            "warnings": [w.to_dict() for w in self.warnings],
            "notes": list(self.notes),
        }

    def summary(self) -> str:
        status = "OK" if self.generated else "DRY-RUN"
        return (
            f"DOC-02 [{self.expediente_id}] {status} — "
            f"{self.heading_count} headings, "
            f"{self.paragraph_count} parrafos, "
            f"{self.table_count} tablas"
            + (f", {self.warning_count()} avisos" if self.warnings else "")
        )


# ---------------------------------------------------------------------------
# Helpers de lectura
# ---------------------------------------------------------------------------


def safe_read_markdown(path: "str | Path") -> str:
    """Lee un archivo Markdown UTF-8.

    Raises:
        FileNotFoundError: si el archivo no existe.
        ValueError: si el archivo esta vacio.
    """
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(f"Archivo Markdown no encontrado: {p}")
    text = p.read_text(encoding="utf-8")
    if not text.strip():
        raise ValueError(f"El archivo Markdown esta vacio: {p}")
    return text


# ---------------------------------------------------------------------------
# Parser Markdown -> bloques
# ---------------------------------------------------------------------------


def _strip_markdown_inline(text: str) -> str:
    """Elimina marcado inline de Markdown (bold, italic, code, links)."""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"__(.+?)__", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"_(.+?)_", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    text = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", text)
    return text.strip()


def _parse_table_lines(lines: list[str]) -> dict:
    """Parsea lineas de tabla Markdown (| col | col |)."""
    if not lines:
        return {"type": "table", "headers": [], "rows": []}

    sep_indices: list[int] = []
    parsed: list[tuple[int, list[str]]] = []

    for idx, line in enumerate(lines):
        raw = line.strip().strip("|")
        cells = [c.strip() for c in raw.split("|")]
        is_sep = bool(cells) and all(
            re.match(r"^[-:]+$", c) for c in cells if c.strip()
        )
        if is_sep:
            sep_indices.append(idx)
        else:
            parsed.append((idx, cells))

    first_sep = min(sep_indices) if sep_indices else None
    headers: list[str] = []
    rows: list[list[str]] = []

    for orig_idx, cells in parsed:
        if first_sep is not None and orig_idx < first_sep:
            headers = cells
        else:
            rows.append(cells)

    return {"type": "table", "headers": headers, "rows": rows}


def parse_markdown_blocks(markdown: str) -> list[dict]:
    """Parsea Markdown a lista de bloques tipados.

    Tipos: heading, paragraph, bullet_list, numbered_list, table,
           blockquote, horizontal_rule, code_block, blank_line.
    Nunca lanza excepcion por contenido inesperado.
    """
    lines = markdown.split("\n")
    blocks: list[dict] = []
    i = 0
    total = len(lines)

    while i < total:
        line = lines[i]
        stripped = line.strip()

        # Bloque de codigo ```
        if stripped.startswith("```"):
            code_lines: list[str] = []
            i += 1
            while i < total and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            if i < total:
                i += 1  # cerrar ```
            blocks.append({"type": "code_block", "text": "\n".join(code_lines)})
            continue

        # Heading
        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            if level > 0 and len(stripped) > level and stripped[level] == " ":
                text = stripped[level:].strip()
                blocks.append({"type": "heading", "level": min(level, 4), "text": text})
                i += 1
                continue

        # Tabla markdown: empieza y termina con |
        if stripped.startswith("|") and stripped.endswith("|"):
            table_lines: list[str] = []
            while (
                i < total
                and lines[i].strip().startswith("|")
                and lines[i].strip().endswith("|")
            ):
                table_lines.append(lines[i].strip())
                i += 1
            blocks.append(_parse_table_lines(table_lines))
            continue

        # Regla horizontal: solo guiones/asteriscos/subrayados
        if stripped and all(c in "-*_ " for c in stripped):
            core = stripped.replace(" ", "")
            if len(core) >= 3 and len(set(core)) == 1:
                blocks.append({"type": "horizontal_rule"})
                i += 1
                continue

        # Blockquote: lineas consecutivas con >
        if stripped.startswith(">"):
            quote_lines: list[str] = []
            while i < total and (
                lines[i].strip().startswith(">") or
                (quote_lines and lines[i].strip() == "")
            ):
                s = lines[i].strip()
                if s.startswith("> "):
                    quote_lines.append(s[2:])
                elif s == ">":
                    quote_lines.append("")
                elif s == "" and quote_lines:
                    # linea en blanco dentro de blockquote — detener
                    break
                i += 1
            text = "\n".join(quote_lines).strip()
            blocks.append({"type": "blockquote", "text": text})
            continue

        # Bullet list: lineas consecutivas con - o *
        if re.match(r"^[-*]\s+", stripped):
            items: list[str] = []
            while i < total and re.match(r"^[-*]\s+", lines[i].strip()):
                s = lines[i].strip()
                items.append(re.sub(r"^[-*]\s+", "", s))
                i += 1
            blocks.append({"type": "bullet_list", "items": items})
            continue

        # Numbered list: lineas consecutivas con N.
        if re.match(r"^\d+\.\s+", stripped):
            items = []
            while i < total and re.match(r"^\d+\.\s+", lines[i].strip()):
                s = lines[i].strip()
                items.append(re.sub(r"^\d+\.\s+", "", s))
                i += 1
            blocks.append({"type": "numbered_list", "items": items})
            continue

        # Linea en blanco
        if not stripped:
            blocks.append({"type": "blank_line"})
            i += 1
            continue

        # Parrafo generico
        blocks.append({"type": "paragraph", "text": stripped})
        i += 1

    return blocks


# ---------------------------------------------------------------------------
# Creacion del documento DOCX
# ---------------------------------------------------------------------------


def create_docx_document(
    title: str = DEFAULT_TITLE,
    subtitle: str = DEFAULT_SUBTITLE,
) -> "docx.Document":
    """Crea un Document() de python-docx con estilos y margenes configurados."""
    from docx import Document
    from docx.enum.text import WD_LINE_SPACING
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor

    doc = Document()

    # Margenes inspirados en la referencia administrativa aportada.
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.3)
        section.header_distance = Cm(1.0)
        section.footer_distance = Cm(1.0)

    def _set_font(style_name: str, name: str, size_pt: float, bold=None, italic=None, color=None) -> None:
        try:
            style = doc.styles[style_name]
            style.font.name = name
            style.font.size = Pt(size_pt)
            if bold is not None:
                style.font.bold = bold
            if italic is not None:
                style.font.italic = italic
            if color is not None:
                style.font.color.rgb = RGBColor(*color)
            rpr = style.element.get_or_add_rPr()
            rfonts = rpr.rFonts
            if rfonts is not None:
                rfonts.set(qn("w:ascii"), name)
                rfonts.set(qn("w:hAnsi"), name)
                rfonts.set(qn("w:cs"), name)
        except Exception:
            pass

    def _set_para_spacing(style_name: str, before: float, after: float, line: float = 1.08) -> None:
        try:
            pf = doc.styles[style_name].paragraph_format
            pf.space_before = Pt(before)
            pf.space_after = Pt(after)
            pf.line_spacing = line
            pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        except Exception:
            pass

    # Estilo Normal
    _set_font("Normal", "Calibri Light", 12)
    _set_para_spacing("Normal", 0, 6)

    # Heading 1
    _set_font("Heading 1", "Calibri Light", 15, bold=True, color=(0x20, 0x20, 0x20))
    _set_para_spacing("Heading 1", 14, 6)

    # Heading 2
    _set_font("Heading 2", "Calibri Light", 13, bold=True, color=(0x25, 0x25, 0x25))
    _set_para_spacing("Heading 2", 10, 4)

    # Heading 3
    _set_font("Heading 3", "Calibri Light", 12, bold=True, color=(0x30, 0x30, 0x30))
    _set_para_spacing("Heading 3", 8, 3)

    # Heading 4
    _set_font("Heading 4", "Calibri Light", 11, bold=True, italic=True, color=(0x40, 0x40, 0x40))
    _set_para_spacing("Heading 4", 6, 2)

    _set_font("Caption", "Calibri Light", 9, bold=True, color=(0x40, 0x40, 0x40))
    _set_font("Header", "Calibri Light", 9, color=(0x70, 0x70, 0x70))
    _set_font("Footer", "Calibri Light", 9, color=(0x70, 0x70, 0x70))

    return doc


def add_cover_page(
    doc: "docx.Document",
    expediente_id: str,
    title: str,
    subtitle: str,
    generated_note: str,
    logo_path: "str | Path | None" = None,
) -> None:
    """Añade portada al documento DOCX."""
    import datetime

    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor

    # Logo opcional
    if logo_path is not None:
        logo = Path(logo_path)
        if logo.exists():
            try:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(str(logo), width=Inches(2.0))
                doc.add_paragraph()
            except Exception:
                pass

    # Titulo principal
    for _ in range(3):
        doc.add_paragraph()

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_after = Pt(2)
    title_run = title_para.add_run(title.upper())
    title_run.bold = True
    title_run.font.name = "Calibri Light"
    title_run.font.size = Pt(20)
    try:
        title_run.font.color.rgb = RGBColor(0x20, 0x20, 0x20)
    except Exception:
        pass

    # Subtitulo
    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_para.paragraph_format.space_after = Pt(18)
    sub_run = sub_para.add_run(subtitle.upper())
    sub_run.font.name = "Calibri Light"
    sub_run.font.size = Pt(15)
    sub_run.bold = True
    try:
        sub_run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
    except Exception:
        pass

    # Expediente ID
    exp_para = doc.add_paragraph()
    exp_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    exp_run = exp_para.add_run(f"Expediente: {expediente_id}")
    exp_run.font.name = "Calibri Light"
    exp_run.font.size = Pt(12)

    # Fecha de generacion
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_str = datetime.date.today().isoformat()
    date_run = date_para.add_run(f"Fecha de generacion: {date_str}")
    date_run.font.name = "Calibri Light"
    date_run.font.size = Pt(10)
    try:
        date_run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    except Exception:
        pass

    # Espacio
    for _ in range(2):
        doc.add_paragraph()

    # Advertencia de alcance
    disc_para = doc.add_paragraph()
    disc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    disc_run = disc_para.add_run(generated_note)
    disc_run.italic = True
    disc_run.font.name = "Calibri Light"
    disc_run.font.size = Pt(9)
    try:
        disc_run.font.color.rgb = RGBColor(0x80, 0x20, 0x20)
    except Exception:
        pass

    # Salto de pagina
    doc.add_page_break()


def add_running_header_footer(doc: "docx.Document", expediente_id: str) -> None:
    """Anade encabezado y pie discretos con paginacion de Word."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    def _page_field(paragraph) -> None:
        run = paragraph.add_run()
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = " PAGE "
        separate = OxmlElement("w:fldChar")
        separate.set(qn("w:fldCharType"), "separate")
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        run._r.append(begin)
        run._r.append(instr)
        run._r.append(separate)
        run._r.append(end)

    for section in doc.sections:
        header = section.header
        footer = section.footer
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.text = f"Documento Ambiental - {expediente_id}"
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.add_run("Pagina ")
        _page_field(fp)
        for para in [hp, fp]:
            for run in para.runs:
                run.font.name = "Calibri Light"
                run.font.size = Pt(9)
                try:
                    run.font.color.rgb = RGBColor(0x70, 0x70, 0x70)
                except Exception:
                    pass


def add_table_of_contents_placeholder(doc: "docx.Document") -> None:
    """Añade indice/TOC al documento. Inserta campo Word si es posible."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt

    doc.add_heading("Indice", level=1)

    p = doc.add_paragraph()
    run = p.add_run()

    try:
        fldChar_begin = OxmlElement("w:fldChar")
        fldChar_begin.set(qn("w:fldCharType"), "begin")

        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")
        instrText.text = r' TOC \o "1-3" \h \z \u '

        fldChar_separate = OxmlElement("w:fldChar")
        fldChar_separate.set(qn("w:fldCharType"), "separate")

        fldChar_end = OxmlElement("w:fldChar")
        fldChar_end.set(qn("w:fldCharType"), "end")

        run._r.append(fldChar_begin)
        run._r.append(instrText)
        run._r.append(fldChar_separate)
        run._r.append(fldChar_end)
    except Exception:
        # Fallback si la manipulacion XML falla
        p.clear()
        p.add_run("[Indice automatico no disponible]")

    note = doc.add_paragraph(
        "Nota: Para actualizar el indice, abrir en Word y presionar Ctrl+A y F9."
    )
    try:
        for run_obj in note.runs:
            run_obj.italic = True
            run_obj.font.size = Pt(9)
    except Exception:
        pass

    doc.add_page_break()


def add_markdown_block_to_docx(doc: "docx.Document", block: dict) -> dict:
    """Añade un bloque parseado al DOCX.

    Returns:
        Dict con conteos: paragraph_added, heading_added, table_added, image_added.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt

    counts = {
        "paragraph_added": 0,
        "heading_added": 0,
        "table_added": 0,
        "image_added": 0,
    }
    btype = block.get("type", "")

    def _style_table_cell(cell, is_header: bool = False) -> None:
        if is_header:
            try:
                tc_pr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:fill"), "EDEDED")
                tc_pr.append(shd)
            except Exception:
                pass
        try:
            tc_pr = cell._tc.get_or_add_tcPr()
            margins = OxmlElement("w:tcMar")
            for side in ["top", "left", "bottom", "right"]:
                node = OxmlElement(f"w:{side}")
                node.set(qn("w:w"), "90")
                node.set(qn("w:type"), "dxa")
                margins.append(node)
            tc_pr.append(margins)
        except Exception:
            pass
        for para in cell.paragraphs:
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            for r in para.runs:
                r.font.name = "Calibri Light"
                r.font.size = Pt(9)
                if is_header:
                    r.bold = True

    if btype == "heading":
        level = max(1, min(int(block.get("level", 1)), 4))
        text = _strip_markdown_inline(block.get("text", ""))
        if text:
            doc.add_heading(text, level=level)
            counts["heading_added"] = 1

    elif btype == "paragraph":
        text = _strip_markdown_inline(block.get("text", ""))
        if text:
            doc.add_paragraph(text)
            counts["paragraph_added"] = 1

    elif btype == "bullet_list":
        for item in block.get("items", []):
            text = _strip_markdown_inline(item)
            if text:
                try:
                    doc.add_paragraph(text, style="List Bullet")
                except Exception:
                    doc.add_paragraph(f"• {text}")
                counts["paragraph_added"] += 1

    elif btype == "numbered_list":
        for item in block.get("items", []):
            text = _strip_markdown_inline(item)
            if text:
                try:
                    doc.add_paragraph(text, style="List Number")
                except Exception:
                    doc.add_paragraph(text)
                counts["paragraph_added"] += 1

    elif btype == "blockquote":
        text = _strip_markdown_inline(block.get("text", ""))
        if text:
            p = doc.add_paragraph()
            try:
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.right_indent = Inches(0.25)
            except Exception:
                pass
            run = p.add_run(text)
            run.italic = True
            counts["paragraph_added"] = 1

    elif btype == "table":
        headers = block.get("headers", [])
        rows = block.get("rows", [])
        all_data = ([headers] if headers else []) + rows
        if not all_data:
            return counts

        n_cols = max(len(r) for r in all_data) if all_data else 1
        n_cols = max(n_cols, 1)
        n_rows = len(all_data)

        try:
            table = doc.add_table(rows=n_rows, cols=n_cols)
            try:
                table.style = "Table Grid"
            except Exception:
                pass

            for row_idx, data_row in enumerate(all_data):
                row = table.rows[row_idx]
                for col_idx, cell_text in enumerate(data_row[:n_cols]):
                    if col_idx < len(row.cells):
                        cell = row.cells[col_idx]
                        clean = _strip_markdown_inline(str(cell_text))
                        cell.text = clean
                        # Cabecera en negrita (primera fila si hay headers)
                        _style_table_cell(cell, is_header=bool(row_idx == 0 and headers))
            counts["table_added"] = 1
        except Exception:
            # Fallback: volcar tabla como parrafos
            for data_row in all_data:
                text = " | ".join(
                    _strip_markdown_inline(str(c)) for c in data_row
                )
                doc.add_paragraph(text)
                counts["paragraph_added"] += 1

    elif btype == "code_block":
        text = block.get("text", "")
        if text:
            p = doc.add_paragraph(text)
            try:
                for run in p.runs:
                    run.font.name = "Courier New"
                    run.font.size = Pt(9)
            except Exception:
                pass
            counts["paragraph_added"] = 1

    elif btype == "horizontal_rule":
        p = doc.add_paragraph()
        try:
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            pPr = p._element.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "auto")
            pBdr.append(bottom)
            pPr.append(pBdr)
        except Exception:
            p.add_run("─" * 40)
        counts["paragraph_added"] = 1

    # blank_line: no añadir nada al DOCX

    return counts


# ---------------------------------------------------------------------------
# Funciones principales de construccion
# ---------------------------------------------------------------------------


def build_docx_from_markdown(
    markdown_path: "str | Path",
    output_docx_path: "str | Path",
    expediente_id: str = "UNKNOWN",
    title: str = DEFAULT_TITLE,
    subtitle: str = DEFAULT_SUBTITLE,
    logo_path: "str | Path | None" = None,
) -> DocumentDocxBuildResult:
    """Genera un DOCX a partir de un archivo Markdown.

    Lee el Markdown, parsea los bloques, crea el documento con portada e
    indice, añade el contenido y guarda el DOCX.
    No modifica el Markdown fuente.
    """
    md_path = Path(markdown_path)
    out_path = Path(output_docx_path)

    markdown = safe_read_markdown(md_path)
    blocks = parse_markdown_blocks(markdown)

    warnings: list[DocxBuildWarning] = []

    doc = create_docx_document(title=title, subtitle=subtitle)
    add_running_header_footer(doc, expediente_id)

    add_cover_page(
        doc,
        expediente_id=expediente_id,
        title=title,
        subtitle=subtitle,
        generated_note=_ADMIN_DISCLAIMER,
        logo_path=logo_path,
    )
    add_table_of_contents_placeholder(doc)

    total_paragraphs = 0
    total_headings = 0
    total_tables = 0
    total_images = 0

    for block in blocks:
        try:
            c = add_markdown_block_to_docx(doc, block)
            total_paragraphs += c.get("paragraph_added", 0)
            total_headings += c.get("heading_added", 0)
            total_tables += c.get("table_added", 0)
            total_images += c.get("image_added", 0)
        except Exception as exc:
            warnings.append(
                DocxBuildWarning(
                    code="BLOCK_ERROR",
                    message=f"Error al procesar bloque {block.get('type','?')}: {exc}",
                    source_line=None,
                    recommendation="Revisar el contenido del bloque en el Markdown.",
                )
            )

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))

    return DocumentDocxBuildResult(
        expediente_id=expediente_id,
        input_markdown_path=str(md_path),
        output_docx_path=str(out_path),
        generated=True,
        paragraph_count=total_paragraphs,
        heading_count=total_headings,
        table_count=total_tables,
        image_count=total_images,
        warnings=warnings,
        notes=[
            "DOCX generado a partir de borrador Markdown DOC-01.",
            _ADMIN_STYLE_NOTE,
            "Requiere revision tecnica y juridica.",
            "No declara aptitud administrativa.",
        ],
    )


def build_docx_from_expediente(
    expediente_path: "str | Path",
    write_outputs: bool = False,
) -> DocumentDocxBuildResult:
    """Genera DOCX desde el borrador Markdown de un expediente.

    Sin write_outputs=True: solo parsea y devuelve resultado con generated=False.
    Con write_outputs=True: genera el DOCX y el JSON de resultado.
    No genera el Markdown si falta; eso es responsabilidad de DOC-01.
    """
    exp = Path(expediente_path)
    md_path = exp / "documento" / "documento_ambiental_borrador.md"

    if not md_path.exists():
        raise FileNotFoundError(
            f"No se encontro el Markdown DOC-01 en: {md_path}. "
            "Ejecute primero 'document-build-md --write'."
        )

    # Leer y parsear (siempre, incluso sin write)
    markdown = safe_read_markdown(md_path)
    blocks = parse_markdown_blocks(markdown)

    para_count = sum(
        1
        for b in blocks
        if b.get("type") in (
            "paragraph", "bullet_list", "numbered_list",
            "blockquote", "code_block",
        )
    )
    heading_count = sum(1 for b in blocks if b.get("type") == "heading")
    table_count = sum(1 for b in blocks if b.get("type") == "table")

    if not write_outputs:
        return DocumentDocxBuildResult(
            expediente_id=exp.name,
            input_markdown_path=str(md_path),
            output_docx_path=None,
            generated=False,
            paragraph_count=para_count,
            heading_count=heading_count,
            table_count=table_count,
            image_count=0,
            warnings=[],
            notes=[
                "write_outputs=False: DOCX no generado.",
                "Pase --write para generar el archivo.",
            ],
        )

    doc_dir = exp / "documento"
    doc_dir.mkdir(parents=True, exist_ok=True)
    docx_path = doc_dir / DOCX_OUTPUT_FILENAME
    json_path = doc_dir / DOCX_BUILD_RESULT_FILENAME

    # Logo (busqueda best-effort relativa a la ubicacion del modulo)
    logo_path: "Path | None" = None
    try:
        module_dir = Path(__file__).parent
        candidate = module_dir.parent.parent.parent / "assets" / "brand" / "logo_ecogestion.png"
        if candidate.exists():
            logo_path = candidate
    except Exception:
        pass

    result = build_docx_from_markdown(
        markdown_path=md_path,
        output_docx_path=docx_path,
        expediente_id=exp.name,
        logo_path=logo_path,
    )

    write_docx_build_result(result, json_path)

    return result


# ---------------------------------------------------------------------------
# Escritura de resultado JSON
# ---------------------------------------------------------------------------


def write_docx_build_result(
    result: DocumentDocxBuildResult,
    output_path: "str | Path",
) -> Path:
    """Escribe el resultado de construccion como JSON UTF-8 indentado."""
    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    out.write_text(
        json.dumps(result.to_dict(), ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    return out


# ---------------------------------------------------------------------------
# Validacion basica del DOCX generado
# ---------------------------------------------------------------------------


def validate_docx_basic(path: "str | Path") -> bool:
    """Validacion basica del archivo DOCX: existe, .docx, size>0, abre sin error."""
    try:
        p = Path(path)
        if not p.exists():
            return False
        if p.suffix.lower() != ".docx":
            return False
        if p.stat().st_size == 0:
            return False
        from docx import Document as _DocxDocument
        _DocxDocument(str(p))
        return True
    except Exception:
        return False
