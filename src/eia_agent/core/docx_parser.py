"""
docx_parser -- IN-01
Extractor de texto, tablas y metadatos de archivos .docx del promotor.

Solo lectura. No usa IA. No escribe en disco.

Uso:
    from eia_agent.core.docx_parser import parse_docx, extract_tables_raw

    content = parse_docx("inputs/memorias/Documento_Ambiental.docx")
    print(content.texto[:500])
    for tabla in content.tablas:
        for fila in tabla:
            print(fila)
"""
from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional


# ---------------------------------------------------------------------------
# Dataclass de resultado
# ---------------------------------------------------------------------------

@dataclass
class DocxContent:
    """Resultado de parsear un archivo .docx.

    texto:
        Texto plano del documento, párrafo a párrafo, unido por saltos de línea.
        Los párrafos vacíos (solo espacios) se omiten.

    tablas:
        Lista de tablas. Cada tabla es una lista de filas.
        Cada fila es un dict {cabecera: valor}.
        La primera fila del DOCX se usa como cabecera; si está vacía,
        se generan claves col_0, col_1, col_2…

    metadatos:
        Dict con propiedades del documento (author, created, modified, title, subject).
        Los campos ausentes o None se incluyen como None.

    num_paginas_estimadas:
        Estimación basada en longitud del texto: max(1, len(texto) // 2500 + 1).
        No usa metadatos de paginación del DOCX; es una aproximación.
    """
    texto: str
    tablas: list[list[dict]] = field(default_factory=list)
    metadatos: dict = field(default_factory=dict)
    num_paginas_estimadas: int = 1


# ---------------------------------------------------------------------------
# Helpers internos
# ---------------------------------------------------------------------------

def _extract_text(doc) -> str:
    """Extrae texto plano de todos los párrafos no vacíos del documento."""
    lines = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n".join(lines)


def _extract_tables(doc) -> list[list[dict]]:
    """Extrae tablas usando la primera fila como cabeceras.

    Si todos los valores de la primera fila están vacíos, usa col_0, col_1…
    """
    result = []
    for table in doc.tables:
        if not table.rows:
            continue

        # Primera fila → cabeceras candidatas
        header_row = [cell.text.strip() for cell in table.rows[0].cells]
        has_headers = any(h for h in header_row)

        if has_headers:
            headers = [h if h else f"col_{i}" for i, h in enumerate(header_row)]
            data_rows = table.rows[1:]
        else:
            headers = [f"col_{i}" for i in range(len(header_row))]
            data_rows = table.rows  # la primera fila también es dato

        rows = []
        for row in data_rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append({headers[i]: cells[i] for i in range(min(len(headers), len(cells)))})

        result.append(rows)
    return result


def _extract_metadatos(doc) -> dict:
    """Extrae las propiedades básicas del documento."""
    props = doc.core_properties
    return {
        "author":   getattr(props, "author",   None),
        "created":  getattr(props, "created",  None),
        "modified": getattr(props, "modified", None),
        "title":    getattr(props, "title",    None),
        "subject":  getattr(props, "subject",  None),
    }


# ---------------------------------------------------------------------------
# API pública
# ---------------------------------------------------------------------------

def parse_docx(ruta: "str | Path") -> DocxContent:
    """Parsea un archivo .docx y devuelve su contenido estructurado.

    Args:
        ruta: Ruta al archivo .docx.

    Returns:
        DocxContent con texto, tablas, metadatos y estimación de páginas.

    Raises:
        FileNotFoundError: si el archivo no existe.
        ValueError: si la extensión no es .docx, o si el archivo no es un DOCX válido.
    """
    import docx as python_docx  # python-docx

    path = Path(ruta)

    if not path.exists():
        raise FileNotFoundError(f"Archivo no encontrado: '{path}'")

    if path.suffix.lower() != ".docx":
        raise ValueError(
            f"Se esperaba un archivo .docx, pero se recibió '{path.suffix}': '{path}'"
        )

    try:
        doc = python_docx.Document(str(path))
    except Exception as exc:
        raise ValueError(
            f"No se pudo abrir '{path}' como documento DOCX válido: {exc}"
        ) from exc

    texto     = _extract_text(doc)
    tablas    = _extract_tables(doc)
    metadatos = _extract_metadatos(doc)
    num_pags  = max(1, len(texto) // 2500 + 1)

    return DocxContent(
        texto=texto,
        tablas=tablas,
        metadatos=metadatos,
        num_paginas_estimadas=num_pags,
    )


def extract_tables_raw(ruta: "str | Path") -> list[list[list[str]]]:
    """Extrae tablas del .docx sin interpretar cabeceras.

    Cada tabla es una lista de filas; cada fila es una lista de strings.
    Útil cuando la primera fila no es cabecera o la estructura es irregular.

    Args:
        ruta: Ruta al archivo .docx.

    Returns:
        Lista de tablas, donde cada tabla es lista[fila[celda]].

    Raises:
        FileNotFoundError: si el archivo no existe.
        ValueError: si la extensión no es .docx o el archivo no es válido.
    """
    import docx as python_docx

    path = Path(ruta)

    if not path.exists():
        raise FileNotFoundError(f"Archivo no encontrado: '{path}'")

    if path.suffix.lower() != ".docx":
        raise ValueError(
            f"Se esperaba un archivo .docx, pero se recibió '{path.suffix}': '{path}'"
        )

    try:
        doc = python_docx.Document(str(path))
    except Exception as exc:
        raise ValueError(
            f"No se pudo abrir '{path}' como documento DOCX válido: {exc}"
        ) from exc

    result = []
    for table in doc.tables:
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        result.append(rows)
    return result
