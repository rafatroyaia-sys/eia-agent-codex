"""
climogram_docx_inserter -- CL-05
Inserta un climograma PNG ya existente en un archivo DOCX, con pie de figura.

No genera climogramas (eso es CL-04).
No llama a AEMET ni a ningún servicio externo.
No usa IA.
No modifica el DOCX original si output_docx es distinto.
"""
from __future__ import annotations

import zipfile
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

_PNG_SIGNATURE = b"\x89PNG\r\n\x1a\n"


# ---------------------------------------------------------------------------
# ClimogramDocxInsertConfig
# ---------------------------------------------------------------------------

@dataclass
class ClimogramDocxInsertConfig:
    """Parámetros de inserción del climograma en el DOCX."""

    heading: str | None = "Climograma"
    caption: str | None = None
    image_width_inches: float = 5.8
    insert_page_break_before: bool = False
    insert_page_break_after: bool = False
    caption_style: str | None = None
    center_image: bool = True

    def to_dict(self) -> dict:
        return {
            "heading": self.heading,
            "caption": self.caption,
            "image_width_inches": self.image_width_inches,
            "insert_page_break_before": self.insert_page_break_before,
            "insert_page_break_after": self.insert_page_break_after,
            "caption_style": self.caption_style,
            "center_image": self.center_image,
        }

    @classmethod
    def from_dict(cls, data: dict) -> "ClimogramDocxInsertConfig":
        return cls(
            heading=data.get("heading", "Climograma"),
            caption=data.get("caption"),
            image_width_inches=float(data.get("image_width_inches", 5.8)),
            insert_page_break_before=bool(data.get("insert_page_break_before", False)),
            insert_page_break_after=bool(data.get("insert_page_break_after", False)),
            caption_style=data.get("caption_style"),
            center_image=bool(data.get("center_image", True)),
        )


# ---------------------------------------------------------------------------
# ClimogramDocxInsertResult
# ---------------------------------------------------------------------------

@dataclass
class ClimogramDocxInsertResult:
    """Resultado de la inserción del climograma en el DOCX."""

    input_docx: str
    output_docx: str
    png_path: str
    inserted: bool
    caption: str | None
    warnings: list[str]
    notes: list[str]

    def to_dict(self) -> dict:
        return {
            "input_docx": self.input_docx,
            "output_docx": self.output_docx,
            "png_path": self.png_path,
            "inserted": self.inserted,
            "caption": self.caption,
            "warnings": list(self.warnings),
            "notes": list(self.notes),
        }

    def summary(self) -> str:
        lines = [
            f"Input DOCX : {self.input_docx}",
            f"Output DOCX: {self.output_docx}",
            f"PNG        : {self.png_path}",
            f"Insertado  : {self.inserted}",
        ]
        if self.caption:
            lines.append(f"Caption    : {self.caption}")
        for w in self.warnings:
            lines.append(f"AVISO: {w}")
        for n in self.notes:
            lines.append(f"Nota: {n}")
        return "\n".join(lines)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _is_valid_png(path: Path) -> bool:
    """Devuelve True si el archivo existe y su cabecera es la firma PNG."""
    if not path.exists() or path.stat().st_size == 0:
        return False
    with open(path, "rb") as f:
        return f.read(8) == _PNG_SIGNATURE


def default_climogram_caption(
    station_name: str | None = None,
    period: str | None = None,
) -> str:
    """Genera un pie de figura estándar para el climograma.

    Ejemplos:
        default_climogram_caption()
            → "Figura. Climograma de la estación climática de referencia."
        default_climogram_caption("Lanzarote Aeropuerto", "1991-2020")
            → "Figura. Climograma de la estación Lanzarote Aeropuerto, periodo 1991-2020."
    """
    if not station_name:
        return "Figura. Climograma de la estación climática de referencia."
    parts = [f"Figura. Climograma de la estación {station_name}"]
    if period:
        parts.append(f", periodo {period}")
    parts.append(".")
    return "".join(parts)


def validate_docx_contains_image(docx_path: "str | Path") -> bool:
    """Devuelve True si el DOCX existe, es un ZIP válido y contiene ≥1 imagen."""
    p = Path(docx_path)
    if not p.exists():
        return False
    try:
        with zipfile.ZipFile(p) as z:
            return any(n.startswith("word/media/") for n in z.namelist())
    except (zipfile.BadZipFile, OSError):
        return False


def count_docx_images(docx_path: "str | Path") -> int:
    """Cuenta el número de archivos en word/media/ del DOCX."""
    p = Path(docx_path)
    if not p.exists():
        return 0
    try:
        with zipfile.ZipFile(p) as z:
            return sum(1 for n in z.namelist() if n.startswith("word/media/"))
    except (zipfile.BadZipFile, OSError):
        return 0


# ---------------------------------------------------------------------------
# insert_climogram_in_docx
# ---------------------------------------------------------------------------

def insert_climogram_in_docx(
    input_docx: "str | Path",
    png_path: "str | Path",
    output_docx: "str | Path",
    config: ClimogramDocxInsertConfig | None = None,
) -> ClimogramDocxInsertResult:
    """Inserta un climograma PNG en un DOCX existente y guarda el resultado.

    Args:
        input_docx:  DOCX de entrada. Debe existir.
        png_path:    PNG del climograma. Debe existir y tener firma PNG válida.
        output_docx: Ruta de salida. Debe terminar en '.docx'.
        config:      Configuración de inserción. Si None, usa ClimogramDocxInsertConfig().

    Returns:
        ClimogramDocxInsertResult con metadatos de la operación.

    Raises:
        FileNotFoundError: Si input_docx o png_path no existen.
        ValueError:        Si png_path no es PNG válido o output_docx no termina en '.docx'.
    """
    inp = Path(input_docx)
    png = Path(png_path)
    out = Path(output_docx)

    if not inp.exists():
        raise FileNotFoundError(f"input_docx no encontrado: {inp}")
    if not png.exists():
        raise FileNotFoundError(f"png_path no encontrado: {png}")
    if not _is_valid_png(png):
        raise ValueError(f"png_path no es un PNG válido (firma incorrecta o archivo vacío): {png}")
    if out.suffix.lower() != ".docx":
        raise ValueError(
            f"output_docx debe terminar en '.docx', recibido: '{out.suffix}'. "
            f"Ruta completa: {out}"
        )

    if config is None:
        config = ClimogramDocxInsertConfig()

    out.parent.mkdir(parents=True, exist_ok=True)

    warnings: list[str] = []
    notes: list[str] = []

    doc = Document(str(inp))

    # ── Salto de página antes ────────────────────────────────────────────────
    if config.insert_page_break_before:
        doc.add_page_break()

    # ── Encabezado ───────────────────────────────────────────────────────────
    if config.heading is not None:
        doc.add_heading(config.heading, level=2)

    # ── Imagen ───────────────────────────────────────────────────────────────
    para = doc.add_paragraph()
    if config.center_image:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(str(png), width=Inches(config.image_width_inches))

    # ── Pie de figura ────────────────────────────────────────────────────────
    effective_caption = config.caption
    if effective_caption is not None:
        if config.caption_style:
            try:
                cap_para = doc.add_paragraph(effective_caption, style=config.caption_style)
            except Exception:
                cap_para = doc.add_paragraph(effective_caption)
                warnings.append(
                    f"Estilo de caption '{config.caption_style}' no encontrado en el DOCX; "
                    "usando estilo Normal."
                )
        else:
            cap_para = doc.add_paragraph(effective_caption)
        if config.center_image:
            cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── Salto de página después ──────────────────────────────────────────────
    if config.insert_page_break_after:
        doc.add_page_break()

    doc.save(str(out))

    return ClimogramDocxInsertResult(
        input_docx=str(inp),
        output_docx=str(out),
        png_path=str(png),
        inserted=True,
        caption=effective_caption,
        warnings=warnings,
        notes=notes,
    )
