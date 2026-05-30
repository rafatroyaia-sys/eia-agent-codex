"""
document_structure_manager.py — EN-02
Validacion y normalizacion de la estructura fisica del DOCX final.

Valida que las secciones principales del Documento Ambiental aparezcan
en orden canonico: PORTADA, INDICE, A-K, ANEXO_GRAFICO, HOJA_FIRMAS.
No reescribe contenido. No regenera el documento. No inventa datos.
No llama a servicios externos. 100 % offline.
"""
from __future__ import annotations

import copy
import json
import re
import shutil
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Optional

# ---------------------------------------------------------------------------
# Orden canonico
# ---------------------------------------------------------------------------

BLOCK_IDS: List[str] = list("ABCDEFGHIJK")

CANONICAL_DOCUMENT_ORDER: List[str] = [
    "PORTADA",
    "INDICE",
    *BLOCK_IDS,
    "ANEXO_GRAFICO",
    "HOJA_FIRMAS",
]

# Titulos esperados (lower-case, para comparacion case-insensitive)
_SECTION_TITLE_HINTS: dict = {
    "PORTADA": [
        "documento ambiental", "borrador", "expediente eia",
        "evaluacion de impacto", "promotor",
    ],
    "INDICE": ["indice", "índice", "table of contents", "contenidos"],
    "A": ["a —", "a–", "a -", "a.", "a:", "bloque a", "a identificacion",
          "a descripcion", "a. "],
    "B": ["b —", "b–", "b -", "b.", "b:", "bloque b", "b descripcion",
          "b alternativas"],
    "C": ["c —", "c–", "c -", "c.", "c:", "bloque c", "c inventario",
          "c situacion"],
    "D": ["d —", "d–", "d -", "d.", "d:", "bloque d"],
    "E": ["e —", "e–", "e -", "e.", "e:", "bloque e"],
    "F": ["f —", "f–", "f -", "f.", "f:", "bloque f"],
    "G": ["g —", "g–", "g -", "g.", "g:", "bloque g", "g alternativas",
          "g justificacion"],
    "H": ["h —", "h–", "h -", "h.", "h:", "bloque h", "h impactos"],
    "I": ["i —", "i–", "i -", "i.", "i:", "bloque i", "i valoracion",
          "i conclusion"],
    "J": ["j —", "j–", "j -", "j.", "j:", "bloque j", "j medidas",
          "j programa"],
    "K": ["k —", "k–", "k -", "k.", "k:", "bloque k", "k plan",
          "k programa de vigilancia"],
    "ANEXO_GRAFICO": [
        "anexo grafico", "anexo gráfico", "anexo cartografico",
        "anexo cartográfico", "anexo y cartografico", "anexo fotografico",
        "annex", "anexo grafico y cartografico",
    ],
    "HOJA_FIRMAS": [
        "hoja de firmas", "firmas y revision", "revision y firmas",
        "firmas", "signature", "hoja firmas", "firmas tecnicas",
        "hoja de firmas y revision",
    ],
}

# Codigos de validacion
_STRUCT_CODES = {
    "EN02-E001": "Portada no encontrada en el DOCX.",
    "EN02-E002": "Indice no encontrado en el DOCX.",
    "EN02-E003": "Bloque requerido ausente en el DOCX.",
    "EN02-E004": "Orden incorrecto de bloques/secciones.",
    "EN02-E005": "Hoja de firmas no esta al final del documento.",
    "EN02-E006": "Anexo grafico aparece antes del bloque K.",
    "EN02-W001": "Bloque/seccion duplicada detectada.",
    "EN02-W002": "Anexo grafico ausente (advertencia — es opcional).",
    "EN02-W003": "Hoja de firmas ausente (advertencia — es opcional).",
    "EN02-I001": "Estructura del documento valida.",
}


# ---------------------------------------------------------------------------
# Dataclasses
# ---------------------------------------------------------------------------

@dataclass
class DocumentSectionPosition:
    """Posicion de una seccion en el DOCX."""

    section_id: str
    title: str
    paragraph_index: Optional[int]
    heading_level: Optional[int]
    found: bool
    notes: List[str] = field(default_factory=list)

    def to_dict(self) -> dict:
        return {
            "section_id": self.section_id,
            "title": self.title,
            "paragraph_index": self.paragraph_index,
            "heading_level": self.heading_level,
            "found": self.found,
            "notes": self.notes,
        }

    def summary(self) -> str:
        status = "OK" if self.found else "NO ENCONTRADA"
        idx = f" (par {self.paragraph_index})" if self.paragraph_index is not None else ""
        return f"[{status}] {self.section_id}: {self.title}{idx}"


@dataclass
class DocumentStructureResult:
    """Resultado de la validacion/normalizacion de estructura del DOCX."""

    input_docx: str
    output_docx: Optional[str]
    sections_found: List[DocumentSectionPosition] = field(default_factory=list)
    expected_order: List[str] = field(default_factory=list)
    detected_order: List[str] = field(default_factory=list)
    errors: List[dict] = field(default_factory=list)
    warnings: List[dict] = field(default_factory=list)
    notes: List[str] = field(default_factory=list)

    def error_count(self) -> int:
        return len(self.errors)

    def warning_count(self) -> int:
        return len(self.warnings)

    def is_valid(self) -> bool:
        return self.error_count() == 0

    def to_dict(self) -> dict:
        return {
            "input_docx": self.input_docx,
            "output_docx": self.output_docx,
            "is_valid": self.is_valid(),
            "sections_found": [s.to_dict() for s in self.sections_found],
            "expected_order": self.expected_order,
            "detected_order": self.detected_order,
            "errors": self.errors,
            "warnings": self.warnings,
            "notes": self.notes,
        }

    def summary(self) -> str:
        status = "VALIDO" if self.is_valid() else "CON ERRORES"
        found_ids = {s.section_id for s in self.sections_found if s.found}
        blocks_found = [b for b in BLOCK_IDS if b in found_ids]
        lines = [
            f"Estructura  : {status}",
            f"Errores     : {self.error_count()}",
            f"Avisos      : {self.warning_count()}",
            f"Bloques A-K : {len(blocks_found)}/11 ({', '.join(blocks_found) if blocks_found else 'ninguno'})",
        ]
        if self.output_docx:
            lines.append(f"Output DOCX : {self.output_docx}")
        if self.notes:
            for n in self.notes[:3]:
                lines.append(f"  [NOTA] {n}")
        return "\n".join(lines)


# ---------------------------------------------------------------------------
# Helpers internos
# ---------------------------------------------------------------------------

def _heading_level(paragraph) -> Optional[int]:
    """Devuelve el nivel del heading (1-9) o None si no es heading."""
    style_name = paragraph.style.name if paragraph.style else ""
    m = re.match(r"[Hh]eading\s+(\d)", style_name)
    if m:
        return int(m.group(1))
    return None


def _para_text(paragraph) -> str:
    """Texto normalizado del parrafo (sin espacios extra, sin acentos normalizados)."""
    import unicodedata
    raw = paragraph.text.strip()
    # Normalizar a ASCII para comparaciones robustas
    normalized = unicodedata.normalize("NFD", raw)
    # Mantener el texto original pero en lower para comparaciones
    return raw


def _matches_section(text: str, section_id: str) -> bool:
    """
    True si el texto del heading corresponde a la seccion indicada.

    Para bloques A-K: solo regex de inicio (mas robusto que hints cortos
    que generan falsos positivos al normalizarse).
    Para secciones no-bloque: hints con palabras clave suficientemente largas.
    """
    import unicodedata

    def _norm(s: str) -> str:
        return unicodedata.normalize("NFD", s.lower()).encode("ascii", "ignore").decode().strip()

    text_clean = text.strip()
    text_norm = _norm(text_clean)

    # --- Bloques A-K: solo regex de inicio, preciso y sin ambigüedad ---
    if section_id in BLOCK_IDS:
        letter = section_id.upper()
        # "A —", "A–", "A -", "A.", "A:", seguidos opcionalmente de espacio
        if re.match(rf"^{letter}\s*[—–\-\.:]", text_clean, re.IGNORECASE):
            return True
        # "A " seguido de letra mayuscula (inicio de titulo)
        if re.match(rf"^{letter}\s+[A-ZÀ-ÿ]", text_clean):
            return True
        # "Bloque A" (cualquier capitalización)
        if re.match(rf"(?i)^bloque\s+{letter}\b", text_clean):
            return True
        return False

    # --- Secciones no-bloque: hints de palabras largas ---
    hints = _SECTION_TITLE_HINTS.get(section_id, [])
    for hint in hints:
        hint_norm = _norm(hint)
        # Solo usar hints con longitud suficiente para evitar falsos positivos
        if len(hint_norm) >= 4 and hint_norm in text_norm:
            return True

    return False


def _is_portada_paragraph(paragraph, idx: int) -> bool:
    """
    True si el parrafo es candidato a formar parte de la portada.
    La portada son los parrafos antes del primer heading.
    """
    text = paragraph.text.strip()
    lev = _heading_level(paragraph)
    return lev is None and len(text) > 3 and idx < 30


def _make_issue(code: str, message: str, section_id: str = "",
                recommendation: str = "") -> dict:
    return {
        "code": code,
        "severity": "ERROR" if code.startswith("EN02-E") else (
            "WARNING" if code.startswith("EN02-W") else "INFO"
        ),
        "section_id": section_id,
        "message": message,
        "recommendation": recommendation,
    }


# ---------------------------------------------------------------------------
# Funcion principal de deteccion
# ---------------------------------------------------------------------------

def detect_document_sections(
    docx_path: "str | Path",
) -> List[DocumentSectionPosition]:
    """
    Detecta secciones del DOCX en orden de aparicion.

    Abre el DOCX con python-docx e identifica:
    PORTADA, INDICE, bloques A-K, ANEXO_GRAFICO, HOJA_FIRMAS.
    No rompe si faltan secciones.
    """
    try:
        import docx as python_docx
    except ImportError:
        import docx as python_docx

    path = Path(docx_path)
    positions: dict[str, DocumentSectionPosition] = {}
    duplicates: dict[str, int] = {}

    # Inicializar todas las secciones como no encontradas
    for sid in CANONICAL_DOCUMENT_ORDER:
        positions[sid] = DocumentSectionPosition(
            section_id=sid,
            title=_SECTION_TITLE_HINTS.get(sid, [sid.lower()])[0],
            paragraph_index=None,
            heading_level=None,
            found=False,
        )

    if not path.exists():
        for sid in CANONICAL_DOCUMENT_ORDER:
            positions[sid].notes.append("DOCX no encontrado en la ruta indicada.")
        return list(positions.values())

    try:
        doc = python_docx.Document(str(path))
    except Exception as exc:
        for sid in CANONICAL_DOCUMENT_ORDER:
            positions[sid].notes.append(f"Error al abrir DOCX: {exc}")
        return list(positions.values())

    paragraphs = doc.paragraphs
    portada_found = False
    first_heading_idx = None

    # Primera pasada: detectar primer heading para saber dónde termina la portada
    for idx, para in enumerate(paragraphs):
        lev = _heading_level(para)
        if lev is not None and para.text.strip():
            first_heading_idx = idx
            break

    # Portada: si hay parrafos con texto antes del primer heading
    if first_heading_idx is None or first_heading_idx > 0:
        # Buscar parrafos con texto antes del primer heading
        limit = first_heading_idx if first_heading_idx is not None else min(30, len(paragraphs))
        for idx in range(limit):
            para = paragraphs[idx]
            if _is_portada_paragraph(para, idx):
                positions["PORTADA"].found = True
                positions["PORTADA"].paragraph_index = idx
                positions["PORTADA"].title = para.text.strip()[:60] or "Portada"
                portada_found = True
                break

    # Segunda pasada: detectar headings
    for idx, para in enumerate(paragraphs):
        lev = _heading_level(para)
        if lev is None:
            continue
        text = para.text.strip()
        if not text:
            continue

        # Comprobar cada seccion pendiente
        for sid in CANONICAL_DOCUMENT_ORDER:
            if sid == "PORTADA":
                continue
            if _matches_section(text, sid):
                if positions[sid].found:
                    # Duplicado
                    duplicates[sid] = duplicates.get(sid, 1) + 1
                    positions[sid].notes.append(
                        f"Duplicado detectado en parrafo {idx}: '{text[:50]}'"
                    )
                else:
                    positions[sid].found = True
                    positions[sid].paragraph_index = idx
                    positions[sid].heading_level = lev
                    positions[sid].title = text[:80]
                break  # Un heading puede pertenecer solo a una seccion

    return list(positions.values())


# ---------------------------------------------------------------------------
# Validacion
# ---------------------------------------------------------------------------

def validate_document_structure(
    docx_path: "str | Path",
) -> "DocumentStructureResult":
    """
    Valida la estructura del DOCX segun el orden canonico.

    Comprueba:
    - PORTADA presente (EN02-E001)
    - INDICE presente (EN02-E002)
    - Bloques A-K presentes (EN02-E003)
    - Orden canonico de bloques respetado (EN02-E004)
    - ANEXO_GRAFICO despues de K si existe (EN02-E006)
    - HOJA_FIRMAS al final si existe (EN02-E005)
    - Warnings por ausencia de ANEXO/FIRMAS y duplicados (EN02-W001..W003)
    """
    path = Path(docx_path)
    sections = detect_document_sections(path)

    # Mapa rapido
    by_id: dict[str, DocumentSectionPosition] = {s.section_id: s for s in sections}

    errors: list[dict] = []
    warnings: list[dict] = []
    notes: list[str] = []

    # -- Orden detectado (solo secciones encontradas, por paragraph_index) --
    found_sections = [s for s in sections if s.found]
    found_sections.sort(key=lambda s: (s.paragraph_index or 999999))
    detected_order = [s.section_id for s in found_sections]

    # -- Regla 1: PORTADA --
    if not by_id["PORTADA"].found:
        errors.append(_make_issue(
            "EN02-E001",
            "Portada no encontrada en el DOCX.",
            "PORTADA",
            "Ejecutar 'document-build-docx --write' para regenerar el DOCX con portada.",
        ))

    # -- Regla 2: INDICE --
    if not by_id["INDICE"].found:
        errors.append(_make_issue(
            "EN02-E002",
            "Indice no encontrado en el DOCX.",
            "INDICE",
            "Ejecutar 'document-build-docx --write' para regenerar el DOCX con indice.",
        ))

    # -- Regla 3: Bloques A-K presentes --
    for bid in BLOCK_IDS:
        if not by_id[bid].found:
            errors.append(_make_issue(
                "EN02-E003",
                f"Bloque requerido '{bid}' no encontrado en el DOCX.",
                bid,
                f"Verificar que DOC-01 genero el bloque {bid} y DOC-02 lo incluyo.",
            ))

    # -- Regla 4: Orden de bloques A-K --
    found_blocks = [b for b in BLOCK_IDS if by_id[b].found]
    if len(found_blocks) > 1:
        # Verificar que el orden de paragraph_index coincide con el orden canonico
        block_positions = [(b, by_id[b].paragraph_index or 0) for b in found_blocks]
        sorted_by_pos = sorted(block_positions, key=lambda x: x[1])
        order_by_pos = [b for b, _ in sorted_by_pos]

        if order_by_pos != found_blocks:
            # Hay al menos un bloque fuera de orden
            errors.append(_make_issue(
                "EN02-E004",
                f"Orden incorrecto de bloques A-K. "
                f"Detectado: {' > '.join(order_by_pos)}. "
                f"Esperado: {' > '.join(found_blocks)}.",
                "",
                "El DOCX tiene bloques fuera del orden canonico A-K.",
            ))

    # -- Regla 5: HOJA_FIRMAS al final --
    firmas = by_id["HOJA_FIRMAS"]
    if firmas.found:
        firmas_idx = firmas.paragraph_index or 0
        # Nada debe aparecer despues de HOJA_FIRMAS (salvo su propio contenido)
        # Verificar que ningun bloque A-K ni ANEXO aparezca despues
        after_firmas = []
        for sid in [*BLOCK_IDS, "ANEXO_GRAFICO"]:
            sec = by_id[sid]
            if sec.found and (sec.paragraph_index or 0) > firmas_idx:
                after_firmas.append(sid)
        if after_firmas:
            errors.append(_make_issue(
                "EN02-E005",
                f"Hoja de firmas no esta al final. "
                f"Estas secciones aparecen despues: {', '.join(after_firmas)}.",
                "HOJA_FIRMAS",
                "La hoja de firmas debe ser la ultima seccion del documento.",
            ))
    else:
        warnings.append(_make_issue(
            "EN02-W003",
            "Hoja de firmas ausente en el DOCX (advertencia — no es error en validacion base).",
            "HOJA_FIRMAS",
            "Ejecutar 'document-prepare-presentation --write' para anadir hoja de firmas.",
        ))

    # -- Regla 6: ANEXO_GRAFICO despues de K --
    anexo = by_id["ANEXO_GRAFICO"]
    if anexo.found:
        anexo_idx = anexo.paragraph_index or 0
        k_sec = by_id["K"]
        if k_sec.found:
            k_idx = k_sec.paragraph_index or 0
            if anexo_idx < k_idx:
                errors.append(_make_issue(
                    "EN02-E006",
                    f"Anexo grafico (par {anexo_idx}) aparece antes del bloque K (par {k_idx}).",
                    "ANEXO_GRAFICO",
                    "El anexo grafico debe ir despues del bloque K.",
                ))
    else:
        warnings.append(_make_issue(
            "EN02-W002",
            "Anexo grafico ausente en el DOCX (advertencia — se anade con 'document-insert-figures').",
            "ANEXO_GRAFICO",
            "Ejecutar 'document-insert-figures --write' para anadir el anexo grafico.",
        ))

    # -- HOJA_FIRMAS despues de ANEXO_GRAFICO si ambos existen --
    if anexo.found and firmas.found:
        anexo_idx = anexo.paragraph_index or 0
        firmas_idx = firmas.paragraph_index or 0
        if firmas_idx < anexo_idx:
            errors.append(_make_issue(
                "EN02-E005",
                f"Hoja de firmas (par {firmas_idx}) aparece antes del anexo grafico (par {anexo_idx}).",
                "HOJA_FIRMAS",
                "La hoja de firmas debe ir despues del anexo grafico.",
            ))

    # -- Duplicados --
    for sid in CANONICAL_DOCUMENT_ORDER:
        sec = by_id[sid]
        if len(sec.notes) > 0 and any("Duplicado" in n for n in sec.notes):
            warnings.append(_make_issue(
                "EN02-W001",
                f"Seccion '{sid}' tiene entradas duplicadas en el DOCX.",
                sid,
                "Revisar la generacion del DOCX para evitar secciones repetidas.",
            ))

    # -- Nota de exito --
    if not errors:
        notes.append("Estructura del documento valida segun orden canonico.")

    notes.append(
        f"Secciones detectadas: {', '.join(detected_order) if detected_order else 'ninguna'}."
    )

    return DocumentStructureResult(
        input_docx=str(path),
        output_docx=None,
        sections_found=sections,
        expected_order=list(CANONICAL_DOCUMENT_ORDER),
        detected_order=detected_order,
        errors=errors,
        warnings=warnings,
        notes=notes,
    )


# ---------------------------------------------------------------------------
# Normalizacion (conservadora)
# ---------------------------------------------------------------------------

def normalize_document_structure(
    input_docx: "str | Path",
    output_docx: "str | Path",
    include_page_breaks: bool = True,
) -> "DocumentStructureResult":
    """
    Crea una copia normalizada del DOCX.

    Alcance conservador (por seguridad con python-docx):
    - Copia el DOCX a output_docx sin tocar el original.
    - Refuerza 'page_break_before=True' en headings principales
      (INDICE, bloques A-K, ANEXO_GRAFICO, HOJA_FIRMAS) cuando falta.
    - No reordena secciones complejas (riesgo de corrupcion).
    - No modifica contenido textual, tablas ni imagenes.
    - Devuelve resultado de validacion sobre el DOCX normalizado.

    Si no es posible normalizar con seguridad, devuelve el resultado
    de validacion del original con notes explicativas.
    """
    try:
        import docx as python_docx
        from docx.oxml.ns import qn
    except ImportError as exc:
        result = validate_document_structure(input_docx)
        result.output_docx = None
        result.notes.append(f"python-docx no disponible: {exc}")
        return result

    in_path = Path(input_docx)
    out_path = Path(output_docx)

    if not in_path.exists():
        result = DocumentStructureResult(
            input_docx=str(in_path),
            output_docx=None,
            errors=[_make_issue("EN02-E001", f"DOCX de entrada no encontrado: {in_path}")],
            notes=["Normalizacion cancelada: archivo de entrada no existe."],
        )
        return result

    # Copiar el DOCX
    try:
        out_path.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(str(in_path), str(out_path))
    except Exception as exc:
        result = validate_document_structure(in_path)
        result.output_docx = None
        result.notes.append(f"Error al copiar DOCX para normalizacion: {exc}")
        return result

    changes_made: list[str] = []

    if include_page_breaks:
        try:
            doc = python_docx.Document(str(out_path))
            paragraphs = doc.paragraphs

            # IDs de secciones principales que deben tener page break before
            main_sections = {"INDICE", *BLOCK_IDS, "ANEXO_GRAFICO", "HOJA_FIRMAS"}

            for idx, para in enumerate(paragraphs):
                lev = _heading_level(para)
                if lev != 1:
                    continue
                text = para.text.strip()
                if not text:
                    continue

                # Identificar si es una seccion principal
                matched_sid = None
                for sid in main_sections:
                    if _matches_section(text, sid):
                        matched_sid = sid
                        break

                if matched_sid is None:
                    continue

                # Comprobar si ya tiene page break before
                try:
                    pf = para.paragraph_format
                    if not pf.page_break_before:
                        pf.page_break_before = True
                        changes_made.append(
                            f"page_break_before anadido antes de '{matched_sid}' (par {idx})"
                        )
                except Exception:
                    pass  # No forzar cambios que puedan corromper

            if changes_made:
                doc.save(str(out_path))
        except Exception as exc:
            changes_made.append(f"Advertencia al reforzar page breaks: {exc}")

    # Validar el DOCX normalizado
    result = validate_document_structure(out_path)
    result.input_docx = str(in_path)
    result.output_docx = str(out_path)

    if changes_made:
        for change in changes_made[:10]:
            result.notes.append(f"Normalizacion: {change}")
    else:
        result.notes.append("Normalizacion: no se requirieron cambios estructurales.")

    return result


# ---------------------------------------------------------------------------
# Outputs
# ---------------------------------------------------------------------------

def build_document_structure_markdown(result: "DocumentStructureResult") -> str:
    """Genera informe Markdown de la validacion/normalizacion estructural."""
    lines: list[str] = []

    status = "VALIDO" if result.is_valid() else "CON ERRORES"
    lines += [
        "# Validacion de estructura del DOCX — EN-02",
        "",
        "## 1. Resumen",
        "",
        f"| Campo | Valor |",
        f"|-------|-------|",
        f"| **Estado** | {status} |",
        f"| **DOCX analizado** | `{Path(result.input_docx).name}` |",
        f"| **Errores** | {result.error_count()} |",
        f"| **Avisos** | {result.warning_count()} |",
    ]
    if result.output_docx:
        lines.append(f"| **DOCX normalizado** | `{Path(result.output_docx).name}` |")
    lines.append("")

    lines += [
        "## 2. Orden esperado",
        "",
        " > ".join(result.expected_order),
        "",
        "## 3. Orden detectado",
        "",
        (" > ".join(result.detected_order) if result.detected_order else "_Ninguna seccion detectada_"),
        "",
        "## 4. Secciones encontradas",
        "",
        "| ID | Titulo detectado | Parrafo | Nivel | Estado |",
        "|----|-----------------|---------|-------|--------|",
    ]
    for sec in result.sections_found:
        idx = str(sec.paragraph_index) if sec.paragraph_index is not None else "—"
        lev = str(sec.heading_level) if sec.heading_level is not None else "—"
        found_str = "✅ OK" if sec.found else "❌ No encontrado"
        title_short = (sec.title[:50] + "...") if len(sec.title) > 50 else sec.title
        lines.append(f"| {sec.section_id} | {title_short} | {idx} | {lev} | {found_str} |")
    lines.append("")

    if result.errors:
        lines += ["## 5. Errores", ""]
        for err in result.errors:
            lines.append(f"- **[{err['code']}]** {err['message']}")
            if err.get("recommendation"):
                lines.append(f"  - _Recomendacion_: {err['recommendation']}")
        lines.append("")
    else:
        lines += ["## 5. Errores", "", "_Sin errores._", ""]

    if result.warnings:
        lines += ["## 6. Avisos", ""]
        for w in result.warnings:
            lines.append(f"- **[{w['code']}]** {w['message']}")
        lines.append("")
    else:
        lines += ["## 6. Avisos", "", "_Sin avisos._", ""]

    if result.notes:
        lines += ["## 7. Notas", ""]
        for n in result.notes:
            lines.append(f"- {n}")
        lines.append("")

    return "\n".join(lines)


def write_document_structure_outputs(
    result: "DocumentStructureResult",
    output_dir: "str | Path",
) -> List[Path]:
    """
    Escribe los outputs de validacion/normalizacion estructural.

    Genera:
    - document_structure_result.json
    - document_structure_result.md

    Devuelve la lista de rutas escritas.
    """
    out_dir = Path(output_dir)
    out_dir.mkdir(parents=True, exist_ok=True)

    json_path = out_dir / "document_structure_result.json"
    md_path = out_dir / "document_structure_result.md"

    json_path.write_text(
        json.dumps(result.to_dict(), ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    md_path.write_text(
        build_document_structure_markdown(result),
        encoding="utf-8",
    )

    return [json_path, md_path]


# ---------------------------------------------------------------------------
# Seleccion automatica del mejor DOCX disponible
# ---------------------------------------------------------------------------

_DOCX_CANDIDATES = [
    "documento/documento_ambiental_final.docx",
    "documento/documento_ambiental_borrador_con_firmas.docx",
    "documento/documento_ambiental_borrador_con_figuras.docx",
    "documento/documento_ambiental_borrador.docx",
]


def find_best_available_docx(expediente_path: "str | Path") -> Optional[Path]:
    """
    Devuelve el mejor DOCX disponible en el expediente, en orden de preferencia.
    Devuelve None si no hay ningun DOCX disponible.
    """
    exp = Path(expediente_path)
    for candidate in _DOCX_CANDIDATES:
        p = exp / candidate
        if p.exists():
            return p
    return None
