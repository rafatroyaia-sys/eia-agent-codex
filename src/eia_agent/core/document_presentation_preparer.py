"""
document_presentation_preparer -- DOC-08
Preparacion del documento y paquete para revision y presentacion administrativa.

Genera elementos formales y metadatos controlados sin alterar el contenido
tecnico de fondo ni declarar aptitud administrativa.

Genera (solo con write_outputs=True):
  - documento/document_metadata.json
  - documento/document_metadata.md
  - documento/hoja_firmas.md
  - documento/checklist_presentacion.json
  - documento/checklist_presentacion.md
  - documento/documento_ambiental_final_revisable.docx  (best-effort, opcional)

Principios no negociables:
  - No usa IA.
  - No consulta fuentes externas.
  - No modifica DOCX, Markdown ni fuentes existentes.
  - No modifica paquete_entrega/.
  - No declara aptitud administrativa (administrative_ready=False siempre).
  - No firma digitalmente.
  - No presenta nada ante la Administracion.
  - Solo escribe si write_outputs=True.
"""
from __future__ import annotations

import json
from dataclasses import dataclass, field
from datetime import datetime, timezone
from pathlib import Path


# ---------------------------------------------------------------------------
# Constantes publicas
# ---------------------------------------------------------------------------

METADATA_JSON = "document_metadata.json"
METADATA_MD = "document_metadata.md"
SIGNATURE_SHEET_MD = "hoja_firmas.md"
PRESENTATION_CHECKLIST_JSON = "checklist_presentacion.json"
PRESENTATION_CHECKLIST_MD = "checklist_presentacion.md"
FINAL_REVIEW_DOCX = "documento_ambiental_final_revisable.docx"


class PRESENTATION_STATUS:
    PREPARADO_PARA_REVISION = "PREPARADO_PARA_REVISION"
    PENDIENTE_REVISION_TECNICA = "PENDIENTE_REVISION_TECNICA"
    PENDIENTE_DOCUMENTACION = "PENDIENTE_DOCUMENTACION"
    NO_PREPARADO = "NO_PREPARADO"


class PRESENTATION_SEVERITY:
    ERROR = "ERROR"
    WARNING = "WARNING"
    INFO = "INFO"


# Frases de aptitud administrativa que no deben aparecer
_ADMIN_READY_PHRASES: tuple[str, ...] = (
    "apto para presentacion",
    "listo para presentar",
    "preparado para presentacion administrativa",
    "expediente apto",
    "procede su presentacion",
)


# ---------------------------------------------------------------------------
# Dataclasses
# ---------------------------------------------------------------------------

@dataclass
class PresentationIssue:
    """Incidencia registrada durante la preparacion para presentacion."""

    severity: str
    code: str
    message: str
    recommendation: str
    evidence: list[str] = field(default_factory=list)

    def to_dict(self) -> dict:
        return {
            "severity": self.severity,
            "code": self.code,
            "message": self.message,
            "recommendation": self.recommendation,
            "evidence": self.evidence,
        }

    def summary(self) -> str:
        return f"[{self.severity}][{self.code}] {self.message}"


@dataclass
class DocumentMetadata:
    """Metadatos documentales del expediente."""

    expediente_id: str
    generated_at: str
    source_docx: "str | None"
    source_markdown: "str | None"
    package_zip: "str | None"
    final_audit_status: "str | None"
    document_qc_status: "str | None"
    package_status: "str | None"
    export_status: "str | None"
    conditional_chain_status: "str | None" = None
    notes: list[str] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)

    @property
    def administrative_ready(self) -> bool:
        """Siempre False. El modulo no declara aptitud administrativa."""
        return False

    def to_dict(self) -> dict:
        return {
            "expediente_id": self.expediente_id,
            "generated_at": self.generated_at,
            "source_docx": self.source_docx,
            "source_markdown": self.source_markdown,
            "package_zip": self.package_zip,
            "final_audit_status": self.final_audit_status,
            "document_qc_status": self.document_qc_status,
            "package_status": self.package_status,
            "export_status": self.export_status,
            "conditional_chain_status": self.conditional_chain_status,
            "administrative_ready": self.administrative_ready,
            "notes": self.notes,
            "warnings": self.warnings,
        }

    def summary(self) -> str:
        lines = [
            f"Metadatos expediente: {self.expediente_id}",
            f"  Generado          : {self.generated_at}",
            f"  Auditoria final   : {self.final_audit_status or 'N/D'}",
            f"  QC documental     : {self.document_qc_status or 'N/D'}",
            f"  Paquete           : {self.package_status or 'N/D'}",
            f"  Exportacion       : {self.export_status or 'N/D'}",
            f"  IM-09 cad.cond.   : {self.conditional_chain_status or 'N/D'}",
            f"  administrative_ready: {self.administrative_ready}",
        ]
        return "\n".join(lines)


@dataclass
class PresentationChecklistItem:
    """Item del checklist de preparacion para presentacion."""

    item_id: str
    description: str
    status: str  # OK / WARNING / ERROR / NO_APLICA
    evidence: list[str] = field(default_factory=list)
    recommendation: str = ""

    def to_dict(self) -> dict:
        return {
            "item_id": self.item_id,
            "description": self.description,
            "status": self.status,
            "evidence": self.evidence,
            "recommendation": self.recommendation,
        }

    def summary(self) -> str:
        return f"[{self.status}] {self.item_id}: {self.description}"


@dataclass
class PresentationPreparationResult:
    """Resultado completo de la preparacion para presentacion."""

    expediente_id: str
    status: str
    metadata: DocumentMetadata
    checklist_items: list[PresentationChecklistItem] = field(default_factory=list)
    issues: list[PresentationIssue] = field(default_factory=list)
    generated_files: list[str] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)
    notes: list[str] = field(default_factory=list)

    def error_count(self) -> int:
        return sum(1 for i in self.issues if i.severity == PRESENTATION_SEVERITY.ERROR)

    def warning_count(self) -> int:
        return sum(1 for i in self.issues if i.severity == PRESENTATION_SEVERITY.WARNING)

    def checklist_ok_count(self) -> int:
        return sum(1 for item in self.checklist_items if item.status == "OK")

    def checklist_error_count(self) -> int:
        return sum(1 for item in self.checklist_items if item.status == "ERROR")

    def is_success(self) -> bool:
        """True si no hay issues con severidad ERROR. PDF y DOCX final no bloquean."""
        return self.error_count() == 0

    def to_dict(self) -> dict:
        return {
            "expediente_id": self.expediente_id,
            "status": self.status,
            "metadata": self.metadata.to_dict(),
            "checklist_items": [item.to_dict() for item in self.checklist_items],
            "issues": [i.to_dict() for i in self.issues],
            "generated_files": self.generated_files,
            "warnings": self.warnings,
            "notes": self.notes,
            "error_count": self.error_count(),
            "warning_count": self.warning_count(),
            "checklist_ok_count": self.checklist_ok_count(),
            "checklist_error_count": self.checklist_error_count(),
            "is_success": self.is_success(),
            "administrative_ready": False,
        }

    def summary(self) -> str:
        lines = [
            f"Preparacion: {self.status}",
            f"  Expediente        : {self.expediente_id}",
            f"  Errores           : {self.error_count()}",
            f"  Advertencias      : {self.warning_count()}",
            f"  Checklist OK      : {self.checklist_ok_count()}/{len(self.checklist_items)}",
            f"  Checklist ERROR   : {self.checklist_error_count()}",
            f"  Archivos generados: {len(self.generated_files)}",
        ]
        if self.issues:
            for issue in self.issues[:5]:
                lines.append(f"  {issue.summary()}")
        result_str = "OK" if self.is_success() else "CON ERRORES"
        lines.append(f"  RESULTADO         : {result_str}")
        lines.append(
            "  AVISO: Esta preparacion no declara el expediente"
            " apto para presentacion administrativa."
        )
        return "\n".join(lines)


# ---------------------------------------------------------------------------
# Funciones de soporte
# ---------------------------------------------------------------------------

def safe_load_json(path: "str | Path") -> "dict | None":
    """Carga JSON de forma tolerante. Devuelve None si no existe o falla."""
    p = Path(path)
    if not p.exists():
        return None
    try:
        with open(p, encoding="utf-8") as fh:
            data = json.load(fh)
        if isinstance(data, dict):
            return data
        return None
    except Exception:
        return None


def build_document_metadata(
    expediente_path: "str | Path",
) -> DocumentMetadata:
    """
    Construye los metadatos documentales del expediente.
    Lee JSONs de resultado si existen; no falla por archivos ausentes.
    administrative_ready es siempre False.
    """
    exp = Path(expediente_path)
    expediente_id = exp.name
    generated_at = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")

    # Leer estados de resultados previos
    audit_data = safe_load_json(exp / "auditoria" / "final_audit_result.json")
    qc_data = safe_load_json(exp / "documento" / "document_quality_result.json")
    pkg_data = safe_load_json(exp / "documento" / "package_build_result.json")
    exp_data = safe_load_json(exp / "documento" / "document_export_result.json")
    cc_data = safe_load_json(exp / "auditoria" / "conditional_chain_result.json")

    final_audit_status = audit_data.get("status") if audit_data else None
    document_qc_status = qc_data.get("status") if qc_data else None
    conditional_chain_status = cc_data.get("status") if cc_data else None
    package_status = (
        "GENERADO" if (pkg_data and pkg_data.get("generated")) else
        ("PENDIENTE" if pkg_data else None)
    )
    export_status = (
        "GENERADO" if (exp_data and exp_data.get("zip_generated")) else
        ("PENDIENTE" if exp_data else None)
    )

    # Detectar rutas de archivos
    enriched_docx = exp / "documento" / "documento_ambiental_borrador_con_figuras.docx"
    base_docx = exp / "documento" / "documento_ambiental_borrador.docx"
    final_docx = exp / "documento" / FINAL_REVIEW_DOCX
    source_md = exp / "documento" / "documento_ambiental_borrador.md"
    zip_path = exp / "documento" / "paquete_entrega.zip"

    if final_docx.exists():
        source_docx_path = str(final_docx)
    elif enriched_docx.exists():
        source_docx_path = str(enriched_docx)
    elif base_docx.exists():
        source_docx_path = str(base_docx)
    else:
        source_docx_path = None

    notes = [
        "Metadatos generados automaticamente. Requieren revision tecnica.",
        "administrative_ready=False: este modulo no declara aptitud administrativa.",
    ]
    warnings: list[str] = []
    if not source_docx_path:
        warnings.append("No se encontro DOCX fuente en documento/.")
    if not (source_md.exists()):
        warnings.append("No se encontro Markdown fuente en documento/.")

    return DocumentMetadata(
        expediente_id=expediente_id,
        generated_at=generated_at,
        source_docx=source_docx_path,
        source_markdown=str(source_md) if source_md.exists() else None,
        package_zip=str(zip_path) if zip_path.exists() else None,
        final_audit_status=final_audit_status,
        document_qc_status=document_qc_status,
        package_status=package_status,
        export_status=export_status,
        conditional_chain_status=conditional_chain_status,
        notes=notes,
        warnings=warnings,
    )


# ---------------------------------------------------------------------------
# Hoja de firmas
# ---------------------------------------------------------------------------

def build_signature_sheet_markdown(metadata: DocumentMetadata) -> str:
    """Genera la hoja de firmas en Markdown."""
    lines: list[str] = [
        "# Hoja de firmas y revision tecnica",
        "",
        "---",
        "",
        "## 1. Expediente",
        "",
        f"**ID de expediente:** {metadata.expediente_id}",
        f"**Fecha de generacion de metadatos:** {metadata.generated_at}",
        "",
        "---",
        "",
        "## 2. Documento revisado",
        "",
        f"**Documento DOCX:** {metadata.source_docx or '_(no disponible)_'}",
        f"**Documento Markdown:** {metadata.source_markdown or '_(no disponible)_'}",
        f"**Paquete ZIP:** {metadata.package_zip or '_(no disponible)_'}",
        "",
        "---",
        "",
        "## 3. Tecnico redactor/revisor",
        "",
        "| Campo | Valor |",
        "|-------|-------|",
        "| Nombre y apellidos | &nbsp; |",
        "| Titulacion | &nbsp; |",
        "| N. colegiado, si procede | &nbsp; |",
        "| Entidad/empresa | &nbsp; |",
        "| Cargo | &nbsp; |",
        "",
        "---",
        "",
        "## 4. Fecha de revision",
        "",
        "| Campo | Valor |",
        "|-------|-------|",
        "| Fecha de revision | &nbsp; |",
        "| Lugar | &nbsp; |",
        "",
        "---",
        "",
        "## 5. Firma",
        "",
        "_Espacio reservado para firma manuscrita o sello._",
        "",
        "&nbsp;",
        "",
        "&nbsp;",
        "",
        "&nbsp;",
        "",
        "---",
        "",
        "## 6. Advertencia",
        "",
        "> **Esta hoja no acredita por si sola la aptitud administrativa del expediente.**",
        "> El presente documento es un borrador tecnico para revision interna.",
        "> La presentacion ante la Administracion requiere firma tecnica/juridica completa",
        "> y validacion por tecnico competente habilitado.",
        "> El promotor es responsable de la presentacion del Documento Ambiental.",
        "",
    ]
    return "\n".join(lines)


# ---------------------------------------------------------------------------
# Checklist de presentacion
# ---------------------------------------------------------------------------

def build_presentation_checklist(
    expediente_path: "str | Path",
    metadata: DocumentMetadata,
) -> list[PresentationChecklistItem]:
    """
    Construye el checklist de preparacion para presentacion/revision.
    No modifica el expediente.
    """
    exp = Path(expediente_path)
    items: list[PresentationChecklistItem] = []

    # CHK-001: DOCX final/revisable existe
    enriched_docx = exp / "documento" / "documento_ambiental_borrador_con_figuras.docx"
    base_docx = exp / "documento" / "documento_ambiental_borrador.docx"
    final_docx = exp / "documento" / FINAL_REVIEW_DOCX

    if final_docx.exists():
        docx_status = "OK"
        docx_evidence = [str(final_docx)]
        docx_rec = ""
    elif enriched_docx.exists():
        docx_status = "OK"
        docx_evidence = [str(enriched_docx)]
        docx_rec = "Se puede generar documento_ambiental_final_revisable.docx con --write."
    elif base_docx.exists():
        docx_status = "WARNING"
        docx_evidence = [str(base_docx)]
        docx_rec = "El DOCX sin figuras es menos completo. Ejecutar document-insert-figures --write."
    else:
        docx_status = "ERROR"
        docx_evidence = []
        docx_rec = "Ejecutar document-build-docx --write y document-insert-figures --write."
    items.append(PresentationChecklistItem(
        item_id="CHK-001",
        description="Documento DOCX final/revisable existe",
        status=docx_status,
        evidence=docx_evidence,
        recommendation=docx_rec,
    ))

    # CHK-002: Markdown fuente existe
    md_src = exp / "documento" / "documento_ambiental_borrador.md"
    items.append(PresentationChecklistItem(
        item_id="CHK-002",
        description="Markdown fuente existe",
        status="OK" if md_src.exists() else "WARNING",
        evidence=[str(md_src)] if md_src.exists() else [],
        recommendation="" if md_src.exists() else
            "Ejecutar document-build-md --write para generar el Markdown.",
    ))

    # CHK-003: QC documental existe
    qc_json = exp / "documento" / "document_quality_result.json"
    items.append(PresentationChecklistItem(
        item_id="CHK-003",
        description="QC documental existe",
        status="OK" if qc_json.exists() else "WARNING",
        evidence=[str(qc_json)] if qc_json.exists() else [],
        recommendation="" if qc_json.exists() else
            "Ejecutar document-qc --write para generar el informe de QC.",
    ))

    # CHK-004: QC documental sin ERROR
    if metadata.document_qc_status is not None:
        if metadata.document_qc_status in ("OK", "CON_OBSERVACIONES"):
            qc_check_status = "OK"
            qc_check_evidence = [f"QC status: {metadata.document_qc_status}"]
            qc_check_rec = ""
        elif metadata.document_qc_status == "NO_CONFORME":
            qc_check_status = "ERROR"
            qc_check_evidence = [f"QC status: {metadata.document_qc_status}"]
            qc_check_rec = "Revisar document_quality_result.md para identificar los errores de QC."
        else:
            qc_check_status = "WARNING"
            qc_check_evidence = [f"QC status: {metadata.document_qc_status}"]
            qc_check_rec = "Estado de QC inesperado. Revisar document_quality_result.md."
    else:
        qc_check_status = "WARNING"
        qc_check_evidence = []
        qc_check_rec = "QC no ejecutado. Ejecutar document-qc --write."
    items.append(PresentationChecklistItem(
        item_id="CHK-004",
        description="QC documental sin ERROR",
        status=qc_check_status,
        evidence=qc_check_evidence,
        recommendation=qc_check_rec,
    ))

    # CHK-005: Auditoria final existe
    audit_json = exp / "auditoria" / "final_audit_result.json"
    items.append(PresentationChecklistItem(
        item_id="CHK-005",
        description="Auditoria final existe",
        status="OK" if audit_json.exists() else "WARNING",
        evidence=[str(audit_json)] if audit_json.exists() else [],
        recommendation="" if audit_json.exists() else
            "Ejecutar audit-final --write para generar la auditoria.",
    ))

    # CHK-006: Auditoria final no oculta estado NO_CONFORME
    audit_data = safe_load_json(audit_json)
    if audit_data is not None:
        audit_status_val = audit_data.get("status", "")
        if audit_status_val == "NO_CONFORME":
            au_check_status = "WARNING"
            au_check_evidence = [f"Auditoria status: {audit_status_val}"]
            au_check_rec = "La auditoria es NO_CONFORME. Revisar final_audit_result.md."
        else:
            au_check_status = "OK"
            au_check_evidence = [f"Auditoria status: {audit_status_val}"]
            au_check_rec = ""
    else:
        au_check_status = "WARNING"
        au_check_evidence = []
        au_check_rec = "Auditoria no disponible. Ejecutar audit-final --write."
    items.append(PresentationChecklistItem(
        item_id="CHK-006",
        description="Auditoria final: estado visible y no oculto",
        status=au_check_status,
        evidence=au_check_evidence,
        recommendation=au_check_rec,
    ))

    # CHK-007: Paquete ZIP existe
    zip_path = exp / "documento" / "paquete_entrega.zip"
    items.append(PresentationChecklistItem(
        item_id="CHK-007",
        description="Paquete ZIP existe",
        status="OK" if zip_path.exists() else "WARNING",
        evidence=[str(zip_path)] if zip_path.exists() else [],
        recommendation="" if zip_path.exists() else
            "Ejecutar document-export --write para generar el ZIP.",
    ))

    # CHK-008: README_ENTREGA existe en paquete
    readme_path = exp / "documento" / "paquete_entrega" / "README_ENTREGA.md"
    items.append(PresentationChecklistItem(
        item_id="CHK-008",
        description="README_ENTREGA existe en paquete",
        status="OK" if readme_path.exists() else "WARNING",
        evidence=[str(readme_path)] if readme_path.exists() else [],
        recommendation="" if readme_path.exists() else
            "Ejecutar document-package --write para regenerar el paquete.",
    ))

    # CHK-009: No consta administrative_ready=True en ningun JSON
    admin_ready_found = False
    admin_ready_sources: list[str] = []
    for json_candidate in [
        exp / "documento" / "document_export_result.json",
        exp / "documento" / "package_build_result.json",
        exp / "documento" / "document_quality_result.json",
    ]:
        data = safe_load_json(json_candidate)
        if data and data.get("administrative_ready") is True:
            admin_ready_found = True
            admin_ready_sources.append(str(json_candidate))

    items.append(PresentationChecklistItem(
        item_id="CHK-009",
        description="No consta administrative_ready=True en ningun JSON",
        status="ERROR" if admin_ready_found else "OK",
        evidence=admin_ready_sources,
        recommendation="" if not admin_ready_found else
            "Ningun modulo del pipeline debe declarar administrative_ready=True.",
    ))

    # CHK-010: Hoja de firmas generable (metadata suficiente)
    sig_generable = metadata.expediente_id and metadata.generated_at
    items.append(PresentationChecklistItem(
        item_id="CHK-010",
        description="Hoja de firmas generable",
        status="OK" if sig_generable else "WARNING",
        evidence=[f"expediente_id={metadata.expediente_id}"],
        recommendation="" if sig_generable else "Revisar metadata del expediente.",
    ))

    # CHK-011: Figuras/captions documentadas si existen
    figs_json = exp / "documento" / "document_figures_result.json"
    figs_data = safe_load_json(figs_json)
    if figs_data is not None:
        if figs_data.get("generated") is True:
            figs_status = "OK"
            figs_evidence = [f"Figuras generadas: {figs_data.get('figures_inserted', 0)}"]
            figs_rec = ""
        else:
            figs_status = "WARNING"
            figs_evidence = [str(figs_json)]
            figs_rec = "document_figures_result indica generated=False. Revisar."
    else:
        figs_status = "NO_APLICA"
        figs_evidence = []
        figs_rec = "No se ejecuto document-insert-figures."
    items.append(PresentationChecklistItem(
        item_id="CHK-011",
        description="Figuras/captions documentadas si existen",
        status=figs_status,
        evidence=figs_evidence,
        recommendation=figs_rec,
    ))

    # CHK-012: No se detectan frases de aptitud administrativa indebida en MD fuente
    admin_phrases_found: list[str] = []
    if md_src.exists():
        try:
            md_text = md_src.read_text(encoding="utf-8", errors="replace").lower()
            for phrase in _ADMIN_READY_PHRASES:
                if phrase in md_text:
                    admin_phrases_found.append(phrase)
        except Exception:
            pass

    items.append(PresentationChecklistItem(
        item_id="CHK-012",
        description="No se detectan frases de aptitud administrativa indebida",
        status="WARNING" if admin_phrases_found else "OK",
        evidence=admin_phrases_found,
        recommendation="" if not admin_phrases_found else
            "Revisar el borrador MD y eliminar afirmaciones de aptitud administrativa.",
    ))

    # CHK-013: Auditoria de cadenas condicionales IM-09 revisada
    # Diseño: WARNING si no existe (no bloquea; es opcional post PIPE-04).
    # WARNING si NO_CONFORME (coherente con CHK-006 que tambien usa WARNING para
    # audit-final NO_CONFORME; el caracter del checklist es advisory, no bloqueante;
    # los errores graves quedan capturados en QC-E009 y en el AVISO del Bloque I).
    cc_json = exp / "auditoria" / "conditional_chain_result.json"
    cc_meta = safe_load_json(cc_json)
    if cc_meta is not None:
        cc_status_val = cc_meta.get("status", "")
        if cc_status_val == "NO_CONFORME":
            chk13_status = "WARNING"
            chk13_evidence = [f"IM-09 status: {cc_status_val}"]
            chk13_rec = (
                "IM-09 NO_CONFORME: revisar cadenas condicionales antes de cerrar."
            )
        else:
            chk13_status = "OK"
            chk13_evidence = [f"IM-09 status: {cc_status_val}"]
            chk13_rec = ""
    else:
        chk13_status = "WARNING"
        chk13_evidence = []
        chk13_rec = (
            "No se encontro auditoria/conditional_chain_result.json. "
            "Ejecutar audit-conditional-chains --write."
        )
    items.append(PresentationChecklistItem(
        item_id="CHK-013",
        description="Auditoria de cadenas condicionales IM-09 revisada",
        status=chk13_status,
        evidence=chk13_evidence,
        recommendation=chk13_rec,
    ))

    return items


# ---------------------------------------------------------------------------
# Generadores de texto
# ---------------------------------------------------------------------------

def build_metadata_markdown(metadata: DocumentMetadata) -> str:
    """Genera el resumen de metadatos en Markdown."""
    lines: list[str] = [
        "# Metadatos documentales",
        "",
        f"**Expediente:** {metadata.expediente_id}",
        f"**Generado:** {metadata.generated_at}",
        f"**administrative_ready:** {metadata.administrative_ready}",
        "",
        "---",
        "",
        "## 1. Archivos fuente",
        "",
        f"- **DOCX:** {metadata.source_docx or '_no disponible_'}",
        f"- **Markdown:** {metadata.source_markdown or '_no disponible_'}",
        f"- **ZIP paquete:** {metadata.package_zip or '_no disponible_'}",
        "",
        "## 2. Estados de auditoria/QC",
        "",
        "| Modulo | Estado |",
        "|--------|--------|",
        f"| Auditoria final | {metadata.final_audit_status or 'N/D'} |",
        f"| QC documental | {metadata.document_qc_status or 'N/D'} |",
        f"| Paquete | {metadata.package_status or 'N/D'} |",
        f"| Exportacion | {metadata.export_status or 'N/D'} |",
        f"| IM-09 cadenas condicionales | {metadata.conditional_chain_status or 'N/D'} |",
        "",
        "## 3. Advertencias",
        "",
    ]
    if metadata.warnings:
        for w in metadata.warnings:
            lines.append(f"- {w}")
    else:
        lines.append("_Sin advertencias._")

    lines += [
        "",
        "## 4. Notas",
        "",
    ]
    for note in metadata.notes:
        lines.append(f"- {note}")

    lines += [
        "",
        "---",
        "",
        "> **Este modulo no declara aptitud administrativa.**",
        "> `administrative_ready=False` siempre.",
        "",
    ]
    return "\n".join(lines)


def build_presentation_checklist_markdown(
    items: list[PresentationChecklistItem],
) -> str:
    """Genera el checklist de presentacion en Markdown."""
    ok = sum(1 for i in items if i.status == "OK")
    warn = sum(1 for i in items if i.status == "WARNING")
    err = sum(1 for i in items if i.status == "ERROR")
    na = sum(1 for i in items if i.status == "NO_APLICA")

    lines: list[str] = [
        "# Checklist de preparacion para revision",
        "",
        f"**Total items:** {len(items)}",
        f"**OK:** {ok} | **WARNING:** {warn} | **ERROR:** {err} | **NO_APLICA:** {na}",
        "",
        "---",
        "",
        "| ID | Descripcion | Estado | Recomendacion |",
        "|----|-------------|--------|---------------|",
    ]

    for item in items:
        rec = item.recommendation or ""
        lines.append(
            f"| {item.item_id} | {item.description} | {item.status} | {rec} |"
        )

    lines += [
        "",
        "---",
        "",
        "## Detalle por item",
        "",
    ]
    for item in items:
        lines.append(f"### {item.item_id} — {item.description}")
        lines.append(f"**Estado:** {item.status}")
        if item.evidence:
            lines.append(f"**Evidencia:** {', '.join(str(e) for e in item.evidence)}")
        if item.recommendation:
            lines.append(f"**Recomendacion:** {item.recommendation}")
        lines.append("")

    lines += [
        "---",
        "",
        "> **Este checklist no declara el expediente apto para presentacion administrativa.**",
        "> Es una herramienta de revision tecnica interna.",
        "",
    ]
    return "\n".join(lines)


# ---------------------------------------------------------------------------
# DOCX final revisable
# ---------------------------------------------------------------------------

def append_signature_sheet_to_docx(
    input_docx_path: "str | Path",
    output_docx_path: "str | Path",
    signature_markdown: str,
) -> bool:
    """
    Crea un DOCX final con la hoja de firmas anadida.
    No modifica el input_docx_path (copia primero, luego anade).
    Devuelve True si el output se genero correctamente.
    """
    try:
        import shutil as _shutil
        from docx import Document
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        return False

    inp = Path(input_docx_path)
    out = Path(output_docx_path)

    if not inp.exists():
        return False

    try:
        out.parent.mkdir(parents=True, exist_ok=True)
        _shutil.copy2(str(inp), str(out))

        doc = Document(str(out))

        # Salto de pagina
        para_break = doc.add_paragraph()
        run_break = para_break.add_run()
        br = OxmlElement("w:br")
        br.set(qn("w:type"), "page")
        run_break._r.append(br)

        # Heading de hoja de firmas
        doc.add_heading("Hoja de firmas y revision tecnica", level=1)

        # Seccion Expediente
        doc.add_heading("1. Expediente", level=2)
        para = doc.add_paragraph()
        para.add_run(f"ID de expediente: ").bold = True
        para.add_run(signature_markdown.split("**ID de expediente:**")[1].split("\n")[0].strip()
                     if "**ID de expediente:**" in signature_markdown else "")

        # Seccion Tecnico redactor
        doc.add_heading("2. Tecnico redactor/revisor", level=2)
        for label in [
            "Nombre y apellidos:",
            "Titulacion:",
            "N. colegiado, si procede:",
            "Entidad/empresa:",
            "Cargo:",
        ]:
            p = doc.add_paragraph()
            p.add_run(f"{label} ").bold = True
            p.add_run("_" * 40)

        # Seccion Fecha
        doc.add_heading("3. Fecha de revision", level=2)
        for label in ["Fecha de revision:", "Lugar:"]:
            p = doc.add_paragraph()
            p.add_run(f"{label} ").bold = True
            p.add_run("_" * 30)

        # Seccion Firma
        doc.add_heading("4. Firma", level=2)
        for _ in range(4):
            doc.add_paragraph("")

        # Advertencia final
        doc.add_heading("5. Advertencia", level=2)
        adv = doc.add_paragraph()
        adv.add_run(
            "Esta hoja no acredita por si sola la aptitud administrativa del expediente. "
            "El presente documento es un borrador tecnico para revision interna. "
            "La presentacion ante la Administracion requiere firma tecnica/juridica "
            "completa y validacion por tecnico competente habilitado."
        ).italic = True

        doc.save(str(out))
        return out.exists() and out.stat().st_size > 0

    except Exception:
        return False


# ---------------------------------------------------------------------------
# Funcion de escritura de outputs
# ---------------------------------------------------------------------------

def write_presentation_outputs(
    result: PresentationPreparationResult,
    output_dir: "str | Path",
) -> list[Path]:
    """
    Escribe todos los outputs de presentacion en output_dir.
    Devuelve lista de rutas escritas.
    """
    out = Path(output_dir)
    out.mkdir(parents=True, exist_ok=True)

    written: list[Path] = []

    # Metadatos JSON
    meta_json = out / METADATA_JSON
    with open(meta_json, "w", encoding="utf-8") as fh:
        json.dump(result.metadata.to_dict(), fh, ensure_ascii=False, indent=2)
    written.append(meta_json)

    # Metadatos MD
    meta_md = out / METADATA_MD
    meta_md.write_text(build_metadata_markdown(result.metadata), encoding="utf-8")
    written.append(meta_md)

    # Hoja de firmas MD
    sig_md = out / SIGNATURE_SHEET_MD
    sig_md.write_text(
        build_signature_sheet_markdown(result.metadata), encoding="utf-8"
    )
    written.append(sig_md)

    # Checklist JSON
    chk_json = out / PRESENTATION_CHECKLIST_JSON
    with open(chk_json, "w", encoding="utf-8") as fh:
        json.dump(
            {"checklist": [i.to_dict() for i in result.checklist_items]},
            fh,
            ensure_ascii=False,
            indent=2,
        )
    written.append(chk_json)

    # Checklist MD
    chk_md = out / PRESENTATION_CHECKLIST_MD
    chk_md.write_text(
        build_presentation_checklist_markdown(result.checklist_items),
        encoding="utf-8",
    )
    written.append(chk_md)

    return written


# ---------------------------------------------------------------------------
# Funcion principal
# ---------------------------------------------------------------------------

def prepare_document_for_presentation(
    expediente_path: "str | Path",
    write_outputs: bool = False,
    create_final_docx: bool = True,
) -> PresentationPreparationResult:
    """
    Prepara el documento para revision y presentacion administrativa.

    Si write_outputs=False (default): analiza, no escribe nada.
    Si write_outputs=True: escribe metadatos, hoja de firmas, checklist.
      Si create_final_docx=True y hay DOCX enriquecido: crea DOCX final revisable.

    No modifica ninguna fuente existente.
    No declara aptitud administrativa (administrative_ready=False siempre).
    """
    exp = Path(expediente_path)
    expediente_id = exp.name
    doc_dir = exp / "documento"
    issues: list[PresentationIssue] = []

    # Construir metadata
    metadata = build_document_metadata(exp)

    # Warnings de metadata como issues INFO
    for w in metadata.warnings:
        issues.append(PresentationIssue(
            severity=PRESENTATION_SEVERITY.WARNING,
            code="PP-W001",
            message=w,
            recommendation="Ejecutar el pipeline documental completo.",
            evidence=[],
        ))

    # Construir checklist
    checklist_items = build_presentation_checklist(exp, metadata)

    # Issues derivados del checklist
    for item in checklist_items:
        if item.status == "ERROR":
            issues.append(PresentationIssue(
                severity=PRESENTATION_SEVERITY.ERROR,
                code=f"PP-E-{item.item_id}",
                message=f"{item.item_id}: {item.description}",
                recommendation=item.recommendation,
                evidence=item.evidence,
            ))
        elif item.status == "WARNING":
            issues.append(PresentationIssue(
                severity=PRESENTATION_SEVERITY.WARNING,
                code=f"PP-W-{item.item_id}",
                message=f"{item.item_id}: {item.description}",
                recommendation=item.recommendation,
                evidence=item.evidence,
            ))

    # Determinar status global
    error_count = sum(1 for i in issues if i.severity == PRESENTATION_SEVERITY.ERROR)
    chk_errors = sum(1 for i in checklist_items if i.status == "ERROR")
    chk_warnings = sum(1 for i in checklist_items if i.status == "WARNING")

    if error_count == 0 and chk_errors == 0 and chk_warnings == 0:
        status = PRESENTATION_STATUS.PREPARADO_PARA_REVISION
    elif error_count > 0 or chk_errors > 0:
        if not metadata.source_docx:
            status = PRESENTATION_STATUS.PENDIENTE_DOCUMENTACION
        else:
            status = PRESENTATION_STATUS.NO_PREPARADO
    else:
        status = PRESENTATION_STATUS.PENDIENTE_REVISION_TECNICA

    generated_files: list[str] = []
    notes = [
        "Este modulo no declara aptitud administrativa.",
        "PREPARADO_PARA_REVISION significa preparado para revision tecnica interna.",
        "La presentacion administrativa requiere firma y validacion tecnica/juridica.",
    ]

    if not write_outputs:
        return PresentationPreparationResult(
            expediente_id=expediente_id,
            status=status,
            metadata=metadata,
            checklist_items=checklist_items,
            issues=issues,
            generated_files=[],
            warnings=[],
            notes=notes + ["Modo dry-run: no se han escrito archivos."],
        )

    # Modo escritura
    written = write_presentation_outputs(
        PresentationPreparationResult(
            expediente_id=expediente_id,
            status=status,
            metadata=metadata,
            checklist_items=checklist_items,
            issues=issues,
            generated_files=[],
            warnings=[],
            notes=notes,
        ),
        doc_dir,
    )
    generated_files = [str(p) for p in written]

    # DOCX final revisable (best-effort)
    final_docx_path = doc_dir / FINAL_REVIEW_DOCX
    if create_final_docx and metadata.source_docx:
        src_docx = Path(metadata.source_docx)
        # No sobreescribir si la fuente ya ES el docx final
        if src_docx.resolve() != final_docx_path.resolve():
            ok = append_signature_sheet_to_docx(
                src_docx,
                final_docx_path,
                build_signature_sheet_markdown(metadata),
            )
            if ok:
                generated_files.append(str(final_docx_path))
            else:
                issues.append(PresentationIssue(
                    severity=PRESENTATION_SEVERITY.WARNING,
                    code="PP-W002",
                    message="No se pudo generar el DOCX final revisable.",
                    recommendation=(
                        "Verificar que python-docx esta instalado y que el DOCX fuente es valido."
                    ),
                    evidence=[str(src_docx)],
                ))

    return PresentationPreparationResult(
        expediente_id=expediente_id,
        status=status,
        metadata=metadata,
        checklist_items=checklist_items,
        issues=issues,
        generated_files=generated_files,
        warnings=[],
        notes=notes,
    )
