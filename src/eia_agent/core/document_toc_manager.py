"""
document_toc_manager.py — EN-05
Gestion del indice automatico (TOC) del DOCX final.

Detecta si un DOCX contiene un campo TOC Word, inserta o reemplaza el TOC
en una copia del DOCX, y marca los campos para actualizacion al abrir.

Principios no negociables:
  - No modifica el DOCX original.
  - No llama a Word COM, LibreOffice, ni ningun conversor externo.
  - No recalcula numeros de pagina (eso lo hace Word/LibreOffice al abrir).
  - No usa IA.
  - No llama a servicios externos.
  - 100 % offline.
"""
from __future__ import annotations

import json
import re
import shutil
import tempfile
import zipfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Optional, Tuple

# ---------------------------------------------------------------------------
# Constantes publicas
# ---------------------------------------------------------------------------

TOC_STATUS = {
    "OK": "OK",
    "CON_OBSERVACIONES": "CON_OBSERVACIONES",
    "NO_CONFORME": "NO_CONFORME",
    "SIN_DATOS": "SIN_DATOS",
}

TOC_SEVERITY = {
    "ERROR": "ERROR",
    "WARNING": "WARNING",
    "INFO": "INFO",
}

TOC_OUTPUT_DOCX = "documento_ambiental_con_toc.docx"
TOC_RESULT_JSON = "document_toc_result.json"
TOC_RESULT_MD = "document_toc_result.md"

DEFAULT_TOC_INSTRUCTION: str = r'TOC \o "1-3" \h \z \u'

# Palabras clave que identifican un parrafo como posible placeholder del TOC
_TOC_PLACEHOLDER_KEYWORDS: List[str] = [
    "índice",
    "indice",
    "tabla de contenido",
    "tabla de contenidos",
    "toc",
]

_DOCX_CANDIDATES: List[str] = [
    "documento/documento_ambiental_numerado.docx",
    "documento/documento_ambiental_final_revisable.docx",
    "documento/documento_ambiental_estructurado.docx",
    "documento/documento_ambiental_borrador_con_figuras.docx",
    "documento/documento_ambiental_borrador.docx",
]

# Longitud maxima de un parrafo para ser candidato a placeholder
_TOC_PLACEHOLDER_MAX_LEN = 80


# ---------------------------------------------------------------------------
# Dataclasses
# ---------------------------------------------------------------------------


@dataclass
class TocIssue:
    severity: str
    code: str
    message: str
    recommendation: str = ""
    evidence: List[str] = field(default_factory=list)

    def to_dict(self) -> dict:
        return {
            "severity": self.severity,
            "code": self.code,
            "message": self.message,
            "recommendation": self.recommendation,
            "evidence": self.evidence,
        }

    def summary(self) -> str:
        return f"[{self.severity}] {self.code}: {self.message}"


@dataclass
class TocDetectionResult:
    docx_path: str
    has_toc: bool
    update_fields_enabled: bool
    toc_paragraph_count: int
    has_settings_xml: bool
    issues: List[TocIssue] = field(default_factory=list)
    notes: List[str] = field(default_factory=list)

    def to_dict(self) -> dict:
        return {
            "docx_path": self.docx_path,
            "has_toc": self.has_toc,
            "update_fields_enabled": self.update_fields_enabled,
            "toc_paragraph_count": self.toc_paragraph_count,
            "has_settings_xml": self.has_settings_xml,
            "issues": [i.to_dict() for i in self.issues],
            "notes": self.notes,
        }

    def summary(self) -> str:
        toc_str = "si" if self.has_toc else "no"
        uf_str = "si" if self.update_fields_enabled else "no"
        return (
            f"TOC detectado: {toc_str} | "
            f"updateFields: {uf_str} | "
            f"parrafos TOC: {self.toc_paragraph_count}"
        )


@dataclass
class TocProcessResult:
    input_docx: str
    output_docx: Optional[str]
    status: str
    toc_inserted: bool
    toc_replaced: bool
    update_fields_set: bool
    placeholder_paragraphs_found: int
    issues: List[TocIssue] = field(default_factory=list)
    warnings: List[str] = field(default_factory=list)
    notes: List[str] = field(default_factory=list)

    def error_count(self) -> int:
        return sum(1 for i in self.issues if i.severity == "ERROR")

    def warning_count(self) -> int:
        return sum(1 for i in self.issues if i.severity == "WARNING") + len(self.warnings)

    def is_valid(self) -> bool:
        return self.error_count() == 0

    def to_dict(self) -> dict:
        return {
            "input_docx": self.input_docx,
            "output_docx": self.output_docx,
            "status": self.status,
            "toc_inserted": self.toc_inserted,
            "toc_replaced": self.toc_replaced,
            "update_fields_set": self.update_fields_set,
            "placeholder_paragraphs_found": self.placeholder_paragraphs_found,
            "issues": [i.to_dict() for i in self.issues],
            "warnings": self.warnings,
            "notes": self.notes,
        }

    def summary(self) -> str:
        action = (
            "reemplazado" if self.toc_replaced
            else ("insertado" if self.toc_inserted else "no insertado")
        )
        uf = "si" if self.update_fields_set else "no"
        return (
            f"[EN-05] TOC: {action} | updateFields: {uf} | "
            f"estado: {self.status} | errores: {self.error_count()} | "
            f"avisos: {self.warning_count()}"
        )


# ---------------------------------------------------------------------------
# Validacion basica del DOCX
# ---------------------------------------------------------------------------


def validate_docx_file(path) -> bool:
    """True si el archivo existe, tiene tamano > 0 y se puede abrir como DOCX."""
    p = Path(path)
    if not p.exists() or not p.is_file():
        return False
    if p.stat().st_size == 0:
        return False
    try:
        from docx import Document
        Document(str(p))
        return True
    except Exception:
        return False


# ---------------------------------------------------------------------------
# Deteccion de TOC en DOCX (inspeccion via zipfile)
# ---------------------------------------------------------------------------


def detect_toc_in_docx(path) -> TocDetectionResult:
    """
    Inspecciona el DOCX via zipfile para detectar campos TOC y updateFields.
    No modifica el archivo.
    """
    p = Path(path)
    issues: List[TocIssue] = []
    notes: List[str] = []

    if not p.exists() or not p.is_file():
        issues.append(TocIssue(
            severity=TOC_SEVERITY["ERROR"],
            code="EN05-E001",
            message=f"DOCX no encontrado: {p}",
        ))
        return TocDetectionResult(
            docx_path=str(p),
            has_toc=False,
            update_fields_enabled=False,
            toc_paragraph_count=0,
            has_settings_xml=False,
            issues=issues,
        )

    has_toc = False
    update_fields_enabled = False
    toc_paragraph_count = 0
    has_settings_xml = False

    try:
        with zipfile.ZipFile(str(p), "r") as zf:
            names = zf.namelist()
            has_settings_xml = "word/settings.xml" in names

            if "word/document.xml" in names:
                doc_xml = zf.read("word/document.xml").decode("utf-8", errors="replace")
                toc_paragraph_count = _count_toc_fields(doc_xml)
                has_toc = toc_paragraph_count > 0

            if has_settings_xml:
                settings_xml = zf.read("word/settings.xml").decode("utf-8", errors="replace")
                update_fields_enabled = _has_update_fields_true(settings_xml)

    except zipfile.BadZipFile:
        issues.append(TocIssue(
            severity=TOC_SEVERITY["ERROR"],
            code="EN05-E001",
            message=f"DOCX corrupto o no es un archivo ZIP valido: {p}",
        ))
        return TocDetectionResult(
            docx_path=str(p),
            has_toc=False,
            update_fields_enabled=False,
            toc_paragraph_count=0,
            has_settings_xml=False,
            issues=issues,
        )
    except Exception as exc:
        issues.append(TocIssue(
            severity=TOC_SEVERITY["ERROR"],
            code="EN05-E001",
            message=f"Error al inspeccionar DOCX: {exc}",
        ))

    if not has_toc:
        notes.append("No se detecto campo TOC en el documento.")
    if not update_fields_enabled:
        notes.append(
            "updateFields no habilitado; el TOC no se actualizara automaticamente al abrir."
        )

    return TocDetectionResult(
        docx_path=str(p),
        has_toc=has_toc,
        update_fields_enabled=update_fields_enabled,
        toc_paragraph_count=toc_paragraph_count,
        has_settings_xml=has_settings_xml,
        issues=issues,
        notes=notes,
    )


def _count_toc_fields(doc_xml: str) -> int:
    """Cuenta referencias a campos TOC en el XML del documento."""
    count = 0
    for match in re.finditer(r"<w:instrText[^>]*>([^<]*)</w:instrText>", doc_xml):
        if "TOC" in match.group(1):
            count += 1
    for match in re.finditer(r'<w:fldSimple[^>]+w:instr="([^"]*)"', doc_xml):
        if "TOC" in match.group(1):
            count += 1
    return count


def _has_update_fields_true(settings_xml: str) -> bool:
    """True si settings.xml contiene w:updateFields con val true/1/on."""
    pattern = r'<w:updateFields[^>]+w:val="(true|1|on)"'
    return bool(re.search(pattern, settings_xml, re.IGNORECASE))


# ---------------------------------------------------------------------------
# Modificacion de settings.xml para updateFields
# ---------------------------------------------------------------------------


def enable_update_fields_on_open(docx_path, output_path) -> bool:
    """
    Crea una copia del DOCX en output_path con w:updateFields w:val="true" en settings.xml.
    Si docx_path == output_path, modifica el archivo en sitio via archivo temporal.
    Retorna True si la operacion tuvo exito.
    """
    src = Path(docx_path)
    dst = Path(output_path)

    if not src.exists() or not src.is_file():
        return False

    same_file = src.resolve() == dst.resolve()
    if same_file:
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tf:
            tmp_path = Path(tf.name)
        try:
            _copy_docx_with_update_fields(src, tmp_path)
            shutil.move(str(tmp_path), str(dst))
        except Exception:
            if tmp_path.exists():
                tmp_path.unlink()
            return False
    else:
        try:
            _copy_docx_with_update_fields(src, dst)
        except Exception:
            return False

    return True


def _copy_docx_with_update_fields(src: Path, dst: Path) -> None:
    """Lee src como ZIP, modifica settings.xml y escribe dst."""
    with zipfile.ZipFile(str(src), "r") as zin:
        names = zin.namelist()
        has_settings = "word/settings.xml" in names

        with zipfile.ZipFile(str(dst), "w") as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "word/settings.xml":
                    data = _inject_update_fields(data)
                zout.writestr(item, data)

            if not has_settings:
                minimal = (
                    "<?xml version='1.0' encoding='UTF-8' standalone='yes'?>"
                    '<w:settings xmlns:w="http://schemas.openxmlformats.org/'
                    'wordprocessingml/2006/main">'
                    '<w:updateFields w:val="true"/>'
                    "</w:settings>"
                )
                zout.writestr("word/settings.xml", minimal.encode("utf-8"))


def _inject_update_fields(xml_bytes: bytes) -> bytes:
    """Modifica el XML de settings.xml para incluir w:updateFields w:val="true"."""
    try:
        text = xml_bytes.decode("utf-8")
    except UnicodeDecodeError:
        text = xml_bytes.decode("latin-1", errors="replace")

    if "w:updateFields" in text:
        text = re.sub(
            r"<w:updateFields[^/]*/?>",
            '<w:updateFields w:val="true"/>',
            text,
        )
        return text.encode("utf-8")

    close_tag = "</w:settings>"
    if close_tag in text:
        text = text.replace(
            close_tag,
            '<w:updateFields w:val="true"/>\n' + close_tag,
        )
        return text.encode("utf-8")

    # Fallback poco probable: sin closing tag
    text += '<w:updateFields w:val="true"/>'
    return text.encode("utf-8")


# ---------------------------------------------------------------------------
# Construccion del campo TOC via python-docx
# ---------------------------------------------------------------------------


def build_toc_field_paragraph(doc):
    """
    Agrega al final del documento un parrafo con el campo TOC Word.
    Retorna el parrafo creado.
    python-docx NO calcula numeros de pagina; Word/LibreOffice deben actualizarlo.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    para = doc.add_paragraph()
    run = para.add_run()

    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")

    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = f" {DEFAULT_TOC_INSTRUCTION} "

    fldChar_separate = OxmlElement("w:fldChar")
    fldChar_separate.set(qn("w:fldCharType"), "separate")

    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")

    run._r.append(fldChar_begin)
    run._r.append(instrText)
    run._r.append(fldChar_separate)
    run._r.append(fldChar_end)

    return para


def _build_toc_run_element():
    """Crea el elemento XML de un run con el campo TOC. Para uso interno en reemplazos."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    run_elem = OxmlElement("w:r")

    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")

    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = f" {DEFAULT_TOC_INSTRUCTION} "

    fldChar_separate = OxmlElement("w:fldChar")
    fldChar_separate.set(qn("w:fldCharType"), "separate")

    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")

    run_elem.append(fldChar_begin)
    run_elem.append(instrText)
    run_elem.append(fldChar_separate)
    run_elem.append(fldChar_end)

    return run_elem


# ---------------------------------------------------------------------------
# Busqueda de placeholders de TOC
# ---------------------------------------------------------------------------


def find_toc_placeholder_paragraphs(doc) -> List[int]:
    """
    Retorna indices de parrafos que parecen ser placeholders del TOC.
    Busca textos cortos que coincidan con keywords de indice.
    """
    candidates = []
    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text or len(text) > _TOC_PLACEHOLDER_MAX_LEN:
            continue
        text_lower = text.lower()
        for kw in _TOC_PLACEHOLDER_KEYWORDS:
            if kw in text_lower:
                candidates.append(idx)
                break
    return candidates


# ---------------------------------------------------------------------------
# Insercion o reemplazo del TOC
# ---------------------------------------------------------------------------


def insert_or_replace_toc(
    input_docx,
    output_docx,
    replace_placeholder: bool = True,
) -> TocProcessResult:
    """
    Crea una copia de input_docx en output_docx con un campo TOC insertado.

    Si replace_placeholder=True y hay paragrafos placeholder, reemplaza el primero.
    Si no hay placeholder o replace_placeholder=False, inserta el TOC al inicio.
    No modifica input_docx.
    """
    from docx import Document
    from docx.oxml import OxmlElement

    src = Path(input_docx)
    dst = Path(output_docx)

    issues: List[TocIssue] = []
    warnings: List[str] = []
    notes: List[str] = []
    toc_inserted = False
    toc_replaced = False
    update_fields_set = False
    placeholder_count = 0

    if not src.exists() or not src.is_file():
        issues.append(TocIssue(
            severity=TOC_SEVERITY["ERROR"],
            code="EN05-E001",
            message=f"DOCX de entrada no encontrado: {src}",
        ))
        return TocProcessResult(
            input_docx=str(src),
            output_docx=None,
            status=TOC_STATUS["NO_CONFORME"],
            toc_inserted=False,
            toc_replaced=False,
            update_fields_set=False,
            placeholder_paragraphs_found=0,
            issues=issues,
        )

    try:
        doc = Document(str(src))
    except Exception as exc:
        issues.append(TocIssue(
            severity=TOC_SEVERITY["ERROR"],
            code="EN05-E001",
            message=f"No se puede abrir el DOCX: {exc}",
        ))
        return TocProcessResult(
            input_docx=str(src),
            output_docx=None,
            status=TOC_STATUS["NO_CONFORME"],
            toc_inserted=False,
            toc_replaced=False,
            update_fields_set=False,
            placeholder_paragraphs_found=0,
            issues=issues,
        )

    placeholder_indices = find_toc_placeholder_paragraphs(doc) if replace_placeholder else []
    placeholder_count = len(placeholder_indices)

    try:
        if replace_placeholder and placeholder_indices:
            target_para = doc.paragraphs[placeholder_indices[0]]
            p_elem = target_para._p
            for child in list(p_elem):
                p_elem.remove(child)
            p_elem.append(_build_toc_run_element())
            toc_replaced = True
            notes.append(f"TOC reemplazado en parrafo indice {placeholder_indices[0]}.")
        else:
            body = doc.element.body
            new_p = OxmlElement("w:p")
            new_p.append(_build_toc_run_element())
            body.insert(0, new_p)
            toc_inserted = True
            if not placeholder_indices:
                notes.append("No se encontro placeholder de TOC; TOC insertado al inicio.")
            else:
                notes.append("TOC insertado al inicio (replace_placeholder=False).")

        dst.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(dst))

    except Exception as exc:
        issues.append(TocIssue(
            severity=TOC_SEVERITY["ERROR"],
            code="EN05-E002",
            message=f"Error al guardar el DOCX con TOC: {exc}",
        ))
        return TocProcessResult(
            input_docx=str(src),
            output_docx=None,
            status=TOC_STATUS["NO_CONFORME"],
            toc_inserted=toc_inserted,
            toc_replaced=toc_replaced,
            update_fields_set=False,
            placeholder_paragraphs_found=placeholder_count,
            issues=issues,
            warnings=warnings,
            notes=notes,
        )

    try:
        ok = enable_update_fields_on_open(dst, dst)
        update_fields_set = ok
        if not ok:
            warnings.append("No se pudo habilitar updateFields en la copia.")
    except Exception as exc:
        warnings.append(f"Error al habilitar updateFields: {exc}")

    status = _compute_toc_status(issues, warnings)

    return TocProcessResult(
        input_docx=str(src),
        output_docx=str(dst),
        status=status,
        toc_inserted=toc_inserted,
        toc_replaced=toc_replaced,
        update_fields_set=update_fields_set,
        placeholder_paragraphs_found=placeholder_count,
        issues=issues,
        warnings=warnings,
        notes=notes,
    )


def _compute_toc_status(issues: List[TocIssue], warnings: List[str]) -> str:
    has_error = any(i.severity == "ERROR" for i in issues)
    has_warning = any(i.severity == "WARNING" for i in issues) or bool(warnings)
    if has_error:
        return TOC_STATUS["NO_CONFORME"]
    if has_warning:
        return TOC_STATUS["CON_OBSERVACIONES"]
    return TOC_STATUS["OK"]


# ---------------------------------------------------------------------------
# Analisis (solo lectura)
# ---------------------------------------------------------------------------


def analyze_toc(path) -> TocProcessResult:
    """
    Analiza el DOCX sin modificarlo.
    Devuelve TocProcessResult con informacion de diagnostico.
    """
    p = Path(path)
    issues: List[TocIssue] = []
    notes: List[str] = []

    if not p.exists() or not p.is_file():
        issues.append(TocIssue(
            severity=TOC_SEVERITY["ERROR"],
            code="EN05-E001",
            message=f"DOCX no encontrado: {p}",
        ))
        return TocProcessResult(
            input_docx=str(p),
            output_docx=None,
            status=TOC_STATUS["NO_CONFORME"],
            toc_inserted=False,
            toc_replaced=False,
            update_fields_set=False,
            placeholder_paragraphs_found=0,
            issues=issues,
        )

    detection = detect_toc_in_docx(p)
    issues.extend(detection.issues)

    placeholder_count = 0
    try:
        from docx import Document
        doc = Document(str(p))
        placeholder_count = len(find_toc_placeholder_paragraphs(doc))
    except Exception as exc:
        issues.append(TocIssue(
            severity=TOC_SEVERITY["ERROR"],
            code="EN05-E001",
            message=f"Error al abrir DOCX para analisis: {exc}",
        ))

    notes.extend(detection.notes)
    if placeholder_count > 0:
        notes.append(f"Se encontraron {placeholder_count} parrafos placeholder candidatos a TOC.")

    status = _compute_toc_status(issues, [])

    return TocProcessResult(
        input_docx=str(p),
        output_docx=None,
        status=status,
        toc_inserted=False,
        toc_replaced=False,
        update_fields_set=detection.update_fields_enabled,
        placeholder_paragraphs_found=placeholder_count,
        issues=issues,
        notes=notes,
    )


# ---------------------------------------------------------------------------
# Funcion principal
# ---------------------------------------------------------------------------


def _find_best_docx(expediente_path) -> Optional[Path]:
    """Selecciona el mejor DOCX disponible segun la lista de candidatos."""
    base = Path(expediente_path)
    for candidate in _DOCX_CANDIDATES:
        p = base / candidate
        if p.exists() and p.is_file() and p.stat().st_size > 0:
            return p
    return None


def process_document_toc(
    expediente_path,
    write_outputs: bool = False,
    apply_toc: bool = False,
    replace_placeholder: bool = True,
) -> TocProcessResult:
    """
    Funcion principal: analiza y opcionalmente aplica TOC al mejor DOCX disponible.

    Args:
        expediente_path: Ruta al directorio del expediente.
        write_outputs: Si True, escribe JSON y MD en documento/.
        apply_toc: Si True, genera copia del DOCX con TOC (documento_ambiental_con_toc.docx).
        replace_placeholder: Si True (defecto), reemplaza placeholder si existe.

    Returns:
        TocProcessResult con el resultado de analisis o aplicacion.
    """
    base = Path(expediente_path)
    docx_path = _find_best_docx(base)

    if docx_path is None:
        result = TocProcessResult(
            input_docx="",
            output_docx=None,
            status=TOC_STATUS["SIN_DATOS"],
            toc_inserted=False,
            toc_replaced=False,
            update_fields_set=False,
            placeholder_paragraphs_found=0,
            issues=[TocIssue(
                severity=TOC_SEVERITY["ERROR"],
                code="EN05-E001",
                message="No se encontro ningun DOCX en documento/.",
            )],
        )
    elif apply_toc:
        output_docx = base / "documento" / TOC_OUTPUT_DOCX
        result = insert_or_replace_toc(
            docx_path,
            output_docx,
            replace_placeholder=replace_placeholder,
        )
    else:
        result = analyze_toc(docx_path)

    if write_outputs and result.status != TOC_STATUS["SIN_DATOS"]:
        out_dir = base / "documento"
        try:
            write_toc_outputs(result, out_dir)
        except Exception as exc:
            result.warnings.append(f"Error al escribir outputs: {exc}")

    return result


# ---------------------------------------------------------------------------
# Informe Markdown
# ---------------------------------------------------------------------------


def build_toc_report_markdown(result: TocProcessResult) -> str:
    """Genera el informe en Markdown del resultado de procesamiento de TOC."""
    lines = [
        "# Informe de TOC — EN-05",
        "",
        "## Resumen",
        "",
        f"- **Estado**: {result.status}",
        f"- **DOCX entrada**: `{result.input_docx}`",
        f"- **DOCX salida**: `{result.output_docx or 'N/A'}`",
        f"- **TOC insertado**: {'si' if result.toc_inserted else 'no'}",
        f"- **TOC reemplazado**: {'si' if result.toc_replaced else 'no'}",
        f"- **updateFields habilitado**: {'si' if result.update_fields_set else 'no'}",
        f"- **Placeholders encontrados**: {result.placeholder_paragraphs_found}",
        f"- **Errores**: {result.error_count()}",
        f"- **Avisos**: {result.warning_count()}",
        "",
    ]

    if result.issues:
        lines += ["## Incidencias", ""]
        for issue in result.issues:
            lines.append(f"- `[{issue.severity}]` **{issue.code}**: {issue.message}")
            if issue.recommendation:
                lines.append(f"  - Recomendacion: {issue.recommendation}")
        lines.append("")

    if result.warnings:
        lines += ["## Avisos adicionales", ""]
        for w in result.warnings:
            lines.append(f"- {w}")
        lines.append("")

    if result.notes:
        lines += ["## Notas", ""]
        for n in result.notes:
            lines.append(f"- {n}")
        lines.append("")

    lines += [
        "## Advertencia",
        "",
        "python-docx inserta el campo TOC pero NO calcula numeros de pagina.",
        "Para actualizar el indice, abra el DOCX en Word o LibreOffice Writer.",
        "Si updateFields=true esta en settings.xml, Word actualiza automaticamente al abrir.",
        "",
    ]

    return "\n".join(lines)


# ---------------------------------------------------------------------------
# Escritura de outputs
# ---------------------------------------------------------------------------


def write_toc_outputs(result: TocProcessResult, output_dir) -> Tuple[Path, Path]:
    """Escribe JSON y MD del resultado en output_dir. Retorna (json_path, md_path)."""
    out = Path(output_dir)
    out.mkdir(parents=True, exist_ok=True)

    json_path = out / TOC_RESULT_JSON
    md_path = out / TOC_RESULT_MD

    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(result.to_dict(), f, ensure_ascii=False, indent=2)

    md_path.write_text(build_toc_report_markdown(result), encoding="utf-8")

    return json_path, md_path
