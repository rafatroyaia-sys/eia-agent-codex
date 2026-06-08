"""
document_figure_inserter -- DOC-03
Insercion de figuras, mapas, climogramas y anexos graficos en DOCX.

Localiza imagenes existentes en el expediente y las inserta en el DOCX
generado por DOC-02, produciendo una version enriquecida.

Principios no negociables:
  - No usa IA.
  - No consulta fuentes externas.
  - No genera mapas nuevos.
  - No genera climogramas nuevos.
  - Solo inserta imagenes ya existentes en el expediente.
  - No modifica el DOCX base (DOC-02).
  - No declara aptitud administrativa.
  - No modifica impactos, medidas, PVA ni auditorias.
  - No genera PDF.
"""
from __future__ import annotations

import json
import re
from dataclasses import dataclass, field
from pathlib import Path

# ---------------------------------------------------------------------------
# Constantes publicas
# ---------------------------------------------------------------------------

FIGURE_OUTPUT_DOCX_FILENAME = "documento_ambiental_borrador_con_figuras.docx"
FIGURE_RESULT_JSON_FILENAME = "document_figures_result.json"
FIGURE_RESULT_MD_FILENAME = "document_figures_result.md"

SUPPORTED_IMAGE_EXTENSIONS: list[str] = [".png", ".jpg", ".jpeg"]

FIGURE_TYPES: list[str] = [
    "MAPA",
    "CLIMOGRAMA",
    "FOTOGRAFIA",
    "LOGO",
    "GRAFICO",
    "OTRO",
]

# Directorios donde buscar figuras (rutas relativas al expediente)
FIGURE_SOURCE_DIRS: list[str] = [
    "mapas",
    "cartografia",
    "cartografia/mapas",
    "clima",
    "documento/figuras",
    "inputs/fotos",
    "inputs/imagenes",
    "assets/brand",
]

# Orden de prioridad para ordenar figuras en el anexo
_FIGURE_TYPE_ORDER: dict[str, int] = {
    "MAPA": 0,
    "CLIMOGRAMA": 1,
    "FOTOGRAFIA": 2,
    "GRAFICO": 3,
    "LOGO": 4,
    "OTRO": 5,
}

# Palabras clave para clasificacion de tipo
_MAPA_KEYWORDS = frozenset({
    "mapa", "map", "cartografia", "situacion", "emplazamiento",
    "parcela", "red_natura", "inundabilidad", "usos_suelo",
    "catastral", "ortofoto", "geologia", "litologia", "espacios",
    "zonificacion", "ubicacion", "ruido", "acustico", "acustica",
    "lden", "sonoro",
})
_CLIMOGRAMA_KEYWORDS = frozenset({
    "climograma", "climate", "clima", "aemet", "koppen", "martonne",
    "precipitacion", "temperatura", "gaussen",
})
_FOTOGRAFIA_KEYWORDS = frozenset({
    "foto", "fotografia", "photo", "imagen", "image", "img",
    "vista", "panoramica",
})
_LOGO_KEYWORDS = frozenset({
    "logo", "brand", "marca", "ecogestion",
})
_GRAFICO_KEYWORDS = frozenset({
    "grafico", "chart", "plot", "figura", "diagrama", "esquema",
})

_ADMIN_NOTE = (
    "Este proceso solo inserta figuras ya existentes en el expediente. "
    "No genera cartografia, no verifica validez oficial y "
    "no declara aptitud administrativa."
)


# ---------------------------------------------------------------------------
# Dataclasses
# ---------------------------------------------------------------------------


@dataclass
class DocumentFigure:
    """Figura localizada en el expediente."""

    figure_id: str
    figure_type: str
    title: str
    source_path: str
    relative_path: str
    caption: str
    section_hint: str
    file_size_bytes: int
    warnings: list[str] = field(default_factory=list)
    notes: list[str] = field(default_factory=list)

    def to_dict(self) -> dict:
        return {
            "figure_id": self.figure_id,
            "figure_type": self.figure_type,
            "title": self.title,
            "source_path": self.source_path,
            "relative_path": self.relative_path,
            "caption": self.caption,
            "section_hint": self.section_hint,
            "file_size_bytes": self.file_size_bytes,
            "warnings": list(self.warnings),
            "notes": list(self.notes),
        }

    def summary(self) -> str:
        kb = self.file_size_bytes // 1024
        return (
            f"[{self.figure_type:12}] {self.figure_id} — {self.title[:50]} "
            f"({kb} KB)"
        )


@dataclass
class FigureInsertionResult:
    """Resultado de la insercion de figuras en el DOCX."""

    expediente_id: str
    input_docx_path: str
    output_docx_path: "str | None" = None
    figures_found: list[DocumentFigure] = field(default_factory=list)
    figures_inserted: list[str] = field(default_factory=list)
    figures_skipped: list[str] = field(default_factory=list)
    generated: bool = False
    warnings: list[str] = field(default_factory=list)
    notes: list[str] = field(default_factory=list)

    def found_count(self) -> int:
        return len(self.figures_found)

    def inserted_count(self) -> int:
        return len(self.figures_inserted)

    def skipped_count(self) -> int:
        return len(self.figures_skipped)

    def warning_count(self) -> int:
        return len(self.warnings)

    def is_success(self) -> bool:
        return self.generated and self.output_docx_path is not None

    def to_dict(self) -> dict:
        return {
            "expediente_id": self.expediente_id,
            "input_docx_path": self.input_docx_path,
            "output_docx_path": self.output_docx_path,
            "generated": self.generated,
            "found_count": self.found_count(),
            "inserted_count": self.inserted_count(),
            "skipped_count": self.skipped_count(),
            "warning_count": self.warning_count(),
            "figures_found": [f.to_dict() for f in self.figures_found],
            "figures_inserted": list(self.figures_inserted),
            "figures_skipped": list(self.figures_skipped),
            "warnings": list(self.warnings),
            "notes": list(self.notes),
        }

    def summary(self) -> str:
        status = "OK" if self.is_success() else ("DRY-RUN" if not self.generated else "ERROR")
        return (
            f"DOC-03 [{self.expediente_id}] {status} — "
            f"{self.found_count()} encontradas, "
            f"{self.inserted_count()} insertadas, "
            f"{self.skipped_count()} omitidas"
            + (f", {self.warning_count()} avisos" if self.warnings else "")
        )


# ---------------------------------------------------------------------------
# Clasificacion y metadatos de figuras
# ---------------------------------------------------------------------------


def _normalize_for_match(s: str) -> str:
    """Normaliza texto para comparacion: minusculas, sin separadores."""
    return re.sub(r"[\W_]+", "", s.lower())


def _contains_keyword(text: str, keywords: frozenset) -> bool:
    """True si alguna keyword (normalizada) es substring del texto normalizado."""
    nt = _normalize_for_match(text)
    for kw in keywords:
        nkw = _normalize_for_match(kw)
        if nkw and nkw in nt:
            return True
    return False


def detect_figure_type(path: "str | Path") -> str:
    """Clasifica el tipo de figura por nombre/ruta."""
    stem = Path(path).stem
    parent = str(Path(path).parent).replace("\\", "/")
    combined = stem + "/" + parent

    if _contains_keyword(combined, _LOGO_KEYWORDS):
        return "LOGO"
    if _contains_keyword(combined, _CLIMOGRAMA_KEYWORDS):
        return "CLIMOGRAMA"
    if _contains_keyword(combined, _MAPA_KEYWORDS):
        return "MAPA"
    if _contains_keyword(combined, _FOTOGRAFIA_KEYWORDS):
        return "FOTOGRAFIA"
    if _contains_keyword(combined, _GRAFICO_KEYWORDS):
        return "GRAFICO"
    return "OTRO"


def build_figure_title(path: "str | Path", figure_type: str) -> str:
    """Crea titulo legible desde el nombre del archivo."""
    stem = Path(path).stem
    # Sustituir separadores
    title = re.sub(r"[_\-]+", " ", stem)
    # Eliminar numeros puros iniciales del tipo MAP001 -> MAP
    title = re.sub(r"^\d+\s*", "", title)
    # Capitalizar
    title = title.strip().capitalize()
    return title if title else figure_type.capitalize()


def build_figure_caption(
    figure_id: str,
    title: str,
    figure_type: str,
) -> str:
    """Genera caption estandar para una figura."""
    return (
        f"Figura {figure_id}. {title}. "
        f"Tipo: {figure_type}. "
        f"Fuente: expediente tecnico."
    )


# ---------------------------------------------------------------------------
# Validacion de imagen
# ---------------------------------------------------------------------------


def validate_image_file(path: "str | Path") -> bool:
    """Validacion basica de imagen: existe, extension soportada, tamanyo > 0."""
    try:
        p = Path(path)
        if not p.exists():
            return False
        if p.suffix.lower() not in SUPPORTED_IMAGE_EXTENSIONS:
            return False
        if p.stat().st_size == 0:
            return False
        return True
    except Exception:
        return False


# ---------------------------------------------------------------------------
# Descubrimiento de figuras
# ---------------------------------------------------------------------------


def discover_document_figures(
    expediente_path: "str | Path",
) -> list[DocumentFigure]:
    """Localiza figuras en los directorios estandar del expediente.

    Busca solo en FIGURE_SOURCE_DIRS. No modifica ningun archivo.
    Ordena por tipo: MAPA > CLIMOGRAMA > FOTOGRAFIA > GRAFICO > LOGO > OTRO.
    """
    exp = Path(expediente_path)
    found: list[tuple[int, Path]] = []  # (order, path)

    _skip_patterns = re.compile(
        r"^[\._~]|thumb|cache|tmp|temp|\btest\b", re.IGNORECASE
    )

    for rel_dir in FIGURE_SOURCE_DIRS:
        dir_path = exp / rel_dir
        if not dir_path.is_dir():
            continue
        for img_path in dir_path.iterdir():
            if not img_path.is_file():
                continue
            if _skip_patterns.search(img_path.name):
                continue
            if img_path.suffix.lower() not in SUPPORTED_IMAGE_EXTENSIONS:
                continue
            if img_path.stat().st_size == 0:
                continue
            ftype = detect_figure_type(img_path)
            order = _FIGURE_TYPE_ORDER.get(ftype, 5)
            found.append((order, img_path))

    # Eliminar duplicados (mismo path absoluto)
    seen: set[str] = set()
    unique: list[tuple[int, Path]] = []
    for order, p in found:
        key = str(p.resolve())
        if key not in seen:
            seen.add(key)
            unique.append((order, p))

    # Ordenar por tipo y luego por nombre
    unique.sort(key=lambda x: (x[0], x[1].name.lower()))

    figures: list[DocumentFigure] = []
    for idx, (_, img_path) in enumerate(unique, start=1):
        ftype = detect_figure_type(img_path)
        fig_id = f"FIG-{idx:03d}"
        title = build_figure_title(img_path, ftype)
        caption = build_figure_caption(fig_id, title, ftype)

        # relative_path desde el expediente
        try:
            rel = str(img_path.relative_to(exp)).replace("\\", "/")
        except ValueError:
            rel = str(img_path)

        # section_hint basado en directorio
        parent_rel = str(img_path.parent.relative_to(exp)).replace("\\", "/") \
            if img_path.parent != exp else ""
        if "cartografia" in parent_rel:
            hint = "Bloque K / Anexo cartografico"
        elif "clima" in parent_rel:
            hint = "Bloque B / Inventario ambiental - FI-001 Clima"
        elif "fotos" in parent_rel or "imagenes" in parent_rel:
            hint = "Bloque K / Anexo fotografico"
        elif "brand" in parent_rel or ftype == "LOGO":
            hint = "Portada"
        else:
            hint = "Bloque K / Anexo general"

        figures.append(
            DocumentFigure(
                figure_id=fig_id,
                figure_type=ftype,
                title=title,
                source_path=str(img_path),
                relative_path=rel,
                caption=caption,
                section_hint=hint,
                file_size_bytes=img_path.stat().st_size,
            )
        )

    return figures


# ---------------------------------------------------------------------------
# Insercion en DOCX
# ---------------------------------------------------------------------------


def add_figures_annex_to_docx(
    input_docx_path: "str | Path",
    output_docx_path: "str | Path",
    figures: list[DocumentFigure],
    title: str = "Anexo grafico y cartografico",
) -> FigureInsertionResult:
    """Abre el DOCX base, anade un anexo de figuras y guarda como nuevo archivo.

    No modifica input_docx_path.
    Si una imagen falla al insertarse, la omite y registra warning.
    """
    from docx import Document
    from docx.shared import Cm

    in_path = Path(input_docx_path)
    out_path = Path(output_docx_path)

    doc = Document(str(in_path))

    # Salto de pagina antes del anexo
    doc.add_page_break()

    # Encabezado del anexo
    doc.add_heading(title, level=1)

    expediente_id = in_path.parent.parent.name if in_path.parent.name == "documento" \
        else in_path.stem

    inserted: list[str] = []
    skipped: list[str] = []
    warnings: list[str] = []

    if not figures:
        doc.add_paragraph(
            "No se localizaron figuras, mapas, climogramas o fotografias "
            "para insertar en este anexo."
        )
    else:
        for fig in figures:
            fig_path = Path(fig.source_path)
            if not validate_image_file(fig_path):
                msg = f"{fig.figure_id}: imagen invalida o no accesible ({fig.relative_path})"
                skipped.append(fig.figure_id)
                warnings.append(msg)
                continue

            try:
                # Agregar caption antes de la imagen
                p_img = doc.add_paragraph()
                run = p_img.add_run()
                run.add_picture(str(fig_path), width=Cm(15))

                # Caption debajo
                cap_para = doc.add_paragraph(fig.caption)
                try:
                    cap_para.runs[0].italic = True
                    from docx.shared import Pt
                    cap_para.runs[0].font.size = Pt(9)
                except Exception:
                    pass

                doc.add_paragraph()  # espacio entre figuras
                inserted.append(fig.figure_id)

            except Exception as exc:
                msg = (
                    f"{fig.figure_id}: error al insertar {fig.relative_path}: {exc}"
                )
                skipped.append(fig.figure_id)
                warnings.append(msg)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))

    return FigureInsertionResult(
        expediente_id=expediente_id,
        input_docx_path=str(in_path),
        output_docx_path=str(out_path),
        figures_found=figures,
        figures_inserted=inserted,
        figures_skipped=skipped,
        generated=True,
        warnings=warnings,
        notes=[
            f"DOCX enriquecido generado con {len(inserted)} figura(s).",
            _ADMIN_NOTE,
        ],
    )


# ---------------------------------------------------------------------------
# Informe Markdown
# ---------------------------------------------------------------------------


def build_figure_result_markdown(result: FigureInsertionResult) -> str:
    """Genera informe Markdown del resultado de insercion."""
    lines: list[str] = []

    lines.append("# Resultado de insercion de figuras")
    lines.append("")

    # 1. Resumen
    lines.append("## 1. Resumen")
    lines.append("")
    lines.append(f"**Expediente:** {result.expediente_id}")
    lines.append(f"**DOCX base:** `{result.input_docx_path}`")
    lines.append(f"**DOCX enriquecido:** `{result.output_docx_path or 'no generado'}`")
    lines.append(f"**Generado:** {'Si' if result.generated else 'No (dry-run)'}")
    lines.append(f"**Figuras encontradas:** {result.found_count()}")
    lines.append(f"**Figuras insertadas:** {result.inserted_count()}")
    lines.append(f"**Figuras omitidas:** {result.skipped_count()}")
    lines.append(f"**Advertencias:** {result.warning_count()}")
    lines.append("")

    # 2. Figuras encontradas
    lines.append("## 2. Figuras encontradas")
    lines.append("")
    if result.figures_found:
        lines.append("| ID | Tipo | Titulo | Tamanyo | Ruta |")
        lines.append("|----|------|--------|---------|------|")
        for fig in result.figures_found:
            kb = fig.file_size_bytes // 1024
            lines.append(
                f"| {fig.figure_id} | {fig.figure_type} | {fig.title[:40]} "
                f"| {kb} KB | `{fig.relative_path}` |"
            )
    else:
        lines.append("_No se localizaron figuras en los directorios de busqueda._")
    lines.append("")

    # 3. Figuras insertadas
    lines.append("## 3. Figuras insertadas")
    lines.append("")
    if result.figures_inserted:
        for fid in result.figures_inserted:
            lines.append(f"- {fid}")
    else:
        lines.append("_Ninguna figura insertada._")
    lines.append("")

    # 4. Figuras omitidas
    lines.append("## 4. Figuras omitidas")
    lines.append("")
    if result.figures_skipped:
        for fid in result.figures_skipped:
            lines.append(f"- {fid}")
    else:
        lines.append("_Sin figuras omitidas._")
    lines.append("")

    # 5. Advertencias
    lines.append("## 5. Advertencias")
    lines.append("")
    if result.warnings:
        for w in result.warnings:
            lines.append(f"- {w}")
    else:
        lines.append("_Sin advertencias._")
    lines.append("")

    # 6. Nota de alcance
    lines.append("## 6. Advertencia de alcance")
    lines.append("")
    lines.append(f"> {_ADMIN_NOTE}")
    lines.append("")

    return "\n".join(lines)


# ---------------------------------------------------------------------------
# Escritura de outputs
# ---------------------------------------------------------------------------


def write_figure_insertion_outputs(
    result: FigureInsertionResult,
    output_dir: "str | Path",
) -> tuple[Path, Path]:
    """Escribe JSON y Markdown del resultado en output_dir."""
    out = Path(output_dir)
    out.mkdir(parents=True, exist_ok=True)

    json_path = out / FIGURE_RESULT_JSON_FILENAME
    md_path = out / FIGURE_RESULT_MD_FILENAME

    json_path.write_text(
        json.dumps(result.to_dict(), ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    md_path.write_text(
        build_figure_result_markdown(result),
        encoding="utf-8",
    )
    return json_path, md_path


# ---------------------------------------------------------------------------
# Funcion principal
# ---------------------------------------------------------------------------


def insert_figures_into_document(
    expediente_path: "str | Path",
    write_outputs: bool = False,
) -> FigureInsertionResult:
    """Localiza figuras del expediente e inserta en el DOCX DOC-02.

    Sin write_outputs=True: descubre figuras y devuelve resultado con generated=False.
    Con write_outputs=True: genera DOCX enriquecido + JSON + MD.
    No modifica el DOCX original de DOC-02.
    No genera mapas ni climogramas nuevos.
    """
    exp = Path(expediente_path)
    doc_dir = exp / "documento"
    base_docx = doc_dir / "documento_ambiental_borrador.docx"

    if not base_docx.exists():
        raise FileNotFoundError(
            f"No se encontro el DOCX DOC-02 en: {base_docx}. "
            "Ejecute primero 'document-build-docx --write'."
        )

    # Descubrir figuras
    all_figures = discover_document_figures(exp)

    # Filtrar logos para el anexo:
    # incluir logos solo si estan en documento/figuras/ o si no hay otras figuras
    non_logo = [f for f in all_figures if f.figure_type != "LOGO"]
    logo_in_doc_figuras = [
        f for f in all_figures
        if f.figure_type == "LOGO" and "documento/figuras" in f.relative_path
    ]
    logo_skipped = [
        f for f in all_figures
        if f.figure_type == "LOGO" and f not in logo_in_doc_figuras
    ]

    if non_logo or logo_in_doc_figuras:
        figures_for_annex = non_logo + logo_in_doc_figuras
        logo_skip_note = (
            f"{len(logo_skipped)} logo(s) de assets/brand excluidos del anexo "
            "(ya aparecen en portada)."
        ) if logo_skipped else ""
    else:
        # Solo hay logos o no hay nada: incluir todo
        figures_for_annex = all_figures
        logo_skip_note = ""

    global_warnings: list[str] = []
    global_notes: list[str] = [
        "Logos de assets/brand excluidos del anexo salvo si estan en documento/figuras/.",
        _ADMIN_NOTE,
    ]
    if logo_skip_note:
        global_warnings.append(logo_skip_note)
    if not all_figures:
        global_warnings.append(
            "No se localizaron figuras en los directorios de busqueda."
        )

    if not write_outputs:
        return FigureInsertionResult(
            expediente_id=exp.name,
            input_docx_path=str(base_docx),
            output_docx_path=None,
            figures_found=all_figures,
            figures_inserted=[],
            figures_skipped=[],
            generated=False,
            warnings=global_warnings,
            notes=global_notes + ["write_outputs=False: DOCX enriquecido no generado."],
        )

    # Generar outputs
    out_docx = doc_dir / FIGURE_OUTPUT_DOCX_FILENAME

    result = add_figures_annex_to_docx(
        input_docx_path=base_docx,
        output_docx_path=out_docx,
        figures=figures_for_annex,
    )

    # Enriquecer result con todas las figuras encontradas
    result.figures_found = all_figures
    result.warnings = list(result.warnings) + global_warnings
    result.notes = list(result.notes) + [n for n in global_notes if n not in result.notes]

    # Escribir JSON y MD
    write_figure_insertion_outputs(result, doc_dir)

    return result
