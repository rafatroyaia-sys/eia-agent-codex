"""
document_numbering_manager.py — EN-04
Analisis y aplicacion de estilos de numeracion en el DOCX final.

Detecta si un DOCX tiene definiciones de numeracion (numbering.xml),
analiza parrafos candidatos a lista (numerada o viñetas) y puede aplicar
estilos de lista Word en una copia conservadora del DOCX.

Principios no negociables:
  - No modifica el DOCX original.
  - No reordena parrafos.
  - No cambia contenido textual.
  - No toca imagenes ni tablas.
  - No usa IA.
  - No llama a servicios externos.
  - 100 % offline.
"""
from __future__ import annotations

import json
import re
import shutil
import zipfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Optional, Tuple

# ---------------------------------------------------------------------------
# Constantes publicas
# ---------------------------------------------------------------------------

NUMBERING_STATUS = {
    "OK": "OK",
    "CON_OBSERVACIONES": "CON_OBSERVACIONES",
    "NO_CONFORME": "NO_CONFORME",
    "SIN_DATOS": "SIN_DATOS",
}

NUMBERING_SEVERITY = {
    "ERROR": "ERROR",
    "WARNING": "WARNING",
    "INFO": "INFO",
}

DEFAULT_NUMBERED_STYLE_CANDIDATES: List[str] = [
    "List Number",
    "List Number 2",
    "List Number 3",
]

DEFAULT_BULLET_STYLE_CANDIDATES: List[str] = [
    "List Bullet",
    "List Bullet 2",
    "List Bullet 3",
]

NUMBERING_OUTPUT_DOCX = "documento_ambiental_numerado.docx"
NUMBERING_RESULT_JSON = "document_numbering_result.json"
NUMBERING_RESULT_MD = "document_numbering_result.md"

_DOCX_CANDIDATES = [
    "documento/documento_ambiental_final_revisable.docx",
    "documento/documento_ambiental_estructurado.docx",
    "documento/documento_ambiental_borrador_con_figuras.docx",
    "documento/documento_ambiental_borrador.docx",
]

# Patrones para listas numeradas (se aplican sobre texto ya strip()ado)
_NUMBERED_PATTERNS = [
    re.compile(r"^\d+\.\s+\S"),           # "1. texto"
    re.compile(r"^\d+\)\s+\S"),           # "1) texto"
    re.compile(r"^\d+\.-\s+\S"),          # "1.- texto"
    re.compile(r"^[a-zA-Z]\)\s+\S"),      # "a) texto", "A) texto", "i) texto"
    re.compile(r"^[a-zA-Z]\.\s+\S"),      # "a. texto", "A. texto", "I. texto"
]

# Patrones para viñetas (se aplican sobre texto ya strip()ado)
_BULLET_PATTERNS = [
    re.compile(r"^-\s+\S"),    # "- texto"
    re.compile(r"^\*\s+\S"),   # "* texto"
    re.compile(r"^•\s+\S"),    # "• texto" (bullet unicode)
    re.compile(r"^–\s+\S"),    # "– texto" (en dash)
    re.compile(r"^—\s+\S"),    # "— texto" (em dash al inicio)
]


# ---------------------------------------------------------------------------
# Dataclasses
# ---------------------------------------------------------------------------


@dataclass
class NumberingIssue:
    """Incidencia detectada durante el analisis o aplicacion de numeracion."""

    severity: str
    code: str
    paragraph_index: "Optional[int]"
    message: str
    recommendation: str = ""
    evidence: List[str] = field(default_factory=list)

    def to_dict(self) -> dict:
        return {
            "severity": self.severity,
            "code": self.code,
            "paragraph_index": self.paragraph_index,
            "message": self.message,
            "recommendation": self.recommendation,
            "evidence": list(self.evidence),
        }

    def summary(self) -> str:
        loc = f" (parrafo {self.paragraph_index})" if self.paragraph_index is not None else ""
        return f"[{self.severity}][{self.code}]{loc} {self.message}"


@dataclass
class ParagraphNumberingStatus:
    """Estado de numeracion de un parrafo especifico."""

    paragraph_index: int
    text_preview: str
    style_name: "Optional[str]"
    detected_as_list_candidate: bool
    applied_style: "Optional[str]"
    warnings: List[str] = field(default_factory=list)
    notes: List[str] = field(default_factory=list)

    def to_dict(self) -> dict:
        return {
            "paragraph_index": self.paragraph_index,
            "text_preview": self.text_preview,
            "style_name": self.style_name,
            "detected_as_list_candidate": self.detected_as_list_candidate,
            "applied_style": self.applied_style,
            "warnings": list(self.warnings),
            "notes": list(self.notes),
        }

    def summary(self) -> str:
        candidate = "CANDIDATO" if self.detected_as_list_candidate else "normal"
        applied = f" -> {self.applied_style}" if self.applied_style else ""
        return f"[{self.paragraph_index}] {candidate}{applied}: {self.text_preview[:60]}"


@dataclass
class NumberingResult:
    """Resultado del analisis/aplicacion de numeracion DOCX."""

    input_docx: str
    output_docx: "Optional[str]"
    status: str
    numbering_definitions_found: bool
    paragraphs_checked: int
    list_candidates_found: int
    styles_applied_count: int
    paragraph_statuses: List[ParagraphNumberingStatus] = field(default_factory=list)
    issues: List[NumberingIssue] = field(default_factory=list)
    warnings: List[str] = field(default_factory=list)
    notes: List[str] = field(default_factory=list)

    def error_count(self) -> int:
        return sum(1 for i in self.issues if i.severity == NUMBERING_SEVERITY["ERROR"])

    def warning_count(self) -> int:
        return sum(1 for i in self.issues if i.severity == NUMBERING_SEVERITY["WARNING"])

    def is_valid(self) -> bool:
        return self.error_count() == 0

    def to_dict(self) -> dict:
        return {
            "input_docx": self.input_docx,
            "output_docx": self.output_docx,
            "status": self.status,
            "numbering_definitions_found": self.numbering_definitions_found,
            "paragraphs_checked": self.paragraphs_checked,
            "list_candidates_found": self.list_candidates_found,
            "styles_applied_count": self.styles_applied_count,
            "error_count": self.error_count(),
            "warning_count": self.warning_count(),
            "is_valid": self.is_valid(),
            "paragraph_statuses": [p.to_dict() for p in self.paragraph_statuses],
            "issues": [i.to_dict() for i in self.issues],
            "warnings": list(self.warnings),
            "notes": list(self.notes),
        }

    def summary(self) -> str:
        lines = [
            f"Resultado numeracion : {self.status}",
            f"DOCX entrada         : {self.input_docx}",
            f"Numbering.xml        : {'SI' if self.numbering_definitions_found else 'NO'}",
            f"Parrafos revisados   : {self.paragraphs_checked}",
            f"Candidatos lista     : {self.list_candidates_found}",
            f"Estilos aplicados    : {self.styles_applied_count}",
            f"ERRORs               : {self.error_count()}",
            f"WARNINGs             : {self.warning_count()}",
        ]
        if self.output_docx:
            lines.insert(2, f"DOCX salida          : {self.output_docx}")
        return "\n".join(lines)


# ---------------------------------------------------------------------------
# Funciones auxiliares
# ---------------------------------------------------------------------------


def validate_docx_file(path: "str | Path") -> bool:
    """Comprueba que el archivo existe, no esta vacio y abre como DOCX."""
    p = Path(path)
    if not p.exists() or not p.is_file():
        return False
    if p.stat().st_size == 0:
        return False
    try:
        import docx
        docx.Document(str(p))
        return True
    except Exception:
        return False


def docx_has_numbering_definitions(path: "str | Path") -> bool:
    """
    Inspecciona el DOCX como ZIP y detecta si contiene word/numbering.xml.
    No modifica el archivo.
    """
    p = Path(path)
    if not p.exists() or not p.is_file():
        return False
    try:
        with zipfile.ZipFile(str(p), "r") as zf:
            names = zf.namelist()
        return "word/numbering.xml" in names
    except Exception:
        return False


def is_numbered_list_candidate(text: str) -> bool:
    """
    Devuelve True si el texto (strip()ado) parece el inicio de un elemento
    de lista numerada.

    Detecta: "1. texto", "1) texto", "1.- texto", "a) texto", "A. texto",
             "I. texto", "i) texto".

    No detecta: numeros decimales, fechas, importes, refs catastrales, coordenadas.
    La clave es el separador (., ), .-) seguido de espacio: "1.5" no tiene
    espacio tras el punto; "01/01/2026" usa /.
    """
    t = text.strip()
    if not t:
        return False
    for pat in _NUMBERED_PATTERNS:
        if pat.match(t):
            return True
    return False


def is_bullet_list_candidate(text: str) -> bool:
    """
    Devuelve True si el texto (strip()ado) parece el inicio de un elemento
    de lista con viñeta.

    Detecta: "- texto", "• texto", "* texto", "– texto", "— texto".
    No detecta guiones internos de frases ("texto-con-guion").
    La clave es que el caracter de lista este al inicio seguido de espacio.
    """
    t = text.strip()
    if not t:
        return False
    for pat in _BULLET_PATTERNS:
        if pat.match(t):
            return True
    return False


def select_numbered_style(level: int = 1) -> str:
    """
    Devuelve el nombre de estilo Word para lista numerada segun el nivel.
    Nivel 1 → List Number, 2 → List Number 2, 3 → List Number 3.
    Nivel fuera de [1,3] se redondea al extremo mas cercano.
    """
    clamped = max(1, min(3, level))
    return DEFAULT_NUMBERED_STYLE_CANDIDATES[clamped - 1]


def select_bullet_style(level: int = 1) -> str:
    """
    Devuelve el nombre de estilo Word para lista de viñetas segun el nivel.
    Nivel 1 → List Bullet, 2 → List Bullet 2, 3 → List Bullet 3.
    Nivel fuera de [1,3] se redondea al extremo mas cercano.
    """
    clamped = max(1, min(3, level))
    return DEFAULT_BULLET_STYLE_CANDIDATES[clamped - 1]


def _available_styles(doc) -> set:
    """Devuelve el conjunto de nombres de estilos disponibles en el documento."""
    try:
        return {s.name for s in doc.styles}
    except Exception:
        return set()


# ---------------------------------------------------------------------------
# Funciones principales
# ---------------------------------------------------------------------------


def analyze_docx_numbering(path: "str | Path") -> NumberingResult:
    """
    Analiza la numeracion de un DOCX sin modificar ningun archivo.

    Detecta numbering.xml, cuenta parrafos y detecta candidatos.
    No escribe outputs.
    """
    p = Path(path)
    base = str(p)

    if not p.exists():
        return NumberingResult(
            input_docx=base,
            output_docx=None,
            status=NUMBERING_STATUS["SIN_DATOS"],
            numbering_definitions_found=False,
            paragraphs_checked=0,
            list_candidates_found=0,
            styles_applied_count=0,
            notes=["DOCX no encontrado."],
        )

    if not validate_docx_file(p):
        return NumberingResult(
            input_docx=base,
            output_docx=None,
            status=NUMBERING_STATUS["NO_CONFORME"],
            numbering_definitions_found=False,
            paragraphs_checked=0,
            list_candidates_found=0,
            styles_applied_count=0,
            issues=[NumberingIssue(
                severity=NUMBERING_SEVERITY["ERROR"],
                code="EN04-E001",
                paragraph_index=None,
                message="El DOCX no puede abrirse o esta corrupto.",
                recommendation="Regenerar el DOCX con DOC-02.",
            )],
        )

    has_numbering = docx_has_numbering_definitions(p)

    try:
        import docx as _docx
        doc = _docx.Document(str(p))
    except Exception as exc:
        return NumberingResult(
            input_docx=base,
            output_docx=None,
            status=NUMBERING_STATUS["NO_CONFORME"],
            numbering_definitions_found=has_numbering,
            paragraphs_checked=0,
            list_candidates_found=0,
            styles_applied_count=0,
            issues=[NumberingIssue(
                severity=NUMBERING_SEVERITY["ERROR"],
                code="EN04-E001",
                paragraph_index=None,
                message=f"Error al abrir DOCX: {exc}",
                recommendation="Verificar integridad del archivo.",
            )],
        )

    para_statuses: List[ParagraphNumberingStatus] = []
    issues: List[NumberingIssue] = []
    warnings_list: List[str] = []
    notes_list: List[str] = []
    candidates = 0

    for idx, para in enumerate(doc.paragraphs):
        text = para.text
        style_name = para.style.name if para.style else None
        is_candidate = is_numbered_list_candidate(text) or is_bullet_list_candidate(text)

        if is_candidate:
            candidates += 1

        para_statuses.append(ParagraphNumberingStatus(
            paragraph_index=idx,
            text_preview=text[:80],
            style_name=style_name,
            detected_as_list_candidate=is_candidate,
            applied_style=None,
        ))

    if not has_numbering:
        notes_list.append(
            "El DOCX no contiene word/numbering.xml. "
            "Los estilos de lista Word pueden crearse al aplicar --apply."
        )

    if candidates == 0:
        notes_list.append("No se detectaron parrafos candidatos a lista.")

    status = _compute_status(issues, warnings_list)

    return NumberingResult(
        input_docx=base,
        output_docx=None,
        status=status,
        numbering_definitions_found=has_numbering,
        paragraphs_checked=len(doc.paragraphs),
        list_candidates_found=candidates,
        styles_applied_count=0,
        paragraph_statuses=para_statuses,
        issues=issues,
        warnings=warnings_list,
        notes=notes_list,
    )


def apply_list_styles_to_docx(
    input_docx: "str | Path",
    output_docx: "str | Path",
    apply_numbered: bool = True,
    apply_bullets: bool = True,
) -> NumberingResult:
    """
    Crea una copia del DOCX con estilos de lista aplicados a parrafos candidatos.

    No modifica el DOCX original. Si un estilo no existe en el documento,
    registra un WARNING y continua sin romper. No reordena parrafos, no
    cambia texto, no toca imagenes ni tablas.
    """
    in_path = Path(input_docx)
    out_path = Path(output_docx)
    base = str(in_path)

    if not in_path.exists():
        return NumberingResult(
            input_docx=base,
            output_docx=str(out_path),
            status=NUMBERING_STATUS["SIN_DATOS"],
            numbering_definitions_found=False,
            paragraphs_checked=0,
            list_candidates_found=0,
            styles_applied_count=0,
            notes=["DOCX de entrada no encontrado."],
        )

    if not validate_docx_file(in_path):
        return NumberingResult(
            input_docx=base,
            output_docx=str(out_path),
            status=NUMBERING_STATUS["NO_CONFORME"],
            numbering_definitions_found=False,
            paragraphs_checked=0,
            list_candidates_found=0,
            styles_applied_count=0,
            issues=[NumberingIssue(
                severity=NUMBERING_SEVERITY["ERROR"],
                code="EN04-E001",
                paragraph_index=None,
                message="El DOCX de entrada no puede abrirse.",
                recommendation="Regenerar con DOC-02.",
            )],
        )

    has_numbering = docx_has_numbering_definitions(in_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)

    # Copiar primero; trabajamos sobre la copia
    shutil.copy2(str(in_path), str(out_path))

    try:
        import docx as _docx
        doc = _docx.Document(str(out_path))
    except Exception as exc:
        return NumberingResult(
            input_docx=base,
            output_docx=str(out_path),
            status=NUMBERING_STATUS["NO_CONFORME"],
            numbering_definitions_found=has_numbering,
            paragraphs_checked=0,
            list_candidates_found=0,
            styles_applied_count=0,
            issues=[NumberingIssue(
                severity=NUMBERING_SEVERITY["ERROR"],
                code="EN04-E001",
                paragraph_index=None,
                message=f"Error al abrir copia DOCX: {exc}",
                recommendation="Verificar integridad del archivo.",
            )],
        )

    available = _available_styles(doc)
    para_statuses: List[ParagraphNumberingStatus] = []
    issues: List[NumberingIssue] = []
    warnings_list: List[str] = []
    notes_list: List[str] = []
    candidates = 0
    applied_count = 0

    for idx, para in enumerate(doc.paragraphs):
        text = para.text
        style_name = para.style.name if para.style else None
        is_numbered = apply_numbered and is_numbered_list_candidate(text)
        is_bullet = apply_bullets and is_bullet_list_candidate(text)
        is_candidate = is_numbered or is_bullet

        applied_style = None
        para_warnings: List[str] = []
        para_notes: List[str] = []

        if is_candidate:
            candidates += 1
            target_style = select_numbered_style(1) if is_numbered else select_bullet_style(1)

            if target_style in available:
                try:
                    para.style = doc.styles[target_style]
                    applied_style = target_style
                    applied_count += 1
                except Exception as exc:
                    msg = f"No se pudo aplicar estilo '{target_style}': {exc}"
                    para_warnings.append(msg)
                    issues.append(NumberingIssue(
                        severity=NUMBERING_SEVERITY["WARNING"],
                        code="EN04-W001",
                        paragraph_index=idx,
                        message=msg,
                        recommendation="Verificar compatibilidad del estilo en el DOCX.",
                        evidence=[text[:60]],
                    ))
            else:
                msg = f"Estilo '{target_style}' no disponible en este DOCX."
                para_warnings.append(msg)
                issues.append(NumberingIssue(
                    severity=NUMBERING_SEVERITY["WARNING"],
                    code="EN04-W002",
                    paragraph_index=idx,
                    message=msg,
                    recommendation="El DOCX no tiene estilos de lista. Aplica estilos manualmente en Word.",
                    evidence=[text[:60]],
                ))

        para_statuses.append(ParagraphNumberingStatus(
            paragraph_index=idx,
            text_preview=text[:80],
            style_name=style_name,
            detected_as_list_candidate=is_candidate,
            applied_style=applied_style,
            warnings=para_warnings,
            notes=para_notes,
        ))

    if applied_count > 0:
        try:
            doc.save(str(out_path))
        except Exception as exc:
            issues.append(NumberingIssue(
                severity=NUMBERING_SEVERITY["ERROR"],
                code="EN04-E002",
                paragraph_index=None,
                message=f"Error al guardar DOCX numerado: {exc}",
                recommendation="Verificar permisos de escritura.",
            ))

    notes_list.append(
        "Este proceso no modifica el contenido tecnico del documento. "
        "Solo analiza o aplica estilos de numeracion en una copia."
    )

    status = _compute_status(issues, warnings_list)

    return NumberingResult(
        input_docx=base,
        output_docx=str(out_path),
        status=status,
        numbering_definitions_found=has_numbering,
        paragraphs_checked=len(para_statuses),
        list_candidates_found=candidates,
        styles_applied_count=applied_count,
        paragraph_statuses=para_statuses,
        issues=issues,
        warnings=warnings_list,
        notes=notes_list,
    )


def _compute_status(
    issues: List[NumberingIssue],
    warnings_list: List[str],
) -> str:
    """Calcula el status a partir de las incidencias."""
    has_errors = any(i.severity == NUMBERING_SEVERITY["ERROR"] for i in issues)
    has_warnings = (
        any(i.severity == NUMBERING_SEVERITY["WARNING"] for i in issues)
        or bool(warnings_list)
    )
    if has_errors:
        return NUMBERING_STATUS["NO_CONFORME"]
    if has_warnings:
        return NUMBERING_STATUS["CON_OBSERVACIONES"]
    return NUMBERING_STATUS["OK"]


# ---------------------------------------------------------------------------
# Markdown y outputs
# ---------------------------------------------------------------------------


def build_numbering_report_markdown(result: NumberingResult) -> str:
    """Genera el informe Markdown del resultado de numeracion."""
    lines: List[str] = []

    lines.append("# Resultado de numeracion DOCX")
    lines.append("")

    # 1. Resumen
    lines.append("## 1. Resumen")
    lines.append("")
    lines.append(f"- **Estado**: {result.status}")
    lines.append(f"- **DOCX entrada**: `{result.input_docx}`")
    if result.output_docx:
        lines.append(f"- **DOCX salida**: `{result.output_docx}`")
    lines.append(f"- **Parrafos revisados**: {result.paragraphs_checked}")
    lines.append(f"- **Candidatos a lista**: {result.list_candidates_found}")
    lines.append(f"- **Estilos aplicados**: {result.styles_applied_count}")
    lines.append(f"- **Errores**: {result.error_count()}")
    lines.append(f"- **Advertencias**: {result.warning_count()}")
    lines.append("")

    # 2. Definiciones de numeracion
    lines.append("## 2. Definiciones de numeracion")
    lines.append("")
    if result.numbering_definitions_found:
        lines.append("El DOCX contiene `word/numbering.xml` con definiciones de numeracion.")
    else:
        lines.append(
            "El DOCX **no** contiene `word/numbering.xml`. "
            "Los estilos de lista Word se crean al aplicar `--apply` si estan disponibles."
        )
    lines.append("")

    # 3. Parrafos candidatos
    lines.append("## 3. Parrafos candidatos")
    lines.append("")
    candidates = [p for p in result.paragraph_statuses if p.detected_as_list_candidate]
    if candidates:
        lines.append(f"Se detectaron {len(candidates)} parrafos candidatos a lista:")
        lines.append("")
        lines.append("| # | Preview | Estilo original | Estilo aplicado |")
        lines.append("|---|---------|----------------|-----------------|")
        for ps in candidates[:30]:
            preview = ps.text_preview[:50].replace("|", "\\|")
            orig = ps.style_name or "—"
            applied = ps.applied_style or "—"
            lines.append(f"| {ps.paragraph_index} | {preview} | {orig} | {applied} |")
        if len(candidates) > 30:
            lines.append(f"| ... | ({len(candidates) - 30} mas) | | |")
    else:
        lines.append("No se detectaron parrafos candidatos a lista numerada o viñetas.")
    lines.append("")

    # 4. Estilos aplicados
    lines.append("## 4. Estilos aplicados")
    lines.append("")
    applied = [p for p in result.paragraph_statuses if p.applied_style]
    if applied:
        style_counts: dict = {}
        for p in applied:
            style_counts[p.applied_style] = style_counts.get(p.applied_style, 0) + 1
        for style, count in sorted(style_counts.items()):
            lines.append(f"- **{style}**: {count} parrafo(s)")
    else:
        lines.append("No se aplicaron estilos en esta ejecucion.")
    lines.append("")

    # 5. Incidencias
    lines.append("## 5. Incidencias")
    lines.append("")
    if result.issues:
        for issue in result.issues:
            lines.append(f"- `[{issue.severity}]` **{issue.code}**: {issue.message}")
            if issue.recommendation:
                lines.append(f"  - *Recomendacion*: {issue.recommendation}")
    else:
        lines.append("Sin incidencias.")
    lines.append("")

    # 6. Advertencia de alcance
    lines.append("## 6. Advertencia de alcance")
    lines.append("")
    lines.append(
        "Este proceso no modifica el contenido tecnico del documento. "
        "Solo analiza o aplica estilos de numeracion en una copia."
    )
    lines.append("")
    lines.append(
        "> **Nota**: La referencia huerfana a `numbering.xml` (OBS-004 del piloto Nave 222) "
        "queda documentada en este modulo. La aplicacion de estilos de lista genera "
        "`word/numbering.xml` de forma controlada en la copia de salida."
    )
    lines.append("")

    return "\n".join(lines)


def write_numbering_outputs(
    result: NumberingResult,
    output_dir: "str | Path",
) -> Tuple[Path, Path]:
    """
    Escribe JSON y MD del resultado en output_dir.
    Devuelve (json_path, md_path).
    """
    out = Path(output_dir)
    out.mkdir(parents=True, exist_ok=True)

    json_path = out / NUMBERING_RESULT_JSON
    md_path = out / NUMBERING_RESULT_MD

    json_path.write_text(
        json.dumps(result.to_dict(), ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    md_path.write_text(
        build_numbering_report_markdown(result),
        encoding="utf-8",
    )

    return json_path, md_path


def _find_best_docx(expediente_path: Path) -> "Optional[Path]":
    """Localiza el mejor DOCX disponible segun orden de preferencia."""
    for candidate in _DOCX_CANDIDATES:
        p = expediente_path / candidate
        if p.exists() and p.stat().st_size > 0:
            return p
    return None


def process_document_numbering(
    expediente_path: "str | Path",
    write_outputs: bool = False,
    apply_styles: bool = False,
) -> NumberingResult:
    """
    Funcion principal: localiza el mejor DOCX, analiza/aplica numeracion
    y opcionalmente escribe outputs.

    Si apply_styles=False: solo analiza (sin crear DOCX numerado).
    Si apply_styles=True: crea documento/documento_ambiental_numerado.docx.
    Si write_outputs=True: escribe JSON/MD en documento/.
    No modifica el DOCX original en ningun caso.
    """
    exp = Path(expediente_path)
    docx_path = _find_best_docx(exp)

    if docx_path is None:
        result = NumberingResult(
            input_docx=str(exp / "documento" / "documento_ambiental_borrador.docx"),
            output_docx=None,
            status=NUMBERING_STATUS["SIN_DATOS"],
            numbering_definitions_found=False,
            paragraphs_checked=0,
            list_candidates_found=0,
            styles_applied_count=0,
            notes=["No se encontro ningun DOCX en documento/."],
        )
        if write_outputs:
            out_dir = exp / "documento"
            write_numbering_outputs(result, out_dir)
        return result

    if apply_styles:
        out_docx = exp / "documento" / NUMBERING_OUTPUT_DOCX
        result = apply_list_styles_to_docx(docx_path, out_docx)
    else:
        result = analyze_docx_numbering(docx_path)

    if write_outputs:
        out_dir = exp / "documento"
        write_numbering_outputs(result, out_dir)

    return result
